"""
Xodimlar — ro'yxat, qo'shish, tahrir, bo'shatish, ishga qabul, davomat, avanslar, ish haqi.
"""
from datetime import datetime, date, timedelta
from typing import Optional, List
from urllib.parse import quote
import io
import uuid

import openpyxl
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from fastapi import APIRouter, Request, Depends, Form, HTTPException, File, UploadFile, Query
from fastapi.responses import HTMLResponse, RedirectResponse, StreamingResponse
from sqlalchemy.orm import Session, joinedload
from sqlalchemy import func, or_, and_

from app.core import templates
from app.models.database import (
    get_db, User, Employee, Department, Position, PieceworkTask,
    Attendance, AttendanceDoc, EmployeeAdvance, EmploymentDoc, DismissalDoc,
    Salary, employee_piecework_tasks,
    ExpenseType, ExpenseDoc, ExpenseDocItem, CashRegister,
    Warehouse, Product, Unit,
)
from app.deps import require_auth, require_admin
from app.utils.production_order import is_qiyom_recipe, recipe_kg_per_unit

router = APIRouter(prefix="/employees", tags=["employees"])


@router.get("", response_class=HTMLResponse)
async def employees_list(
    request: Request,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
    show_dismissed: bool = False,
):
    """Xodimlar ro'yxati — odatiy holda faqat faol xodimlar."""
    q = db.query(Employee).order_by(Employee.full_name)
    if not show_dismissed:
        q = q.filter(Employee.is_active == True)
    employees = q.all()
    piecework_tasks = db.query(PieceworkTask).filter(PieceworkTask.is_active == True).order_by(PieceworkTask.name).all()
    departments = db.query(Department).filter(Department.is_active == True).order_by(Department.name).all()
    positions = db.query(Position).filter(Position.is_active == True).order_by(Position.name).all()
    return templates.TemplateResponse("employees/list.html", {
        "request": request,
        "employees": employees,
        "piecework_tasks": piecework_tasks,
        "departments": departments,
        "positions": positions,
        "current_user": current_user,
        "page_title": "Xodimlar",
        "show_dismissed": show_dismissed,
    })


@router.post("/add")
async def employee_add(
    request: Request,
    full_name: str = Form(...),
    code: str = Form(""),
    position: str = Form(""),
    department: str = Form(""),
    phone: str = Form(""),
    salary: float = Form(0),
    salary_type: str = Form(""),
    piecework_task_ids: List[int] = Form([]),
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    """Xodim qo'shish. Kod bo'sh qolsa saqlashda avtomatik yaratiladi (EMP-<id>)."""
    import uuid
    code_val = (code or "").strip()
    if not code_val:
        code_val = f"_auto_{uuid.uuid4().hex[:12]}"
    st = (salary_type or "").strip() or None
    if st and st not in ("oylik", "soatlik", "bo'lak", "bo'lak_oylik"):
        st = None
    task_ids = [int(x) for x in (piecework_task_ids or []) if str(x).strip().isdigit()]
    task_ids = list(dict.fromkeys(task_ids))
    employee = Employee(
        full_name=full_name,
        code=code_val,
        position=position,
        department=department,
        phone=phone,
        salary=salary,
        salary_type=st,
        piecework_task_id=task_ids[0] if task_ids else None,  # legacy
    )
    db.add(employee)
    db.flush()
    if not (code or "").strip():
        employee.code = f"EMP-{employee.id}"
    if st in ("bo'lak", "bo'lak_oylik") and task_ids:
        tasks = db.query(PieceworkTask).filter(PieceworkTask.id.in_(task_ids)).all()
        employee.piecework_tasks = tasks
    db.commit()
    return RedirectResponse(url="", status_code=303)


@router.get("/edit/{employee_id}", response_class=HTMLResponse)
async def employee_edit_page(
    request: Request,
    employee_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    """Xodim tahrirlash sahifasi"""
    emp = db.query(Employee).filter(Employee.id == employee_id).first()
    if not emp:
        return RedirectResponse(url="/employees?error=Xodim topilmadi", status_code=303)
    piecework_tasks = db.query(PieceworkTask).filter(PieceworkTask.is_active == True).order_by(PieceworkTask.name).all()
    return templates.TemplateResponse("employees/edit.html", {
        "request": request,
        "emp": emp,
        "piecework_tasks": piecework_tasks,
        "current_user": current_user,
        "page_title": "Xodimni tahrirlash"
    })


@router.post("/update/{employee_id}")
async def employee_update(
    employee_id: int,
    full_name: str = Form(...),
    code: str = Form(""),
    position: str = Form(""),
    department: str = Form(""),
    phone: str = Form(""),
    salary: float = Form(0),
    salary_type: str = Form(""),
    piecework_task_ids: List[int] = Form([]),
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    """Xodim ma'lumotlarini yangilash. Kod bo'sh qolsa avtomatik EMP-<id> qo'yiladi."""
    from urllib.parse import quote
    emp = db.query(Employee).filter(Employee.id == employee_id).first()
    if not emp:
        return RedirectResponse(url="/employees?error=Xodim topilmadi", status_code=303)
    code_val = (code or "").strip() or f"EMP-{employee_id}"
    duplicate = db.query(Employee).filter(Employee.code == code_val, Employee.id != employee_id).first()
    if duplicate:
        return RedirectResponse(url="/employees?error=" + quote("Bunday kod boshqa xodimda mavjud: " + code_val), status_code=303)
    emp.full_name = full_name
    emp.code = code_val
    emp.position = position
    emp.department = department
    emp.phone = phone
    emp.salary = salary
    st = (salary_type or "").strip() or None
    if st and st not in ("oylik", "soatlik", "bo'lak", "bo'lak_oylik"):
        st = None
    emp.salary_type = st
    task_ids = [int(x) for x in (piecework_task_ids or []) if str(x).strip().isdigit()]
    task_ids = list(dict.fromkeys(task_ids))
    emp.piecework_task_id = task_ids[0] if task_ids else None  # legacy
    if st in ("bo'lak", "bo'lak_oylik") and task_ids:
        tasks = db.query(PieceworkTask).filter(PieceworkTask.id.in_(task_ids)).all()
        emp.piecework_tasks = tasks
    else:
        emp.piecework_tasks = []
    db.commit()
    return RedirectResponse(url="/employees?updated=1", status_code=303)


@router.post("/delete/{employee_id}")
async def employee_delete(
    employee_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    """Xodimni o'chirish. Bog'liq hujjatlar bo'lsa (ishga qabul, avans, oylik, davomat va h.k.) DB xatolik beradi — foydalanuvchiga xabar."""
    from urllib.parse import quote
    emp = db.query(Employee).filter(Employee.id == employee_id).first()
    if not emp:
        return RedirectResponse(url="/employees?error=Xodim topilmadi", status_code=303)
    try:
        db.delete(emp)
        db.commit()
        return RedirectResponse(url="/employees?deleted=1", status_code=303)
    except IntegrityError:
        db.rollback()
        return RedirectResponse(
            url="/employees?error=" + quote(
                "Xodimni o'chirib bo'lmaydi: unga bog'liq yozuvlar mavjud (ishga qabul hujjati, avans, oylik, davomat va h.k.). "
                "Xodimni o'chirmasdan «Faol emas» deb belgilang yoki avval bog'liq hujjatlarni olib tashlang."
            ),
            status_code=303,
        )


# --- ISHDAN BO'SHATISH ---
DISMISSAL_REASONS = [
    ("own_will", "O'z ixtiyori bilan"),
    ("contract_end", "Shartnoma muddati tugadi"),
    ("discipline", "Mehnat intizomini buzgani"),
    ("reduction", "Loyihadan (shtatdan) qisqartirish"),
    ("agreement", "O'zaro kelishuv"),
    ("other", "Boshqa"),
]


@router.get("/dismissal/create", response_class=HTMLResponse)
async def dismissal_create_page(
    request: Request,
    employee_id: int = Query(..., description="Xodim ID"),
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    """Ishdan bo'shatish hujjati yaratish — forma."""
    emp = db.query(Employee).filter(Employee.id == employee_id).first()
    if not emp:
        return RedirectResponse(url="/employees?error=Xodim topilmadi", status_code=303)
    if not emp.is_active:
        return RedirectResponse(url="/employees?error=Xodim allaqachon ishdan bo'shatilgan", status_code=303)
    default_date = datetime.now().date().strftime("%Y-%m-%d")
    return templates.TemplateResponse("employees/dismissal_form.html", {
        "request": request,
        "employee": emp,
        "reasons": DISMISSAL_REASONS,
        "default_date": default_date,
        "current_user": current_user,
        "page_title": "Ishdan bo'shatish",
    })


@router.post("/dismissal/create", response_class=RedirectResponse)
async def dismissal_create_submit(
    request: Request,
    employee_id: int = Form(...),
    doc_date: str = Form(...),
    reason: str = Form(""),
    note: str = Form(""),
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    """Ishdan bo'shatish hujjatini yaratadi, xodimni faol emas qiladi."""
    from urllib.parse import quote
    emp = db.query(Employee).filter(Employee.id == employee_id).first()
    if not emp:
        return RedirectResponse(url="/employees?error=Xodim topilmadi", status_code=303)
    if not emp.is_active:
        return RedirectResponse(url="/employees?error=Xodim allaqachon ishdan bo'shatilgan", status_code=303)
    try:
        doc_d = datetime.strptime(doc_date.strip()[:10], "%Y-%m-%d").date()
    except (ValueError, TypeError):
        return RedirectResponse(url=f"/dismissal/create?employee_id={employee_id}&error=Noto%27g%27ri sana", status_code=303)
    reason_label = next((r[1] for r in DISMISSAL_REASONS if r[0] == reason), reason or "—")
    count = db.query(DismissalDoc).filter(DismissalDoc.doc_date >= doc_d.replace(day=1)).count()
    number = f"IB-{doc_d.strftime('%Y%m%d')}-{count + 1:04d}"
    doc = DismissalDoc(
        number=number,
        employee_id=emp.id,
        doc_date=doc_d,
        reason=reason_label,
        note=(note or "").strip() or None,
        user_id=current_user.id if current_user else None,
    )
    db.add(doc)
    db.flush()
    emp.is_active = False
    db.commit()
    return RedirectResponse(url=f"/dismissal/{doc.id}?created=1", status_code=303)


@router.get("/dismissal/{doc_id}", response_class=HTMLResponse)
async def dismissal_doc_view(
    doc_id: int,
    request: Request,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    """Ishdan bo'shatish hujjati ko'rinishi."""
    doc = (
        db.query(DismissalDoc)
        .options(joinedload(DismissalDoc.employee), joinedload(DismissalDoc.user))
        .filter(DismissalDoc.id == doc_id)
        .first()
    )
    if not doc:
        raise HTTPException(status_code=404, detail="Hujjat topilmadi")
    return templates.TemplateResponse("employees/dismissal_doc.html", {
        "request": request,
        "doc": doc,
        "current_user": current_user,
        "page_title": f"Ishdan bo'shatish {doc.number}",
    })


def _build_dismissal_docx(doc, company_name: str, employer_rep_name: str):
    """Ishdan bo'shatish hujjatini Word (.docx) sifatida qaytaradi (BytesIO)."""
    d = Document()
    style = d.styles["Normal"]
    style.font.size = Pt(11)
    style.font.name = "Times New Roman"
    d.add_heading("ISHDAN BO'SHATISH HAQIDA BUYRUQ", level=0)
    h = d.paragraphs[-1]
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    d.add_paragraph()
    p = d.add_paragraph()
    p.add_run(f"№ {doc.number}").bold = True
    p.add_run(f"   Sana: {doc.doc_date.strftime('%d.%m.%Y') if doc.doc_date else '—'}")
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    d.add_paragraph()
    d.add_paragraph(f"Joy: ____________________________")
    d.add_paragraph(f"Korxona: {company_name}")
    d.add_paragraph()
    emp = doc.employee
    d.add_paragraph(
        f"1. {emp.full_name} (xodim kodi: {emp.code or '—'}), "
        f"{doc.doc_date.strftime('%d.%m.%Y')} sanadan boshlab ishdan bo'shatiladi."
    )
    d.add_paragraph(f"2. Ishdan bo'shatish sababi: {doc.reason or '—'}.")
    if doc.note:
        d.add_paragraph(f"3. Izoh: {doc.note}")
    d.add_paragraph()
    d.add_paragraph("Ish beruvchi:")
    d.add_paragraph(f"Korxona: {company_name}")
    d.add_paragraph(f"Rahbar: {employer_rep_name}")
    d.add_paragraph("Imzo: ______________________")
    d.add_paragraph()
    d.add_paragraph("Xodim bilan tanishtirildi:")
    d.add_paragraph(f"F.I.O: {emp.full_name}")
    d.add_paragraph("Imzo: ______________________")
    buf = io.BytesIO()
    d.save(buf)
    buf.seek(0)
    return buf


@router.get("/dismissal/{doc_id}/export-word")
async def dismissal_doc_export_word(
    doc_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    """Ishdan bo'shatish hujjatini Word (.docx) formatida yuklab olish."""
    doc = (
        db.query(DismissalDoc)
        .options(joinedload(DismissalDoc.employee))
        .filter(DismissalDoc.id == doc_id)
        .first()
    )
    if not doc:
        raise HTTPException(status_code=404, detail="Hujjat topilmadi")
    company_name = "TOTLI HOLVA SWEETS"
    employer_rep_name = "Rahimov D.A."
    buf = _build_dismissal_docx(doc, company_name, employer_rep_name)
    safe_number = (doc.number or "ib").replace("/", "-").replace("\\", "-")
    filename = f"Ishdan_bo'shatish_{safe_number}.docx"
    return Response(
        content=buf.read(),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename*=UTF-8\'\'{quote(filename)}'},
    )


# --- ISHGA QABUL QILISH HUJJATI ---
@router.get("/hiring-docs", response_class=HTMLResponse)
async def employment_docs_list(
    request: Request,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    """Ishga qabul qilish hujjatlari ro'yxati — barcha foydalanuvchilar barcha hujjatlarni ko'radi"""
    docs = (
        db.query(EmploymentDoc)
        .options(joinedload(EmploymentDoc.employee))
        .order_by(EmploymentDoc.created_at.desc())
        .all()
    )
    return templates.TemplateResponse("employees/hiring_docs_list.html", {
        "request": request,
        "docs": docs,
        "current_user": current_user,
        "page_title": "Ishga qabul qilish hujjatlari"
    })


@router.get("/hiring-doc/create", response_class=HTMLResponse)
async def employment_doc_create_page(
    request: Request,
    employee_id: Optional[int] = None,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    """Ishga qabul hujjati yaratish (xodim tanlash). Har bir xodim faqat bir marta ishga qabul qilinadi."""
    employees = db.query(Employee).filter(Employee.is_active == True).order_by(Employee.full_name).all()
    emp = db.query(Employee).filter(Employee.id == employee_id).first() if employee_id else None
    # Agar tanlangan xodimda allaqachon ishga qabul hujjati bo'lsa — yangi yaratishga ruxsat yo'q
    if emp:
        existing = db.query(EmploymentDoc).filter(EmploymentDoc.employee_id == emp.id).order_by(EmploymentDoc.doc_date.desc()).first()
        if existing:
            return RedirectResponse(
                url="/hiring-docs?error=" + quote(f"«{emp.full_name}» allaqachon ishga qabul qilingan. Yangi hujjat yaratib bo'lmaydi — mavjud hujjatni ko'ring yoki tahrirlang.")
                + "&existing_doc_id=" + str(existing.id),
                status_code=303,
            )
    today_str = date.today().isoformat()
    departments = db.query(Department).filter(Department.is_active == True).order_by(Department.name).all()
    positions = db.query(Position).filter(Position.is_active == True).order_by(Position.name).all()
    piecework_tasks = db.query(PieceworkTask).filter(PieceworkTask.is_active == True).order_by(PieceworkTask.name).all()
    return templates.TemplateResponse("employees/hiring_doc_form.html", {
        "request": request,
        "employees": employees,
        "emp": emp,
        "today_str": today_str,
        "departments": departments,
        "positions": positions,
        "piecework_tasks": piecework_tasks,
        "current_user": current_user,
        "page_title": "Ishga qabul hujjati yaratish"
    })


@router.post("/hiring-doc/create")
async def employment_doc_create(
    employee_id: int = Form(...),
    doc_date: str = Form(...),
    hire_date: str = Form(None),
    position: str = Form(""),
    department: str = Form(""),
    salary: float = Form(0),
    salary_type: str = Form(""),
    piecework_task_ids: List[int] = Form([]),
    rest_days: List[str] = Form([]),
    probation: str = Form(""),
    contract_type: str = Form("indefinite"),
    contract_end_date: str = Form(None),
    note: str = Form(""),
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    """Ishga qabul hujjati yaratish (O'zR Mehnat kodeksi, gov.uz tamoyillari asosida). Har bir xodim faqat bir marta ishga qabul qilinadi."""
    from urllib.parse import quote
    emp = db.query(Employee).filter(Employee.id == employee_id).first()
    if not emp:
        return RedirectResponse(url="/hiring-docs?error=" + quote("Xodim topilmadi"), status_code=303)
    existing = db.query(EmploymentDoc).filter(EmploymentDoc.employee_id == emp.id).first()
    if existing:
        return RedirectResponse(
            url="/hiring-docs?error=" + quote(f"«{emp.full_name}» allaqachon ishga qabul qilingan. Yangi hujjat yaratib bo'lmaydi.")
            + "&existing_doc_id=" + str(existing.id),
            status_code=303,
        )
    try:
        doc_d = datetime.strptime(doc_date.strip(), "%Y-%m-%d").date()
    except (ValueError, TypeError):
        return RedirectResponse(url="/hiring-doc/create?employee_id=" + str(employee_id) + "&error=" + quote("Noto'g'ri sana"), status_code=303)
    hire_d = None
    if hire_date and hire_date.strip():
        try:
            hire_d = datetime.strptime(hire_date.strip(), "%Y-%m-%d").date()
        except (ValueError, TypeError):
            pass
    end_d = None
    if contract_end_date and contract_end_date.strip() and (contract_type or "").strip() == "fixed":
        try:
            end_d = datetime.strptime(contract_end_date.strip(), "%Y-%m-%d").date()
        except (ValueError, TypeError):
            pass
    count = db.query(EmploymentDoc).filter(EmploymentDoc.doc_date >= doc_d.replace(day=1)).count()
    number = f"IQ-{doc_d.strftime('%Y%m%d')}-{count + 1:04d}"
    doc_salary = float(salary) if salary else (emp.salary or 0)
    doc_department = (department or "").strip() or (emp.department or "").strip() or None
    st = (salary_type or "").strip() or None
    if st and st not in ("oylik", "soatlik", "bo'lak", "bo'lak_oylik"):
        st = None
    task_ids = [int(x) for x in (piecework_task_ids or []) if str(x).strip().isdigit()]
    task_ids = list(dict.fromkeys(task_ids))
    rest_days_clean = [d for d in (rest_days or []) if d in ("mon","tue","wed","thu","fri","sat","sun")]
    probation_clean = (probation or "").strip() or None
    ct = (contract_type or "").strip() or "indefinite"
    if ct not in ("indefinite", "fixed", "task"):
        ct = "indefinite"
    doc = EmploymentDoc(
        number=number,
        employee_id=emp.id,
        doc_date=doc_d,
        hire_date=hire_d,
        position=(position or "").strip() or (emp.position or "").strip() or None,
        department=doc_department,
        salary=doc_salary,
        salary_type=st,
        piecework_task_ids=",".join(str(x) for x in task_ids) if (st in ("bo'lak", "bo'lak_oylik") and task_ids) else None,
        rest_days=",".join(rest_days_clean) if rest_days_clean else None,
        probation=probation_clean,
        contract_type=ct,
        contract_end_date=end_d,
        note=note or None,
        user_id=current_user.id,
        confirmed_at=datetime.now(),  # Hujjat yaratilganda avtomatik tasdiqlanadi
    )
    db.add(doc)
    db.flush()
    emp.salary = doc_salary
    if st:
        emp.salary_type = st
    if st in ("bo'lak", "bo'lak_oylik"):
        if task_ids:
            tasks = db.query(PieceworkTask).filter(PieceworkTask.id.in_(task_ids)).all()
            emp.piecework_tasks = tasks
            emp.piecework_task_id = task_ids[0]  # legacy
        else:
            emp.piecework_tasks = []
            emp.piecework_task_id = None
    if hire_d:
        emp.hire_date = hire_d
    if (position or "").strip():
        emp.position = (position or "").strip()
    if doc_department:
        emp.department = doc_department
    db.commit()
    return RedirectResponse(url=f"/hiring-doc/{doc.id}?created=1", status_code=303)


@router.get("/hiring-doc/{doc_id}", response_class=HTMLResponse)
async def employment_doc_view(
    request: Request,
    doc_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    """Ishga qabul hujjati ko'rish / chop etish — barcha maydonlar to'liq ko'rsatiladi."""
    doc = (
        db.query(EmploymentDoc)
        .options(joinedload(EmploymentDoc.employee))
        .filter(EmploymentDoc.id == doc_id)
        .first()
    )
    if not doc:
        return RedirectResponse(url="/hiring-docs?error=Hujjat topilmadi", status_code=303)
    # Bo'lim: hujjatdagi yoki xodimdagi (matn) yoki xodimning department_id orqali
    display_department = (doc.department or "").strip() or None
    if not display_department and doc.employee:
        display_department = (doc.employee.department or "").strip() or None
        if not display_department and getattr(doc.employee, "department_id", None):
            dept = db.query(Department).filter(Department.id == doc.employee.department_id).first()
            if dept:
                display_department = dept.name
    if not display_department:
        display_department = "—"

    # Bo'lak ishlar (snapshot) — hujjatda saqlangan ro'yxat
    piecework_task_names = []
    try:
        raw = (doc.piecework_task_ids or "").strip()
        ids = [int(x) for x in raw.split(",") if x.strip().isdigit()] if raw else []
        if ids:
            tasks = db.query(PieceworkTask).filter(PieceworkTask.id.in_(ids)).order_by(PieceworkTask.name).all()
            for t in tasks:
                nm = (t.name or t.code or str(t.id))
                piecework_task_names.append(nm)
    except Exception:
        piecework_task_names = []
    return templates.TemplateResponse("employees/hiring_doc.html", {
        "request": request,
        "doc": doc,
        "display_department": display_department,
        "piecework_task_names": piecework_task_names,
        "current_user": current_user,
        "page_title": f"Ishga qabul {doc.number}"
    })


@router.get("/hiring-doc/{doc_id}/contract", response_class=HTMLResponse)
async def employment_doc_contract(
    request: Request,
    doc_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    """Mehnat shartnomasi (to'liq) — namuna asosida chop etish."""
    doc = (
        db.query(EmploymentDoc)
        .options(joinedload(EmploymentDoc.employee), joinedload(EmploymentDoc.user))
        .filter(EmploymentDoc.id == doc_id)
        .first()
    )
    if not doc:
        return RedirectResponse(url="/hiring-docs?error=Hujjat topilmadi", status_code=303)

    # Bo'lim ko'rsatish
    display_department = (doc.department or "").strip() or None
    if not display_department and doc.employee:
        display_department = (doc.employee.department or "").strip() or None
        if not display_department and getattr(doc.employee, "department_id", None):
            dept = db.query(Department).filter(Department.id == doc.employee.department_id).first()
            if dept:
                display_department = dept.name
    if not display_department:
        display_department = "—"

    # Tanlangan bo'lak ishlar (snapshot) — stavkalari bilan
    selected_piecework_tasks = []
    try:
        raw = (doc.piecework_task_ids or "").strip()
        ids = [int(x) for x in raw.split(",") if x.strip().isdigit()] if raw else []
        if ids:
            selected_piecework_tasks = db.query(PieceworkTask).filter(PieceworkTask.id.in_(ids)).order_by(PieceworkTask.name).all()
    except Exception:
        selected_piecework_tasks = []

    # Dam olish kunlari matni
    rest_days_display = ""
    try:
        raw_rest = (doc.rest_days or "").strip()
        codes = [x for x in raw_rest.split(",") if x]
        name_map = {
            "mon": "dushanba",
            "tue": "seshanba",
            "wed": "chorshanba",
            "thu": "payshanba",
            "fri": "juma",
            "sat": "shanba",
            "sun": "yakshanba",
        }
        names = [name_map.get(c, c) for c in codes]
        if names:
            rest_days_display = ", ".join(names)
    except Exception:
        rest_days_display = ""

    company_name = "TOTLI HOLVA SWEETS"
    employer_rep_name = "Rahimov D.A."

    return templates.TemplateResponse("employees/labor_contract.html", {
        "request": request,
        "doc": doc,
        "display_department": display_department,
        "selected_piecework_tasks": selected_piecework_tasks,
        "company_name": company_name,
        "employer_rep_name": employer_rep_name,
        "rest_days_display": rest_days_display,
        "current_user": current_user,
        "page_title": f"Mehnat shartnomasi {doc.number}",
    })


def _build_labor_contract_docx(doc, display_department, selected_piecework_tasks, rest_days_display, company_name, employer_rep_name):
    """Shartnoma matnini Word hujjati (.docx) sifatida qaytaradi (BytesIO)."""
    d = Document()
    style = d.styles["Normal"]
    style.font.size = Pt(11)
    style.font.name = "Times New Roman"

    # Sarlavha
    h = d.add_heading("MEHNAT SHARTNOMASI", level=0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    d.add_paragraph()
    p = d.add_paragraph()
    p.add_run(f"№ {doc.number}").bold = True
    p.add_run(f"   Sana: {doc.doc_date.strftime('%d.%m.%Y') if doc.doc_date else '—'}")
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    d.add_paragraph()

    d.add_paragraph("Joy: ____________________________")
    d.add_paragraph(f"Korxona: {company_name}")
    d.add_paragraph()

    d.add_paragraph(
        f"{company_name} (keyingi o'rinlarda Ish beruvchi) va fuqaro {doc.employee.full_name} "
        "(keyingi o'rinlarda Xodim), mazkur mehnat shartnomasini quyidagilar haqida tuzdilar."
    )
    d.add_paragraph()

    d.add_heading("1. UMUMIY QOIDALAR", level=1)
    hire_date_str = doc.hire_date.strftime("%d.%m.%Y") if doc.hire_date else "________________"
    pos = doc.position or (doc.employee.position if doc.employee else "") or "________________"
    d.add_paragraph(
        f"1.1. Xodim {hire_date_str} sanadan boshlab {display_department} bo'limida {pos} lavozimiga ishga qabul qilinadi."
    )
    d.add_paragraph(f"1.2. Xodimning ish joyi: {display_department}.")
    if doc.contract_type == "fixed":
        end = f" ({doc.contract_end_date.strftime('%d.%m.%Y')} gacha)" if doc.contract_end_date else ""
        d.add_paragraph(f"1.3. Mazkur shartnomaning amal qilish muddati: muayyan muddatga{end}.")
    elif doc.contract_type == "task":
        d.add_paragraph("1.3. Mazkur shartnomaning amal qilish muddati: muayyan ishni bajarish davriga.")
    else:
        d.add_paragraph("1.3. Mazkur shartnomaning amal qilish muddati: nomuayyan muddatga.")
    prob = doc.probation if doc.probation else "sinovsiz"
    d.add_paragraph(f"1.4. Sinov muddati: {prob}.")
    d.add_paragraph("1.5. Xodim lavozim yo'riqnomasi va amaldagi qonunchilikka muvofiq mehnat majburiyatlarini bajaradi.")
    d.add_paragraph()

    d.add_heading("2. TOMONLARNING HUQUQ VA MAJBURIYATLARI", level=1)
    d.add_paragraph("2.1. Ish beruvchining majburiyatlari:")
    d.add_paragraph("  • Xodimga xavfsiz va samarali mehnat qilish uchun shart-sharoitlar yaratish.")
    d.add_paragraph("  • Ichki mehnat tartibi qoidalari va lavozim yo'riqnomasi bilan tanishtirish.")
    d.add_paragraph("  • Ish haqini o'z vaqtida to'lash.")
    d.add_paragraph("2.2. Xodimning majburiyatlari:")
    d.add_paragraph("  • Mehnat intizomi va ichki tartib qoidalariga rioya qilish.")
    d.add_paragraph("  • Ish beruvchining qonuniy topshiriqlarini o'z vaqtida va aniq bajarish.")
    d.add_paragraph("  • Mehnat muhofazasi va texnika xavfsizligi talablariga rioya qilish.")
    d.add_paragraph()

    d.add_heading("3. ISH VAQTI VA DAM OLISH VAQTI", level=1)
    d.add_paragraph("3.1. Ish kuni vaqti: 09:00 dan 18:00 gacha (to'liq ish kuni asosida).")
    rest = rest_days_display if rest_days_display else "shanba va yakshanba"
    d.add_paragraph(f"3.2. Dam olish kunlari: {rest}.")
    d.add_paragraph("3.3. Qonunchilikda belgilangan tartibda dam olish/bayram kunlari ishga jalb etilishi mumkin.")
    d.add_paragraph()

    d.add_heading("4. MEHNATGA HAQ TO'LASH", level=1)
    salary_type_map = {"oylik": "Oylik", "soatlik": "Soatlik", "bo'lak": "Bo'lak", "bo'lak_oylik": "Bo'lak + oylik"}
    st = salary_type_map.get(doc.salary_type, "________________")
    d.add_paragraph(f"4.1. Ish haqi turi: {st}.")
    if doc.salary_type in ("bo'lak", "bo'lak_oylik") and selected_piecework_tasks:
        d.add_paragraph("Bo'lak ishlar va stavkalar:")
        for t in selected_piecework_tasks:
            name = t.name or t.code or str(t.id)
            price = f"{t.price_per_unit:,.0f}" if t.price_per_unit is not None else "0"
            unit = t.unit_name or "birlik"
            d.add_paragraph(f"  • {name} — {price} so'm/{unit}")
    salary_val = f"{doc.salary:,.0f}" if doc.salary else "0"
    d.add_paragraph(f"4.2. Mehnat haqi miqdori: {salary_val} so'm.")
    d.add_paragraph("4.3. Ish haqi har oyda kamida ikki marta to'lanadi.")
    d.add_paragraph()

    d.add_heading("5. XIZMAT SAFARLARI", level=1)
    d.add_paragraph("5.1. Ish zaruriyatiga ko'ra Xodim xizmat safariga yuborilishi mumkin. Xarajatlar amaldagi qonunchilikka muvofiq qoplanadi.")
    d.add_paragraph()

    d.add_heading("6. MEHNAT SHARTNOMASINI BEKOR QILISH", level=1)
    d.add_paragraph("6.1. Mehnat shartnomasi O'zbekiston Respublikasi Mehnat kodeksida belgilangan tartibda bekor qilinishi mumkin.")
    d.add_paragraph()

    d.add_heading("7. MEHNAT NIZOLARI", level=1)
    d.add_paragraph("7.1. Mehnat nizolari qonun hujjatlarida belgilangan tartibda hal qilinadi.")
    d.add_paragraph()

    d.add_heading("8. TOMONLAR REKVIZITLARI VA IMZOLARI", level=1)
    d.add_paragraph("Ish beruvchi:")
    d.add_paragraph(f"Korxona: {company_name}")
    d.add_paragraph("Manzil: O'zbekiston Respublikasi, Qo'qon shahri, Jasorat ko'chasi, 52-uy")
    d.add_paragraph("STIR: 311469106")
    d.add_paragraph("Hisob raqam: 202088409071067110001")
    d.add_paragraph('Bank: "Asaka" banki Qo\'qon filiali')
    d.add_paragraph("MFO: 00873")
    d.add_paragraph(f"Rahbar: {employer_rep_name}")
    d.add_paragraph("Imzo: ______________________")
    d.add_paragraph()
    d.add_paragraph("Xodim:")
    d.add_paragraph(f"F.I.O: {doc.employee.full_name}")
    d.add_paragraph(f"Kodi: {doc.employee.code or '—'}")
    d.add_paragraph(f"Telefon: {doc.employee.phone or '—'}")
    d.add_paragraph("Manzil: ____________________________")
    d.add_paragraph("Pasport: ____________________________")
    d.add_paragraph("Imzo: ______________________")

    buf = io.BytesIO()
    d.save(buf)
    buf.seek(0)
    return buf


@router.get("/hiring-doc/{doc_id}/contract/export-word")
async def employment_doc_contract_export_word(
    doc_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    """Mehnat shartnomasini Word (.docx) formatida yuklab olish."""
    doc = (
        db.query(EmploymentDoc)
        .options(joinedload(EmploymentDoc.employee), joinedload(EmploymentDoc.user))
        .filter(EmploymentDoc.id == doc_id)
        .first()
    )
    if not doc:
        raise HTTPException(status_code=404, detail="Hujjat topilmadi")

    display_department = (doc.department or "").strip() or None
    if not display_department and doc.employee:
        display_department = (doc.employee.department or "").strip() or None
        if not display_department and getattr(doc.employee, "department_id", None):
            dept = db.query(Department).filter(Department.id == doc.employee.department_id).first()
            if dept:
                display_department = dept.name
    if not display_department:
        display_department = "—"

    selected_piecework_tasks = []
    try:
        raw = (doc.piecework_task_ids or "").strip()
        ids = [int(x) for x in raw.split(",") if x.strip().isdigit()] if raw else []
        if ids:
            selected_piecework_tasks = db.query(PieceworkTask).filter(PieceworkTask.id.in_(ids)).order_by(PieceworkTask.name).all()
    except Exception:
        selected_piecework_tasks = []

    rest_days_display = ""
    try:
        raw_rest = (doc.rest_days or "").strip()
        codes = [x for x in raw_rest.split(",") if x]
        name_map = {"mon": "dushanba", "tue": "seshanba", "wed": "chorshanba", "thu": "payshanba", "fri": "juma", "sat": "shanba", "sun": "yakshanba"}
        names = [name_map.get(c, c) for c in codes]
        if names:
            rest_days_display = ", ".join(names)
    except Exception:
        rest_days_display = ""

    company_name = "TOTLI HOLVA SWEETS"
    employer_rep_name = "Rahimov D.A."

    buf = _build_labor_contract_docx(
        doc, display_department, selected_piecework_tasks, rest_days_display, company_name, employer_rep_name
    )
    safe_number = (doc.number or "shartnoma").replace("/", "-").replace("\\", "-")
    filename = f"Mehnat_shartnomasi_{safe_number}.docx"
    return Response(
        content=buf.read(),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename*=UTF-8\'\'{quote(filename)}'},
    )


@router.post("/hiring-docs/bulk-confirm")
async def employment_docs_bulk_confirm(
    request: Request,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_admin),
):
    """Tanlangan ishga qabul hujjatlarini tasdiqlash"""
    from urllib.parse import quote
    form = await request.form()
    doc_ids_raw = form.getlist("doc_ids")
    try:
        doc_ids = [int(x) for x in doc_ids_raw if str(x).strip().isdigit()]
    except (ValueError, TypeError):
        doc_ids = []
    if not doc_ids:
        return RedirectResponse(url="/hiring-docs?error=" + quote("Hech qanday hujjat tanlanmagan."), status_code=303)
    confirmed = 0
    for did in doc_ids:
        doc = db.query(EmploymentDoc).filter(EmploymentDoc.id == did).first()
        if doc and not doc.confirmed_at:
            doc.confirmed_at = datetime.now()
            confirmed += 1
    db.commit()
    return RedirectResponse(url=f"/hiring-docs?confirmed=1&count={confirmed}", status_code=303)


@router.post("/hiring-docs/bulk-cancel-confirm")
async def employment_docs_bulk_cancel_confirm(
    request: Request,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_admin),
):
    """Tanlangan hujjatlarda tasdiqlashni bekor qilish"""
    from urllib.parse import quote
    form = await request.form()
    doc_ids_raw = form.getlist("doc_ids")
    try:
        doc_ids = [int(x) for x in doc_ids_raw if str(x).strip().isdigit()]
    except (ValueError, TypeError):
        doc_ids = []
    if not doc_ids:
        return RedirectResponse(url="/hiring-docs?error=" + quote("Hech qanday hujjat tanlanmagan."), status_code=303)
    unconfirmed = 0
    for did in doc_ids:
        doc = db.query(EmploymentDoc).filter(EmploymentDoc.id == did).first()
        if doc and doc.confirmed_at:
            doc.confirmed_at = None
            unconfirmed += 1
    db.commit()
    return RedirectResponse(url=f"/hiring-docs?unconfirmed=1&count={unconfirmed}", status_code=303)


@router.post("/hiring-doc/{doc_id}/confirm")
async def employment_doc_confirm(
    doc_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_admin),
):
    """Ishga qabul hujjatini tasdiqlash"""
    doc = db.query(EmploymentDoc).filter(EmploymentDoc.id == doc_id).first()
    if not doc:
        return RedirectResponse(url="/hiring-docs?error=Hujjat topilmadi", status_code=303)
    doc.confirmed_at = datetime.now()
    db.commit()
    return RedirectResponse(url="/hiring-docs?confirmed=1", status_code=303)


@router.post("/hiring-doc/{doc_id}/cancel-confirm")
async def employment_doc_cancel_confirm(
    doc_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_admin),
):
    """Ishga qabul hujjati tasdiqlashni bekor qilish"""
    doc = db.query(EmploymentDoc).filter(EmploymentDoc.id == doc_id).first()
    if not doc:
        return RedirectResponse(url="/hiring-docs?error=Hujjat topilmadi", status_code=303)
    doc.confirmed_at = None
    db.commit()
    return RedirectResponse(url="/hiring-docs?unconfirmed=1", status_code=303)


@router.post("/hiring-doc/{doc_id}/delete")
async def employment_doc_delete(
    doc_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_admin),
):
    """Ishga qabul hujjatini o'chirish — faqat tasdiqlanmagan hujjatni o'chirish mumkin."""
    from urllib.parse import quote
    doc = db.query(EmploymentDoc).filter(EmploymentDoc.id == doc_id).first()
    if not doc:
        return RedirectResponse(url="/hiring-docs?error=Hujjat topilmadi", status_code=303)
    if doc.confirmed_at:
        return RedirectResponse(
            url="/hiring-docs?error=" + quote("Tasdiqlangan hujjatni o'chirish mumkin emas. Avval «Bekor qilish» orqali tasdiqlashni bekor qiling."),
            status_code=303
        )
    db.delete(doc)
    db.commit()
    return RedirectResponse(url="/hiring-docs?deleted=1", status_code=303)


@router.get("/hiring-doc/{doc_id}/edit", response_class=HTMLResponse)
async def employment_doc_edit_page(
    request: Request,
    doc_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_admin),
):
    """Ishga qabul hujjatini tahrirlash"""
    doc = (
        db.query(EmploymentDoc)
        .options(joinedload(EmploymentDoc.employee))
        .filter(EmploymentDoc.id == doc_id)
        .first()
    )
    if not doc:
        return RedirectResponse(url="/hiring-docs?error=Hujjat topilmadi", status_code=303)
    if doc.confirmed_at:
        from urllib.parse import quote
        return RedirectResponse(
            url="/hiring-docs?error=" + quote("Tasdiqlangan hujjatni tahrirlash mumkin emas. Avval «Bekor qilish» orqali tasdiqlashni bekor qiling."),
            status_code=303,
        )
    departments = db.query(Department).filter(Department.is_active == True).order_by(Department.name).all()
    positions = db.query(Position).filter(Position.is_active == True).order_by(Position.name).all()
    piecework_tasks = db.query(PieceworkTask).filter(PieceworkTask.is_active == True).order_by(PieceworkTask.name).all()
    # Tanlangan bo'lak ishlar IDlari (checkbox checked uchun)
    selected_piecework_ids = []
    if doc.piecework_task_ids:
        for x in (doc.piecework_task_ids or "").split(","):
            if str(x).strip().isdigit():
                selected_piecework_ids.append(int(x.strip()))
    display_department = (doc.department or "").strip() or (getattr(doc.employee, "department", "") or "").strip() or "—"
    return templates.TemplateResponse("employees/hiring_doc_edit.html", {
        "request": request,
        "doc": doc,
        "departments": departments,
        "positions": positions,
        "piecework_tasks": piecework_tasks,
        "selected_piecework_ids": selected_piecework_ids,
        "display_department": display_department,
        "current_user": current_user,
        "page_title": f"Ishga qabul {doc.number} — tahrirlash",
    })


@router.post("/hiring-doc/{doc_id}/edit")
async def employment_doc_edit_save(
    doc_id: int,
    doc_date: str = Form(...),
    hire_date: str = Form(None),
    position: str = Form(""),
    department: str = Form(""),
    salary: float = Form(0),
    salary_type: str = Form(""),
    piecework_task_ids: List[int] = Form([]),
    rest_days: List[str] = Form([]),
    probation: str = Form(""),
    contract_type: str = Form("indefinite"),
    contract_end_date: str = Form(None),
    note: str = Form(""),
    db: Session = Depends(get_db),
    current_user: User = Depends(require_admin),
):
    """Ishga qabul hujjatini saqlash (tahrirlash) — faqat tasdiqlanmagan hujjatni tahrirlash mumkin."""
    from urllib.parse import quote
    doc = db.query(EmploymentDoc).filter(EmploymentDoc.id == doc_id).first()
    if not doc:
        return RedirectResponse(url="/hiring-docs?error=Hujjat topilmadi", status_code=303)
    if doc.confirmed_at:
        return RedirectResponse(
            url="/hiring-docs?error=" + quote("Tasdiqlangan hujjatni tahrirlash mumkin emas. Avval «Bekor qilish» orqali tasdiqlashni bekor qiling."),
            status_code=303,
        )
    emp = db.query(Employee).filter(Employee.id == doc.employee_id).first()
    if not emp:
        return RedirectResponse(url="/hiring-docs?error=" + quote("Xodim topilmadi"), status_code=303)
    try:
        doc_d = datetime.strptime(doc_date.strip(), "%Y-%m-%d").date()
    except (ValueError, TypeError):
        return RedirectResponse(url=f"/hiring-doc/{doc_id}/edit?error=" + quote("Noto'g'ri sana"), status_code=303)
    hire_d = None
    if hire_date and hire_date.strip():
        try:
            hire_d = datetime.strptime(hire_date.strip(), "%Y-%m-%d").date()
        except (ValueError, TypeError):
            pass
    end_d = None
    if contract_end_date and contract_end_date.strip() and (contract_type or "").strip() == "fixed":
        try:
            end_d = datetime.strptime(contract_end_date.strip(), "%Y-%m-%d").date()
        except (ValueError, TypeError):
            pass
    st = (salary_type or "").strip() or None
    if st and st not in ("oylik", "soatlik", "bo'lak", "bo'lak_oylik"):
        st = None
    task_ids = [int(x) for x in (piecework_task_ids or []) if str(x).strip().isdigit()]
    task_ids = list(dict.fromkeys(task_ids))
    rest_days_clean = [d for d in (rest_days or []) if d in ("mon","tue","wed","thu","fri","sat","sun")]
    probation_clean = (probation or "").strip() or None
    ct = (contract_type or "").strip() or "indefinite"
    if ct not in ("indefinite", "fixed", "task"):
        ct = "indefinite"

    doc.doc_date = doc_d
    doc.hire_date = hire_d
    doc.position = (position or "").strip() or None
    doc.department = (department or "").strip() or None
    doc.salary = float(salary or 0)
    doc.salary_type = st
    doc.piecework_task_ids = ",".join(str(x) for x in task_ids) if (st in ("bo'lak", "bo'lak_oylik") and task_ids) else None
    doc.contract_type = ct
    doc.contract_end_date = end_d
    doc.note = (note or "").strip() or None
    doc.probation = probation_clean
    doc.rest_days = ",".join(rest_days_clean) if rest_days_clean else None

    # Employee snapshot yangilash (o'ylik hisoblash uchun)
    emp.salary = doc.salary
    if st:
        emp.salary_type = st
    if hire_d:
        emp.hire_date = hire_d
    if doc.position:
        emp.position = doc.position
    if doc.department:
        emp.department = doc.department
    if st in ("bo'lak", "bo'lak_oylik"):
        if task_ids:
            tasks = db.query(PieceworkTask).filter(PieceworkTask.id.in_(task_ids)).all()
            emp.piecework_tasks = tasks
            emp.piecework_task_id = task_ids[0]  # legacy
        else:
            emp.piecework_tasks = []
            emp.piecework_task_id = None

    db.commit()
    return RedirectResponse(url=f"/hiring-doc/{doc.id}?edited=1", status_code=303)


# --- EMPLOYEES EXCEL OPERATIONS ---
@router.get("/export")
async def export_employees(db: Session = Depends(get_db), current_user: User = Depends(require_auth)):
    employees = db.query(Employee).all()
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Employees"
    ws.append(["ID", "Kod", "F.I.SH", "Lavozim", "Bo'lim", "Telefon", "Oylik"])
    for e in employees:
        ws.append([e.id, e.code, e.full_name, e.position, e.department, e.phone, e.salary])
    stream = io.BytesIO()
    wb.save(stream)
    stream.seek(0)
    return StreamingResponse(stream, media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", headers={"Content-Disposition": "attachment; filename=xodimlar.xlsx"})

@router.get("/template")
async def template_employees():
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Template"
    ws.append(["Kod", "F.I.SH", "Lavozim", "Bo'lim", "Telefon", "Oylik"])
    ws.append(["X001", "Aliyev Vali", "Ishchi", "Ishlab chiqarish", "+998901234567", 3000000])
    stream = io.BytesIO()
    wb.save(stream)
    stream.seek(0)
    return StreamingResponse(stream, media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", headers={"Content-Disposition": "attachment; filename=xodim_andoza.xlsx"})

@router.post("/import")
async def import_employees(file: UploadFile = File(...), db: Session = Depends(get_db), current_user: User = Depends(require_auth)):
    contents = await file.read()
    wb = openpyxl.load_workbook(io.BytesIO(contents))
    ws = wb.active
    rows = list(ws.iter_rows(min_row=2, values_only=True))
    for row in rows:
        if not row[0]: continue
        code, full_name, position, department, phone, salary = row[0:6]
        employee = db.query(Employee).filter(Employee.code == code).first()
        if not employee:
            employee = Employee(
                code=code, 
                full_name=full_name, 
                position=position, 
                department=department, 
                phone=phone, 
                salary=salary
            )
            db.add(employee)
        else:
            employee.full_name = full_name
            employee.position = position
            employee.department = department
            employee.phone = phone
            employee.salary = salary
        db.commit()
    return RedirectResponse(url="", status_code=303)


@router.post("/import-from-hikvision-preview")
async def employees_import_from_hikvision_preview(
    request: Request,
    hikvision_host: str = Form(...),
    hikvision_port: str = Form("443"),
    hikvision_username: str = Form("admin"),
    hikvision_password: str = Form(""),
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    """Hikvision ulanishi va yuklanadigan shaxslar ro'yxatini ko'rsatadi; tanlanganlarni keyin yuklash mumkin."""
    from urllib.parse import quote
    try:
        port = int((hikvision_port or "").strip() or "443")
    except (ValueError, TypeError):
        port = 443
    try:
        from app.utils.hikvision import HikvisionAPI
        api = HikvisionAPI(
            host=(hikvision_host or "").strip(),
            port=port,
            username=(hikvision_username or "admin").strip(),
            password=(hikvision_password or ""),
        )
        if not api.test_connection():
            return RedirectResponse(
                url="/employees?error=" + quote(api._last_error or "Qurilma bilan bog'lanib bo'lmadi."),
                status_code=303
            )
        persons = api.get_person_list()
    except Exception as e:
        return RedirectResponse(url="/employees?error=" + quote("Hikvision: " + str(e)[:150]), status_code=303)
    return templates.TemplateResponse("employees/hikvision_import_preview.html", {
        "request": request,
        "persons": persons or [],
        "hikvision_host": (hikvision_host or "").strip(),
        "hikvision_port": str(port),
        "hikvision_username": (hikvision_username or "admin").strip(),
        "hikvision_password": hikvision_password or "",
        "current_user": current_user,
        "page_title": "Hikvision — xodimlarni tanlash"
    })


@router.get("/import-from-hikvision-preview", response_class=HTMLResponse)
async def employees_import_from_hikvision_preview_get(
    request: Request,
    current_user: User = Depends(require_auth),
):
    """Preview sahifasiga to'g'ridan-to'g'ri kirilsa xodimlar ro'yxatiga yo'naltiradi."""
    return RedirectResponse(url="", status_code=303)


@router.post("/import-from-hikvision")
async def employees_import_from_hikvision(
    hikvision_host: str = Form(...),
    hikvision_port: str = Form("443"),
    hikvision_username: str = Form("admin"),
    hikvision_password: str = Form(""),
    employee_no: Optional[List[str]] = Form(None),
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    """Hikvision qurilmasidan tanlangan (yoki barcha) xodimlarni Employee jadvaliga qo'shadi."""
    from urllib.parse import quote
    try:
        port = int((hikvision_port or "").strip() or "443")
    except (ValueError, TypeError):
        port = 443
    employee_nos = employee_no if isinstance(employee_no, list) and employee_no else None
    try:
        from app.utils.hikvision import import_employees_from_hikvision
        result = import_employees_from_hikvision(
            (hikvision_host or "").strip(),
            port,
            (hikvision_username or "admin").strip(),
            (hikvision_password or ""),
            db,
            employee_nos=employee_nos,
        )
        err_list = result.get("errors") or []
        imported = result.get("imported", 0)
        updated = result.get("updated", 0)
        if err_list:
            msg = f"Qo'shildi: {imported}, yangilandi: {updated}. Xato: " + "; ".join(str(e) for e in err_list[:3])
            return RedirectResponse(url="/employees?warning=" + quote(msg), status_code=303)
        msg = f"Qo'shildi: {imported}, yangilandi: {updated}."
        return RedirectResponse(url="/employees?imported=1&msg=" + quote(msg), status_code=303)
    except Exception as e:
        return RedirectResponse(url="/employees?error=" + quote("Hikvision: " + str(e)[:150]), status_code=303)


# ==========================================
# DAVOMAT (KUNLIK TABELLAR)
# ==========================================

@router.get("/attendance", response_class=HTMLResponse)
async def attendance_docs_list(
    request: Request,
    start_date: Optional[str] = None,
    end_date: Optional[str] = None,
    sort: Optional[str] = None,
    order: Optional[str] = None,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    """Kunlik tabel hujjatlari ro'yxati — saralash: number, date, count, confirmed_at"""
    today = date.today()
    start_date = start_date or (today - timedelta(days=30)).strftime("%Y-%m-%d")
    end_date = end_date or today.strftime("%Y-%m-%d")
    sort = (sort or "date").strip().lower()
    order = (order or "asc").strip().lower()
    if order not in ("asc", "desc"):
        order = "asc"
    query = (
        db.query(AttendanceDoc)
        .filter(AttendanceDoc.date >= start_date, AttendanceDoc.date <= end_date)
    )
    if sort == "number":
        query = query.order_by(AttendanceDoc.number.desc() if order == "desc" else AttendanceDoc.number.asc())
    elif sort == "date":
        query = query.order_by(AttendanceDoc.date.desc() if order == "desc" else AttendanceDoc.date.asc())
    elif sort == "confirmed_at":
        query = query.order_by(
            AttendanceDoc.confirmed_at.desc() if order == "desc" else AttendanceDoc.confirmed_at.asc()
        )
    else:
        query = query.order_by(AttendanceDoc.date.desc())
    docs = query.all()
    count_by_doc = {}
    for doc in docs:
        count_by_doc[doc.id] = db.query(Attendance).filter(Attendance.date == doc.date).count()
    if sort == "count":
        reverse = order == "desc"
        docs = sorted(docs, key=lambda d: count_by_doc.get(d.id, 0), reverse=reverse)
    return templates.TemplateResponse("employees/attendance_docs_list.html", {
        "request": request,
        "docs": docs,
        "count_by_doc": count_by_doc,
        "start_date": start_date,
        "end_date": end_date,
        "sort": sort,
        "order": order,
        "current_user": current_user,
        "page_title": "Kunlik tabellar",
    })


@router.get("/attendance/form", response_class=HTMLResponse)
async def attendance_form(
    request: Request,
    date_param: Optional[str] = Query(None, alias="date"),
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    """Tabel formasi — sana tanlash, shu kundagi yozuvlar, Hikvision yuklash. Sana query ?date=YYYY-MM-DD orqali olinadi."""
    today = date.today()
    form_date_str = (date_param or "").strip() or today.strftime("%Y-%m-%d")
    try:
        form_date = datetime.strptime(form_date_str, "%Y-%m-%d").date()
    except ValueError:
        form_date = today
        form_date_str = form_date.strftime("%Y-%m-%d")
    attendances = (
        db.query(Attendance)
        .filter(Attendance.date == form_date)
        .order_by(Attendance.employee_id)
        .all()
    )
    attendance_by_employee = {a.employee_id: a for a in attendances}
    # Barcha faol xodimlar + shu kunda davomat bo‘lgan (lekin ro‘yxatda bo‘lmagan) xodimlar
    employees_active = (
        db.query(Employee)
        .filter(Employee.is_active == True)
        .order_by(Employee.full_name)
        .all()
    )
    employee_ids_in_rows = {e.id for e in employees_active}
    attendance_rows = [{"employee": e, "attendance": attendance_by_employee.get(e.id)} for e in employees_active]
    for att in attendances:
        if att.employee_id not in employee_ids_in_rows:
            emp = db.query(Employee).filter(Employee.id == att.employee_id).first()
            if emp:
                attendance_rows.append({"employee": emp, "attendance": att})
                employee_ids_in_rows.add(emp.id)
    doc = db.query(AttendanceDoc).filter(AttendanceDoc.date == form_date).first()
    return templates.TemplateResponse("employees/attendance_form.html", {
        "request": request,
        "form_date": form_date,
        "form_date_str": form_date_str,
        "attendances": attendances,
        "attendance_rows": attendance_rows,
        "doc": doc,
        "current_user": current_user,
        "page_title": "Tabel formasi",
    })


@router.post("/attendance/sync-hikvision")
async def attendance_sync_hikvision(
    request: Request,
    start_date: str = Form(...),
    end_date: str = Form(...),
    hikvision_host: str = Form(...),
    hikvision_port: str = Form("443"),
    hikvision_username: str = Form("admin"),
    hikvision_password: str = Form(""),
    redirect_url: str = Form("/attendance/form"),
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    """Hikvision'dan davomat yuklash"""
    from urllib.parse import quote
    sep = "&" if "?" in (redirect_url or "") else "?"
    base_redirect = (redirect_url or "/attendance/form").strip()
    try:
        start_d = datetime.strptime((start_date or "").strip(), "%Y-%m-%d").date()
        end_d = datetime.strptime((end_date or "").strip(), "%Y-%m-%d").date()
    except (ValueError, TypeError):
        return RedirectResponse(url=base_redirect + sep + "error=" + quote("Noto'g'ri sana"), status_code=303)
    try:
        port = int(hikvision_port.strip() or "443")
    except (ValueError, TypeError):
        port = 443
    try:
        from app.utils.hikvision import sync_hikvision_attendance
        result = sync_hikvision_attendance(
            (hikvision_host or "").strip(),
            port,
            (hikvision_username or "admin").strip(),
            (hikvision_password or ""),
            start_d,
            end_d,
            db,
        )
        err_list = result.get("errors") or []
        events_count = result.get("events_count", 0)
        imported = result.get("imported", 0)
        msg = f"Hodisa: {events_count} ta, yuklangan: {imported} ta. Xato: {len(err_list)} ta."
        if err_list:
            msg += " " + "; ".join(str(e) for e in err_list[:3])
        return RedirectResponse(url=base_redirect + sep + "synced=1&msg=" + quote(msg), status_code=303)
    except Exception as e:
        err_msg = str(e)[:200] if e else "Noma'lum xato"
        traceback.print_exc()
        return RedirectResponse(url=base_redirect + sep + "error=" + quote("Hikvision yuklash: " + err_msg), status_code=303)


def _parse_time(s: str):
    """'09:00' yoki '09:00:00' dan time object qaytaradi, bo'sh bo'lsa None."""
    if not s or not str(s).strip():
        return None
    s = str(s).strip()
    for fmt in ("%H:%M", "%H:%M:%S"):
        try:
            from datetime import time as dt_time
            t = datetime.strptime(s, fmt).time()
            return t
        except ValueError:
            continue
    return None


@router.post("/attendance/form/bulk-time")
async def attendance_form_bulk_time(
    request: Request,
    date_param: str = Form(..., alias="date"),
    check_in_time: str = Form("09:00"),
    check_out_time: str = Form("18:00"),
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    """Barcha faol xodimlarga tanlangan kun uchun Keldi/Ketdi/Soat (9:00–18:00) yuklash va saqlash."""
    from urllib.parse import quote
    try:
        doc_date = datetime.strptime(date_param.strip(), "%Y-%m-%d").date()
    except (ValueError, AttributeError):
        return RedirectResponse(url="/attendance/form?error=" + quote("Noto'g'ri sana"), status_code=303)
    t_in = _parse_time((check_in_time or "09:00").strip())
    t_out = _parse_time((check_out_time or "18:00").strip())
    if not t_in:
        t_in = datetime.strptime("09:00", "%H:%M").time()
    if not t_out:
        t_out = datetime.strptime("18:00", "%H:%M").time()
    check_in_dt = datetime.combine(doc_date, t_in)
    check_out_dt = datetime.combine(doc_date, t_out)
    delta = check_out_dt - check_in_dt
    if delta.total_seconds() < 0:
        delta += timedelta(days=1)
    hours_worked = round(delta.total_seconds() / 3600 * 2) / 2
    form = await request.form()
    employee_ids_param = form.getlist("employee_ids")
    if employee_ids_param:
        try:
            emp_ids = [int(x) for x in employee_ids_param if str(x).strip().isdigit()]
        except (ValueError, TypeError):
            emp_ids = []
        employees = db.query(Employee).filter(Employee.id.in_(emp_ids), Employee.is_active == True).all() if emp_ids else []
    else:
        employees = db.query(Employee).filter(Employee.is_active == True).all()
    saved = 0
    for emp in employees:
        att = db.query(Attendance).filter(Attendance.employee_id == emp.id, Attendance.date == doc_date).first()
        if not att:
            att = Attendance(employee_id=emp.id, date=doc_date)
            db.add(att)
        att.check_in = check_in_dt
        att.check_out = check_out_dt
        att.hours_worked = hours_worked
        att.status = "present"
        saved += 1
    db.commit()
    msg = f"{saved} ta xodimga vaqt yuklandi (Keldi {check_in_time or '09:00'}, Ketdi {check_out_time or '18:00'})."
    return RedirectResponse(
        url=f"/attendance/form?date={doc_date.strftime('%Y-%m-%d')}&saved={saved}&msg=" + quote(msg),
        status_code=303,
    )


@router.post("/attendance/form/save")
async def attendance_form_save(
    request: Request,
    date_param: str = Form(..., alias="date"),
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    """Tabelni qo'lda to'ldirish — har bir xodim uchun Keldi/Ketdi/Soat/Holat/Izoh saqlanadi."""
    from urllib.parse import quote
    try:
        doc_date = datetime.strptime(date_param.strip(), "%Y-%m-%d").date()
    except (ValueError, AttributeError):
        return RedirectResponse(url="/attendance/form?error=" + quote("Noto'g'ri sana"), status_code=303)
    form = await request.form()
    # employee_ids = form.getlist("employee_id"), check_in_1=..., check_out_1=..., hours_1=..., status_1=..., note_1=...
    employee_ids = form.getlist("employee_id")
    saved = 0
    for i, emp_id_str in enumerate(employee_ids):
        try:
            emp_id = int(emp_id_str)
        except (ValueError, TypeError):
            continue
        emp = db.query(Employee).filter(Employee.id == emp_id).first()
        if not emp:
            continue
        check_in_str = (form.get(f"check_in_{emp_id}") or "").strip()
        check_out_str = (form.get(f"check_out_{emp_id}") or "").strip()
        hours_str = (form.get(f"hours_{emp_id}") or "").strip().replace(",", ".")
        status_val = (form.get(f"status_{emp_id}") or "").strip() or "present"
        note_val = (form.get(f"note_{emp_id}") or "").strip() or None
        if status_val not in ("present", "absent", "leave"):
            status_val = "present"
        if not check_in_str and not check_out_str:
            status_val = "absent"
        try:
            hours_worked = float(hours_str) if hours_str else None
        except ValueError:
            hours_worked = None
        check_in_time = _parse_time(check_in_str)
        check_out_time = _parse_time(check_out_str)
        check_in_dt = datetime.combine(doc_date, check_in_time) if check_in_time else None
        check_out_dt = datetime.combine(doc_date, check_out_time) if check_out_time else None
        if hours_worked is None and check_in_dt and check_out_dt:
            delta = check_out_dt - check_in_dt
            if delta.total_seconds() < 0:
                delta += timedelta(days=1)
            hours_worked = round(delta.total_seconds() / 3600 * 2) / 2
        att = db.query(Attendance).filter(Attendance.employee_id == emp_id, Attendance.date == doc_date).first()
        if not att:
            att = Attendance(employee_id=emp_id, date=doc_date)
            db.add(att)
        att.check_in = check_in_dt
        att.check_out = check_out_dt
        att.hours_worked = hours_worked if hours_worked is not None else (att.hours_worked if att.hours_worked is not None else 0)
        att.status = status_val
        att.note = note_val
        saved += 1
    db.commit()
    return RedirectResponse(
        url=f"/attendance/form?date={doc_date.strftime('%Y-%m-%d')}&saved={saved}&msg=" + quote("Tabel qo'lda saqlandi."),
        status_code=303,
    )


@router.post("/attendance/form/confirm")
async def attendance_form_confirm(
    request: Request,
    date_param: str = Form(..., alias="date"),
    db: Session = Depends(get_db),
    current_user: User = Depends(require_admin),
):
    """Kunni tasdiqlash — AttendanceDoc yaratiladi"""
    try:
        doc_date = datetime.strptime(date_param, "%Y-%m-%d").date()
    except ValueError:
        return RedirectResponse(url="/attendance/form?error=Noto'g'ri sana", status_code=303)
    existing = db.query(AttendanceDoc).filter(AttendanceDoc.date == doc_date).first()
    if existing:
        if existing.confirmed_at:
            return RedirectResponse(url="/attendance?already=1", status_code=303)
        existing.confirmed_at = datetime.now()
        existing.user_id = current_user.id
        db.commit()
        return RedirectResponse(url="/attendance?confirmed=1", status_code=303)
    count = db.query(AttendanceDoc).filter(AttendanceDoc.date >= doc_date.replace(day=1)).count()
    number = f"TBL-{doc_date.strftime('%Y%m%d')}-{count + 1:04d}"
    doc = AttendanceDoc(number=number, date=doc_date, user_id=current_user.id, confirmed_at=datetime.now())
    db.add(doc)
    db.commit()
    return RedirectResponse(url="/attendance?confirmed=1", status_code=303)


@router.get("/attendance/doc/{doc_id}", response_class=HTMLResponse)
async def attendance_doc_view(
    request: Request,
    doc_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    """Kunlik tabel hujjati ko'rinishi"""
    doc = db.query(AttendanceDoc).filter(AttendanceDoc.id == doc_id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Hujjat topilmadi")
    rows = db.query(Attendance).filter(Attendance.date == doc.date).order_by(Attendance.employee_id).all()
    return templates.TemplateResponse("employees/attendance_doc.html", {
        "request": request,
        "doc": doc,
        "rows": rows,
        "current_user": current_user,
        "page_title": f"Tabel {doc.number}",
    })


@router.get("/attendance/records", response_class=HTMLResponse)
async def attendance_records(
    request: Request,
    start_date: Optional[str] = None,
    end_date: Optional[str] = None,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    """Barcha davomat yozuvlari (sana oralig'i) — qo'lda qo'shish/tahrirlash"""
    today = date.today()
    start_date = start_date or (today - timedelta(days=7)).strftime("%Y-%m-%d")
    end_date = end_date or today.strftime("%Y-%m-%d")
    records = (
        db.query(Attendance)
        .filter(Attendance.date >= start_date, Attendance.date <= end_date)
        .order_by(Attendance.date.desc(), Attendance.employee_id)
        .all()
    )
    employees = db.query(Employee).filter(Employee.is_active == True).order_by(Employee.full_name).all()
    return templates.TemplateResponse("employees/attendance_records.html", {
        "request": request,
        "records": records,
        "employees": employees,
        "start_date": start_date,
        "end_date": end_date,
        "current_user": current_user,
        "page_title": "Davomat yozuvlari",
    })


@router.post("/attendance/doc/{doc_id}/delete")
async def attendance_doc_delete(
    doc_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_admin),
):
    """Tabel hujjatini ro'yxatdan butunlay o'chirish (AttendanceDoc jadvaldan o'chiriladi; davomat yozuvlari saqlanadi)."""
    doc = db.query(AttendanceDoc).filter(AttendanceDoc.id == doc_id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Hujjat topilmadi")
    db.delete(doc)
    db.commit()
    return RedirectResponse(url="/attendance?deleted=1", status_code=303)


@router.post("/attendance/doc/{doc_id}/cancel-confirm")
async def attendance_doc_cancel_confirm(
    doc_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_admin),
):
    """Tasdiqlashni bekor qilish"""
    doc = db.query(AttendanceDoc).filter(AttendanceDoc.id == doc_id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Hujjat topilmadi")
    doc.confirmed_at = None
    db.commit()
    return RedirectResponse(url="/attendance?unconfirmed=1", status_code=303)


@router.post("/attendance/records/add")
async def attendance_record_add(
    request: Request,
    employee_id: int = Form(...),
    att_date: str = Form(...),
    check_in: Optional[str] = Form(None),
    check_out: Optional[str] = Form(None),
    hours_worked: float = Form(0),
    note: str = Form(""),
    start_date: str = Form(""),
    end_date: str = Form(""),
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    """Davomat yozuvi qo'shish (qo'lda)"""
    try:
        att_d = datetime.strptime(att_date, "%Y-%m-%d").date()
    except ValueError:
        return RedirectResponse(url=f"/attendance/records?start_date={start_date}&end_date={end_date}&error=Noto'g'ri sana", status_code=303)
    check_in_dt = None
    if check_in:
        try:
            check_in_dt = datetime.strptime(f"{att_date} {check_in}", "%Y-%m-%d %H:%M")
        except ValueError:
            pass
    check_out_dt = None
    if check_out:
        try:
            check_out_dt = datetime.strptime(f"{att_date} {check_out}", "%Y-%m-%d %H:%M")
        except ValueError:
            pass
    att = Attendance(
        employee_id=employee_id,
        date=att_d,
        check_in=check_in_dt,
        check_out=check_out_dt,
        hours_worked=hours_worked or 0,
        status="present",
        note=note or None,
    )
    db.add(att)
    db.commit()
    return RedirectResponse(url=f"/attendance/records?start_date={start_date}&end_date={end_date}&added=1", status_code=303)


@router.get("/attendance/records/edit/{record_id}", response_class=HTMLResponse)
async def attendance_record_edit_page(
    request: Request,
    record_id: int,
    start_date: str = Query(""),
    end_date: str = Query(""),
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    """Davomat yozuvini tahrirlash sahifasi"""
    att = db.query(Attendance).filter(Attendance.id == record_id).first()
    if not att:
        raise HTTPException(status_code=404, detail="Yozuv topilmadi")
    return templates.TemplateResponse("employees/attendance_record_edit.html", {
        "request": request,
        "record": att,
        "start_date": start_date or att.date.strftime("%Y-%m-%d"),
        "end_date": end_date or att.date.strftime("%Y-%m-%d"),
        "current_user": current_user,
        "page_title": "Davomat yozuvini tahrirlash",
    })


@router.post("/attendance/records/edit/{record_id}")
async def attendance_record_edit_save(
    record_id: int,
    check_in: Optional[str] = Form(None),
    check_out: Optional[str] = Form(None),
    hours_worked: float = Form(None),
    status: str = Form("present"),
    note: str = Form(""),
    start_date: str = Form(""),
    end_date: str = Form(""),
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    """Davomat yozuvini saqlash"""
    att = db.query(Attendance).filter(Attendance.id == record_id).first()
    if not att:
        raise HTTPException(status_code=404, detail="Yozuv topilmadi")
    att_date_str = att.date.strftime("%Y-%m-%d")
    check_in_dt = None
    if check_in and str(check_in).strip():
        try:
            check_in_dt = datetime.strptime(f"{att_date_str} {check_in.strip()}", "%Y-%m-%d %H:%M")
        except ValueError:
            pass
    check_out_dt = None
    if check_out and str(check_out).strip():
        try:
            check_out_dt = datetime.strptime(f"{att_date_str} {check_out.strip()}", "%Y-%m-%d %H:%M")
        except ValueError:
            pass
    att.check_in = check_in_dt
    att.check_out = check_out_dt
    if hours_worked is not None:
        att.hours_worked = float(hours_worked)
    elif check_in_dt and check_out_dt:
        delta = check_out_dt - check_in_dt
        if delta.total_seconds() < 0:
            delta += timedelta(days=1)
        att.hours_worked = round(delta.total_seconds() / 3600 * 2) / 2
    if status in ("present", "absent", "leave"):
        att.status = status
    att.note = (note or "").strip() or None
    db.commit()
    return RedirectResponse(url=f"/attendance/records?start_date={start_date}&end_date={end_date}&updated=1", status_code=303)


@router.post("/attendance/records/bulk-time")
async def attendance_records_bulk_time(
    request: Request,
    start_date: str = Form(""),
    end_date: str = Form(""),
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    """Tanlangan yozuvlarga Keldi 9:00, Ketdi 18:00, Soat 9 qo'llash"""
    from urllib.parse import quote
    form = await request.form()
    record_ids_raw = form.getlist("record_ids")
    try:
        record_ids = [int(x) for x in record_ids_raw if str(x).strip().isdigit()]
    except (ValueError, TypeError):
        record_ids = []
    if not record_ids:
        return RedirectResponse(url=f"/attendance/records?start_date={start_date}&end_date={end_date}&error=" + quote("Hech qanday yozuv tanlanmagan"), status_code=303)
    check_in_dt = datetime.strptime("09:00", "%H:%M").time()
    check_out_dt = datetime.strptime("18:00", "%H:%M").time()
    hours_worked = 9.0
    updated = 0
    for rid in record_ids:
        att = db.query(Attendance).filter(Attendance.id == rid).first()
        if not att:
            continue
        att.check_in = datetime.combine(att.date, check_in_dt)
        att.check_out = datetime.combine(att.date, check_out_dt)
        att.hours_worked = hours_worked
        att.status = "present"
        updated += 1
    db.commit()
    return RedirectResponse(url=f"/attendance/records?start_date={start_date}&end_date={end_date}&updated={updated}&msg=" + quote("Vaqt yuklandi (9:00–18:00)."), status_code=303)


@router.post("/attendance/records/bulk-time-all")
async def attendance_records_bulk_time_all(
    start_date: str = Form(...),
    end_date: str = Form(...),
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    """Tanlangan sana oralig'idagi har bir kun uchun barcha faol xodimlarga Keldi 9:00, Ketdi 18:00, Soat 9 yuklash."""
    from urllib.parse import quote
    try:
        d_start = datetime.strptime(start_date.strip()[:10], "%Y-%m-%d").date()
        d_end = datetime.strptime(end_date.strip()[:10], "%Y-%m-%d").date()
    except (ValueError, TypeError):
        return RedirectResponse(url=f"/attendance/records?error=" + quote("Noto'g'ri sana"), status_code=303)
    if d_end < d_start:
        d_end = d_start
    check_in_t = datetime.strptime("09:00", "%H:%M").time()
    check_out_t = datetime.strptime("18:00", "%H:%M").time()
    hours_worked = 9.0
    employees = db.query(Employee).filter(Employee.is_active == True).all()
    saved = 0
    d = d_start
    while d <= d_end:
        check_in_dt = datetime.combine(d, check_in_t)
        check_out_dt = datetime.combine(d, check_out_t)
        for emp in employees:
            att = db.query(Attendance).filter(Attendance.employee_id == emp.id, Attendance.date == d).first()
            if not att:
                att = Attendance(employee_id=emp.id, date=d)
                db.add(att)
            att.check_in = check_in_dt
            att.check_out = check_out_dt
            att.hours_worked = hours_worked
            att.status = "present"
            saved += 1
        d += timedelta(days=1)
    db.commit()
    msg = quote(f"Barcha xodimlar uchun vaqt yuklandi: {saved} ta yozuv (9:00–18:00).")
    return RedirectResponse(url=f"/attendance/records?start_date={start_date}&end_date={end_date}&updated={saved}&msg={msg}", status_code=303)


@router.post("/attendance/records/delete/{record_id}")
async def attendance_record_delete(
    record_id: int,
    start_date: str = Form(""),
    end_date: str = Form(""),
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    """Davomat yozuvini o'chirish"""
    att = db.query(Attendance).filter(Attendance.id == record_id).first()
    if att:
        db.delete(att)
        db.commit()
    return RedirectResponse(url=f"/attendance/records?start_date={start_date}&end_date={end_date}", status_code=303)


@router.post("/attendance/records/bulk-delete")
async def attendance_records_bulk_delete(
    request: Request,
    start_date: str = Form(""),
    end_date: str = Form(""),
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    """Tanlangan davomat yozuvlarini o'chirish."""
    from urllib.parse import quote
    form = await request.form()
    record_ids_raw = form.getlist("record_ids")
    try:
        record_ids = [int(x) for x in record_ids_raw if str(x).strip().isdigit()]
    except (ValueError, TypeError):
        record_ids = []
    if not record_ids:
        return RedirectResponse(url=f"/attendance/records?start_date={start_date}&end_date={end_date}&error=" + quote("Hech qanday yozuv tanlanmagan."), status_code=303)
    deleted = 0
    for rid in record_ids:
        att = db.query(Attendance).filter(Attendance.id == rid).first()
        if att:
            db.delete(att)
            deleted += 1
    db.commit()
    msg = quote(f"Tanlangan {deleted} ta yozuv o'chirildi.")
    return RedirectResponse(url=f"/attendance/records?start_date={start_date}&end_date={end_date}&deleted={deleted}&msg={msg}", status_code=303)


# ==========================================
# AVANS BERISH
# ==========================================

def _advances_list_redirect_params(form_or_params, key_from="date_from", key_to="date_to"):
    """Filtr parametrlarini redirect URL ga qo'shish."""
    parts = []
    if hasattr(form_or_params, "get"):
        df, dt = form_or_params.get(key_from) or "", form_or_params.get(key_to) or ""
    else:
        df = form_or_params.get(key_from, "") or ""
        dt = form_or_params.get(key_to, "") or ""
    if (df or "").strip():
        parts.append("date_from=" + quote(str(df).strip()[:10]))
    if (dt or "").strip():
        parts.append("date_to=" + quote(str(dt).strip()[:10]))
    return "&".join(parts)


@router.get("/advances", response_class=HTMLResponse)
async def employee_advances_list(
    request: Request,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
    date_from: Optional[str] = None,
    date_to: Optional[str] = None,
):
    """Xodim avanslari ro'yxati — sana bo'yicha filtrlash."""
    q = db.query(EmployeeAdvance).options(
        joinedload(EmployeeAdvance.cash_register),
        joinedload(EmployeeAdvance.employee),
    ).order_by(EmployeeAdvance.advance_date.desc())
    if (date_from or "").strip():
        try:
            df = datetime.strptime(str(date_from).strip()[:10], "%Y-%m-%d").date()
            q = q.filter(EmployeeAdvance.advance_date >= df)
        except ValueError:
            pass
    if (date_to or "").strip():
        try:
            dt = datetime.strptime(str(date_to).strip()[:10], "%Y-%m-%d").date()
            q = q.filter(EmployeeAdvance.advance_date <= dt)
        except ValueError:
            pass
    advances = q.all()
    employees = db.query(Employee).filter(Employee.is_active == True).order_by(Employee.full_name).all()
    cash_registers = db.query(CashRegister).filter(CashRegister.is_active == True).order_by(CashRegister.name).all()
    default_date = date.today().strftime("%Y-%m-%d")
    filter_date_from = str(date_from or "").strip()[:10] if date_from else ""
    filter_date_to = str(date_to or "").strip()[:10] if date_to else ""
    return templates.TemplateResponse("employees/advances_list.html", {
        "request": request,
        "advances": advances,
        "employees": employees,
        "cash_registers": cash_registers,
        "default_date": default_date,
        "filter_date_from": filter_date_from,
        "filter_date_to": filter_date_to,
        "current_user": current_user,
        "page_title": "Avans berish",
    })


@router.post("/advances/add")
async def employee_advance_add(
    request: Request,
    employee_id: int = Form(...),
    amount: float = Form(...),
    advance_date: str = Form(...),
    cash_register_id: Optional[int] = Form(None),
    note: str = Form(""),
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    """Avans qo'shish; tanlangan kassadan chiqim yoziladi."""
    try:
        adv_date = datetime.strptime(advance_date, "%Y-%m-%d").date()
    except ValueError:
        return RedirectResponse(url="/advances?error=Noto'g'ri sana", status_code=303)
    if amount <= 0:
        return RedirectResponse(url="/advances?error=Summa 0 dan katta bo'lishi kerak", status_code=303)
    emp = db.query(Employee).filter(Employee.id == employee_id).first()
    if not emp:
        return RedirectResponse(url="/advances?error=Xodim topilmadi", status_code=303)
    cash = None
    if cash_register_id:
        cash = db.query(CashRegister).filter(CashRegister.id == cash_register_id, CashRegister.is_active == True).first()
    if not cash:
        return RedirectResponse(url="/advances?error=Kassani tanlang", status_code=303)
    today = datetime.now()
    # Status ustuni startup da ta'minlangan; bitta commit da saqlaymiz
    adv = EmployeeAdvance(
        employee_id=employee_id,
        amount=amount,
        advance_date=adv_date,
        cash_register_id=cash.id,
        note=note or None,
        confirmed_at=today,
    )
    db.add(adv)
    db.flush()
    pay_count = db.query(Payment).filter(Payment.created_at >= today.replace(hour=0, minute=0, second=0)).count()
    pay_number = f"PAY-{today.strftime('%Y%m%d')}-{pay_count + 1:04d}"
    emp_name = (emp.full_name or f"Xodim {employee_id}")[:100]
    db.add(Payment(
        number=pay_number,
        type="expense",
        cash_register_id=cash.id,
        partner_id=None,
        order_id=None,
        amount=amount,
        payment_type="cash",
        category="other",
        description=f"Avans: {emp_name}",
        user_id=current_user.id if current_user else None,
        status="confirmed",
    ))
    _sync_cash_balance(db, cash.id)
    try:
        db.commit()
    except Exception as e:
        db.rollback()
        err_msg = str(e).replace("'", " ").replace("%", "")[:200]
        return RedirectResponse(url="/advances?error=" + quote(f"Saqlash xatosi: {err_msg}"), status_code=303)
    # Filtrsiz qaytamiz — ro'yxat sana bo'yicha kamayishda, yangi avans birinchi qatorda
    return RedirectResponse(url="/advances?added=1", status_code=303)


@router.get("/advances/edit/{advance_id}", response_class=HTMLResponse)
async def employee_advance_edit_page(
    advance_id: int,
    request: Request,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    """Avansni tahrirlash sahifasi — faqat tasdiqlanmagan avanslar."""
    adv = db.query(EmployeeAdvance).options(
        joinedload(EmployeeAdvance.employee),
        joinedload(EmployeeAdvance.cash_register),
    ).filter(EmployeeAdvance.id == advance_id).first()
    if not adv:
        return RedirectResponse(url="/advances?error=Avans topilmadi", status_code=303)
    if adv.confirmed_at:
        return RedirectResponse(
            url="/advances?error=" + quote("Tasdiqlangan avansni tahrirlash mumkin emas. Avval tasdiqni bekor qiling."),
            status_code=303,
        )
    employees = db.query(Employee).filter(Employee.is_active == True).order_by(Employee.full_name).all()
    if adv.employee and not any(e.id == adv.employee_id for e in employees):
        employees = [adv.employee] + list(employees)
    cash_registers = db.query(CashRegister).filter(CashRegister.is_active == True).order_by(CashRegister.name).all()
    next_ids = (request.query_params.get("next_ids") or "").strip()
    next_count = len([x for x in next_ids.split(",") if x.strip()]) if next_ids else 0
    return templates.TemplateResponse("employees/advance_edit.html", {
        "request": request,
        "advance": adv,
        "employees": employees,
        "cash_registers": cash_registers,
        "current_user": current_user,
        "page_title": "Avansni tahrirlash",
        "next_ids": next_ids,
        "next_count": next_count,
    })


@router.post("/advances/edit/{advance_id}")
async def employee_advance_edit_save(
    advance_id: int,
    request: Request,
    employee_id: int = Form(...),
    amount: float = Form(...),
    advance_date: str = Form(...),
    cash_register_id: Optional[int] = Form(None),
    note: str = Form(""),
    next_ids: str = Form(""),
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    """Avansni saqlash (tahrirlash) — faqat tasdiqlanmagan avanslar."""
    adv = db.query(EmployeeAdvance).filter(EmployeeAdvance.id == advance_id).first()
    if not adv:
        return RedirectResponse(url="/advances?error=Avans topilmadi", status_code=303)
    if adv.confirmed_at:
        return RedirectResponse(
            url="/advances?error=" + quote("Tasdiqlangan avansni tahrirlash mumkin emas."),
            status_code=303,
        )
    try:
        adv_date = datetime.strptime(advance_date, "%Y-%m-%d").date()
    except ValueError:
        return RedirectResponse(url=f"/advances/edit/{advance_id}?error=Noto'g'ri sana", status_code=303)
    if amount <= 0:
        return RedirectResponse(url=f"/advances/edit/{advance_id}?error=Summa 0 dan katta bo'lishi kerak", status_code=303)
    emp = db.query(Employee).filter(Employee.id == employee_id).first()
    if not emp:
        return RedirectResponse(url=f"/advances/edit/{advance_id}?error=Xodim topilmadi", status_code=303)
    if cash_register_id:
        cash = db.query(CashRegister).filter(CashRegister.id == cash_register_id, CashRegister.is_active == True).first()
        adv.cash_register_id = cash.id if cash else adv.cash_register_id
    else:
        adv.cash_register_id = None
    adv.employee_id = employee_id
    adv.amount = amount
    adv.advance_date = adv_date
    adv.note = note or None
    adv.confirmed_at = datetime.now()  # Tahrirlashda saqlash = tasdiqlash
    db.commit()
    # Ketma-ket tahrirlash: keyingi avansga yo'naltirish
    next_param = (next_ids or "").strip()
    if next_param:
        rest = [x.strip() for x in next_param.split(",") if x.strip()]
        if rest:
            try:
                next_id = int(rest[0])
                remaining = ",".join(rest[1:])
                url = f"/advances/edit/{next_id}"
                if remaining:
                    url += "?next_ids=" + remaining
                return RedirectResponse(url=url, status_code=303)
            except (ValueError, TypeError):
                pass
    return RedirectResponse(url="/advances?edited=1", status_code=303)


@router.post("/advances/confirm/{advance_id}")
async def employee_advance_confirm(
    advance_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    """Avansni tasdiqlash"""
    adv = db.query(EmployeeAdvance).filter(EmployeeAdvance.id == advance_id).first()
    if not adv:
        return RedirectResponse(url="/advances?error=Avans topilmadi", status_code=303)
    adv.confirmed_at = datetime.now()
    db.commit()
    return RedirectResponse(url="/advances?confirmed=1", status_code=303)


@router.post("/advances/unconfirm/{advance_id}")
async def employee_advance_unconfirm(
    advance_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    """Avans tasdiqini bekor qilish"""
    adv = db.query(EmployeeAdvance).filter(EmployeeAdvance.id == advance_id).first()
    if not adv:
        return RedirectResponse(url="/advances?error=Avans topilmadi", status_code=303)
    adv.confirmed_at = None
    db.commit()
    return RedirectResponse(url="/advances?unconfirmed=1", status_code=303)


@router.post("/advances/delete/{advance_id}")
async def employee_advance_delete(
    advance_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    """Avansni ro'yxatdan o'chirish — faqat tasdiqlanmagan (tasdiq bekor qilingan) avanslar."""
    adv = db.query(EmployeeAdvance).filter(EmployeeAdvance.id == advance_id).first()
    if not adv:
        return RedirectResponse(url="/advances?error=Avans topilmadi", status_code=303)
    if adv.confirmed_at:
        return RedirectResponse(
            url="/advances?error=" + quote("Tasdiqlangan avansni o'chirish mumkin emas. Avval tasdiqni bekor qiling."),
            status_code=303,
        )
    db.delete(adv)
    db.commit()
    return RedirectResponse(url="/advances?deleted=1", status_code=303)


@router.post("/advances/bulk-edit", response_class=RedirectResponse)
async def employee_advances_bulk_edit(
    request: Request,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    """Tanlangan tasdiqlanmagan avanslarni ketma-ket tahrirlash — birinchisiga yo'naltiradi."""
    form = await request.form()
    raw = form.getlist("advance_ids")
    ids = []
    for x in raw:
        try:
            ids.append(int(x))
        except (TypeError, ValueError):
            pass
    if not ids:
        return RedirectResponse(url="/advances?error=" + quote("Hech qaysi avans tanlanmagan."), status_code=303)
    unconfirmed = (
        db.query(EmployeeAdvance.id)
        .filter(EmployeeAdvance.id.in_(ids), EmployeeAdvance.confirmed_at.is_(None))
        .order_by(EmployeeAdvance.id)
        .all()
    )
    unconfirmed_ids = [r[0] for r in unconfirmed]
    if not unconfirmed_ids:
        return RedirectResponse(url="/advances?error=" + quote("Tanlangan avanslar tasdiqlangan. Faqat tasdiqlanmagan avanslarni tahrirlash mumkin."), status_code=303)
    first_id = unconfirmed_ids[0]
    next_ids = unconfirmed_ids[1:]
    next_param = ",".join(str(i) for i in next_ids) if next_ids else ""
    url = f"/advances/edit/{first_id}"
    if next_param:
        url += "?next_ids=" + next_param
    return RedirectResponse(url=url, status_code=303)


@router.post("/advances/bulk-unconfirm", response_class=RedirectResponse)
async def employee_advances_bulk_unconfirm(
    request: Request,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    """Tanlangan tasdiqlangan avanslarning tasdiqini bekor qilish"""
    form = await request.form()
    raw = form.getlist("advance_ids")
    ids = []
    for x in raw:
        try:
            ids.append(int(x))
        except (TypeError, ValueError):
            pass
    if not ids:
        return RedirectResponse(url="/advances?error=" + quote("Hech qaysi avans tanlanmagan."), status_code=303)
    updated = db.query(EmployeeAdvance).filter(EmployeeAdvance.id.in_(ids), EmployeeAdvance.confirmed_at.isnot(None)).update({EmployeeAdvance.confirmed_at: None}, synchronize_session=False)
    db.commit()
    base = "/advances?bulk_unconfirmed=" + str(updated)
    extra = _advances_list_redirect_params(form)
    return RedirectResponse(url=base + ("&" + extra if extra else ""), status_code=303)


@router.post("/advances/bulk-confirm", response_class=RedirectResponse)
async def employee_advances_bulk_confirm(
    request: Request,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    """Tanlangan tasdiqlanmagan avanslarni tasdiqlash"""
    form = await request.form()
    raw = form.getlist("advance_ids")
    ids = []
    for x in raw:
        try:
            ids.append(int(x))
        except (TypeError, ValueError):
            pass
    if not ids:
        return RedirectResponse(url="/advances?error=" + quote("Hech qaysi avans tanlanmagan."), status_code=303)
    now = datetime.now()
    updated = db.query(EmployeeAdvance).filter(EmployeeAdvance.id.in_(ids), EmployeeAdvance.confirmed_at.is_(None)).update({EmployeeAdvance.confirmed_at: now}, synchronize_session=False)
    db.commit()
    base = "/advances?bulk_confirmed=" + str(updated)
    extra = _advances_list_redirect_params(form)
    return RedirectResponse(url=base + ("&" + extra if extra else ""), status_code=303)


@router.post("/advances/bulk-delete", response_class=RedirectResponse)
async def employee_advances_bulk_delete(
    request: Request,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    """Tanlangan tasdiqlanmagan avanslarni o'chirish (tasdiqlanganlarni o'chirish mumkin emas)."""
    form = await request.form()
    raw = form.getlist("advance_ids")
    ids = []
    for x in raw:
        try:
            ids.append(int(x))
        except (TypeError, ValueError):
            pass
    if not ids:
        base = "/advances?error=" + quote("Hech qaysi avans tanlanmagan.")
        extra = _advances_list_redirect_params(form)
        return RedirectResponse(url=base + ("&" + extra if extra else ""), status_code=303)
    deleted = db.query(EmployeeAdvance).filter(EmployeeAdvance.id.in_(ids), EmployeeAdvance.confirmed_at.is_(None)).delete(synchronize_session=False)
    db.commit()
    if ids and deleted == 0:
        base = "/advances?error=" + quote("Tanlangan avanslar tasdiqlangan. O'chirish uchun avval tasdiqni bekor qiling.")
    else:
        base = "/advances?bulk_deleted=" + str(deleted)
    extra = _advances_list_redirect_params(form)
    return RedirectResponse(url=base + ("&" + extra if extra else ""), status_code=303)


# ==========================================
# OYLIK HISOBLASH
# ==========================================

@router.get("/salary", response_class=HTMLResponse)
async def employee_salary_page(
    request: Request,
    year: Optional[int] = None,
    month: Optional[int] = None,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    """Oylik hisoblash — oy tanlash, xodimlar ro'yxati (base, bonus, deduction, avans, total)"""
    today = date.today()
    year = year or today.year
    month = month or today.month
    if not (1 <= month <= 12):
        month = today.month
    if year < 2020 or year > 2030:
        year = today.year
    # Faqat ishga qabul hujjati bor xodimlar ro'yxatda ko'rinadi
    hired_employee_ids = db.query(EmploymentDoc.employee_id).distinct().all()
    hired_ids = [r[0] for r in hired_employee_ids if r[0]]
    if not hired_ids:
        employees = []
    else:
        employees = (
            db.query(Employee)
            .filter(Employee.is_active == True, Employee.id.in_(hired_ids))
            .order_by(Employee.full_name)
            .all()
        )
    salaries = {s.employee_id: s for s in db.query(Salary).filter(Salary.year == year, Salary.month == month).all()}
    # Ishga qabul hujjatidagi oylik — avval tasdiqlangan, keyin har qanday oxirgi hujjat (qadoqlovchilar va b. uchun)
    emp_ids = [e.id for e in employees]
    latest_doc_salary = {}
    if emp_ids:
        # Har bir xodim uchun tasdiqlangan hujjatlar ichidan eng oxirgi (max doc_date) dagi oylik
        subq_conf = (
            db.query(EmploymentDoc.employee_id, func.max(EmploymentDoc.doc_date).label("max_date"))
            .filter(EmploymentDoc.employee_id.in_(emp_ids), EmploymentDoc.confirmed_at.isnot(None))
            .group_by(EmploymentDoc.employee_id)
        ).subquery()
        docs_confirmed = (
            db.query(EmploymentDoc.employee_id, EmploymentDoc.salary)
            .join(subq_conf, (EmploymentDoc.employee_id == subq_conf.c.employee_id) & (EmploymentDoc.doc_date == subq_conf.c.max_date))
            .filter(EmploymentDoc.employee_id.in_(emp_ids), EmploymentDoc.confirmed_at.isnot(None))
            .all()
        )
        for row in docs_confirmed:
            if (row.salary or 0) > 0:
                latest_doc_salary[row.employee_id] = float(row.salary)
        # Asos hali 0 bo'lgan xodimlar uchun: har bir xodimning eng oxirgi hujjatidagi oylik (max doc_date)
        missing = [eid for eid in emp_ids if eid not in latest_doc_salary]
        if missing:
            subq = (
                db.query(EmploymentDoc.employee_id, func.max(EmploymentDoc.doc_date).label("max_date"))
                .filter(EmploymentDoc.employee_id.in_(missing))
                .group_by(EmploymentDoc.employee_id)
            ).subquery()
            docs_latest = (
                db.query(EmploymentDoc.employee_id, EmploymentDoc.salary)
                .join(subq, (EmploymentDoc.employee_id == subq.c.employee_id) & (EmploymentDoc.doc_date == subq.c.max_date))
                .filter(EmploymentDoc.employee_id.in_(missing))
                .all()
            )
            for row in docs_latest:
                if (row.salary or 0) > 0:
                    latest_doc_salary[row.employee_id] = float(row.salary)
    # Avanslar (shu oy berilgan) — hisoblash uchun
    advance_sums = {}
    from calendar import monthrange
    _, last_day = monthrange(year, month)
    start_d = date(year, month, 1)
    end_d = date(year, month, last_day)
    for a in db.query(EmployeeAdvance).filter(
        EmployeeAdvance.advance_date >= start_d,
        EmployeeAdvance.advance_date <= end_d,
        EmployeeAdvance.confirmed_at.isnot(None),
    ).all():
        advance_sums[a.employee_id] = advance_sums.get(a.employee_id, 0) + (a.amount or 0)
    # Oylik turi xodimlar uchun: tabel bo'yicha ishlagan kunlar — kunlik = oylik/oydagi kunlar, asos = kunlik * ishlagan kun
    worked_days_by_emp = {}
    if emp_ids:
        try:
            worked_rows = (
                db.query(Attendance.employee_id, func.count(func.distinct(Attendance.date)).label("days"))
                .filter(
                    Attendance.employee_id.in_(emp_ids),
                    Attendance.date >= start_d,
                    Attendance.date <= end_d,
                    or_(
                        Attendance.status == "present",
                        Attendance.check_in.isnot(None),
                    ),
                )
                .group_by(Attendance.employee_id)
                .all()
            )
            for r in worked_rows:
                worked_days_by_emp[r.employee_id] = int(r.days or 0)
        except Exception:
            pass
    days_in_month = last_day
    # Bo'lak ish haqi: ishlab chiqarilgan miqdor * (bitta bo'lak narxi). Bitta stavka ishlatiladi (min), yig'indi emas.
    piecework_calculated = {}
    emp_by_id = {e.id: e for e in employees}
    group_member_ids = set()
    # Ishlab chiqarish guruhlari (qiyomchilar): operator ishi kunlik tabel bo'yicha kелgan a'zolar orasida teng bo'linadi
    production_groups = (
        db.query(ProductionGroup)
        .options(joinedload(ProductionGroup.members), joinedload(ProductionGroup.piecework_task))
        .filter(ProductionGroup.is_active == True, ProductionGroup.operator_id.in_(emp_ids) if emp_ids else False)
        .all()
    )
    for gr in production_groups:
        member_ids = [m.id for m in gr.members if m.id in emp_ids] if hasattr(gr, "members") and gr.members else []
        if not member_ids or gr.operator_id not in emp_ids:
            continue
        group_member_ids.update(member_ids)
        rate = float(gr.piecework_task.price_per_unit or 0) if gr.piecework_task else 0
        if rate <= 0:
            continue
        prod_list = (
            db.query(Production)
            .options(joinedload(Production.recipe))
            .filter(
                Production.operator_id == gr.operator_id,
                Production.status == "completed",
                func.date(Production.date) >= start_d,
                func.date(Production.date) <= end_d,
            )
            .all()
        )
        day_kg = {}
        for p in prod_list:
            if not getattr(gr, "include_qiyom", True) and is_qiyom_recipe(p.recipe):
                continue
            kg = (float(p.quantity or 0) * recipe_kg_per_unit(p.recipe)) if p.recipe else 0
            if kg <= 0:
                continue
            d = p.date.date() if hasattr(p.date, "date") else p.date
            day_kg[d] = day_kg.get(d, 0) + kg
        attendances = (
            db.query(Attendance.employee_id, Attendance.date, Attendance.status, Attendance.check_in)
            .filter(Attendance.employee_id.in_(member_ids), Attendance.date >= start_d, Attendance.date <= end_d)
            .all()
        )
        present_by_date = {}
        for row in attendances:
            d = row.date
            if d not in present_by_date:
                present_by_date[d] = set()
            if (row.status or "").strip() == "present" or (getattr(row, "check_in", None) is not None):
                present_by_date[d].add(row.employee_id)
        member_kg = {mid: 0.0 for mid in member_ids}
        for d, kg in day_kg.items():
            present = present_by_date.get(d, set()) & set(member_ids)
            cnt = len(present)
            if cnt <= 0:
                continue
            per_person = kg / cnt
            for mid in present:
                member_kg[mid] = member_kg.get(mid, 0) + per_person
        for mid in member_ids:
            piecework_calculated[mid] = member_kg.get(mid, 0) * rate
    # Har bir xodim uchun bitta bo'lak narxi (min stavka — bir xil ish uchun bitta narx)
    piece_rate_sum = {}
    if emp_ids:
        rows_rates = (
            db.query(employee_piecework_tasks.c.employee_id, func.min(PieceworkTask.price_per_unit).label("rate"))
            .join(PieceworkTask, PieceworkTask.id == employee_piecework_tasks.c.task_id)
            .filter(employee_piecework_tasks.c.employee_id.in_(emp_ids))
            .filter(PieceworkTask.price_per_unit > 0)
            .group_by(employee_piecework_tasks.c.employee_id)
            .all()
        )
        for eid, rate in rows_rates:
            piece_rate_sum[int(eid)] = float(rate or 0)
    # Bo'lak va Bo'lak+oylik xodimlar, lekin employee_piecework_tasks da yo'q: stavkani ishga qabul hujjatidagi piecework_task_ids dan olaymiz
    boalak_emp_ids = [e.id for e in employees if getattr(e, "salary_type", None) in ("bo'lak", "bo'lak_oylik")]
    if boalak_emp_ids:
        docs_with_tasks = (
            db.query(EmploymentDoc.employee_id, EmploymentDoc.piecework_task_ids)
            .filter(EmploymentDoc.employee_id.in_(boalak_emp_ids), EmploymentDoc.confirmed_at.isnot(None))
            .order_by(EmploymentDoc.doc_date.desc())
            .all()
        )
        for row in docs_with_tasks:
            eid = row.employee_id
            if piece_rate_sum.get(eid, 0) > 0:
                continue
            raw = (row.piecework_task_ids or "").strip()
            ids = [int(x) for x in raw.split(",") if x.strip().isdigit()] if raw else []
            if not ids:
                continue
            first_task = db.query(PieceworkTask).filter(PieceworkTask.id == ids[0], PieceworkTask.price_per_unit > 0).first()
            if first_task:
                piece_rate_sum[eid] = float(first_task.price_per_unit)
    boalak_employees = [e for e in employees if getattr(e, "salary_type", None) in ("bo'lak", "bo'lak_oylik") and piece_rate_sum.get(e.id, 0) > 0]
    emp_by_id = {e.id: e for e in employees}
    user_to_employee_id = {}
    for e in employees:
        if e.user_id:
            user_to_employee_id[e.user_id] = e.id
    # Bo'lak ish haqi: Operator bo'yicha ishlab chiqarishdagi kabi KG hisoblanadi (qiyom hisobga olinmaydi), keyin kg * bo'lak narxi
    if boalak_employees:
        productions_for_salary = (
            db.query(Production)
            .options(joinedload(Production.recipe))
            .filter(
                Production.status == "completed",
                func.date(Production.date) >= start_d,
                func.date(Production.date) <= end_d,
            )
            .all()
        )
        total_kg_by_emp_id = {}
        for p in productions_for_salary:
            if is_qiyom_recipe(p.recipe):
                continue
            kg = (float(p.quantity or 0) * recipe_kg_per_unit(p.recipe)) if p.recipe else 0
            if kg <= 0:
                continue
            emp_id = None
            if p.operator_id and p.operator_id in emp_by_id:
                emp_id = p.operator_id
            elif p.user_id and p.user_id in user_to_employee_id:
                emp_id = user_to_employee_id[p.user_id]
            if emp_id and emp_id not in group_member_ids:
                total_kg_by_emp_id[emp_id] = total_kg_by_emp_id.get(emp_id, 0) + kg
        for emp in boalak_employees:
            if emp.id in group_member_ids:
                continue
            total_kg = total_kg_by_emp_id.get(emp.id, 0)
            rate = piece_rate_sum.get(emp.id, 0)
            if total_kg > 0 and rate > 0:
                piecework_calculated[emp.id] = total_kg * rate
    rows = []
    for emp in employees:
        s = salaries.get(emp.id)
        piecework_amount = float(piecework_calculated.get(emp.id, 0) or 0)
        base_source = ""  # "oylik" | "bo'lak" — asos qaysi manbadan olingan
        # Guruh a'zosi (qiyomchilar): guruh bo'lak ulushi bo'yicha, asos = max(mehnat haqi, bo'lak)
        if emp.id in group_member_ids and emp.id in piecework_calculated:
            mehnat_haqi = float(latest_doc_salary.get(emp.id, 0) or 0) or float(emp.salary or 0)
            piece_total = piecework_amount
            base = max(mehnat_haqi, piece_total)
            base_source = "bo'lak" if piece_total >= mehnat_haqi and piece_total > 0 else "oylik"
        elif getattr(emp, "salary_type", None) == "bo'lak":
            base = piecework_amount
            base_source = "bo'lak" if piecework_amount > 0 else ""
        elif getattr(emp, "salary_type", None) == "bo'lak_oylik":
            mehnat_haqi = float(latest_doc_salary.get(emp.id, 0) or 0) or float(emp.salary or 0)
            piece_total = piecework_amount
            base = max(mehnat_haqi, piece_total)
            base_source = "bo'lak" if piece_total >= mehnat_haqi and piece_total > 0 else "oylik"
        else:
            base = (s.base_salary if s else 0) or (emp.salary or 0) or latest_doc_salary.get(emp.id, 0)
            if not base and emp.id in piecework_calculated:
                base = piecework_calculated[emp.id]
            base = float(base or 0)
            if getattr(emp, "salary_type", None) in ("oylik", "soatlik") or not getattr(emp, "salary_type", None):
                base_source = "oylik" if base > 0 else ""
        base = float(base or 0)
        # Hisoblangan oylik (tabel bo'yicha): Oylik turi — doim; Bo'lak+oylik / guruh a'zosi — faqat asos "oylikdan" bo'lsa
        calculated_base = None
        if days_in_month and days_in_month > 0:
            contract_monthly = float(latest_doc_salary.get(emp.id, 0) or 0) or float(emp.salary or 0)
            worked_days = worked_days_by_emp.get(emp.id, 0) or 0
            if getattr(emp, "salary_type", None) == "oylik":
                calculated_base = round((contract_monthly / days_in_month) * worked_days, 2)
            elif base_source == "oylik" and contract_monthly > 0:
                # Bo'lak+oylik yoki guruh a'zosi, asos oylikdan — tabel bo'yicha hisoblangan oylik
                calculated_base = round((contract_monthly / days_in_month) * worked_days, 2)
        amount_for_total = calculated_base if calculated_base is not None else base
        bonus = float(s.bonus if s and s.bonus is not None else 0) or 0
        deduction = float(s.deduction if s and s.deduction is not None else 0) or 0
        # Avans: avval shu oydagi tasdiqlangan avanslar yig'indisi (EmployeeAdvance), bo'sh bo'lsa saqlangan qiymat
        adv_ded = float(advance_sums.get(emp.id, 0) or 0)
        if adv_ded == 0 and s and getattr(s, "advance_deduction", None) is not None:
            adv_ded = float(s.advance_deduction)
        total = amount_for_total + bonus - deduction - adv_ded
        total = round(total, 2)
        paid = float(s.paid if s and s.paid is not None else 0) or 0
        status = (s.status if s else "pending") or "pending"
        if total == 0 and paid == 0:
            status = "pending"
        elif (total or 0) > 0 and paid >= total:
            status = "paid"
        elif (total or 0) > 0:
            status = "pending"
        rows.append({
            "employee": emp,
            "salary_row": s,
            "base_salary": base,
            "calculated_base": calculated_base,
            "piecework_amount": piecework_amount,
            "base_source": base_source,
            "bonus": bonus,
            "deduction": deduction,
            "advance_deduction": adv_ded,
            "total": total,
            "paid": paid,
            "status": status,
            "worked_days": worked_days_by_emp.get(emp.id, 0) or 0,
            "days_in_month": days_in_month,
        })
    cash_doc_id = request.query_params.get("cash_doc")
    try:
        cash_doc_id = int(cash_doc_id) if cash_doc_id else None
    except (TypeError, ValueError):
        cash_doc_id = None
    cash_register_id = request.query_params.get("cash_id")
    try:
        cash_register_id = int(cash_register_id) if cash_register_id else None
    except (TypeError, ValueError):
        cash_register_id = None
    expense_doc_id = request.query_params.get("expense_doc_id")
    try:
        expense_doc_id = int(expense_doc_id) if expense_doc_id else None
    except (TypeError, ValueError):
        expense_doc_id = None
    no_cash_warn = request.query_params.get("no_cash") == "1"
    cash_registers = db.query(CashRegister).filter(CashRegister.is_active == True).order_by(CashRegister.name).all()
    try:
        last_day = calendar.monthrange(year, month)[1]
        harajatlar_date_from = f"{year}-{month:02d}-01"
        harajatlar_date_to = f"{year}-{month:02d}-{last_day}"
    except (ValueError, TypeError):
        harajatlar_date_from = harajatlar_date_to = ""
    return templates.TemplateResponse("employees/salary_list.html", {
        "request": request,
        "year": year,
        "month": month,
        "rows": rows,
        "current_user": current_user,
        "page_title": "Oylik hisoblash",
        "cash_doc_id": cash_doc_id,
        "cash_register_id": cash_register_id,
        "expense_doc_id": expense_doc_id,
        "no_cash_warn": no_cash_warn,
        "cash_registers": cash_registers,
        "harajatlar_date_from": harajatlar_date_from,
        "harajatlar_date_to": harajatlar_date_to,
    })


@router.post("/salary/save")
async def employee_salary_save(
    request: Request,
    year: int = Form(...),
    month: int = Form(...),
    cash_register_id: Optional[int] = Form(None),
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    """Oylik yozuvlarini saqlash; tanlangan kassadan chiqim hujjati (Payment) va qoldiq hujjati (CashBalanceDoc) yaratiladi. Faqat ishga qabul qilingan xodimlar."""
    if not (1 <= month <= 12) or year < 2020 or year > 2030:
        return RedirectResponse(url="/salary?error=Noto'g'ri oy yoki yil", status_code=303)
    form = await request.form()
    hired_ids = [r[0] for r in db.query(EmploymentDoc.employee_id).distinct().all() if r[0]]
    if not hired_ids:
        employees = []
    else:
        employees = db.query(Employee).filter(Employee.is_active == True, Employee.id.in_(hired_ids)).all()
    total_payroll = 0.0
    for emp in employees:
        base = float(form.get(f"base_{emp.id}", 0) or 0)
        bonus = float(form.get(f"bonus_{emp.id}", 0) or 0)
        deduction = float(form.get(f"deduction_{emp.id}", 0) or 0)
        advance_deduction = float(form.get(f"advance_{emp.id}", 0) or 0)
        total = base + bonus - deduction - advance_deduction
        total_payroll += max(0, float(total))  # Kassadan faqat musbat to'lovlar chiqadi
        s = db.query(Salary).filter(Salary.employee_id == emp.id, Salary.year == year, Salary.month == month).first()
        if not s:
            s = Salary(employee_id=emp.id, year=year, month=month)
            db.add(s)
        s.base_salary = base
        s.bonus = bonus
        s.deduction = deduction
        s.advance_deduction = advance_deduction
        s.total = total
        if s.paid is None:
            s.paid = 0
        s.status = "paid" if (s.paid or 0) >= total else "pending"
    db.commit()
    # Oylik to'lovi — avval Harajatlar jurnalida hujjat (qoralama) yaratiladi; tasdiqlashdan keyin kassadan chiqim yoziladi
    expense_doc_id = None
    no_cash_warn = False
    if total_payroll > 0:
        cash = None
        if cash_register_id:
            cash = db.query(CashRegister).filter(CashRegister.id == int(cash_register_id), CashRegister.is_active == True).first()
        if not cash:
            cash = db.query(CashRegister).filter(CashRegister.is_active == True).order_by(CashRegister.id).first()
        if not cash:
            no_cash_warn = True
        else:
            # "ish haqqi" harajat turi
            et = db.query(ExpenseType).filter(ExpenseType.is_active == True, ExpenseType.name.ilike("%ish haqqi%")).first()
            if not et:
                et = db.query(ExpenseType).filter(ExpenseType.is_active == True).order_by(ExpenseType.id).first()
            if not et:
                # Harajat turi yo'q bo'lsa yaratamiz
                et = ExpenseType(name="ish haqqi", category="Ishlab chiqarish xarajatlari", is_active=True)
                db.add(et)
                db.flush()
            try:
                last_day = calendar.monthrange(year, month)[1]
                doc_date = datetime(year, month, last_day)
            except (ValueError, TypeError):
                doc_date = datetime.now()
            doc_number = _next_expense_doc_number(db)
            doc = ExpenseDoc(
                number=doc_number,
                date=doc_date,
                cash_register_id=cash.id,
                direction_id=None,
                department_id=None,
                status="draft",
                total_amount=total_payroll,
                payment_id=None,
                user_id=current_user.id if current_user else None,
            )
            db.add(doc)
            db.flush()
            db.add(ExpenseDocItem(
                expense_doc_id=doc.id,
                expense_type_id=et.id,
                amount=total_payroll,
                description=f"Oylik to'lovi {year}-{month:02d}",
            ))
            expense_doc_id = doc.id
            db.commit()
    params = f"year={year}&month={month}&saved=1"
    if expense_doc_id:
        params += f"&expense_doc_id={expense_doc_id}"
    if no_cash_warn:
        params += "&no_cash=1"
    return RedirectResponse(url=f"/salary?{params}", status_code=303)


@router.post("/salary/mark-paid/{employee_id}")
async def employee_salary_mark_paid(
    request: Request,
    employee_id: int,
    year: int = Form(...),
    month: int = Form(...),
    paid_amount: float = Form(...),
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    """Oylik to'langanligini belgilash"""
    s = db.query(Salary).filter(
        Salary.employee_id == employee_id,
        Salary.year == year,
        Salary.month == month,
    ).first()
    if not s:
        s = Salary(employee_id=employee_id, year=year, month=month, base_salary=0, total=0, paid=0)
        db.add(s)
    s.paid = paid_amount
    s.status = "paid" if paid_amount >= (s.total or 0) else "pending"
    db.commit()
    return RedirectResponse(url=f"/salary?year={year}&month={month}", status_code=303)


