
# --- Barcha importlar ---

from fastapi import FastAPI, Request, Depends, HTTPException, Form, Cookie, File, UploadFile, Query
from fastapi.staticfiles import StaticFiles
from fastapi.templating import Jinja2Templates
from fastapi.responses import HTMLResponse, RedirectResponse, FileResponse, Response, StreamingResponse, JSONResponse
from sqlalchemy.orm import Session, joinedload
from sqlalchemy import func, or_, and_, text
from sqlalchemy.exc import IntegrityError, OperationalError
from datetime import datetime, timedelta, date
import uvicorn
import base64
import barcode
from barcode.writer import ImageWriter
from PIL import Image
import os
import traceback
from typing import Optional, List
import openpyxl
import io
from urllib.parse import quote
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from app.models.database import (
    get_db, init_db, SessionLocal,
    User, Product, Category, Unit, Warehouse, Stock, StockMovement,
    WarehouseTransfer, WarehouseTransferItem,
    Partner, Order, OrderItem, Payment, CashRegister, CashTransfer, PosDraft,
    Recipe, RecipeItem, RecipeStage, Production, ProductionItem, ProductionStage, PRODUCTION_STAGE_NAMES, Machine, Employee, Salary,
    PieceworkTask,
    employee_piecework_tasks,
    Attendance, AttendanceDoc, EmployeeAdvance, EmploymentDoc, DismissalDoc,
    ProductionGroup, production_group_members,
    Agent, AgentLocation, Route, RoutePoint, Visit,
    Driver, DriverLocation, Delivery, PartnerLocation,
    Purchase, PurchaseItem, PurchaseExpense, Department, Direction, Region, Position,
    PriceType, ProductPrice, ProductPriceHistory,
    StockAdjustmentDoc, StockAdjustmentDocItem,
    CashBalanceDoc, CashBalanceDocItem,
    PartnerBalanceDoc, PartnerBalanceDocItem,
)
from app.utils.auth import (
    hash_password, get_user_from_token,
    generate_csrf_token, verify_csrf_token,
)
from app.utils.notifications import check_low_stock_and_notify, get_unread_count, get_user_notifications
from app.utils.user_scope import get_warehouses_for_user
from app.deps import get_current_user, require_auth, require_admin
from app.core import templates
from app.routes import auth as auth_routes
from app.routes import dashboard as dashboard_routes
from app.routes import home as home_routes
from app.routes import reports as reports_routes
from app.routes import info as info_routes

app = FastAPI(title="TOTLI HOLVA", description="Biznes boshqaruv tizimi", version="1.0")
app.mount("/static", StaticFiles(directory="app/static"), name="static")

# Routerlar (auth, dashboard, home, reports, info)
app.include_router(auth_routes.router)
app.include_router(home_routes.router)
app.include_router(reports_routes.router)
app.include_router(info_routes.router)
app.include_router(dashboard_routes.router)


# ==========================================
# 404 – sahifa topilmadi (HTML)
# ==========================================
_HTML_404 = """
<!DOCTYPE html>
<html lang="uz">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>404 - Sahifa topilmadi - TOTLI HOLVA</title>
    <link href="https://cdn.jsdelivr.net/npm/bootstrap@5.3.2/dist/css/bootstrap.min.css" rel="stylesheet">
</head>
<body class="bg-light d-flex align-items-center justify-content-center min-vh-100">
    <div class="text-center p-5">
        <h1 class="display-1 text-muted">404</h1>
        <h2 class="text-secondary">Sahifa topilmadi</h2>
        <p class="lead text-muted">So'ralgan sahifa mavjud emas yoki ko'chirilgan.</p>
        <a href="/" class="btn btn-success mt-3">Bosh sahifaga</a>
        <a href="/login" class="btn btn-outline-secondary mt-3 ms-2">Kirish</a>
    </div>
</body>
</html>
"""

_HTML_500 = """
<!DOCTYPE html>
<html lang="uz">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>500 - Server xatosi - TOTLI HOLVA</title>
    <link href="https://cdn.jsdelivr.net/npm/bootstrap@5.3.2/dist/css/bootstrap.min.css" rel="stylesheet">
</head>
<body class="bg-light d-flex align-items-center justify-content-center min-vh-100">
    <div class="text-center p-5">
        <h1 class="display-1 text-danger">500</h1>
        <h2 class="text-secondary">Server xatosi</h2>
        <p class="lead text-muted">Iltimos, keyinroq urinib ko'ring yoki administrator bilan bog'laning.</p>
        <a href="/" class="btn btn-success mt-3">Bosh sahifaga</a>
        <a href="/login" class="btn btn-outline-secondary mt-3 ms-2">Kirish</a>
    </div>
</body>
</html>
"""


@app.exception_handler(404)
async def not_found_handler(request: Request, exc):
    """404 – sahifa topilmadi (HTML)."""
    return HTMLResponse(content=_HTML_404, status_code=404)


# ==========================================
# GLOBAL FALLBACK - istalgan xatoda 500 o'rniga login (HTML)
# ==========================================
@app.middleware("http")
async def global_safe_middleware(request: Request, call_next):
    try:
        response = await call_next(request)
        try:
            response.headers["X-Server-Source"] = "pwp"
        except Exception:
            pass
        return response
    except (KeyboardInterrupt, SystemExit):
        raise
    except BaseException as e:
        tb = traceback.format_exc()
        traceback.print_exc()
        for _dir in [os.path.dirname(os.path.abspath(__file__)), os.getcwd(), r"C:\Users\ELYOR\.cursor\worktrees\business_system\pwp"]:
            try:
                if _dir:
                    log_path = os.path.join(_dir, "server_error.log")
                    with open(log_path, "a", encoding="utf-8") as f:
                        f.write("\n--- [global_safe] %s ---\n%s\n" % (datetime.now().isoformat(), tb))
                    break
            except Exception:
                continue
        try:
            path = (getattr(request, "url", None) and getattr(request.url, "path", None)) or getattr(request, "path", None) or "/"
        except Exception:
            path = "/"
        if path == "/login" or path == "/favicon.ico":
            r = JSONResponse(status_code=500, content={"detail": "Server xatosi"})
        else:
            try:
                accept = getattr(request, "headers", None) and (request.headers.get("accept") or "")
            except Exception:
                accept = ""
            if "text/html" in (accept or ""):
                # 500 sahifasini ko'rsatamiz, session o'chirilmaydi (logout hissi bermaslik)
                r = HTMLResponse(content=_HTML_500, status_code=500)
            else:
                r = JSONResponse(status_code=500, content={"detail": "Server xatosi"})
        try:
            r.headers["X-Server-Source"] = "pwp"
        except Exception:
            pass
        return r


# ==========================================
# CSRF MIDDLEWARE - POST/PUT/DELETE so'rovlarni himoya qilish
# ==========================================
@app.middleware("http")
async def csrf_middleware(request: Request, call_next):
    try:
        return await _csrf_middleware_impl(request, call_next)
    except (KeyboardInterrupt, SystemExit):
        raise
    except Exception as e:
        import traceback
        traceback.print_exc()
        return await call_next(request)


async def _csrf_middleware_impl(request: Request, call_next):
    from urllib.parse import parse_qs
    from starlette.requests import Request as StarletteRequest

    try:
        path = (getattr(request, "url", None) and getattr(request.url, "path", None)) or getattr(request, "path", None) or "/"
    except Exception:
        path = "/"
    method = (getattr(request, "method", None) or "GET")
    if not isinstance(method, str):
        method = "GET"
    method = method.upper()
    # GET, HEAD, OPTIONS da CSRF tekshiruvi yo'q
    if method in ("GET", "HEAD", "OPTIONS"):
        token = request.cookies.get("csrf_token")
        if not token:
            token = generate_csrf_token()
        try:
            setattr(request.state, "csrf_token", token)
        except Exception:
            pass
        response = await call_next(request)
        if not request.cookies.get("csrf_token"):
            response.set_cookie("csrf_token", token, path="/", httponly=False, samesite="lax", max_age=86400 * 7)
        return response

    # Himoyalanmaydigan yo'llar (API login, static, PWA location)
    if path in ("/login", "/api/agent/login", "/api/driver/login") or path.startswith("/static"):
        try:
            setattr(request.state, "csrf_token", request.cookies.get("csrf_token") or generate_csrf_token())
        except Exception:
            pass
        return await call_next(request)

    # Cookie dan yoki yangi token
    token = request.cookies.get("csrf_token")
    if not token:
        token = generate_csrf_token()
    try:
        setattr(request.state, "csrf_token", token)
    except Exception:
        pass

    # POST/PUT/PATCH/DELETE da token tekshirish
    received_token = request.headers.get("X-CSRF-Token")
    content_type = request.headers.get("content-type", "")
    if not received_token and "application/x-www-form-urlencoded" in content_type:
        body = await request.body()
        parsed = parse_qs(body.decode("utf-8", errors="replace"))
        received_token = (parsed.get("csrf_token") or [None])[0]
        async def receive():
            return {"type": "http.request", "body": body}
        request = StarletteRequest(request.scope, receive)
        try:
            setattr(request.state, "csrf_token", token)
        except Exception:
            pass
    elif "multipart/form-data" in content_type and not received_token:
        body = await request.body()
        # multipart dan csrf_token ni qidirish (name="csrf_token" dan keyingi qiymat)
        idx = body.find(b'name="csrf_token"')
        if idx != -1:
            start = body.find(b"\r\n\r\n", idx) + 4
            end = body.find(b"\r\n", start)
            if start != 3 and end != -1:
                received_token = body[start:end].decode("utf-8", errors="replace")
        async def receive():
            return {"type": "http.request", "body": body}
        request = StarletteRequest(request.scope, receive)
        try:
            setattr(request.state, "csrf_token", token)
        except Exception:
            pass

    if not verify_csrf_token(received_token, token):
        if "text/html" in request.headers.get("accept", ""):
            return RedirectResponse(url=f"/?error=csrf", status_code=303)
        return JSONResponse(status_code=403, content={"detail": "CSRF token noto'g'ri yoki yo'q"})

    response = await call_next(request)
    if not request.cookies.get("csrf_token"):
        response.set_cookie("csrf_token", token, path="/", httponly=False, samesite="lax", max_age=86400 * 7)
    return response


# ==========================================
# AUTH MIDDLEWARE - barcha sahifa va API ni himoyalash
# ==========================================
@app.middleware("http")
async def auth_middleware(request: Request, call_next):
    try:
        return await _auth_middleware_impl(request, call_next)
    except (KeyboardInterrupt, SystemExit):
        raise
    except BaseException as e:
        # ExceptionGroup va boshqa BaseException (traceback brauzerga chiqmasin)
        tb = traceback.format_exc()
        traceback.print_exc()
        for _dir in [os.path.dirname(os.path.abspath(__file__)), os.getcwd(), r"C:\Users\ELYOR\.cursor\worktrees\business_system\pwp"]:
            try:
                if _dir:
                    with open(os.path.join(_dir, "server_error.log"), "a", encoding="utf-8") as f:
                        f.write("\n--- [auth_middleware] %s ---\n%s\n" % (datetime.now().isoformat(), tb))
                    break
            except Exception:
                continue
        try:
            path = (getattr(request, "url", None) and getattr(request.url, "path", None)) or getattr(request, "path", None) or "/"
        except Exception:
            path = "/"
        if path == "/login" or path == "/favicon.ico":
            return JSONResponse(status_code=500, content={"detail": "Server xatosi"})
        try:
            accept = (request.headers.get("accept") or "") if getattr(request, "headers", None) else ""
        except Exception:
            accept = ""
        if "text/html" in accept:
            resp = RedirectResponse(url="/login?error=please_retry", status_code=303)
            try:
                resp.delete_cookie("session_token", path="/")
            except Exception:
                pass
            return resp
        return JSONResponse(status_code=500, content={"detail": "Server xatosi"})


async def _auth_middleware_impl(request: Request, call_next):
    path = (getattr(request, "url", None) and getattr(request.url, "path", None)) or getattr(request, "path", "/") or "/"
    method = (getattr(request, "method", None) or "GET").upper() if isinstance(getattr(request, "method", None), str) else "GET"
    # Login, logout, static, favicon, ping - himoya kerak emas
    if path in ("/login", "/logout", "/favicon.ico", "/ping"):
        return await call_next(request)
    if path.startswith("/static"):
        return await call_next(request)
    # Mobil/PWA agent va haydovchi API (alohida token bilan)
    if path in ("/api/agent/login", "/api/driver/login"):
        return await call_next(request)
    if (path == "/api/agent/location" or path == "/api/driver/location") and method == "POST":
        return await call_next(request)
    if path in ("/api/agent/orders", "/api/agent/partners"):
        return await call_next(request)
    # Session tekshiruvi
    token = request.cookies.get("session_token")
    if not token:
        if path.startswith("/api/"):
            return JSONResponse(status_code=401, content={"detail": "Login talab qilindi"})
        return RedirectResponse(url="/login", status_code=303)
    user_data = get_user_from_token(token)
    if not user_data:
        if path.startswith("/api/"):
            return JSONResponse(status_code=401, content={"detail": "Session muddati tugadi"})
        resp = RedirectResponse(url="/login", status_code=303)
        resp.delete_cookie("session_token")
        return resp
    db = SessionLocal()
    try:
        user = db.query(User).filter(User.id == user_data["user_id"]).first()
        if not user or not user.is_active:
            if path.startswith("/api/"):
                return JSONResponse(status_code=401, content={"detail": "Foydalanuvchi faol emas"})
            resp = RedirectResponse(url="/login", status_code=303)
            resp.delete_cookie("session_token")
            return resp
        return await call_next(request)
    finally:
        db.close()


# ==========================================
# AUTENTIFIKATSIYA — app.routes.auth da
# ==========================================

@app.exception_handler(403)
async def forbidden_handler(request: Request, exc: HTTPException):
    """403 - brauzer so'rovida bosh sahifaga yo'naltirish"""
    if "text/html" in request.headers.get("accept", ""):
        return RedirectResponse(url="/?error=admin_required", status_code=303)
    return JSONResponse(status_code=403, content={"detail": exc.detail})


@app.exception_handler(Exception)
async def debug_500_handler(request: Request, exc: Exception):
    """500 da: brauzer uchun login ga yo'naltirish, traceback konsolda va server_error.log da."""
    tb = traceback.format_exc()
    traceback.print_exc()
    for _dir in [os.path.dirname(os.path.abspath(__file__)), os.getcwd(), r"C:\Users\ELYOR\.cursor\worktrees\business_system\pwp"]:
        try:
            if _dir:
                with open(os.path.join(_dir, "server_error.log"), "a", encoding="utf-8") as f:
                    f.write("\n--- [exception_handler] %s ---\n%s\n" % (datetime.now().isoformat(), tb))
                break
        except Exception:
            continue
    try:
        path = (getattr(request, "url", None) and getattr(request.url, "path", None)) or getattr(request, "path", None) or "/"
    except Exception:
        path = "/"
    if path == "/login" or path == "/favicon.ico":
        return JSONResponse(status_code=500, content={"detail": "Server xatosi"})
    try:
        accept = (request.headers.get("accept") or "") if getattr(request, "headers", None) else ""
    except Exception:
        accept = ""
    if "text/html" in accept:
        resp = RedirectResponse(url="/login?error=please_retry", status_code=303)
        try:
            resp.delete_cookie("session_token", path="/")
        except Exception:
            pass
        return resp
    return JSONResponse(status_code=500, content={"detail": "Server xatosi"})


@app.get("/ping", include_in_schema=False)
async def ping():
    """Qaysi main.py ishlayotganini tekshirish (auth kerak emas)."""
    return {"ok": True, "main_py": os.path.abspath(__file__)}


@app.get("/favicon.ico", include_in_schema=False)
async def favicon():
    """Brauzer uchun favicon (logo) — 404 oldini olish"""
    try:
        root = os.path.dirname(os.path.abspath(__file__))
        favicon_path = os.path.join(root, "app", "static", "images", "logo.png")
        if os.path.isfile(favicon_path):
            return FileResponse(os.path.abspath(favicon_path), media_type="image/png")
    except Exception:
        pass
    return Response(status_code=204)


# ==========================================
# DASHBOARDS
# ==========================================

# Test route without authentication
@app.get("/test/dashboard/executive", response_class=HTMLResponse)
async def executive_dashboard_test(request: Request, db: Session = Depends(get_db), current_user: User = Depends(require_admin)):
    """Rahbariyat Dashboard - Test (fake data), faqat admin"""
    
    # Fake user for testing
    fake_user = {"username": "test", "role": "admin"}
    
    # Fake statistika
    stats = {
        'today_sales': 5500000,
        'sales_growth': 12.5,
        'today_orders': 45,
        'completed_orders': 38,
        'active_agents': 12,
        'total_agents': 15,
        'warehouse_value': 25000000,
        'low_stock_count': 3
    }
    
    # Fake 7 kunlik savdo
    sales_trend = {
        "labels": ["26.01", "27.01", "28.01", "29.01", "30.01", "31.01", "01.02"],
        "data": [4200000, 4500000, 4800000, 5100000, 5300000, 5200000, 5500000]
    }
    
    # Fake top mahsulotlar
    top_products = {
        "labels": ["Shokolad tort", "Medovik", "Napoleon", "Tiramisu", "Eclair"],
        "data": [150, 120, 100, 85, 70]
    }
    
    # Fake top agentlar
    top_agents = [
        {'name': 'Alisher Karimov', 'sales': 1200000, 'orders': 25},
        {'name': 'Dilshod Rahimov', 'sales': 980000, 'orders': 20},
        {'name': 'Sardor Usmonov', 'sales': 850000, 'orders': 18},
        {'name': 'Jasur Toshmatov', 'sales': 720000, 'orders': 15},
        {'name': 'Bobur Sharipov', 'sales': 650000, 'orders': 12}
    ]
    
    # Fake ogohlantirishlar
    alerts = [
        {
            'title': 'Past qoldiq',
            'message': '3 ta mahsulot qoldig\'i past darajada'
        },
        {
            'title': 'Yangi buyurtma',
            'message': 'Toshkent filialidan yangi buyurtma keldi'
        }
    ]
    
    return templates.TemplateResponse("dashboards/executive.html", {
        "request": request,
        "page_title": "Rahbariyat Dashboard",
        "current_user": None,
        "user": fake_user,
        "stats": stats,
        "sales_trend": sales_trend,
        "top_products": top_products,
        "top_agents": top_agents,
        "alerts": alerts
    })


@app.get("/dashboard/executive", response_class=HTMLResponse)
async def executive_dashboard(request: Request, db: Session = Depends(get_db), current_user: User = Depends(require_auth)):
    """Rahbariyat Dashboard - Real Data"""
    from datetime import datetime, timedelta
    from sqlalchemy import func
    from app.models.database import Order, OrderItem, Agent, Stock, Product
    
    if not current_user:
        return RedirectResponse(url="/login", status_code=303)
    
    # Bugungi sana
    today = datetime.now().date()
    yesterday = today - timedelta(days=1)
    week_ago = today - timedelta(days=7)
    
    # Bugungi savdo (completed orders)
    today_sales = db.query(func.sum(Order.total)).filter(
        func.date(Order.created_at) == today,
        Order.status == 'completed'
    ).scalar() or 0
    
    # Kechagi savdo
    yesterday_sales = db.query(func.sum(Order.total)).filter(
        func.date(Order.created_at) == yesterday,
        Order.status == 'completed'
    ).scalar() or 0
    
    # O'sish foizi
    sales_growth = 0
    if yesterday_sales > 0:
        sales_growth = ((today_sales - yesterday_sales) / yesterday_sales) * 100
    
    # Bugungi buyurtmalar
    today_orders = db.query(func.count(Order.id)).filter(
        func.date(Order.created_at) == today
    ).scalar() or 0
    
    # Bajarilgan buyurtmalar
    completed_orders = db.query(func.count(Order.id)).filter(
        func.date(Order.created_at) == today,
        Order.status == 'completed'
    ).scalar() or 0
    
    # Faol agentlar
    active_agents = db.query(func.count(Agent.id)).filter(
        Agent.is_active == True
    ).scalar() or 0
    
    # Jami agentlar
    total_agents = db.query(func.count(Agent.id)).scalar() or 0
    
    # Ombor qiymati
    warehouse_value = db.query(
        func.sum(Stock.quantity * Product.purchase_price)
    ).join(
        Product, Stock.product_id == Product.id
    ).scalar() or 0
    
    # Past qoldiq mahsulotlar
    low_stock_count = db.query(func.count(Stock.id)).filter(
        Stock.quantity < 10
    ).scalar() or 0
    
    # 7 kunlik savdo dinamikasi
    sales_trend_labels = []
    sales_trend_data = []
    for i in range(6, -1, -1):
        date = today - timedelta(days=i)
        sales = db.query(func.sum(Order.total)).filter(
            func.date(Order.created_at) == date,
            Order.status == 'completed'
        ).scalar() or 0
        sales_trend_labels.append(date.strftime('%d.%m'))
        sales_trend_data.append(float(sales))
    
    # Top 5 mahsulotlar
    top_products_query = db.query(
        Product.name,
        func.sum(OrderItem.quantity).label('total_qty')
    ).join(
        OrderItem, Product.id == OrderItem.product_id
    ).join(
        Order, OrderItem.order_id == Order.id
    ).filter(
        func.date(Order.created_at) >= week_ago,
        Order.status == 'completed'
    ).group_by(Product.id, Product.name).order_by(
        func.sum(OrderItem.quantity).desc()
    ).limit(5).all()
    
    top_products_labels = [p.name for p in top_products_query] or ['Ma\'lumot yo\'q']
    top_products_data = [float(p.total_qty) for p in top_products_query] or [0]
    
    # Top 5 agentlar
    top_agents_query = db.query(
        Agent.name,
        func.sum(Order.total).label('total_sales'),
        func.count(Order.id).label('order_count')
    ).join(
        Order, Agent.id == Order.partner_id  # Assuming agent is partner
    ).filter(
        func.date(Order.created_at) >= week_ago,
        Order.status == 'completed'
    ).group_by(Agent.id, Agent.name).order_by(
        func.sum(Order.total).desc()
    ).limit(5).all()
    
    top_agents = [
        {
            'name': a.name,
            'sales': float(a.total_sales or 0),
            'orders': a.order_count
        }
        for a in top_agents_query
    ] or [{'name': 'Ma\'lumot yo\'q', 'sales': 0, 'orders': 0}]
    
    # Ogohlantirishlar
    alerts = []
    
    # Past qoldiq ogohlantirishlari
    if low_stock_count > 0:
        alerts.append({
            'title': 'Past qoldiq',
            'message': f'{low_stock_count} ta mahsulot qoldig\'i past darajada'
        })
    
    # Bugungi savdo past bo'lsa
    if yesterday_sales > 0 and sales_growth < -10:
        alerts.append({
            'title': 'Savdo pasaygan',
            'message': f'Bugungi savdo kechaga nisbatan {abs(sales_growth):.1f}% kamaygan'
        })
    
    # Statistika
    stats = {
        'today_sales': float(today_sales),
        'sales_growth': round(sales_growth, 1),
        'today_orders': today_orders,
        'completed_orders': completed_orders,
        'active_agents': active_agents,
        'total_agents': total_agents,
        'warehouse_value': float(warehouse_value),
        'low_stock_count': low_stock_count
    }
    
    return templates.TemplateResponse("dashboards/executive.html", {
        "request": request,
        "page_title": "Rahbariyat Dashboard",
        "current_user": current_user,
        "user": current_user,
        "stats": stats,
        "sales_trend": {
            "labels": sales_trend_labels,
            "data": sales_trend_data
        },
        "top_products": {
            "labels": top_products_labels,
            "data": top_products_data
        },
        "top_agents": top_agents,
        "alerts": alerts
    })


# Executive Dashboard - Export to Excel
@app.get("/dashboard/executive/export")
async def executive_export(request: Request, db: Session = Depends(get_db), current_user: User = Depends(require_auth)):
    """Export Executive Dashboard to Excel"""
    return await export_executive_dashboard(request, db)


# Live Data Endpoints
@app.get("/dashboard/executive/live")
async def executive_live(request: Request, db: Session = Depends(get_db)):
    """Live data for Executive Dashboard"""
    return await executive_live_data(request, db)

@app.get("/dashboard/warehouse/live")
async def warehouse_live(request: Request, db: Session = Depends(get_db)):
    """Live data for Warehouse Dashboard"""
    return await warehouse_live_data(request, db)

@app.get("/dashboard/delivery/live")
async def delivery_live(request: Request, db: Session = Depends(get_db)):
    """Live data for Delivery Dashboard"""
    return await delivery_live_data(request, db)


# Sales Dashboard - Real Data
@app.get("/dashboard/sales", response_class=HTMLResponse)
async def sales_dashboard(request: Request, db: Session = Depends(get_db), current_user: User = Depends(require_auth)):
    """Savdo Dashboard - Real Data"""
    from datetime import datetime, timedelta
    from sqlalchemy import func
    from app.models.database import Order, Partner
    
    if not current_user:
        return RedirectResponse(url="/login", status_code=303)
    
    today = datetime.now().date()
    yesterday = today - timedelta(days=1)
    week_ago = today - timedelta(days=7)
    month_ago = today - timedelta(days=30)
    
    # Today's sales
    today_sales = db.query(func.sum(Order.total)).filter(
        func.date(Order.created_at) == today,
        Order.status == 'completed'
    ).scalar() or 0
    
    yesterday_sales = db.query(func.sum(Order.total)).filter(
        func.date(Order.created_at) == yesterday,
        Order.status == 'completed'
    ).scalar() or 0
    
    sales_growth = 0
    if yesterday_sales > 0:
        sales_growth = ((today_sales - yesterday_sales) / yesterday_sales) * 100
    
    # Orders
    total_orders = db.query(func.count(Order.id)).filter(
        func.date(Order.created_at) == today
    ).scalar() or 0
    
    completed_orders = db.query(func.count(Order.id)).filter(
        func.date(Order.created_at) == today,
        Order.status == 'completed'
    ).scalar() or 0
    
    # Customers
    active_customers = db.query(func.count(func.distinct(Order.partner_id))).filter(
        func.date(Order.created_at) >= month_ago
    ).scalar() or 0
    
    new_customers = db.query(func.count(func.distinct(Order.partner_id))).filter(
        func.date(Order.created_at) >= week_ago
    ).scalar() or 0
    
    # Average check
    avg_check = today_sales / total_orders if total_orders > 0 else 0
    
    metrics = {
        'today_sales': float(today_sales),
        'sales_growth': round(sales_growth, 1),
        'total_orders': total_orders,
        'completed_orders': completed_orders,
        'active_customers': active_customers,
        'new_customers': new_customers,
        'avg_check': float(avg_check)
    }
    
    # Order Status
    status_counts = db.query(
        Order.status,
        func.count(Order.id)
    ).filter(
        func.date(Order.created_at) >= week_ago
    ).group_by(Order.status).all()
    
    status_map = {'draft': 'Yangi', 'confirmed': 'Jarayonda', 'completed': 'Bajarilgan', 'cancelled': 'Bekor qilingan'}
    order_status = {
        "labels": [status_map.get(s[0], s[0]) for s in status_counts] or ['Ma\'lumot yo\'q'],
        "data": [s[1] for s in status_counts] or [0]
    }
    
    # Weekly Sales
    weekly_labels = []
    weekly_data = []
    for i in range(6, -1, -1):
        date = today - timedelta(days=i)
        sales = db.query(func.sum(Order.total)).filter(
            func.date(Order.created_at) == date,
            Order.status == 'completed'
        ).scalar() or 0
        weekly_labels.append(['Yak', 'Dush', 'Sesh', 'Chor', 'Pay', 'Juma', 'Shan'][date.weekday()])
        weekly_data.append(float(sales))
    
    weekly_sales = {"labels": weekly_labels, "data": weekly_data}
    
    # Recent Orders
    recent = db.query(Order, Partner.name).join(
        Partner, Order.partner_id == Partner.id, isouter=True
    ).filter(
        func.date(Order.created_at) >= week_ago
    ).order_by(Order.created_at.desc()).limit(5).all()
    
    status_text_map = {'draft': 'Yangi', 'confirmed': 'Jarayonda', 'completed': 'Bajarilgan', 'cancelled': 'Bekor qilingan'}
    recent_orders = [
        {
            'number': o.number,
            'customer': p_name or 'Noma\'lum',
            'total': float(o.total),
            'status': o.status,
            'status_text': status_text_map.get(o.status, o.status)
        }
        for o, p_name in recent
    ] or [{'number': '-', 'customer': 'Ma\'lumot yo\'q', 'total': 0, 'status': 'draft', 'status_text': '-'}]
    
    # Top Customers
    top = db.query(
        Partner.name,
        func.count(Order.id).label('order_count'),
        func.sum(Order.total).label('total_sales')
    ).join(
        Order, Partner.id == Order.partner_id
    ).filter(
        func.date(Order.created_at) >= month_ago,
        Order.status == 'completed'
    ).group_by(Partner.id, Partner.name).order_by(
        func.sum(Order.total).desc()
    ).limit(5).all()
    
    top_customers = [
        {'name': t.name, 'orders': t.order_count, 'total': float(t.total_sales)}
        for t in top
    ] or [{'name': 'Ma\'lumot yo\'q', 'orders': 0, 'total': 0}]
    
    # Fake funnel (not in database yet)
    funnel = {
        "labels": ["Tashrif", "Qiziqish", "Taklif", "Buyurtma", "To'lov"],
        "data": [250, 180, 120, total_orders, completed_orders]
    }
    
    return templates.TemplateResponse("dashboards/sales.html", {
        "request": request,
        "page_title": "Savdo Dashboard",
        "current_user": current_user,
        "user": current_user,
        "metrics": metrics,
        "order_status": order_status,
        "funnel": funnel,
        "weekly_sales": weekly_sales,
        "recent_orders": recent_orders,
        "top_customers": top_customers
    })


# Sales Dashboard - Test (fake data)
@app.get("/test/dashboard/sales", response_class=HTMLResponse)
async def sales_dashboard_test(request: Request, db: Session = Depends(get_db), current_user: User = Depends(require_admin)):
    """Savdo Dashboard - Test (fake data), faqat admin"""
    
    # Fake user
    fake_user = {"username": "test", "role": "sales"}
    
    # Metrics
    metrics = {
        'today_sales': 5500000,
        'sales_growth': 15.3,
        'total_orders': 45,
        'completed_orders': 38,
        'active_customers': 128,
        'new_customers': 12,
        'avg_check': 122222
    }
    
    # Order Status
    order_status = {
        "labels": ["Yangi", "Jarayonda", "Bajarilgan", "Bekor qilingan"],
        "data": [7, 12, 38, 3]
    }
    
    # Sales Funnel
    funnel = {
        "labels": ["Tashrif", "Qiziqish", "Taklif", "Buyurtma", "To'lov"],
        "data": [250, 180, 120, 60, 45]
    }
    
    # Weekly Sales
    weekly_sales = {
        "labels": ["Dush", "Sesh", "Chor", "Pay", "Juma", "Shan", "Yak"],
        "data": [720000, 850000, 920000, 880000, 1100000, 950000, 780000]
    }
    
    # Recent Orders
    recent_orders = [
        {'number': 'ORD-1234', 'customer': 'Anvar Toshmatov', 'total': 450000, 'status': 'completed', 'status_text': 'Bajarilgan'},
        {'number': 'ORD-1235', 'customer': 'Dilshod Karimov', 'total': 320000, 'status': 'processing', 'status_text': 'Jarayonda'},
        {'number': 'ORD-1236', 'customer': 'Sardor Usmonov', 'total': 180000, 'status': 'new', 'status_text': 'Yangi'},
        {'number': 'ORD-1237', 'customer': 'Jasur Rahimov', 'total': 520000, 'status': 'completed', 'status_text': 'Bajarilgan'},
        {'number': 'ORD-1238', 'customer': 'Bobur Sharipov', 'total': 280000, 'status': 'processing', 'status_text': 'Jarayonda'}
    ]
    
    # Top Customers
    top_customers = [
        {'name': 'Anvar Toshmatov', 'orders': 45, 'total': 5200000},
        {'name': 'Dilshod Karimov', 'orders': 38, 'total': 4100000},
        {'name': 'Sardor Usmonov', 'orders': 32, 'total': 3500000},
        {'name': 'Jasur Rahimov', 'orders': 28, 'total': 2900000},
        {'name': 'Bobur Sharipov', 'orders': 25, 'total': 2400000}
    ]
    
    return templates.TemplateResponse("dashboards/sales.html", {
        "request": request,
        "page_title": "Savdo Dashboard",
        "user": fake_user,
        "metrics": metrics,
        "order_status": order_status,
        "funnel": funnel,
        "weekly_sales": weekly_sales,
        "recent_orders": recent_orders,
        "top_customers": top_customers
    })


# Agent Dashboard - Real Data
@app.get("/dashboard/agent", response_class=HTMLResponse)
async def agent_dashboard(request: Request, db: Session = Depends(get_db), current_user: User = Depends(require_auth)):
    """Agent Dashboard - Real Data"""
    from datetime import datetime, timedelta
    from sqlalchemy import func
    from app.models.database import Agent, Visit, Route, RoutePoint, Partner, Order, AgentLocation
    
    if not current_user:
        return RedirectResponse(url="/login", status_code=303)
    
    # Get agent for current user (assuming user has agent_id or we use first agent)
    agent = db.query(Agent).filter(Agent.is_active == True).first()
    if not agent:
        # No agent found - show empty dashboard
        agent = {'name': 'Agent topilmadi', 'location': '-'}
        return templates.TemplateResponse("dashboards/agent.html", {
            "request": request,
            "page_title": "Agent Dashboard",
            "current_user": current_user,
            "user": current_user,
            "agent": agent,
            "kpi": {'visits_completed': 0, 'visits_total': 0, 'visits_percent': 0, 'today_sales': 0, 'orders': 0, 'orders_completed': 0, 'target_achieved': 0, 'target_total': 25000000, 'target_percent': 0},
            "schedule": [],
            "recent_orders": [],
            "customers": [],
            "performance": {'labels': [], 'sales': [], 'target': []}
        })
    
    today = datetime.now().date()
    month_ago = today - timedelta(days=30)
    
    # Agent info with location
    latest_location = db.query(AgentLocation).filter(
        AgentLocation.agent_id == agent.id
    ).order_by(AgentLocation.recorded_at.desc()).first()
    
    agent_info = {
        'name': agent.full_name,
        'location': latest_location.address if latest_location and latest_location.address else agent.region or 'Noma\'lum'
    }
    
    # Today's visits
    today_visits = db.query(func.count(Visit.id)).filter(
        Visit.agent_id == agent.id,
        func.date(Visit.visit_date) == today
    ).scalar() or 0
    
    completed_visits = db.query(func.count(Visit.id)).filter(
        Visit.agent_id == agent.id,
        func.date(Visit.visit_date) == today,
        Visit.status == 'visited'
    ).scalar() or 0
    
    visits_percent = int((completed_visits / today_visits * 100)) if today_visits > 0 else 0
    
    # Today's sales (orders created by agent)
    today_sales = db.query(func.sum(Order.total)).filter(
        func.date(Order.created_at) == today,
        Order.status == 'completed'
    ).scalar() or 0
    
    today_orders = db.query(func.count(Order.id)).filter(
        func.date(Order.created_at) == today
    ).scalar() or 0
    
    completed_orders = db.query(func.count(Order.id)).filter(
        func.date(Order.created_at) == today,
        Order.status == 'completed'
    ).scalar() or 0
    
    # Monthly target (placeholder)
    target_total = 25000000
    month_sales = db.query(func.sum(Order.total)).filter(
        func.date(Order.created_at) >= month_ago,
        Order.status == 'completed'
    ).scalar() or 0
    
    target_percent = int((month_sales / target_total * 100)) if target_total > 0 else 0
    
    kpi = {
        'visits_completed': completed_visits,
        'visits_total': today_visits,
        'visits_percent': visits_percent,
        'today_sales': float(today_sales),
        'orders': today_orders,
        'orders_completed': completed_orders,
        'target_achieved': float(month_sales),
        'target_total': target_total,
        'target_percent': target_percent
    }
    
    # Today's schedule from visits
    schedule_visits = db.query(Visit, Partner).join(
        Partner, Visit.partner_id == Partner.id
    ).filter(
        Visit.agent_id == agent.id,
        func.date(Visit.visit_date) == today
    ).order_by(Visit.check_in_time).all()
    
    schedule = []
    for visit, partner in schedule_visits:
        schedule.append({
            'customer': partner.name,
            'time': visit.check_in_time.strftime('%H:%M') if visit.check_in_time else '-',
            'address': partner.address or '-',
            'completed': visit.status == 'visited'
        })
    
    if not schedule:
        schedule = [{'customer': 'Bugun tashrif rejalashtirilmagan', 'time': '-', 'address': '-', 'completed': False}]
    
    # Recent orders
    recent = db.query(Order, Partner).join(
        Partner, Order.partner_id == Partner.id
    ).filter(
        func.date(Order.created_at) >= today - timedelta(days=7)
    ).order_by(Order.created_at.desc()).limit(5).all()
    
    status_map = {'draft': ('Yangi', 'primary'), 'confirmed': ('Jarayonda', 'warning'), 'completed': ('Bajarilgan', 'success'), 'cancelled': ('Bekor qilingan', 'danger')}
    recent_orders = []
    for order, partner in recent:
        status_text, status_color = status_map.get(order.status, ('Noma\'lum', 'secondary'))
        recent_orders.append({
            'number': order.number,
            'customer': partner.name,
            'total': float(order.total),
            'status_color': status_color,
            'status_text': status_text
        })
    
    if not recent_orders:
        recent_orders = [{'number': '-', 'customer': 'Ma\'lumot yo\'q', 'total': 0, 'status_color': 'secondary', 'status_text': '-'}]
    
    # My customers (partners with recent orders)
    customers_query = db.query(
        Partner,
        func.max(Order.total).label('last_order')
    ).join(
        Order, Partner.id == Order.partner_id
    ).filter(
        func.date(Order.created_at) >= month_ago
    ).group_by(Partner.id).order_by(func.max(Order.created_at).desc()).limit(5).all()
    
    customers = []
    for partner, last_order in customers_query:
        customers.append({
            'name': partner.name,
            'phone': partner.phone or '-',
            'address': partner.address or '-',
            'last_order': float(last_order) if last_order else 0
        })
    
    if not customers:
        customers = [{'name': 'Ma\'lumot yo\'q', 'phone': '-', 'address': '-', 'last_order': 0}]
    
    # 30-day performance
    performance_labels = []
    performance_sales = []
    performance_target = []
    
    daily_target = target_total / 30
    cumulative_sales = 0
    
    for i in range(0, 30, 5):
        date = month_ago + timedelta(days=i)
        sales = db.query(func.sum(Order.total)).filter(
            func.date(Order.created_at) >= month_ago,
            func.date(Order.created_at) <= date,
            Order.status == 'completed'
        ).scalar() or 0
        
        cumulative_sales = float(sales)
        performance_labels.append(f'{i+1}-kun')
        performance_sales.append(cumulative_sales)
        performance_target.append(daily_target * (i + 1))
    
    performance = {
        'labels': performance_labels,
        'sales': performance_sales,
        'target': performance_target
    }
    
    return templates.TemplateResponse("dashboards/agent.html", {
        "request": request,
        "page_title": "Agent Dashboard",
        "current_user": current_user,
        "user": current_user,
        "agent": agent_info,
        "kpi": kpi,
        "schedule": schedule,
        "recent_orders": recent_orders,
        "customers": customers,
        "performance": performance
    })


# Agent Dashboard - Test (fake data)
@app.get("/test/dashboard/agent", response_class=HTMLResponse)
async def agent_dashboard_test(request: Request, db: Session = Depends(get_db), current_user: User = Depends(require_admin)):
    """Agent Dashboard - Test (fake data), faqat admin"""
    
    # Fake user
    fake_user = {"username": "agent1", "role": "agent"}
    
    # Agent info
    agent = {
        'name': 'Alisher Karimov',
        'location': 'Chilonzor tumani, Toshkent'
    }
    
    # KPI
    kpi = {
        'visits_completed': 8,
        'visits_total': 12,
        'visits_percent': 67,
        'today_sales': 1250000,
        'orders': 15,
        'orders_completed': 12,
        'target_achieved': 18500000,
        'target_total': 25000000,
        'target_percent': 74
    }
    
    # Today's Schedule
    schedule = [
        {'customer': 'Anvar Toshmatov', 'time': '09:00', 'address': 'Chilonzor 12-kv', 'completed': True},
        {'customer': 'Dilshod Karimov', 'time': '10:30', 'address': 'Yunusobod 5-kv', 'completed': True},
        {'customer': 'Sardor Usmonov', 'time': '11:45', 'address': 'Mirzo Ulug\'bek 8-kv', 'completed': True},
        {'customer': 'Jasur Rahimov', 'time': '13:00', 'address': 'Yakkasaroy 3-kv', 'completed': False},
        {'customer': 'Bobur Sharipov', 'time': '14:30', 'address': 'Sergeli 7-kv', 'completed': False},
        {'customer': 'Otabek Normatov', 'time': '16:00', 'address': 'Uchtepa 4-kv', 'completed': False}
    ]
    
    # Recent Orders
    recent_orders = [
        {'number': 'ORD-1234', 'customer': 'Anvar Toshmatov', 'total': 450000, 'status_color': 'success', 'status_text': 'Bajarilgan'},
        {'number': 'ORD-1235', 'customer': 'Dilshod Karimov', 'total': 320000, 'status_color': 'warning', 'status_text': 'Jarayonda'},
        {'number': 'ORD-1236', 'customer': 'Sardor Usmonov', 'total': 180000, 'status_color': 'primary', 'status_text': 'Yangi'},
        {'number': 'ORD-1237', 'customer': 'Jasur Rahimov', 'total': 520000, 'status_color': 'success', 'status_text': 'Bajarilgan'},
        {'number': 'ORD-1238', 'customer': 'Bobur Sharipov', 'total': 280000, 'status_color': 'warning', 'status_text': 'Jarayonda'}
    ]
    
    # My Customers
    customers = [
        {'name': 'Anvar Toshmatov', 'phone': '+998 90 123 45 67', 'address': 'Chilonzor 12-kv', 'last_order': 450000},
        {'name': 'Dilshod Karimov', 'phone': '+998 91 234 56 78', 'address': 'Yunusobod 5-kv', 'last_order': 320000},
        {'name': 'Sardor Usmonov', 'phone': '+998 93 345 67 89', 'address': 'Mirzo Ulug\'bek 8-kv', 'last_order': 180000},
        {'name': 'Jasur Rahimov', 'phone': '+998 94 456 78 90', 'address': 'Yakkasaroy 3-kv', 'last_order': 520000},
        {'name': 'Bobur Sharipov', 'phone': '+998 95 567 89 01', 'address': 'Sergeli 7-kv', 'last_order': 280000}
    ]
    
    # Performance (30 days)
    performance = {
        'labels': ['1-kun', '5-kun', '10-kun', '15-kun', '20-kun', '25-kun', '30-kun'],
        'sales': [500000, 1200000, 2100000, 3500000, 5200000, 7100000, 8500000],
        'target': [833333, 1666666, 2500000, 3333333, 4166666, 5000000, 5833333]
    }
    
    return templates.TemplateResponse("dashboards/agent.html", {
        "request": request,
        "page_title": "Agent Dashboard",
        "current_user": None,
        "user": fake_user,
        "agent": agent,
        "kpi": kpi,
        "schedule": schedule,
        "recent_orders": recent_orders,
        "customers": customers,
        "performance": performance
    })



# Production Dashboard - Real Data
@app.get("/dashboard/production", response_class=HTMLResponse)
async def production_dashboard(request: Request, db: Session = Depends(get_db), current_user: User = Depends(require_auth)):
    """Ishlab chiqarish Dashboard - Real Data"""
    from datetime import datetime, timedelta
    from sqlalchemy import func
    from app.models.database import Production, Recipe, Product, Employee
    
    if not current_user:
        return RedirectResponse(url="/login", status_code=303)
    
    today = datetime.now().date()
    week_ago = today - timedelta(days=7)
    
    # Today's production
    today_production = db.query(func.sum(Production.quantity)).filter(
        func.date(Production.date) == today,
        Production.status == 'completed'
    ).scalar() or 0
    
    # Daily plan (placeholder - could be from a Plan table)
    plan = 3000
    efficiency = int((today_production / plan * 100)) if plan > 0 else 0
    
    # Active workers
    active_workers = db.query(func.count(Employee.id)).filter(
        Employee.is_active == True
    ).scalar() or 0
    
    # Raw materials stock percentage (placeholder)
    raw_materials = 72  # Placeholder
    
    metrics = {
        'today_production': int(today_production),
        'plan': plan,
        'active_machines': 0,  # Placeholder - no Machine model
        'total_machines': 0,   # Placeholder
        'efficiency': efficiency,
        'workers': active_workers,
        'shifts': 3,  # Placeholder
        'raw_materials': raw_materials
    }
    
    # Production orders (from Production table)
    production_query = db.query(Production, Recipe, Product).join(
        Recipe, Production.recipe_id == Recipe.id
    ).join(
        Product, Recipe.product_id == Product.id
    ).filter(
        func.date(Production.date) >= today - timedelta(days=1),
        Production.status.in_(['draft', 'completed'])
    ).order_by(Production.date.desc()).limit(10).all()
    
    production_orders = []
    for prod, recipe, product in production_query:
        # Calculate progress based on status
        progress = 100 if prod.status == 'completed' else 50
        deadline = prod.date.strftime('%H:%M') if prod.date else '-'
        
        production_orders.append({
            'product': product.name,
            'quantity': int(prod.quantity),
            'deadline': deadline,
            'progress': progress
        })
    
    if not production_orders:
        production_orders = [{'product': 'Ma\'lumot yo\'q', 'quantity': 0, 'deadline': '-', 'progress': 0}]
    
    # Machines (placeholder - no Machine model)
    machines = [
        {'name': 'Uskunalar ma\'lumoti', 'status': 'unknown', 'operator': '-', 'badge_color': 'secondary', 'status_text': 'Ma\'lumot yo\'q'}
    ]
    
    # Weekly production chart
    chart_labels = []
    chart_data = []
    chart_plan = []
    
    for i in range(6, -1, -1):
        date = today - timedelta(days=i)
        production = db.query(func.sum(Production.quantity)).filter(
            func.date(Production.date) == date,
            Production.status == 'completed'
        ).scalar() or 0
        
        chart_labels.append(['Yak', 'Dush', 'Sesh', 'Chor', 'Pay', 'Juma', 'Shan'][date.weekday()])
        chart_data.append(int(production))
        chart_plan.append(plan)
    
    chart = {
        'labels': chart_labels,
        'data': chart_data,
        'plan': chart_plan
    }
    
    return templates.TemplateResponse("dashboards/production.html", {
        "request": request,
        "page_title": "Ishlab chiqarish Dashboard",
        "current_user": current_user,
        "user": current_user,
        "metrics": metrics,
        "production_orders": production_orders,
        "machines": machines,
        "chart": chart
    })


# Production Dashboard - Test (fake data)
@app.get("/test/dashboard/production", response_class=HTMLResponse)
async def production_dashboard_test(request: Request, db: Session = Depends(get_db), current_user: User = Depends(require_admin)):
    """Ishlab chiqarish Dashboard - Test (fake data), faqat admin"""
    fake_user = {"username": "prod_manager", "role": "production"}
    
    metrics = {
        'today_production': 2500,
        'plan': 3000,
        'active_machines': 8,
        'total_machines': 10,
        'efficiency': 85,
        'workers': 45,
        'shifts': 3,
        'raw_materials': 72
    }
    
    production_orders = [
        {'product': 'Shokolad tort', 'quantity': 500, 'deadline': '16:00', 'progress': 75},
        {'product': 'Medovik', 'quantity': 300, 'deadline': '14:00', 'progress': 100},
        {'product': 'Napoleon', 'quantity': 400, 'deadline': '18:00', 'progress': 45},
        {'product': 'Tiramisu', 'quantity': 200, 'deadline': '15:00', 'progress': 90}
    ]
    
    machines = [
        {'name': 'Mixer #1', 'status': 'active', 'operator': 'Alisher', 'badge_color': 'success', 'status_text': 'Ishlayapti'},
        {'name': 'Mixer #2', 'status': 'active', 'operator': 'Dilshod', 'badge_color': 'success', 'status_text': 'Ishlayapti'},
        {'name': 'Oven #1', 'status': 'active', 'operator': 'Sardor', 'badge_color': 'success', 'status_text': 'Ishlayapti'},
        {'name': 'Oven #2', 'status': 'idle', 'operator': '-', 'badge_color': 'warning', 'status_text': 'Dam olishda'},
        {'name': 'Packaging #1', 'status': 'maintenance', 'operator': 'Texnik', 'badge_color': 'danger', 'status_text': 'Ta\'mirda'}
    ]
    
    chart = {
        'labels': ['Dush', 'Sesh', 'Chor', 'Pay', 'Juma', 'Shan', 'Yak'],
        'data': [2200, 2400, 2600, 2300, 2800, 2500, 2100],
        'plan': [3000, 3000, 3000, 3000, 3000, 2500, 2000]
    }
    
    return templates.TemplateResponse("dashboards/production.html", {
        "request": request,
        "page_title": "Ishlab chiqarish Dashboard",
        "current_user": None,
        "user": fake_user,
        "metrics": metrics,
        "production_orders": production_orders,
        "machines": machines,
        "chart": chart
    })


# Warehouse Dashboard - Real Data
@app.get("/dashboard/warehouse", response_class=HTMLResponse)
async def warehouse_dashboard(request: Request, db: Session = Depends(get_db), current_user: User = Depends(require_auth)):
    """Ombor Dashboard - Real Data"""
    from datetime import datetime, timedelta
    from sqlalchemy import func
    from app.models.database import Stock, Product, Category, Purchase, PurchaseItem
    
    if not current_user:
        return RedirectResponse(url="/login", status_code=303)
    
    today = datetime.now().date()
    week_ago = today - timedelta(days=7)
    
    # Total warehouse value
    total_value = db.query(
        func.sum(Stock.quantity * Product.purchase_price)
    ).join(
        Product, Stock.product_id == Product.id
    ).scalar() or 0
    
    # Total products
    total_products = db.query(func.count(Stock.id)).scalar() or 0
    
    # Categories
    categories = db.query(func.count(Category.id)).scalar() or 0
    
    # Today's incoming (purchases)
    today_in = db.query(func.sum(PurchaseItem.quantity)).join(
        Purchase, PurchaseItem.purchase_id == Purchase.id
    ).filter(
        func.date(Purchase.date) == today
    ).scalar() or 0
    
    # Today's outgoing (from orders - we'll use a simple count for now)
    today_out = db.query(func.count(Stock.id)).filter(
        Stock.quantity > 0
    ).scalar() or 0  # Placeholder
    
    metrics = {
        'total_value': float(total_value),
        'total_products': total_products,
        'categories': categories,
        'today_in': int(today_in),
        'today_out': 0  # Placeholder - need stock movement tracking
    }
    
    # Low stock items
    low_stock_items = db.query(Stock, Product).join(
        Product, Stock.product_id == Product.id
    ).filter(
        Stock.quantity < 20
    ).order_by(Stock.quantity).limit(10).all()
    
    low_stock = []
    for stock, product in low_stock_items:
        level = 'critical' if stock.quantity < 10 else 'low'
        badge = 'danger' if stock.quantity < 10 else 'warning'
        low_stock.append({
            'name': product.name,
            'quantity': int(stock.quantity),
            'min_quantity': 20,
            'level': level,
            'badge': badge
        })
    
    if not low_stock:
        low_stock = [{'name': 'Barcha mahsulotlar yetarli', 'quantity': 0, 'min_quantity': 0, 'level': 'ok', 'badge': 'success'}]
    
    # Recent movements (using purchases as proxy)
    recent_purchases = db.query(Purchase, PurchaseItem, Product).join(
        PurchaseItem, Purchase.id == PurchaseItem.purchase_id
    ).join(
        Product, PurchaseItem.product_id == Product.id
    ).filter(
        func.date(Purchase.date) >= week_ago
    ).order_by(Purchase.date.desc()).limit(5).all()
    
    recent_moves = []
    for purchase, item, product in recent_purchases:
        recent_moves.append({
            'product': product.name,
            'quantity': int(item.quantity),
            'type_text': 'Kirim',
            'type_color': 'success',
            'time': purchase.date.strftime('%H:%M')
        })
    
    if not recent_moves:
        recent_moves = [{'product': 'Ma\'lumot yo\'q', 'quantity': 0, 'type_text': '-', 'type_color': 'secondary', 'time': '-'}]
    
    # Weekly movement chart (placeholder with real structure)
    chart_labels = []
    chart_incoming = []
    chart_outgoing = []
    
    for i in range(6, -1, -1):
        date = today - timedelta(days=i)
        incoming = db.query(func.sum(PurchaseItem.quantity)).join(
            Purchase, PurchaseItem.purchase_id == Purchase.id
        ).filter(
            func.date(Purchase.date) == date
        ).scalar() or 0
        
        chart_labels.append(['Yak', 'Dush', 'Sesh', 'Chor', 'Pay', 'Juma', 'Shan'][date.weekday()])
        chart_incoming.append(int(incoming))
        chart_outgoing.append(0)  # Placeholder
    
    chart = {
        'labels': chart_labels,
        'incoming': chart_incoming,
        'outgoing': chart_outgoing
    }
    
    return templates.TemplateResponse("dashboards/warehouse.html", {
        "request": request,
        "page_title": "Ombor Dashboard",
        "current_user": current_user,
        "user": current_user,
        "metrics": metrics,
        "low_stock": low_stock,
        "recent_moves": recent_moves,
        "chart": chart
    })


# Warehouse Dashboard - Test (fake data)
@app.get("/test/dashboard/warehouse", response_class=HTMLResponse)
async def warehouse_dashboard_test(request: Request, db: Session = Depends(get_db), current_user: User = Depends(require_admin)):
    """Ombor Dashboard - Test (fake data), faqat admin"""
    fake_user = {"username": "warehouse_manager", "role": "warehouse"}
    
    metrics = {
        'total_value': 25000000,
        'total_products': 156,
        'categories': 12,
        'today_in': 45,
        'today_out': 38
    }
    
    low_stock = [
        {'name': 'Shokolad', 'quantity': 5, 'min_quantity': 20, 'level': 'critical', 'badge': 'danger'},
        {'name': 'Un', 'quantity': 15, 'min_quantity': 50, 'level': 'low', 'badge': 'warning'},
        {'name': 'Shakar', 'quantity': 25, 'min_quantity': 40, 'level': 'low', 'badge': 'warning'},
        {'name': 'Yog\'', 'quantity': 8, 'min_quantity': 30, 'level': 'critical', 'badge': 'danger'},
        {'name': 'Tuxum', 'quantity': 35, 'min_quantity': 50, 'level': 'low', 'badge': 'warning'}
    ]
    
    recent_moves = [
        {'product': 'Shokolad tort', 'quantity': 50, 'type_text': 'Chiqim', 'type_color': 'danger', 'time': '14:30'},
        {'product': 'Un', 'quantity': 100, 'type_text': 'Kirim', 'type_color': 'success', 'time': '13:15'},
        {'product': 'Medovik', 'quantity': 30, 'type_text': 'Chiqim', 'type_color': 'danger', 'time': '12:45'},
        {'product': 'Shakar', 'quantity': 50, 'type_text': 'Kirim', 'type_color': 'success', 'time': '11:20'},
        {'product': 'Napoleon', 'quantity': 25, 'type_text': 'Chiqim', 'type_color': 'danger', 'time': '10:30'}
    ]
    
    chart = {
        'labels': ['Dush', 'Sesh', 'Chor', 'Pay', 'Juma', 'Shan', 'Yak'],
        'incoming': [120, 150, 180, 140, 200, 160, 130],
        'outgoing': [100, 130, 150, 120, 170, 140, 110]
    }
    
    return templates.TemplateResponse("dashboards/warehouse.html", {
        "request": request,
        "page_title": "Ombor Dashboard",
        "user": fake_user,
        "metrics": metrics,
        "low_stock": low_stock,
        "recent_moves": recent_moves,
        "chart": chart
    })


# Delivery Dashboard - Real Data
@app.get("/dashboard/delivery", response_class=HTMLResponse)
async def delivery_dashboard(request: Request, db: Session = Depends(get_db), current_user: User = Depends(require_auth)):
    """Yetkazib berish Dashboard - Real Data"""
    from datetime import datetime, timedelta
    from sqlalchemy import func, case
    from app.models.database import Delivery, Driver, DriverLocation, Order, Partner
    
    if not current_user:
        return RedirectResponse(url="/login", status_code=303)
    
    today = datetime.now().date()
    week_ago = today - timedelta(days=7)
    
    # Today's deliveries
    total_deliveries = db.query(func.count(Delivery.id)).filter(
        func.date(Delivery.planned_date) == today
    ).scalar() or 0
    
    completed_deliveries = db.query(func.count(Delivery.id)).filter(
        func.date(Delivery.planned_date) == today,
        Delivery.status == 'delivered'
    ).scalar() or 0
    
    percent = int((completed_deliveries / total_deliveries * 100)) if total_deliveries > 0 else 0
    
    # Active drivers
    active_drivers = db.query(func.count(Driver.id)).filter(
        Driver.is_active == True
    ).scalar() or 0
    
    total_drivers = db.query(func.count(Driver.id)).scalar() or 0
    
    # Average delivery time (placeholder - need delivery duration tracking)
    avg_time = 45  # Placeholder
    
    # Delays (deliveries not completed on time)
    delays = db.query(func.count(Delivery.id)).filter(
        func.date(Delivery.planned_date) < today,
        Delivery.status.in_(['pending', 'in_progress'])
    ).scalar() or 0
    
    metrics = {
        'completed': completed_deliveries,
        'total': total_deliveries,
        'percent': percent,
        'active_drivers': active_drivers,
        'total_drivers': total_drivers,
        'avg_time': avg_time,
        'delays': delays
    }
    
    # Today's deliveries with details
    deliveries_query = db.query(Delivery, Driver, Partner).join(
        Driver, Delivery.driver_id == Driver.id
    ).outerjoin(
        Order, Delivery.order_id == Order.id
    ).outerjoin(
        Partner, Order.partner_id == Partner.id
    ).filter(
        func.date(Delivery.planned_date) >= week_ago
    ).order_by(Delivery.planned_date.desc()).limit(20).all()
    
    status_map = {
        'pending': ('Kutilmoqda', 'secondary'),
        'in_progress': ('Yo\'lda', 'warning'),
        'delivered': ('Yetkazilgan', 'success'),
        'failed': ('Bekor qilingan', 'danger')
    }
    
    deliveries = []
    for delivery, driver, partner in deliveries_query:
        status_text, badge_color = status_map.get(delivery.status, ('Noma\'lum', 'secondary'))
        deliveries.append({
            'customer': partner.name if partner else 'Noma\'lum',
            'address': delivery.delivery_address or '-',
            'driver': driver.full_name,
            'status': delivery.status,
            'badge_color': badge_color,
            'status_text': status_text,
            'time': delivery.planned_date.strftime('%H:%M') if delivery.planned_date else '-'
        })
    
    if not deliveries:
        deliveries = [{'customer': 'Ma\'lumot yo\'q', 'address': '-', 'driver': '-', 'status': 'pending', 'badge_color': 'secondary', 'status_text': '-', 'time': '-'}]
    
    # Drivers with their stats
    drivers_query = db.query(
        Driver,
        func.count(Delivery.id).label('delivery_count')
    ).outerjoin(
        Delivery, 
        (Driver.id == Delivery.driver_id) & (func.date(Delivery.planned_date) == today)
    ).group_by(Driver.id).order_by(func.count(Delivery.id).desc()).limit(10).all()
    
    drivers = []
    for driver, delivery_count in drivers_query:
        # Get latest location
        latest_location = db.query(DriverLocation).filter(
            DriverLocation.driver_id == driver.id
        ).order_by(DriverLocation.timestamp.desc()).first()
        
        location_text = 'Noma\'lum'
        if latest_location and latest_location.address:
            location_text = latest_location.address[:30] + '...' if len(latest_location.address) > 30 else latest_location.address
        
        status_color = 'success' if driver.is_active else 'secondary'
        status_text = 'Faol' if driver.is_active else 'Faol emas'
        
        drivers.append({
            'name': driver.full_name,
            'deliveries': delivery_count,
            'location': location_text,
            'status_color': status_color,
            'status_text': status_text
        })
    
    if not drivers:
        drivers = [{'name': 'Ma\'lumot yo\'q', 'deliveries': 0, 'location': '-', 'status_color': 'secondary', 'status_text': '-'}]
    
    # Weekly delivery chart
    chart_labels = []
    chart_completed = []
    chart_delayed = []
    
    for i in range(6, -1, -1):
        date = today - timedelta(days=i)
        
        completed = db.query(func.count(Delivery.id)).filter(
            func.date(Delivery.planned_date) == date,
            Delivery.status == 'delivered'
        ).scalar() or 0
        
        delayed = db.query(func.count(Delivery.id)).filter(
            func.date(Delivery.planned_date) == date,
            Delivery.status.in_(['pending', 'in_progress', 'failed'])
        ).scalar() or 0
        
        chart_labels.append(['Yak', 'Dush', 'Sesh', 'Chor', 'Pay', 'Juma', 'Shan'][date.weekday()])
        chart_completed.append(completed)
        chart_delayed.append(delayed)
    
    chart = {
        'labels': chart_labels,
        'completed': chart_completed,
        'delayed': chart_delayed
    }
    
    return templates.TemplateResponse("dashboards/delivery.html", {
        "request": request,
        "page_title": "Yetkazib berish Dashboard",
        "current_user": current_user,
        "user": current_user,
        "metrics": metrics,
        "deliveries": deliveries,
        "drivers": drivers,
        "chart": chart
    })


# Delivery Dashboard - Test (fake data)
@app.get("/test/dashboard/delivery", response_class=HTMLResponse)
async def delivery_dashboard_test(request: Request, db: Session = Depends(get_db), current_user: User = Depends(require_admin)):
    """Yetkazib berish Dashboard - Test (fake data), faqat admin"""
    fake_user = {"username": "delivery_manager", "role": "delivery"}
    
    metrics = {
        'completed': 28,
        'total': 35,
        'percent': 80,
        'active_drivers': 8,
        'total_drivers': 10,
        'avg_time': 45,
        'delays': 3
    }
    
    deliveries = [
        {'customer': 'Anvar Toshmatov', 'address': 'Chilonzor 12-kv', 'driver': 'Alisher', 'status': 'completed', 'badge_color': 'success', 'status_text': 'Yetkazilgan', 'time': '09:30'},
        {'customer': 'Dilshod Karimov', 'address': 'Yunusobod 5-kv', 'driver': 'Sardor', 'status': 'in-progress', 'badge_color': 'warning', 'status_text': 'Yo\'lda', 'time': '10:15'},
        {'customer': 'Jasur Rahimov', 'address': 'Yakkasaroy 3-kv', 'driver': 'Bobur', 'status': 'in-progress', 'badge_color': 'warning', 'status_text': 'Yo\'lda', 'time': '11:00'},
        {'customer': 'Otabek Normatov', 'address': 'Uchtepa 4-kv', 'driver': 'Jasur', 'status': 'pending', 'badge_color': 'secondary', 'status_text': 'Kutilmoqda', 'time': '13:00'},
        {'customer': 'Sardor Usmonov', 'address': 'Mirzo Ulug\'bek 8-kv', 'driver': 'Dilshod', 'status': 'completed', 'badge_color': 'success', 'status_text': 'Yetkazilgan', 'time': '08:45'}
    ]
    
    drivers = [
        {'name': 'Alisher Karimov', 'deliveries': 5, 'location': 'Chilonzor', 'status_color': 'success', 'status_text': 'Faol'},
        {'name': 'Sardor Usmonov', 'deliveries': 4, 'location': 'Yunusobod', 'status_color': 'success', 'status_text': 'Faol'},
        {'name': 'Bobur Sharipov', 'deliveries': 3, 'location': 'Yakkasaroy', 'status_color': 'success', 'status_text': 'Faol'},
        {'name': 'Jasur Rahimov', 'deliveries': 2, 'location': 'Uchtepa', 'status_color': 'warning', 'status_text': 'Dam olishda'},
        {'name': 'Dilshod Karimov', 'deliveries': 4, 'location': 'Sergeli', 'status_color': 'success', 'status_text': 'Faol'}
    ]
    
    chart = {
        'labels': ['Dush', 'Sesh', 'Chor', 'Pay', 'Juma', 'Shan', 'Yak'],
        'completed': [32, 35, 38, 34, 40, 36, 28],
        'delayed': [3, 2, 4, 3, 5, 2, 3]
    }
    
    return templates.TemplateResponse("dashboards/delivery.html", {
        "request": request,
        "page_title": "Yetkazib berish Dashboard",
        "user": fake_user,
        "metrics": metrics,
        "deliveries": deliveries,
        "drivers": drivers,
        "chart": chart
    })


# O'lchov birliklari bo'limi (moved to app/routes/info.py)
# @app.get("/info/units", response_class=HTMLResponse)
async def info_units(request: Request, db: Session = Depends(get_db), current_user: User = Depends(require_auth)):
    units = db.query(Unit).all()
    return templates.TemplateResponse("info/units.html", {"request": request, "units": units, "current_user": current_user, "page_title": "O'lchov birliklari"})

@app.post("/info/units/add")
async def info_units_add(
    request: Request,
    code: str = Form(...),
    name: str = Form(...),
    db: Session = Depends(get_db)
):
    existing = db.query(Unit).filter(Unit.code == code).first()
    if existing:
        raise HTTPException(status_code=400, detail=f"'{code}' kodli o'lchov birligi allaqachon mavjud!")
    
    unit = Unit(code=code, name=name)
    db.add(unit)
    db.commit()
    return RedirectResponse(url="/info/units", status_code=303)

@app.post("/info/units/edit/{unit_id}")
async def info_units_edit(
    unit_id: int,
    code: str = Form(...),
    name: str = Form(...),
    db: Session = Depends(get_db)
):
    unit = db.query(Unit).filter(Unit.id == unit_id).first()
    if not unit:
        raise HTTPException(status_code=404, detail="O'lchov birligi topilmadi")
    
    existing = db.query(Unit).filter(Unit.code == code, Unit.id != unit_id).first()
    if existing:
        raise HTTPException(status_code=400, detail=f"'{code}' kodli o'lchov birligi allaqachon mavjud!")
    
    unit.code = code
    unit.name = name
    db.commit()
    return RedirectResponse(url="/info/units", status_code=303)

@app.post("/info/units/delete/{unit_id}")
async def info_units_delete(unit_id: int, db: Session = Depends(get_db), current_user: User = Depends(require_auth)):
    unit = db.query(Unit).filter(Unit.id == unit_id).first()
    if not unit:
        raise HTTPException(status_code=404, detail="O'lchov birligi topilmadi")
    
    db.delete(unit)
    db.commit()
    return RedirectResponse(url="/info/units", status_code=303)

# --- UNITS EXCEL OPERATIONS ---
@app.get("/info/units/export")
async def export_units(db: Session = Depends(get_db), current_user: User = Depends(require_auth)):
    units = db.query(Unit).all()
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Units"
    ws.append(["ID", "Kod", "Nomi"])
    for u in units:
        ws.append([u.id, u.code, u.name])
    stream = io.BytesIO()
    wb.save(stream)
    stream.seek(0)
    return StreamingResponse(stream, media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", headers={"Content-Disposition": "attachment; filename=olchov_birliklari.xlsx"})

@app.get("/info/units/template")
async def template_units(current_user: User = Depends(require_auth)):
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Template"
    ws.append(["Kod", "Nomi"])
    ws.append(["kg", "Kilogramm"])
    stream = io.BytesIO()
    wb.save(stream)
    stream.seek(0)
    return StreamingResponse(stream, media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", headers={"Content-Disposition": "attachment; filename=birlik_andoza.xlsx"})

@app.post("/info/units/import")
async def import_units(file: UploadFile = File(...), db: Session = Depends(get_db), current_user: User = Depends(require_auth)):
    contents = await file.read()
    wb = openpyxl.load_workbook(io.BytesIO(contents))
    ws = wb.active
    rows = list(ws.iter_rows(min_row=2, values_only=True))
    for row in rows:
        if not row[0]: continue
        code, name = row[0], row[1]
        unit = db.query(Unit).filter(Unit.code == code).first()
        if not unit:
            unit = Unit(code=code, name=name)
            db.add(unit)
        else:
            unit.name = name
        db.commit()
    return RedirectResponse(url="/info/units", status_code=303)

# Kategoriyalar bo'limi
@app.get("/info/categories", response_class=HTMLResponse)
async def info_categories(request: Request, db: Session = Depends(get_db), current_user: User = Depends(require_auth)):
    categories = db.query(Category).all()
    return templates.TemplateResponse("info/categories.html", {"request": request, "categories": categories, "current_user": current_user, "page_title": "Kategoriyalar"})

@app.post("/info/categories/add")
async def info_categories_add(
    request: Request,
    code: str = Form(...),
    name: str = Form(...),
    type: str = Form(...),
    db: Session = Depends(get_db)
):
    existing = db.query(Category).filter(Category.code == code).first()
    if existing:
        raise HTTPException(status_code=400, detail=f"'{code}' kodli kategoriya allaqachon mavjud!")
    
    category = Category(code=code, name=name, type=type)
    db.add(category)
    db.commit()
    return RedirectResponse(url="/info/categories", status_code=303)

@app.post("/info/categories/edit/{category_id}")
async def info_categories_edit(
    category_id: int,
    code: str = Form(...),
    name: str = Form(...),
    type: str = Form(...),
    db: Session = Depends(get_db)
):
    category = db.query(Category).filter(Category.id == category_id).first()
    if not category:
        raise HTTPException(status_code=404, detail="Kategoriya topilmadi")
    
    existing = db.query(Category).filter(Category.code == code, Category.id != category_id).first()
    if existing:
        raise HTTPException(status_code=400, detail=f"'{code}' kodli kategoriya allaqachon mavjud!")
    
    category.code = code
    category.name = name
    category.type = type
    db.commit()
    return RedirectResponse(url="/info/categories", status_code=303)

@app.post("/info/categories/delete/{category_id}")
async def info_categories_delete(category_id: int, db: Session = Depends(get_db), current_user: User = Depends(require_auth)):
    category = db.query(Category).filter(Category.id == category_id).first()
    if not category:
        raise HTTPException(status_code=404, detail="Kategoriya topilmadi")
    
    db.delete(category)
    db.commit()
    return RedirectResponse(url="/info/categories", status_code=303)

# --- CATEGORIES EXCEL OPERATIONS ---
@app.get("/info/categories/export")
async def export_categories(db: Session = Depends(get_db), current_user: User = Depends(require_auth)):
    categories = db.query(Category).all()
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Categories"
    ws.append(["ID", "Kod", "Nomi", "Turi"])
    for c in categories:
        ws.append([c.id, c.code, c.name, c.type])
    stream = io.BytesIO()
    wb.save(stream)
    stream.seek(0)
    return StreamingResponse(stream, media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", headers={"Content-Disposition": "attachment; filename=kategoriyalar.xlsx"})

@app.get("/info/categories/template")
async def template_categories(current_user: User = Depends(require_auth)):
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Template"
    ws.append(["Kod", "Nomi", "Turi"])
    ws.append(["CAT001", "Shirinliklar", "tayyor"])
    stream = io.BytesIO()
    wb.save(stream)
    stream.seek(0)
    return StreamingResponse(stream, media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", headers={"Content-Disposition": "attachment; filename=kategoriya_andoza.xlsx"})

@app.post("/info/categories/import")
async def import_categories(file: UploadFile = File(...), db: Session = Depends(get_db), current_user: User = Depends(require_auth)):
    contents = await file.read()
    wb = openpyxl.load_workbook(io.BytesIO(contents))
    ws = wb.active
    rows = list(ws.iter_rows(min_row=2, values_only=True))
    for row in rows:
        if not row[0]: continue
        code, name, type_ = row[0], row[1], row[2]
        category = db.query(Category).filter(Category.code == code).first()
        if not category:
            category = Category(code=code, name=name, type=type_)
            db.add(category)
        else:
            category.name = name
            category.type = type_
        db.commit()
    return RedirectResponse(url="/info/categories", status_code=303)

# Narx turlari (Chakana, Ulgurji, VIP va h.k.) — Ma'lumotnomalar
@app.get("/info/price-types", response_class=HTMLResponse)
async def info_price_types(request: Request, db: Session = Depends(get_db), current_user: User = Depends(require_auth)):
    """Narx turlari ro'yxati"""
    if not current_user:
        return RedirectResponse(url="/login", status_code=303)
    price_types = db.query(PriceType).filter(PriceType.is_active == True).order_by(PriceType.name).all()
    return templates.TemplateResponse("info/price_types.html", {
        "request": request,
        "price_types": price_types,
        "current_user": current_user,
        "page_title": "Narx turlari"
    })


@app.post("/info/price-types/add")
async def info_price_types_add(
    name: str = Form(...),
    code: str = Form(""),
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth)
):
    if not current_user:
        return RedirectResponse(url="/login", status_code=303)
    code = (code or "").strip() or None
    if code and db.query(PriceType).filter(PriceType.code == code).first():
        raise HTTPException(status_code=400, detail=f"'{code}' kodli narx turi allaqachon mavjud!")
    pt = PriceType(name=name, code=code, is_active=True)
    db.add(pt)
    db.commit()
    return RedirectResponse(url="/info/price-types", status_code=303)


@app.post("/info/price-types/edit/{price_type_id}")
async def info_price_types_edit(
    price_type_id: int,
    name: str = Form(...),
    code: str = Form(""),
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth)
):
    if not current_user:
        return RedirectResponse(url="/login", status_code=303)
    pt = db.query(PriceType).filter(PriceType.id == price_type_id).first()
    if not pt:
        raise HTTPException(status_code=404, detail="Narx turi topilmadi")
    code = (code or "").strip() or None
    if code and db.query(PriceType).filter(PriceType.code == code, PriceType.id != price_type_id).first():
        raise HTTPException(status_code=400, detail=f"'{code}' kodli narx turi allaqachon mavjud!")
    pt.name = name
    pt.code = code
    db.commit()
    return RedirectResponse(url="/info/price-types", status_code=303)


@app.post("/info/price-types/delete/{price_type_id}")
async def info_price_types_delete(price_type_id: int, db: Session = Depends(get_db), current_user: User = Depends(require_auth)):
    if not current_user:
        return RedirectResponse(url="/login", status_code=303)
    pt = db.query(PriceType).filter(PriceType.id == price_type_id).first()
    if not pt:
        raise HTTPException(status_code=404, detail="Narx turi topilmadi")
    db.query(ProductPrice).filter(ProductPrice.price_type_id == price_type_id).delete()
    pt.is_active = False
    db.commit()
    return RedirectResponse(url="/info/price-types", status_code=303)


# Narxni o'rnatish (mahsulot narxlari narx turi bo'yicha) — Ma'lumotnomalar
@app.get("/info/prices", response_class=HTMLResponse)
async def info_prices(
    request: Request,
    price_type_id: Optional[int] = None,
    search: Optional[str] = None,
    type_filter: Optional[str] = None,
    price_status: Optional[str] = None,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth)
):
    """Mahsulot tannarxi va narx turi bo'yicha sotuv narxini o'rnatish"""
    if not current_user:
        return RedirectResponse(url="/login", status_code=303)
    price_types = db.query(PriceType).filter(PriceType.is_active == True).order_by(PriceType.name).all()
    if not price_types:
        return templates.TemplateResponse("info/prices.html", {
            "request": request,
            "products": [],
            "price_types": [],
            "current_price_type_id": None,
            "product_prices_by_type": {},
            "product_tannarx": {},
            "current_user": current_user,
            "page_title": "Narxni o'rnatish",
            "filter_search": "",
            "filter_type": "",
            "filter_price_status": "all",
        })
    current_pt_id = price_type_id or (price_types[0].id if price_types else None)
    products = db.query(Product).options(joinedload(Product.unit)).filter(
        Product.is_active == True
    ).order_by(Product.name).all()
    # Mahsulot × narx turi bo'yicha sotuv narxi
    product_prices = db.query(ProductPrice).filter(ProductPrice.price_type_id == current_pt_id).all()
    product_prices_by_type = {pp.product_id: pp.sale_price for pp in product_prices}

    # Tannarx: avvalo Product.purchase_price, bo'lmasa so'nggi kirim yoki qoldiq hujjatidan
    product_tannarx = {p.id: float(p.purchase_price or 0) for p in products}
    last_purchase_cost = {}
    for row in (
        db.query(PurchaseItem.product_id, PurchaseItem.price)
        .join(Purchase, PurchaseItem.purchase_id == Purchase.id)
        .filter(Purchase.status == "confirmed", PurchaseItem.price.isnot(None))
        .order_by(Purchase.date.desc())
        .all()
    ):
        pid, price = (row[0], row[1]) if hasattr(row, "__getitem__") else (row.product_id, row.price)
        if pid not in last_purchase_cost and (price or 0) > 0:
            last_purchase_cost[pid] = float(price)
    # Qoldiq hujjati (StockAdjustmentDoc) tasdiqlangan bo'lsa, shu hujjat qatorlaridagi cost_price dan tannarx olamiz
    confirmed_doc_ids = [r[0] for r in db.query(StockAdjustmentDoc.id).filter(StockAdjustmentDoc.status == "confirmed").all()]
    last_qoldiq_cost = {}
    if confirmed_doc_ids:
        # So'nggi tasdiqlangan hujjatlar bo'yicha (sana kamayish tartibida) qatorlarni olish
        qoldiq_rows = (
            db.query(StockAdjustmentDocItem.product_id, StockAdjustmentDocItem.cost_price, StockAdjustmentDoc.date, StockAdjustmentDoc.id)
            .join(StockAdjustmentDoc, StockAdjustmentDocItem.doc_id == StockAdjustmentDoc.id)
            .filter(StockAdjustmentDoc.id.in_(confirmed_doc_ids))
            .order_by(StockAdjustmentDoc.date.desc(), StockAdjustmentDoc.id.desc())
            .all()
        )
        for row in qoldiq_rows:
            pid = row[0] if hasattr(row, "__getitem__") else row.product_id
            cost = row[1] if hasattr(row, "__getitem__") else row.cost_price
            if pid not in last_qoldiq_cost and (cost is not None and float(cost) > 0):
                last_qoldiq_cost[pid] = float(cost)
    for pid in product_tannarx:
        if product_tannarx[pid] <= 0 and pid in last_purchase_cost:
            product_tannarx[pid] = last_purchase_cost[pid]
        if product_tannarx[pid] <= 0 and pid in last_qoldiq_cost:
            product_tannarx[pid] = last_qoldiq_cost[pid]
    # Qoldiqdan olgan tannarxni Product jadvaliga bir marta yozish (Narxni o'rnatishda 0 ko'rinmasin)
    for pid, cost in last_qoldiq_cost.items():
        if product_tannarx.get(pid, 0) != cost:
            continue
        prod = db.query(Product).filter(Product.id == pid).first()
        if prod and (prod.purchase_price or 0) <= 0 and cost > 0:
            prod.purchase_price = cost
    try:
        db.commit()
    except Exception:
        db.rollback()

    # Filtrlar: mahsulot nomi, turi, sotuv narxi holati
    search_q = (search or "").strip().lower()
    type_filter_val = (type_filter or "").strip() or None
    price_status_val = (price_status or "all").strip() or "all"
    if search_q or type_filter_val or price_status_val != "all":
        filtered_products = []
        for p in products:
            if search_q and search_q not in (p.name or "").lower() and search_q not in (p.barcode or "").lower():
                continue
            if type_filter_val and (p.type or "") != type_filter_val:
                continue
            sale_val = product_prices_by_type.get(p.id)
            if sale_val is None:
                sale_val = getattr(p, "sale_price", None)
            has_sale_price = sale_val is not None and float(sale_val or 0) > 0
            if price_status_val == "set" and not has_sale_price:
                continue
            if price_status_val == "not_set" and has_sale_price:
                continue
            filtered_products.append(p)
        products = filtered_products

    return templates.TemplateResponse("info/prices.html", {
        "request": request,
        "products": products,
        "price_types": price_types,
        "current_price_type_id": current_pt_id,
        "product_prices_by_type": product_prices_by_type,
        "product_tannarx": product_tannarx,
        "current_user": current_user,
        "page_title": "Narxni o'rnatish",
        "filter_search": search or "",
        "filter_type": type_filter or "",
        "filter_price_status": price_status_val,
    })


def _next_price_history_doc_number(db: Session) -> str:
    """Narx o'zgarishi hujjati raqami: PN-YYYYMMDD-NNN"""
    from datetime import datetime
    prefix = f"PN-{datetime.now().strftime('%Y%m%d')}-"
    last = db.query(ProductPriceHistory).filter(ProductPriceHistory.doc_number.like(f"{prefix}%")).order_by(ProductPriceHistory.id.desc()).first()
    if not last or not last.doc_number:
        num = 1
    else:
        try:
            num = int(last.doc_number.rsplit("-", 1)[-1]) + 1
        except (ValueError, IndexError):
            num = 1
    return f"{prefix}{num:03d}"


@app.post("/info/prices/edit/{product_id}")
async def info_prices_edit(
    product_id: int,
    purchase_price: float = Form(0),
    sale_price: float = Form(0),
    price_type_id: Optional[int] = Form(None),
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth)
):
    """Mahsulot tannarxi va narx turi bo'yicha sotuv narxini yangilash; har bir o'zgarish hujjat sifatida tarixga yoziladi."""
    if not current_user:
        return RedirectResponse(url="/login", status_code=303)
    product = db.query(Product).filter(Product.id == product_id).first()
    if not product:
        raise HTTPException(status_code=404, detail="Mahsulot topilmadi")
    old_purchase = float(product.purchase_price or 0)
    if price_type_id:
        pp = db.query(ProductPrice).filter(
            ProductPrice.product_id == product_id,
            ProductPrice.price_type_id == price_type_id
        ).first()
        old_sale = float(pp.sale_price if pp else 0)
        product.purchase_price = purchase_price
        if pp:
            pp.sale_price = sale_price
        else:
            db.add(ProductPrice(product_id=product_id, price_type_id=price_type_id, sale_price=sale_price))
    else:
        old_sale = float(product.sale_price or 0)
        product.purchase_price = purchase_price
        product.sale_price = sale_price
    # Harakat tarixi — hujjat sifatida saqlash (avvalgi va yangi narx)
    doc_number = _next_price_history_doc_number(db)
    db.add(ProductPriceHistory(
        doc_number=doc_number,
        product_id=product_id,
        price_type_id=price_type_id,
        old_purchase_price=old_purchase,
        new_purchase_price=float(purchase_price or 0),
        old_sale_price=old_sale,
        new_sale_price=float(sale_price or 0),
        changed_by_id=current_user.id,
    ))
    db.commit()
    redirect_url = f"/info/prices?price_type_id={price_type_id}" if price_type_id else "/info/prices"
    return RedirectResponse(url=redirect_url, status_code=303)


# Kassalar bo'limi
@app.get("/info/cash", response_class=HTMLResponse)
async def info_cash(request: Request, db: Session = Depends(get_db), current_user: User = Depends(require_auth)):
    cash_registers = db.query(CashRegister).all()
    departments = db.query(Department).filter(Department.is_active == True).order_by(Department.name).all()
    return templates.TemplateResponse("info/cash.html", {
        "request": request,
        "cash_registers": cash_registers,
        "departments": departments,
        "current_user": current_user,
        "page_title": "Kassalar",
    })


@app.post("/info/cash/add")
async def info_cash_add(
    request: Request,
    name: str = Form(...),
    balance: float = Form(0),
    department_id: Optional[int] = Form(None),
    payment_type: Optional[str] = Form(None),
    db: Session = Depends(get_db),
):
    pt = (payment_type or "").strip() or None
    if pt and pt not in ("naqd", "plastik", "click", "terminal"):
        pt = None
    cash = CashRegister(
        name=name,
        balance=balance,
        department_id=department_id if department_id and department_id > 0 else None,
        payment_type=pt,
        is_active=True,
    )
    db.add(cash)
    db.commit()
    return RedirectResponse(url="/info/cash", status_code=303)


@app.post("/info/cash/edit/{cash_id}")
async def info_cash_edit(
    cash_id: int,
    name: str = Form(...),
    balance: float = Form(0),
    department_id: Optional[int] = Form(None),
    payment_type: Optional[str] = Form(None),
    db: Session = Depends(get_db),
):
    cash = db.query(CashRegister).filter(CashRegister.id == cash_id).first()
    if not cash:
        raise HTTPException(status_code=404, detail="Kassa topilmadi")
    cash.name = name
    cash.balance = balance
    cash.department_id = department_id if department_id and department_id > 0 else None
    pt = (payment_type or "").strip() or None
    cash.payment_type = pt if pt in ("naqd", "plastik", "click", "terminal") else None
    db.commit()
    return RedirectResponse(url="/info/cash", status_code=303)

@app.post("/info/cash/delete/{cash_id}")
async def info_cash_delete(cash_id: int, db: Session = Depends(get_db), current_user: User = Depends(require_auth)):
    cash = db.query(CashRegister).filter(CashRegister.id == cash_id).first()
    if not cash:
        raise HTTPException(status_code=404, detail="Kassa topilmadi")
    
    db.delete(cash)
    db.commit()
    return RedirectResponse(url="/info/cash", status_code=303)


@app.post("/info/cash/recalculate/{cash_id}")
async def info_cash_recalculate(
    cash_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_admin),
):
    """Kassa balansini Payment yozuvlaridan qayta hisoblaydi (faqat admin)."""
    cash = db.query(CashRegister).filter(CashRegister.id == cash_id).first()
    if not cash:
        raise HTTPException(status_code=404, detail="Kassa topilmadi")
    total = (
        db.query(func.coalesce(func.sum(Payment.amount), 0))
        .filter(Payment.cash_register_id == cash_id)
        .scalar()
    )
    if total is None:
        total = 0
    cash.balance = float(total)
    db.commit()
    from urllib.parse import quote
    return RedirectResponse(
        url="/info/cash?recalculated=1&balance=" + quote(str(cash.balance)),
        status_code=303,
    )


# ==========================================
# QOLDİQLAR (bitta oyna: kassa, tovar, kontragent)
# ==========================================

@app.get("/qoldiqlar", response_class=HTMLResponse)
async def qoldiqlar_page(request: Request, db: Session = Depends(get_db), current_user: User = Depends(require_auth)):
    """Qoldiqlar sahifasi: kassa, tovar (forma spiska 1C), kontragent qoldiqlarini kiritish"""
    cash_registers = db.query(CashRegister).filter(CashRegister.is_active == True).all()
    warehouses = db.query(Warehouse).filter(Warehouse.is_active == True).all()
    products = db.query(Product).filter(Product.is_active == True).order_by(Product.name).all()
    stocks = db.query(Stock).join(Warehouse).join(Product).order_by(Stock.updated_at.desc()).limit(300).all()
    partners = db.query(Partner).filter(Partner.is_active == True).order_by(Partner.name).all()
    tovar_docs = (
        db.query(StockAdjustmentDoc)
        .order_by(StockAdjustmentDoc.id.desc())
        .limit(500)
        .all()
    )
    cash_docs = (
        db.query(CashBalanceDoc)
        .order_by(CashBalanceDoc.created_at.desc())
        .limit(200)
        .all()
    )
    kontragent_docs = (
        db.query(PartnerBalanceDoc)
        .order_by(PartnerBalanceDoc.created_at.desc())
        .limit(200)
        .all()
    )
    return templates.TemplateResponse("qoldiqlar/index.html", {
        "request": request,
        "cash_registers": cash_registers,
        "warehouses": warehouses,
        "products": products,
        "stocks": stocks,
        "partners": partners,
        "tovar_docs": tovar_docs,
        "cash_docs": cash_docs,
        "kontragent_docs": kontragent_docs,
        "current_user": current_user,
        "page_title": "Qoldiqlar",
    })


def _tarix_doc_type_label(doc_type: str) -> str:
    """Hujjat turi uchun o'qiladigan nom (tarix sahifasi)."""
    labels = {
        "Purchase": "Kirim",
        "Production": "Ishlab chiqarish",
        "WarehouseTransfer": "Ombordan omborga",
        "StockAdjustmentDoc": "Qoldiq tuzatish",
        "Sale": "Sotuv",
        "SaleReturn": "Qaytish",
    }
    return labels.get(doc_type or "", doc_type or "—")


@app.get("/qoldiqlar/tarix", response_class=HTMLResponse)
async def qoldiqlar_tarix(
    request: Request,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
    warehouse_id: Optional[str] = None,
    product_id: Optional[str] = None,
):
    """Mahsulot harakati tarixi — tanlanmasa barcha harakatlar; filtr ixtiyoriy (ombor/mahsulot)."""
    try:
        warehouses = db.query(Warehouse).filter(Warehouse.is_active == True).all()
        products = db.query(Product).filter(Product.is_active == True).order_by(Product.name).all()
        try:
            selected_warehouse_id = int(warehouse_id) if (warehouse_id and str(warehouse_id).strip()) else None
        except (ValueError, TypeError):
            selected_warehouse_id = None
        try:
            selected_product_id = int(product_id) if (product_id and str(product_id).strip()) else None
        except (ValueError, TypeError):
            selected_product_id = None

        q = db.query(StockMovement)
        if selected_warehouse_id:
            q = q.filter(StockMovement.warehouse_id == selected_warehouse_id)
        if selected_product_id:
            q = q.filter(StockMovement.product_id == selected_product_id)
        q = q.order_by(StockMovement.created_at.desc()).limit(500)
        movements = q.all()

        movement_rows = []
        warehouse_ids = [m.warehouse_id for m in movements if m.warehouse_id is not None]
        product_ids = [m.product_id for m in movements if m.product_id is not None]
        wh_by_id = {}
        if warehouse_ids:
            for w in db.query(Warehouse).filter(Warehouse.id.in_(list(set(warehouse_ids)))).all():
                wh_by_id[w.id] = w
        prod_by_id = {}
        if product_ids:
            for p in db.query(Product).filter(Product.id.in_(list(set(product_ids)))).all():
                prod_by_id[p.id] = p
        for m in movements:
            wh = wh_by_id.get(m.warehouse_id) if m.warehouse_id is not None else None
            pr = prod_by_id.get(m.product_id) if m.product_id is not None else None
            movement_rows.append({
                "date": m.created_at.strftime("%d.%m.%Y %H:%M") if m.created_at else "—",
                "warehouse_name": (getattr(wh, "name", None) if wh else None) or (f"#{m.warehouse_id}" if m.warehouse_id is not None else "—"),
                "product_name": (getattr(pr, "name", None) if pr else None) or (f"#{m.product_id}" if m.product_id is not None else "—"),
                "product_code": (getattr(pr, "code", None) or "") if pr else "",
                "doc_type_label": _tarix_doc_type_label(m.document_type or ""),
                "doc_number": m.document_number or (f"{m.document_type or ''}-{m.document_id}" if m.document_id else "—"),
                "quantity_change": float(m.quantity_change or 0),
                "warehouse_id": m.warehouse_id,
                "product_id": m.product_id,
            })

        return templates.TemplateResponse("qoldiqlar/tarix.html", {
            "request": request,
            "current_user": current_user,
            "warehouses": warehouses,
            "products": products,
            "selected_product_id": selected_product_id,
            "selected_warehouse_id": selected_warehouse_id,
            "movements": movement_rows,
            "page_title": "Qoldiqlar — Mahsulot harakati tarixi",
        })
    except Exception as e:
        import logging
        logging.getLogger(__name__).exception("qoldiqlar_tarix: %s", e)
        try:
            _wh = db.query(Warehouse).filter(Warehouse.is_active == True).all()
            _pr = db.query(Product).filter(Product.is_active == True).order_by(Product.name).all()
        except Exception:
            _wh = []
            _pr = []
        return templates.TemplateResponse("qoldiqlar/tarix.html", {
            "request": request,
            "current_user": current_user,
            "warehouses": _wh,
            "products": _pr,
            "selected_product_id": None,
            "selected_warehouse_id": None,
            "movements": [],
            "page_title": "Qoldiqlar — Mahsulot harakati tarixi",
            "error_message": str(e),
        }, status_code=500)


@app.post("/qoldiqlar/kassa/{cash_id}")
async def qoldiqlar_kassa_save(
    cash_id: int,
    balance: float = Form(...),
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    """Kassa qoldig'ini yangilash (eski tezkor forma uchun qolgan)"""
    cash = db.query(CashRegister).filter(CashRegister.id == cash_id).first()
    if not cash:
        raise HTTPException(status_code=404, detail="Kassa topilmadi")
    cash.balance = balance
    db.commit()
    return RedirectResponse(url="/qoldiqlar#kassa", status_code=303)


# --- Kassa qoldiq HUJJATLARI (1C uslubida) ---
@app.get("/qoldiqlar/kassa/hujjat/new", response_class=HTMLResponse)
async def qoldiqlar_kassa_hujjat_new(
    request: Request,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    """Yangi kassa qoldiq hujjati"""
    cash_registers = db.query(CashRegister).filter(CashRegister.is_active == True).all()
    return templates.TemplateResponse("qoldiqlar/kassa_hujjat_form.html", {
        "request": request,
        "doc": None,
        "cash_registers": cash_registers,
        "current_user": current_user,
        "page_title": "Kassa qoldiqlari — yangi hujjat",
    })


@app.post("/qoldiqlar/kassa/hujjat")
async def qoldiqlar_kassa_hujjat_create(
    request: Request,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    """Kassa qoldiq hujjatini yaratish (qoralama)"""
    form = await request.form()
    cash_ids = form.getlist("cash_register_id")
    balances = form.getlist("balance")

    items_data = []
    for i, cid in enumerate(cash_ids):
        if not cid:
            continue
        try:
            bid = int(cid)
            bal = float(balances[i]) if i < len(balances) and balances[i] != "" else None
        except (TypeError, ValueError):
            continue
        if bal is not None:
            items_data.append((bid, bal))

    if not items_data:
        return RedirectResponse(url="/qoldiqlar/kassa/hujjat/new", status_code=303)

    today = datetime.now()
    count = db.query(CashBalanceDoc).filter(
        CashBalanceDoc.date >= today.replace(hour=0, minute=0, second=0)
    ).count()
    number = f"KLD-{today.strftime('%Y%m%d')}-{str(count + 1).zfill(4)}"

    doc = CashBalanceDoc(
        number=number,
        date=today,
        user_id=current_user.id if current_user else None,
        status="draft",
    )
    db.add(doc)
    db.flush()
    for cid, bal in items_data:
        db.add(CashBalanceDocItem(doc_id=doc.id, cash_register_id=cid, balance=bal))
    db.commit()
    return RedirectResponse(url=f"/qoldiqlar/kassa/hujjat/{doc.id}", status_code=303)


@app.get("/qoldiqlar/kassa/hujjat/{doc_id}", response_class=HTMLResponse)
async def qoldiqlar_kassa_hujjat_view(
    request: Request,
    doc_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    """Kassa qoldiq hujjatini ko'rish"""
    doc = db.query(CashBalanceDoc).filter(CashBalanceDoc.id == doc_id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Hujjat topilmadi")
    cash_registers = db.query(CashRegister).filter(CashRegister.is_active == True).all()
    return templates.TemplateResponse("qoldiqlar/kassa_hujjat_form.html", {
        "request": request,
        "doc": doc,
        "cash_registers": cash_registers,
        "current_user": current_user,
        "page_title": f"Kassa qoldiqlari {doc.number}",
    })


@app.post("/qoldiqlar/kassa/hujjat/{doc_id}/tasdiqlash")
async def qoldiqlar_kassa_hujjat_tasdiqlash(
    doc_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    """Kassa hujjatini tasdiqlash — kassa balanslarini yangilash"""
    doc = db.query(CashBalanceDoc).filter(CashBalanceDoc.id == doc_id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Hujjat topilmadi")
    if doc.status != "draft":
        raise HTTPException(status_code=400, detail="Hujjat allaqachon tasdiqlangan")
    if not doc.items:
        raise HTTPException(status_code=400, detail="Kamida bitta kassa qatori bo'lishi kerak")
    for item in doc.items:
        cash = db.query(CashRegister).filter(CashRegister.id == item.cash_register_id).first()
        if cash:
            item.previous_balance = cash.balance
            cash.balance = item.balance
    doc.status = "confirmed"
    db.commit()
    return RedirectResponse(url=f"/qoldiqlar/kassa/hujjat/{doc_id}", status_code=303)


@app.post("/qoldiqlar/kassa/hujjat/{doc_id}/revert")
async def qoldiqlar_kassa_hujjat_revert(
    doc_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_admin),
):
    """Kassa hujjati tasdiqini bekor qilish (faqat admin)"""
    doc = db.query(CashBalanceDoc).filter(CashBalanceDoc.id == doc_id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Hujjat topilmadi")
    if doc.status != "confirmed":
        raise HTTPException(status_code=400, detail="Faqat tasdiqlangan hujjatning tasdiqini bekor qilish mumkin")
    for item in doc.items:
        cash = db.query(CashRegister).filter(CashRegister.id == item.cash_register_id).first()
        if cash and item.previous_balance is not None:
            cash.balance = item.previous_balance
    doc.status = "draft"
    db.commit()
    return RedirectResponse(url=f"/qoldiqlar/kassa/hujjat/{doc_id}", status_code=303)


@app.post("/qoldiqlar/kassa/hujjat/{doc_id}/delete")
async def qoldiqlar_kassa_hujjat_delete(
    doc_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_admin),
):
    """Kassa hujjatini o'chirish (faqat qoralama, faqat admin)"""
    doc = db.query(CashBalanceDoc).filter(CashBalanceDoc.id == doc_id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Hujjat topilmadi")
    if doc.status != "draft":
        raise HTTPException(status_code=400, detail="Faqat qoralama holatidagi hujjatni o'chirish mumkin. Avval tasdiqni bekor qiling.")
    db.delete(doc)
    db.commit()
    return RedirectResponse(url="/qoldiqlar#kassa", status_code=303)


# ==========================================
# KASSADAN KASSAGA O'TKAZISH (jo'natuvchi yuboradi — qabul qiluvchi tasdiqlaydi)
# ==========================================

@app.get("/cash/transfiers")
async def cash_transfiers_redirect():
    """Yozuv xatosi: transfiers -> transfers (ro'yxatga yo'naltirish)."""
    return RedirectResponse(url="/cash/transfers", status_code=301)


@app.get("/cash/transfers", response_class=HTMLResponse)
async def cash_transfers_list(
    request: Request,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    """Kassadan kassaga o'tkazish hujjatlari ro'yxati"""
    try:
        transfers = (
            db.query(CashTransfer)
            .options(
                joinedload(CashTransfer.from_cash),
                joinedload(CashTransfer.to_cash),
                joinedload(CashTransfer.user),
                joinedload(CashTransfer.approved_by),
            )
            .order_by(CashTransfer.created_at.desc())
            .limit(100)
            .all()
        )
    except Exception as e:
        err = str(e).lower()
        if "payment_type" in err or "no such column" in err or "operationalerror" in err:
            return HTMLResponse(
                "<!DOCTYPE html><html><head><meta charset='utf-8'><title>Migratsiya kerak</title></head><body style='font-family:sans-serif;padding:2rem;'>"
                "<h2>Bazada yangi ustun yo'q</h2><p>Kassalar jadvaliga <code>payment_type</code> qo'shilishi kerak. "
                "Loyiha ildizida terminalda bajariladi:</p><pre>alembic upgrade head</pre>"
                "<p><a href='/cash/transfers'>Qayta urinish</a> &nbsp; <a href='/'>Bosh sahifa</a></p></body></html>",
                status_code=500,
            )
        raise
    return templates.TemplateResponse("cash/transfers_list.html", {
        "request": request,
        "transfers": transfers,
        "current_user": current_user,
        "page_title": "Kassadan kassaga o'tkazish",
    })


@app.get("/cash/transfers/new", response_class=HTMLResponse)
async def cash_transfer_new(
    request: Request,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    """Yangi kassadan kassaga o'tkazish (qoralama)"""
    cash_registers = db.query(CashRegister).filter(CashRegister.is_active == True).order_by(CashRegister.name).all()
    return templates.TemplateResponse("cash/transfer_form.html", {
        "request": request,
        "transfer": None,
        "cash_registers": cash_registers,
        "current_user": current_user,
        "page_title": "Kassadan kassaga o'tkazish (yaratish)",
    })


@app.post("/cash/transfers/create")
async def cash_transfer_create(
    request: Request,
    from_cash_id: int = Form(...),
    to_cash_id: int = Form(...),
    amount: float = Form(...),
    note: str = Form(""),
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    from urllib.parse import quote
    if from_cash_id == to_cash_id:
        return RedirectResponse(url="/cash/transfers/new?error=" + quote("Qayerdan va qayerga kassa bir xil bo'lmasin."), status_code=303)
    if amount <= 0:
        return RedirectResponse(url="/cash/transfers/new?error=" + quote("Summa 0 dan katta bo'lishi kerak."), status_code=303)
    from_cash = db.query(CashRegister).filter(CashRegister.id == from_cash_id).first()
    if not from_cash or (from_cash.balance or 0) < amount:
        return RedirectResponse(url="/cash/transfers/new?error=" + quote("Kassada yetarli mablag' yo'q."), status_code=303)
    last_t = db.query(CashTransfer).order_by(CashTransfer.id.desc()).first()
    num = f"KK-{datetime.now().strftime('%Y%m%d')}-{(last_t.id + 1) if last_t else 1:04d}"
    t = CashTransfer(
        number=num,
        from_cash_id=from_cash_id,
        to_cash_id=to_cash_id,
        amount=amount,
        status="draft",
        user_id=current_user.id if current_user else None,
        note=note or None,
    )
    db.add(t)
    db.commit()
    return RedirectResponse(url=f"/cash/transfers/{t.id}", status_code=303)


@app.get("/cash/transfers/{transfer_id}", response_class=HTMLResponse)
async def cash_transfer_view(
    request: Request,
    transfer_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    transfer = db.query(CashTransfer).filter(CashTransfer.id == transfer_id).first()
    if not transfer:
        raise HTTPException(status_code=404, detail="Hujjat topilmadi")
    cash_registers = db.query(CashRegister).filter(CashRegister.is_active == True).order_by(CashRegister.name).all()
    return templates.TemplateResponse("cash/transfer_form.html", {
        "request": request,
        "transfer": transfer,
        "cash_registers": cash_registers,
        "current_user": current_user,
        "page_title": f"Kassadan kassaga {transfer.number}",
    })


@app.post("/cash/transfers/{transfer_id}/send")
async def cash_transfer_send(
    transfer_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    """Jo'natuvchi yuboradi — hujjat tasdiqlash kutilmoqdaga o'tadi"""
    t = db.query(CashTransfer).filter(CashTransfer.id == transfer_id).first()
    if not t:
        raise HTTPException(status_code=404, detail="Hujjat topilmadi")
    if t.status != "draft":
        from urllib.parse import quote
        return RedirectResponse(url=f"/cash/transfers/{transfer_id}?error=" + quote("Faqat qoralamani yuborish mumkin."), status_code=303)
    from_cash = db.query(CashRegister).filter(CashRegister.id == t.from_cash_id).first()
    if not from_cash or (from_cash.balance or 0) < (t.amount or 0):
        from urllib.parse import quote
        return RedirectResponse(url=f"/cash/transfers/{transfer_id}?error=" + quote("Kassada yetarli mablag' yo'q."), status_code=303)
    t.status = "pending_approval"
    db.commit()
    return RedirectResponse(url=f"/cash/transfers/{transfer_id}?sent=1", status_code=303)


@app.post("/cash/transfers/{transfer_id}/confirm")
async def cash_transfer_confirm(
    transfer_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    """Qabul qiluvchi tasdiqlaydi — from_cash dan ayiriladi, to_cash ga qo'shiladi"""
    t = db.query(CashTransfer).filter(CashTransfer.id == transfer_id).first()
    if not t:
        raise HTTPException(status_code=404, detail="Hujjat topilmadi")
    if t.status != "pending_approval":
        from urllib.parse import quote
        return RedirectResponse(url=f"/cash/transfers/{transfer_id}?error=" + quote("Faqat tasdiqlash kutilayotgan hujjatni tasdiqlash mumkin."), status_code=303)
    from_cash = db.query(CashRegister).filter(CashRegister.id == t.from_cash_id).first()
    to_cash = db.query(CashRegister).filter(CashRegister.id == t.to_cash_id).first()
    if not from_cash or not to_cash:
        from urllib.parse import quote
        return RedirectResponse(url=f"/cash/transfers/{transfer_id}?error=" + quote("Kassa topilmadi."), status_code=303)
    amount = t.amount or 0
    if (from_cash.balance or 0) < amount:
        from urllib.parse import quote
        return RedirectResponse(url=f"/cash/transfers/{transfer_id}?error=" + quote("Jo'natuvchi kassada yetarli mablag' yo'q."), status_code=303)
    from_cash.balance = (from_cash.balance or 0) - amount
    to_cash.balance = (to_cash.balance or 0) + amount
    t.status = "confirmed"
    t.approved_by_user_id = current_user.id if current_user else None
    t.approved_at = datetime.now()
    db.commit()
    return RedirectResponse(url=f"/cash/transfers/{transfer_id}?confirmed=1", status_code=303)


@app.post("/cash/transfers/{transfer_id}/revert")
async def cash_transfer_revert(
    transfer_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_admin),
):
    """Tasdiqni bekor qilish (faqat admin): balanslarni qaytarish"""
    t = db.query(CashTransfer).filter(CashTransfer.id == transfer_id).first()
    if not t or t.status != "confirmed":
        from urllib.parse import quote
        return RedirectResponse(url=f"/cash/transfers/{transfer_id}?error=" + quote("Faqat tasdiqlangan hujjatning tasdiqini bekor qilish mumkin."), status_code=303)
    amount = t.amount or 0
    from_cash = db.query(CashRegister).filter(CashRegister.id == t.from_cash_id).first()
    to_cash = db.query(CashRegister).filter(CashRegister.id == t.to_cash_id).first()
    if from_cash:
        from_cash.balance = (from_cash.balance or 0) + amount
    if to_cash:
        to_cash.balance = max(0, (to_cash.balance or 0) - amount)
    t.status = "pending_approval"
    t.approved_by_user_id = None
    t.approved_at = None
    db.commit()
    return RedirectResponse(url=f"/cash/transfers/{transfer_id}?reverted=1", status_code=303)


@app.post("/cash/transfers/{transfer_id}/delete")
async def cash_transfer_delete(
    transfer_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    t = db.query(CashTransfer).filter(CashTransfer.id == transfer_id).first()
    if not t:
        raise HTTPException(status_code=404, detail="Hujjat topilmadi")
    if t.status != "draft":
        from urllib.parse import quote
        return RedirectResponse(url=f"/cash/transfers?error=" + quote("Faqat qoralamani o'chirish mumkin."), status_code=303)
    db.delete(t)
    db.commit()
    return RedirectResponse(url="/cash/transfers?deleted=1", status_code=303)


# --- Kontragent qoldiq HUJJATLARI (1C uslubida) ---
@app.get("/qoldiqlar/kontragent/hujjat/new", response_class=HTMLResponse)
async def qoldiqlar_kontragent_hujjat_new(
    request: Request,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    """Yangi kontragent balans hujjati"""
    partners = db.query(Partner).filter(Partner.is_active == True).order_by(Partner.name).all()
    return templates.TemplateResponse("qoldiqlar/kontragent_hujjat_form.html", {
        "request": request,
        "doc": None,
        "partners": partners,
        "current_user": current_user,
        "page_title": "Kontragent qoldiqlari — yangi hujjat",
    })


@app.post("/qoldiqlar/kontragent/hujjat")
async def qoldiqlar_kontragent_hujjat_create(
    request: Request,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    """Kontragent balans hujjatini yaratish (qoralama)"""
    form = await request.form()
    partner_ids = form.getlist("partner_id")
    balances = form.getlist("balance")

    items_data = []
    for i, pid in enumerate(partner_ids):
        if not pid:
            continue
        try:
            pid_int = int(pid)
            bal_str = (balances[i] if i < len(balances) else "").strip()
            if not bal_str:
                continue
            bal = float(bal_str)
        except (TypeError, ValueError):
            continue
        items_data.append((pid_int, bal))

    if not items_data:
        return RedirectResponse(url="/qoldiqlar/kontragent/hujjat/new", status_code=303)

    today = datetime.now()
    count = db.query(PartnerBalanceDoc).filter(
        PartnerBalanceDoc.date >= today.replace(hour=0, minute=0, second=0)
    ).count()
    number = f"KNT-{today.strftime('%Y%m%d')}-{str(count + 1).zfill(4)}"

    doc = PartnerBalanceDoc(
        number=number,
        date=today,
        user_id=current_user.id if current_user else None,
        status="draft",
    )
    db.add(doc)
    db.flush()
    for pid, bal in items_data:
        db.add(PartnerBalanceDocItem(doc_id=doc.id, partner_id=pid, balance=bal))
    db.commit()
    return RedirectResponse(url=f"/qoldiqlar/kontragent/hujjat/{doc.id}", status_code=303)


@app.get("/qoldiqlar/kontragent/hujjat/{doc_id}", response_class=HTMLResponse)
async def qoldiqlar_kontragent_hujjat_view(
    request: Request,
    doc_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    """Kontragent balans hujjatini ko'rish"""
    doc = db.query(PartnerBalanceDoc).filter(PartnerBalanceDoc.id == doc_id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Hujjat topilmadi")
    partners = db.query(Partner).filter(Partner.is_active == True).order_by(Partner.name).all()
    return templates.TemplateResponse("qoldiqlar/kontragent_hujjat_form.html", {
        "request": request,
        "doc": doc,
        "partners": partners,
        "current_user": current_user,
        "page_title": f"Kontragent qoldiqlari {doc.number}",
    })


@app.post("/qoldiqlar/kontragent/hujjat/{doc_id}/tasdiqlash")
async def qoldiqlar_kontragent_hujjat_tasdiqlash(
    doc_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    """Kontragent hujjatini tasdiqlash — kontragent balanslarini yangilash"""
    doc = db.query(PartnerBalanceDoc).filter(PartnerBalanceDoc.id == doc_id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Hujjat topilmadi")
    if doc.status != "draft":
        raise HTTPException(status_code=400, detail="Hujjat allaqachon tasdiqlangan")
    if not doc.items:
        raise HTTPException(status_code=400, detail="Kamida bitta kontragent qatori bo'lishi kerak")
    for item in doc.items:
        partner = db.query(Partner).filter(Partner.id == item.partner_id).first()
        if partner:
            item.previous_balance = partner.balance
            partner.balance = item.balance
    doc.status = "confirmed"
    db.commit()
    return RedirectResponse(url=f"/qoldiqlar/kontragent/hujjat/{doc_id}", status_code=303)


@app.post("/qoldiqlar/kontragent/hujjat/{doc_id}/revert")
async def qoldiqlar_kontragent_hujjat_revert(
    doc_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_admin),
):
    """Kontragent hujjati tasdiqini bekor qilish (faqat admin)"""
    doc = db.query(PartnerBalanceDoc).filter(PartnerBalanceDoc.id == doc_id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Hujjat topilmadi")
    if doc.status != "confirmed":
        raise HTTPException(status_code=400, detail="Faqat tasdiqlangan hujjatning tasdiqini bekor qilish mumkin")
    for item in doc.items:
        partner = db.query(Partner).filter(Partner.id == item.partner_id).first()
        if partner and item.previous_balance is not None:
            partner.balance = item.previous_balance
    doc.status = "draft"
    db.commit()
    return RedirectResponse(url=f"/qoldiqlar/kontragent/hujjat/{doc_id}", status_code=303)


@app.post("/qoldiqlar/kontragent/hujjat/{doc_id}/delete")
async def qoldiqlar_kontragent_hujjat_delete(
    doc_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_admin),
):
    """Kontragent hujjatini o'chirish (faqat qoralama, faqat admin)"""
    doc = db.query(PartnerBalanceDoc).filter(PartnerBalanceDoc.id == doc_id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Hujjat topilmadi")
    if doc.status != "draft":
        raise HTTPException(status_code=400, detail="Faqat qoralama holatidagi hujjatni o'chirish mumkin. Avval tasdiqni bekor qiling.")
    db.delete(doc)
    db.commit()
    return RedirectResponse(url="/qoldiqlar#kontragent", status_code=303)


@app.post("/qoldiqlar/tovar")
async def qoldiqlar_tovar_save(
    warehouse_id: int = Form(...),
    product_id: int = Form(...),
    quantity: float = Form(...),
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    """Tovar qoldig'ini kiritish yoki qo'shish (omborda mavjud bo'lsa qo'shiladi)"""
    if quantity < 0:
        return RedirectResponse(url="/qoldiqlar#tovar", status_code=303)
    stock = db.query(Stock).filter(
        Stock.warehouse_id == warehouse_id,
        Stock.product_id == product_id,
    ).first()
    if stock:
        stock.quantity = (stock.quantity or 0) + quantity
        stock.updated_at = datetime.now()
    else:
        stock = Stock(warehouse_id=warehouse_id, product_id=product_id, quantity=quantity)
        db.add(stock)
    db.commit()
    return RedirectResponse(url="/qoldiqlar#tovar", status_code=303)


@app.post("/qoldiqlar/kontragent/{partner_id}")
async def qoldiqlar_kontragent_save(
    partner_id: int,
    balance: str = Form(""),
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    """Kontragent balansini yangilash"""
    partner = db.query(Partner).filter(Partner.id == partner_id).first()
    if not partner:
        raise HTTPException(status_code=404, detail="Kontragent topilmadi")
    balance_str = (balance or "").strip()
    if not balance_str:
        return RedirectResponse(url="/qoldiqlar#kontragent", status_code=303)
    try:
        partner.balance = float(balance_str)
    except (TypeError, ValueError):
        return RedirectResponse(url="/qoldiqlar#kontragent", status_code=303)
    db.commit()
    return RedirectResponse(url="/qoldiqlar#kontragent", status_code=303)


@app.get("/qoldiqlar/export")
async def qoldiqlar_export(db: Session = Depends(get_db), current_user: User = Depends(require_auth)):
    """Tovar qoldiqlari hisoboti — Excel hujjat sifatida yuklab olish"""
    stocks = (
        db.query(Stock)
        .join(Warehouse, Stock.warehouse_id == Warehouse.id)
        .join(Product, Stock.product_id == Product.id)
        .order_by(Warehouse.name, Product.name)
        .all()
    )
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Tovar qoldiqlari"
    ws.append(["Ombor", "Mahsulot", "Kod", "Miqdor"])
    for s in stocks:
        ws.append([
            s.warehouse.name if s.warehouse else "-",
            s.product.name if s.product else "-",
            (s.product.code or "") if s.product else "",
            float(s.quantity) if s.quantity is not None else 0,
        ])
    stream = io.BytesIO()
    wb.save(stream)
    stream.seek(0)
    return StreamingResponse(
        stream,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": "attachment; filename=tovar_qoldiqlari.xlsx"},
    )


# --- Tovar qoldiq HUJJATLARI (1C uslubida: ro'yxat + hujjat + qatorlar) ---
@app.get("/qoldiqlar/tovar/hujjat", response_class=HTMLResponse)
async def qoldiqlar_tovar_hujjat_list(
    request: Request,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    """Tovar qoldiqlari hujjatlari ro'yxati"""
    docs = (
        db.query(StockAdjustmentDoc)
        .order_by(StockAdjustmentDoc.created_at.desc())
        .limit(200)
        .all()
    )
    reverted = request.query_params.get("reverted") == "1"
    return templates.TemplateResponse("qoldiqlar/hujjat_list.html", {
        "request": request,
        "docs": docs,
        "current_user": current_user,
        "page_title": "Tovar qoldiqlari hujjatlari",
        "reverted": reverted,
    })


@app.get("/qoldiqlar/tovar/hujjat/new", response_class=HTMLResponse)
async def qoldiqlar_tovar_hujjat_new(
    request: Request,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    """Yangi tovar qoldiq hujjati (qoralama)"""
    warehouses = db.query(Warehouse).filter(Warehouse.is_active == True).all()
    products = db.query(Product).filter(Product.is_active == True).order_by(Product.name).all()
    return templates.TemplateResponse("qoldiqlar/hujjat_form.html", {
        "request": request,
        "doc": None,
        "warehouses": warehouses,
        "products": products,
        "current_user": current_user,
        "page_title": "Tovar qoldiqlari — yangi hujjat",
    })


@app.post("/qoldiqlar/tovar/hujjat")
async def qoldiqlar_tovar_hujjat_create(
    request: Request,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    """Tovar qoldiq hujjatini yaratish (qoralama)"""
    form = await request.form()
    product_ids = form.getlist("product_id")
    warehouse_ids = form.getlist("warehouse_id")
    quantities = form.getlist("quantity")
    cost_prices = form.getlist("cost_price")
    sale_prices = form.getlist("sale_price")

    items_data = []
    for i, pid in enumerate(product_ids):
        if not pid or not str(pid).strip():
            continue
        try:
            wid = int(warehouse_ids[i]) if i < len(warehouse_ids) and warehouse_ids[i] else None
            qty = float(quantities[i]) if i < len(quantities) and str(quantities[i]).strip() else 0
            _cp = cost_prices[i] if i < len(cost_prices) else ""
            _sp = sale_prices[i] if i < len(sale_prices) else ""
            cp = float(_cp) if str(_cp).strip() else 0
            sp = float(_sp) if str(_sp).strip() else 0
        except (TypeError, ValueError):
            continue
        if wid and qty > 0:
            try:
                items_data.append((int(pid), wid, qty, cp, sp))
            except ValueError:
                continue

    today = datetime.now()
    count = db.query(StockAdjustmentDoc).filter(
        StockAdjustmentDoc.date >= today.replace(hour=0, minute=0, second=0)
    ).count()
    number = f"QLD-{today.strftime('%Y%m%d')}-{str(count + 1).zfill(4)}"

    total_tannarx = sum(qty * cp for _, _, qty, cp, _ in items_data)
    total_sotuv = sum(qty * sp for _, _, qty, _, sp in items_data)

    doc = StockAdjustmentDoc(
        number=number,
        date=today,
        user_id=current_user.id if current_user else None,
        status="draft",
        total_tannarx=total_tannarx,
        total_sotuv=total_sotuv,
    )
    db.add(doc)
    db.flush()

    for pid, wid, qty, cp, sp in items_data:
        db.add(StockAdjustmentDocItem(
            doc_id=doc.id,
            product_id=pid,
            warehouse_id=wid,
            quantity=qty,
            cost_price=cp,
            sale_price=sp,
        ))
    db.commit()
    return RedirectResponse(url=f"/qoldiqlar/tovar/hujjat/{doc.id}", status_code=303)


@app.post("/qoldiqlar/tovar/import-excel")
async def qoldiqlar_tovar_import_excel(
    request: Request,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    """Exceldan tovar qoldiqlarini yuklash — hujjat (QLD-...) yaratiladi, jadvalda ko'rinadi."""
    from urllib.parse import quote
    form = await request.form()
    file = form.get("file") or form.get("excel_file")
    if not file or not getattr(file, "filename", None):
        return RedirectResponse(url="/qoldiqlar?error=import&detail=" + quote("Excel fayl tanlang") + "#tovar", status_code=303)
    try:
        contents = await file.read()
        if not contents:
            return RedirectResponse(url="/qoldiqlar?error=import&detail=" + quote("Fayl bo'sh") + "#tovar", status_code=303)
        wb = openpyxl.load_workbook(io.BytesIO(contents), read_only=False, data_only=True)
        ws = wb.active
        rows = list(ws.iter_rows(min_row=2, values_only=True))
        items_data = []
        skip_no_wh = 0
        skip_no_prod = 0
        skip_empty = 0
        missing_products = []
        missing_warehouses = []
        for row in rows:
            if not row or (row[0] is None and (len(row) < 2 or row[1] is None)):
                skip_empty += 1
                continue
            wh_key = str(row[0] or "").strip() if len(row) > 0 else ""
            raw_prod = row[1] if len(row) > 1 else None
            if raw_prod is not None and isinstance(raw_prod, (int, float)) and float(raw_prod) == int(float(raw_prod)):
                prod_key = str(int(float(raw_prod)))
            else:
                prod_key = str(raw_prod or "").strip()
            try:
                qty = float(row[2]) if len(row) > 2 and row[2] is not None else 0
            except (TypeError, ValueError):
                qty = 0
            cp = 0.0
            sp = 0.0
            if len(row) > 3 and row[3] is not None and row[3] != "":
                try:
                    cp = float(row[3])
                except (TypeError, ValueError):
                    pass
            if len(row) > 4 and row[4] is not None and row[4] != "":
                try:
                    sp = float(row[4])
                except (TypeError, ValueError):
                    pass
            if not wh_key or not prod_key:
                skip_empty += 1
                continue
            warehouse = db.query(Warehouse).filter(
                (func.lower(Warehouse.name) == wh_key.lower()) | (Warehouse.code == wh_key)
            ).first()
            product = db.query(Product).filter(
                or_(
                    and_(Product.code.isnot(None), Product.code != "", func.lower(func.trim(Product.code)) == prod_key.lower()),
                    and_(Product.barcode.isnot(None), Product.barcode != "", func.lower(func.trim(Product.barcode)) == prod_key.lower()),
                )
            ).first()
            if not product and prod_key:
                product = db.query(Product).filter(
                    Product.name.isnot(None),
                    func.lower(func.trim(Product.name)) == prod_key.strip().lower()
                ).first()
            if not warehouse:
                if wh_key and wh_key not in missing_warehouses:
                    missing_warehouses.append(wh_key)
                skip_no_wh += 1
                continue
            if not product:
                if prod_key and prod_key not in missing_products:
                    missing_products.append(prod_key)
                skip_no_prod += 1
                continue
            items_data.append((product.id, warehouse.id, qty, cp, sp))
        if not items_data:
            detail = "Hech qanday to'g'ri qator topilmadi. Ombor va mahsulot nomi/kodi to'g'ri ekanligini tekshiring."
            if missing_products:
                detail += " Mahsulot topilmadi: " + ", ".join(missing_products[:10])
                if len(missing_products) > 10:
                    detail += f" va yana {len(missing_products) - 10} ta"
            if missing_warehouses:
                detail += ". Ombor topilmadi: " + ", ".join(missing_warehouses[:5])
            return RedirectResponse(
                url="/qoldiqlar?error=import&detail=" + quote(detail) + "#tovar",
                status_code=303,
            )
        today = datetime.now()
        count = db.query(StockAdjustmentDoc).filter(
            StockAdjustmentDoc.date >= today.replace(hour=0, minute=0, second=0)
        ).count()
        number = f"QLD-{today.strftime('%Y%m%d')}-{str(count + 1).zfill(4)}"
        total_tannarx = sum(qty * cp for _, _, qty, cp, _ in items_data)
        total_sotuv = sum(qty * sp for _, _, qty, _, sp in items_data)
        doc = StockAdjustmentDoc(
            number=number,
            date=today,
            user_id=current_user.id if current_user else None,
            status="draft",
            total_tannarx=total_tannarx,
            total_sotuv=total_sotuv,
        )
        db.add(doc)
        db.flush()
        for pid, wid, qty, cp, sp in items_data:
            db.add(StockAdjustmentDocItem(
                doc_id=doc.id,
                product_id=pid,
                warehouse_id=wid,
                quantity=qty,
                cost_price=cp,
                sale_price=sp,
            ))
        db.commit()
        detail = f"Yuklandi: {len(items_data)} ta"
        if skip_no_prod or skip_no_wh:
            detail += f", o'tkazib yuborildi: {skip_no_prod + skip_no_wh} ta"
        if missing_products:
            detail += ". Mahsulot topilmadi: " + ", ".join(missing_products[:8])
            if len(missing_products) > 8:
                detail += f" va yana {len(missing_products) - 8} ta"
        if missing_warehouses:
            detail += ". Ombor topilmadi: " + ", ".join(missing_warehouses[:3])
        return RedirectResponse(
            url="/qoldiqlar?success=import&detail=" + quote(detail) + "&doc_id=" + str(doc.id) + "&doc_number=" + quote(doc.number) + "#tovar",
            status_code=303,
        )
    except Exception as e:
        traceback.print_exc()
        return RedirectResponse(
            url="/qoldiqlar?error=import&detail=" + quote(str(e)[:180]) + "#tovar",
            status_code=303,
        )


@app.get("/qoldiqlar/tovar/hujjat/{doc_id}", response_class=HTMLResponse)
async def qoldiqlar_tovar_hujjat_view(
    request: Request,
    doc_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    """Tovar qoldiq hujjatini ko'rish/tahrirlash"""
    doc = db.query(StockAdjustmentDoc).filter(StockAdjustmentDoc.id == doc_id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Hujjat topilmadi")
    warehouses = db.query(Warehouse).filter(Warehouse.is_active == True).all()
    products = db.query(Product).filter(Product.is_active == True).order_by(Product.name).all()
    return templates.TemplateResponse("qoldiqlar/hujjat_form.html", {
        "request": request,
        "doc": doc,
        "warehouses": warehouses,
        "products": products,
        "current_user": current_user,
        "page_title": f"Tovar qoldiqlari {doc.number}",
    })


@app.post("/qoldiqlar/tovar/hujjat/{doc_id}/add-row")
async def qoldiqlar_tovar_hujjat_add_row(
    doc_id: int,
    product_id: int = Form(...),
    warehouse_id: int = Form(...),
    quantity: float = Form(...),
    cost_price: float = Form(0),
    sale_price: float = Form(0),
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    """Hujjatga qator qo'shish (faqat qoralama)"""
    doc = db.query(StockAdjustmentDoc).filter(StockAdjustmentDoc.id == doc_id).first()
    if not doc or doc.status != "draft":
        raise HTTPException(status_code=400, detail="Faqat qoralamani tahrirlash mumkin")
    if quantity <= 0:
        return RedirectResponse(url=f"/qoldiqlar/tovar/hujjat/{doc_id}", status_code=303)
    doc.total_tannarx = (doc.total_tannarx or 0) + quantity * (cost_price or 0)
    doc.total_sotuv = (doc.total_sotuv or 0) + quantity * (sale_price or 0)
    db.add(StockAdjustmentDocItem(
        doc_id=doc_id,
        product_id=product_id,
        warehouse_id=warehouse_id,
        quantity=quantity,
        cost_price=cost_price or 0,
        sale_price=sale_price or 0,
    ))
    db.commit()
    return RedirectResponse(url=f"/qoldiqlar/tovar/hujjat/{doc_id}", status_code=303)


@app.post("/qoldiqlar/tovar/hujjat/{doc_id}/delete-row/{item_id}")
async def qoldiqlar_tovar_hujjat_delete_row(
    doc_id: int,
    item_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    """Hujjatdan qatorni o'chirish (faqat qoralama)"""
    doc = db.query(StockAdjustmentDoc).filter(StockAdjustmentDoc.id == doc_id).first()
    if not doc or doc.status != "draft":
        raise HTTPException(status_code=400, detail="Faqat qoralamani tahrirlash mumkin")
    item = db.query(StockAdjustmentDocItem).filter(
        StockAdjustmentDocItem.id == item_id,
        StockAdjustmentDocItem.doc_id == doc_id,
    ).first()
    if not item:
        raise HTTPException(status_code=404, detail="Qator topilmadi")
    doc.total_tannarx = (doc.total_tannarx or 0) - (item.quantity * (item.cost_price or 0))
    doc.total_sotuv = (doc.total_sotuv or 0) - (item.quantity * (item.sale_price or 0))
    if doc.total_tannarx < 0:
        doc.total_tannarx = 0
    if doc.total_sotuv < 0:
        doc.total_sotuv = 0
    db.delete(item)
    db.commit()
    return RedirectResponse(url=f"/qoldiqlar/tovar/hujjat/{doc_id}", status_code=303)


@app.post("/qoldiqlar/tovar/hujjat/{doc_id}/edit-row/{item_id}")
async def qoldiqlar_tovar_hujjat_edit_row(
    doc_id: int,
    item_id: int,
    quantity: float = Form(...),
    warehouse_id: int = Form(...),
    cost_price: float = Form(0),
    sale_price: float = Form(0),
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    """Hujjat qatorini tahrirlash (faqat qoralama): soni, ombor, tannarx, sotuv narx"""
    doc = db.query(StockAdjustmentDoc).filter(StockAdjustmentDoc.id == doc_id).first()
    if not doc or doc.status != "draft":
        raise HTTPException(status_code=400, detail="Faqat qoralamani tahrirlash mumkin")
    item = db.query(StockAdjustmentDocItem).filter(
        StockAdjustmentDocItem.id == item_id,
        StockAdjustmentDocItem.doc_id == doc_id,
    ).first()
    if not item:
        raise HTTPException(status_code=404, detail="Qator topilmadi")
    if quantity <= 0:
        return RedirectResponse(url=f"/qoldiqlar/tovar/hujjat/{doc_id}", status_code=303)
    # Eski summalarni hisobdan chiqarish
    doc.total_tannarx = (doc.total_tannarx or 0) - (item.quantity * (item.cost_price or 0))
    doc.total_sotuv = (doc.total_sotuv or 0) - (item.quantity * (item.sale_price or 0))
    item.quantity = quantity
    item.warehouse_id = warehouse_id
    item.cost_price = cost_price or 0
    item.sale_price = sale_price or 0
    doc.total_tannarx = (doc.total_tannarx or 0) + item.quantity * item.cost_price
    doc.total_sotuv = (doc.total_sotuv or 0) + item.quantity * item.sale_price
    db.commit()
    return RedirectResponse(url=f"/qoldiqlar/tovar/hujjat/{doc_id}", status_code=303)


@app.post("/qoldiqlar/tovar/hujjat/{doc_id}/tasdiqlash")
async def qoldiqlar_tovar_hujjat_tasdiqlash(
    doc_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    """Hujjatni tasdiqlash — ombor qoldiqlariga qo'shiladi"""
    doc = db.query(StockAdjustmentDoc).filter(StockAdjustmentDoc.id == doc_id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Hujjat topilmadi")
    if doc.status != "draft":
        raise HTTPException(status_code=400, detail="Hujjat allaqachon tasdiqlangan")
    if not doc.items:
        raise HTTPException(status_code=400, detail="Kamida bitta qator bo'lishi kerak")

    for item in doc.items:
        # Eski qoldiqni olish
        stock = db.query(Stock).filter(
            Stock.warehouse_id == item.warehouse_id,
            Stock.product_id == item.product_id,
        ).first()
        old_quantity = stock.quantity if stock else 0
        
        # Yangi qoldiqni hisoblash
        new_quantity = item.quantity
        quantity_change = new_quantity - old_quantity
        
        if stock:
            stock.quantity = new_quantity
            stock.updated_at = datetime.now()
        else:
            db.add(Stock(
                warehouse_id=item.warehouse_id,
                product_id=item.product_id,
                quantity=item.quantity,
            ))
        
        # StockMovement yozuvini yaratish (adjustment)
        if quantity_change != 0:
            create_stock_movement(
                db=db,
                warehouse_id=item.warehouse_id,
                product_id=item.product_id,
                quantity_change=quantity_change,  # O'zgarish (+ yoki -)
                operation_type="adjustment",
                document_type="StockAdjustmentDoc",
                document_id=doc.id,
                document_number=doc.number,
                user_id=current_user.id if current_user else None,
                note=f"Qoldiq tuzatish: {doc.number}"
            )
        # Tannarxni mahsulotga yozish — Narxni o'rnatish sahifasida ko'rinsin
        if (item.cost_price or 0) > 0:
            prod = db.query(Product).filter(Product.id == item.product_id).first()
            if prod:
                prod.purchase_price = item.cost_price
    doc.status = "confirmed"
    db.commit()
    return RedirectResponse(url=f"/qoldiqlar/tovar/hujjat/{doc_id}", status_code=303)


@app.post("/qoldiqlar/tovar/hujjat/{doc_id}/revert")
async def qoldiqlar_tovar_hujjat_revert(
    doc_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_admin),
):
    """Tovar qoldiq hujjati tasdiqini bekor qilish (faqat admin) — ombor qoldig'ini kamaytirish"""
    doc = db.query(StockAdjustmentDoc).filter(StockAdjustmentDoc.id == doc_id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Hujjat topilmadi")
    if doc.status != "confirmed":
        raise HTTPException(status_code=400, detail="Faqat tasdiqlangan hujjatning tasdiqini bekor qilish mumkin")
    for item in doc.items:
        stock = db.query(Stock).filter(
            Stock.warehouse_id == item.warehouse_id,
            Stock.product_id == item.product_id,
        ).first()
        if stock:
            stock.quantity = (stock.quantity or 0) - item.quantity
            if stock.quantity < 0:
                stock.quantity = 0
            stock.updated_at = datetime.now()
    delete_stock_movements_for_document(db, "StockAdjustmentDoc", doc_id)
    doc.status = "draft"
    db.commit()
    return RedirectResponse(url="/qoldiqlar/tovar/hujjat?reverted=1", status_code=303)


@app.post("/qoldiqlar/tovar/hujjat/{doc_id}/delete")
async def qoldiqlar_tovar_hujjat_delete(
    doc_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_admin),
):
    """Tovar qoldiq hujjatini o'chirish (faqat qoralama, faqat admin)"""
    doc = db.query(StockAdjustmentDoc).filter(StockAdjustmentDoc.id == doc_id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Hujjat topilmadi")
    if doc.status != "draft":
        raise HTTPException(status_code=400, detail="Faqat qoralama holatidagi hujjatni o'chirish mumkin. Avval tasdiqni bekor qiling.")
    db.delete(doc)
    db.commit()
    return RedirectResponse(url="/qoldiqlar#tovar", status_code=303)


# Bo'limlar bo'limi
@app.get("/info/departments", response_class=HTMLResponse)
async def info_departments(request: Request, db: Session = Depends(get_db), current_user: User = Depends(require_auth)):
    departments = db.query(Department).all()
    return templates.TemplateResponse("info/departments.html", {"request": request, "departments": departments, "current_user": current_user, "page_title": "Bo'limlar"})

@app.post("/info/departments/add")
async def info_departments_add(
    request: Request,
    code: str = Form(...),
    name: str = Form(...),
    description: str = Form(""),
    db: Session = Depends(get_db)
):
    existing = db.query(Department).filter(Department.code == code).first()
    if existing:
        raise HTTPException(status_code=400, detail=f"'{code}' kodli bo'lim allaqachon mavjud!")
    
    department = Department(code=code, name=name, description=description, is_active=True)
    db.add(department)
    db.commit()
    return RedirectResponse(url="/info/departments", status_code=303)

@app.post("/info/departments/edit/{department_id}")
async def info_departments_edit(
    department_id: int,
    code: str = Form(...),
    name: str = Form(...),
    description: str = Form(""),
    db: Session = Depends(get_db)
):
    department = db.query(Department).filter(Department.id == department_id).first()
    if not department:
        raise HTTPException(status_code=404, detail="Bo'lim topilmadi")
    
    existing = db.query(Department).filter(Department.code == code, Department.id != department_id).first()
    if existing:
        raise HTTPException(status_code=400, detail=f"'{code}' kodli bo'lim allaqachon mavjud!")
    
    department.code = code
    department.name = name
    department.description = description
    db.commit()
    return RedirectResponse(url="/info/departments", status_code=303)

@app.post("/info/departments/delete/{department_id}")
async def info_departments_delete(department_id: int, db: Session = Depends(get_db), current_user: User = Depends(require_auth)):
    department = db.query(Department).filter(Department.id == department_id).first()
    if not department:
        raise HTTPException(status_code=404, detail="Bo'lim topilmadi")
    
    db.delete(department)
    db.commit()
    return RedirectResponse(url="/info/departments", status_code=303)

# --- DEPARTMENTS EXCEL OPERATIONS ---
@app.get("/info/departments/export")
async def export_departments(db: Session = Depends(get_db), current_user: User = Depends(require_auth)):
    departments = db.query(Department).all()
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Departments"
    ws.append(["ID", "Kod", "Nomi", "Izoh"])
    for d in departments:
        ws.append([d.id, d.code, d.name, d.description])
    stream = io.BytesIO()
    wb.save(stream)
    stream.seek(0)
    return StreamingResponse(stream, media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", headers={"Content-Disposition": "attachment; filename=bolimlar.xlsx"})

@app.get("/info/departments/template")
async def template_departments(current_user: User = Depends(require_auth)):
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Template"
    ws.append(["Kod", "Nomi", "Izoh"])
    ws.append(["DEP001", "Ishlab chiqarish", "Asosiy tsex"])
    stream = io.BytesIO()
    wb.save(stream)
    stream.seek(0)
    return StreamingResponse(stream, media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", headers={"Content-Disposition": "attachment; filename=bolim_andoza.xlsx"})

@app.post("/info/departments/import")
async def import_departments(file: UploadFile = File(...), db: Session = Depends(get_db), current_user: User = Depends(require_auth)):
    contents = await file.read()
    wb = openpyxl.load_workbook(io.BytesIO(contents))
    ws = wb.active
    rows = list(ws.iter_rows(min_row=2, values_only=True))
    for row in rows:
        if not row[0]: continue
        code, name, description = row[0], row[1], row[2]
        department = db.query(Department).filter(Department.code == code).first()
        if not department:
            department = Department(code=code, name=name, description=description)
            db.add(department)
        else:
            department.name = name
            department.description = description
        db.commit()
    return RedirectResponse(url="/info/departments", status_code=303)

# Yo'nalishlar bo'limi
@app.get("/info/directions", response_class=HTMLResponse)
async def info_directions(request: Request, db: Session = Depends(get_db), current_user: User = Depends(require_auth)):
    directions = db.query(Direction).all()
    return templates.TemplateResponse("info/directions.html", {"request": request, "directions": directions, "current_user": current_user, "page_title": "Yo'nalishlar"})

@app.post("/info/directions/add")
async def info_directions_add(
    request: Request,
    code: str = Form(...),
    name: str = Form(...),
    description: str = Form(""),
    db: Session = Depends(get_db)
):
    existing = db.query(Direction).filter(Direction.code == code).first()
    if existing:
        raise HTTPException(status_code=400, detail=f"'{code}' kodli yo'nalish allaqachon mavjud!")
    
    direction = Direction(code=code, name=name, description=description, is_active=True)
    db.add(direction)
    db.commit()
    return RedirectResponse(url="/info/directions", status_code=303)

@app.post("/info/directions/edit/{direction_id}")
async def info_directions_edit(
    direction_id: int,
    code: str = Form(...),
    name: str = Form(...),
    description: str = Form(""),
    db: Session = Depends(get_db)
):
    direction = db.query(Direction).filter(Direction.id == direction_id).first()
    if not direction:
        raise HTTPException(status_code=404, detail="Yo'nalish topilmadi")
    
    existing = db.query(Direction).filter(Direction.code == code, Direction.id != direction_id).first()
    if existing:
        raise HTTPException(status_code=400, detail=f"'{code}' kodli yo'nalish allaqachon mavjud!")
    
    direction.code = code
    direction.name = name
    direction.description = description
    db.commit()
    return RedirectResponse(url="/info/directions", status_code=303)

@app.post("/info/directions/delete/{direction_id}")
async def info_directions_delete(direction_id: int, db: Session = Depends(get_db), current_user: User = Depends(require_auth)):
    direction = db.query(Direction).filter(Direction.id == direction_id).first()
    if not direction:
        raise HTTPException(status_code=404, detail="Yo'nalish topilmadi")
    
    db.delete(direction)
    db.commit()
    return RedirectResponse(url="/info/directions", status_code=303)

# --- DIRECTIONS EXCEL OPERATIONS ---
@app.get("/info/directions/export")
async def export_directions(db: Session = Depends(get_db), current_user: User = Depends(require_auth)):
    directions = db.query(Direction).all()
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Directions"
    ws.append(["ID", "Kod", "Nomi", "Izoh"])
    for d in directions:
        ws.append([d.id, d.code, d.name, d.description])
    stream = io.BytesIO()
    wb.save(stream)
    stream.seek(0)
    return StreamingResponse(stream, media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", headers={"Content-Disposition": "attachment; filename=yonalishlar.xlsx"})

@app.get("/info/directions/template")
async def template_directions(current_user: User = Depends(require_auth)):
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Template"
    ws.append(["Kod", "Nomi", "Izoh"])
    ws.append(["DIR001", "Halva", "Halva mahsulotlari"])
    stream = io.BytesIO()
    wb.save(stream)
    stream.seek(0)
    return StreamingResponse(stream, media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", headers={"Content-Disposition": "attachment; filename=yonalish_andoza.xlsx"})

@app.post("/info/directions/import")
async def import_directions(file: UploadFile = File(...), db: Session = Depends(get_db), current_user: User = Depends(require_auth)):
    contents = await file.read()
    wb = openpyxl.load_workbook(io.BytesIO(contents))
    ws = wb.active
    rows = list(ws.iter_rows(min_row=2, values_only=True))
    for row in rows:
        if not row[0]: continue
        code, name, description = row[0], row[1], row[2]
        direction = db.query(Direction).filter(Direction.code == code).first()
        if not direction:
            direction = Direction(code=code, name=name, description=description)
            db.add(direction)
        else:
            direction.name = name
            direction.description = description
        db.commit()
    return RedirectResponse(url="/info/directions", status_code=303)

# Foydalanuvchilar bo'limi (faqat admin)
@app.get("/info/users", response_class=HTMLResponse)
async def info_users(request: Request, db: Session = Depends(get_db), current_user: User = Depends(require_admin)):
    """Foydalanuvchilar ro'yxati — hodimlar, bo'lim, ombor, kassa ro'yxatlari bilan."""
    users = (
        db.query(User)
        .options(
            joinedload(User.department),
            joinedload(User.warehouse),
            joinedload(User.cash_register),
            joinedload(User.departments_list),
            joinedload(User.warehouses_list),
            joinedload(User.cash_registers_list),
        )
        .order_by(User.id)
        .all()
    )
    employees = (
        db.query(Employee)
        .filter(Employee.is_active == True)
        .order_by(Employee.full_name)
        .all()
    )
    user_to_employee = {e.user_id: e for e in db.query(Employee).filter(Employee.user_id != None).all()}
    departments = db.query(Department).filter(Department.is_active == True).order_by(Department.name).all()
    warehouses = db.query(Warehouse).filter(Warehouse.is_active == True).order_by(Warehouse.name).all()
    cash_registers = db.query(CashRegister).filter(CashRegister.is_active == True).order_by(CashRegister.name).all()
    error = request.query_params.get("error", "").strip()
    return templates.TemplateResponse("info/users.html", {
        "request": request,
        "users": users,
        "employees": employees,
        "user_to_employee": user_to_employee,
        "departments": departments,
        "warehouses": warehouses,
        "cash_registers": cash_registers,
        "current_user": current_user,
        "page_title": "Foydalanuvchilar",
        "error": error,
    })

@app.post("/info/users/add")
async def info_users_add(
    request: Request,
    username: str = Form(...),
    password: str = Form(...),
    full_name: str = Form(...),
    role: str = Form("user"),
    is_active: bool = Form(True),
    db: Session = Depends(get_db),
    current_user: User = Depends(require_admin)
):
    """Yangi foydalanuvchi qo'shish (faqat admin)"""
    # Username dublikat tekshiruvi
    existing = db.query(User).filter(User.username == username).first()
    if existing:
        raise HTTPException(status_code=400, detail=f"'{username}' login bilan foydalanuvchi allaqachon mavjud!")
    
    # Yangi foydalanuvchi yaratish
    user = User(
        username=username,
        password_hash=hash_password(password),
        full_name=full_name,
        role=role,
        is_active=is_active
    )
    db.add(user)
    db.commit()
    return RedirectResponse(url="/info/users", status_code=303)

@app.post("/info/users/edit/{user_id}")
async def info_users_edit(
    user_id: int,
    username: str = Form(...),
    full_name: str = Form(...),
    role: str = Form("user"),
    is_active: bool = Form(True),
    db: Session = Depends(get_db),
    current_user: User = Depends(require_admin)
):
    """Foydalanuvchini tahrirlash (faqat admin)"""
    user = db.query(User).filter(User.id == user_id).first()
    if not user:
        raise HTTPException(status_code=404, detail="Foydalanuvchi topilmadi")
    
    # Username dublikat tekshiruvi (o'zidan boshqa)
    existing = db.query(User).filter(
        User.username == username,
        User.id != user_id
    ).first()
    if existing:
        raise HTTPException(status_code=400, detail=f"'{username}' login bilan foydalanuvchi allaqachon mavjud!")
    
    user.username = username
    user.full_name = full_name
    user.role = role
    user.is_active = is_active
    db.commit()
    return RedirectResponse(url="/info/users", status_code=303)

@app.post("/info/users/change-password/{user_id}")
async def info_users_change_password(
    user_id: int,
    new_password: str = Form(...),
    db: Session = Depends(get_db),
    current_user: User = Depends(require_admin)
):
    """Foydalanuvchi parolini o'zgartirish (faqat admin)"""
    user = db.query(User).filter(User.id == user_id).first()
    if not user:
        raise HTTPException(status_code=404, detail="Foydalanuvchi topilmadi")
    
    user.password_hash = hash_password(new_password)
    db.commit()
    return RedirectResponse(url="/info/users", status_code=303)

@app.post("/info/users/delete/{user_id}")
async def info_users_delete(user_id: int, db: Session = Depends(get_db), current_user: User = Depends(require_admin)):
    """Foydalanuvchini o'chirish (faqat admin)"""
    # O'zini o'chirishga ruxsat bermaslik
    if user_id == current_user.id:
        raise HTTPException(status_code=400, detail="O'zingizni o'chira olmaysiz!")
    
    user = db.query(User).filter(User.id == user_id).first()
    if not user:
        raise HTTPException(status_code=404, detail="Foydalanuvchi topilmadi")
    
    db.delete(user)
    db.commit()
    return RedirectResponse(url="/info/users", status_code=303)


# ==========================================
# LAVOZIMLAR
# ==========================================
@app.get("/info/positions", response_class=HTMLResponse)
async def info_positions(request: Request, db: Session = Depends(get_db), current_user: User = Depends(require_auth)):
    """Lavozimlar ro'yxati"""
    if not current_user:
        return RedirectResponse(url="/login", status_code=303)
    positions = db.query(Position).filter(Position.is_active == True).all()
    return templates.TemplateResponse("info/positions.html", {
        "request": request,
        "current_user": current_user,
        "positions": positions,
        "page_title": "Lavozimlar"
    })


@app.post("/info/positions/add")
async def info_positions_add(
    code: str = Form(...),
    name: str = Form(...),
    description: str = Form(""),
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth)
):
    """Lavozim qo'shish"""
    existing = db.query(Position).filter(Position.code == code).first()
    if existing:
        raise HTTPException(status_code=400, detail=f"'{code}' kodli lavozim allaqachon mavjud!")
    position = Position(code=code, name=name, description=description or None)
    db.add(position)
    db.commit()
    return RedirectResponse(url="/info/positions", status_code=303)


@app.post("/info/positions/edit/{position_id}")
async def info_positions_edit(
    position_id: int,
    code: str = Form(...),
    name: str = Form(...),
    description: str = Form(""),
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth)
):
    """Lavozimni tahrirlash"""
    position = db.query(Position).filter(Position.id == position_id).first()
    if not position:
        raise HTTPException(status_code=404, detail="Lavozim topilmadi")
    existing = db.query(Position).filter(Position.code == code, Position.id != position_id).first()
    if existing:
        raise HTTPException(status_code=400, detail=f"'{code}' kodli lavozim allaqachon mavjud!")
    position.code = code
    position.name = name
    position.description = description or None
    db.commit()
    return RedirectResponse(url="/info/positions", status_code=303)


@app.post("/info/positions/delete/{position_id}")
async def info_positions_delete(position_id: int, db: Session = Depends(get_db), current_user: User = Depends(require_auth)):
    """Lavozimni o'chirish (soft - is_active=False)"""
    position = db.query(Position).filter(Position.id == position_id).first()
    if not position:
        raise HTTPException(status_code=404, detail="Lavozim topilmadi")
    position.is_active = False
    db.commit()
    return RedirectResponse(url="/info/positions", status_code=303)


# --- ISHLAB CHIQARISH GURUHLARI (QIYOMCHILAR) ---
@app.get("/info/production-groups", response_class=HTMLResponse)
async def info_production_groups(
    request: Request,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    """Ishlab chiqarish guruhlari — operator + a'zolar, kunlik tabel bo'yicha bo'lak taqsimlanadi."""
    groups = (
        db.query(ProductionGroup)
        .options(joinedload(ProductionGroup.operator), joinedload(ProductionGroup.piecework_task), joinedload(ProductionGroup.members))
        .filter(ProductionGroup.is_active == True)
        .order_by(ProductionGroup.name)
        .all()
    )
    employees = db.query(Employee).filter(Employee.is_active == True).order_by(Employee.full_name).all()
    piecework_tasks = db.query(PieceworkTask).filter(PieceworkTask.is_active == True).order_by(PieceworkTask.name).all()
    return templates.TemplateResponse("info/production_groups.html", {
        "request": request,
        "groups": groups,
        "employees": employees,
        "piecework_tasks": piecework_tasks,
        "current_user": current_user,
        "page_title": "Ishlab chiqarish guruhlari (qiyomchilar)",
    })


@app.post("/info/production-groups/add", response_class=RedirectResponse)
async def info_production_groups_add(
    request: Request,
    name: str = Form(...),
    operator_id: int = Form(...),
    piecework_task_id: Optional[int] = Form(None),
    include_qiyom: str = Form("1"),
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    """Guruh qo'shish."""
    form = await request.form()
    member_ids_raw = form.getlist("member_ids") if hasattr(form, "getlist") else []
    member_ids = list(dict.fromkeys([int(x) for x in member_ids_raw if str(x).strip().isdigit()]))
    gr = ProductionGroup(
        name=name.strip(),
        operator_id=operator_id,
        piecework_task_id=int(piecework_task_id) if piecework_task_id else None,
        include_qiyom=(include_qiyom == "1"),
    )
    db.add(gr)
    db.flush()
    for eid in member_ids:
        db.execute(production_group_members.insert().values(group_id=gr.id, employee_id=eid))
    db.commit()
    return RedirectResponse(url="/info/production-groups?added=1", status_code=303)


@app.get("/info/production-groups/edit/{group_id}", response_class=HTMLResponse)
async def info_production_groups_edit_page(
    group_id: int,
    request: Request,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    """Guruhni tahrirlash sahifasi."""
    gr = (
        db.query(ProductionGroup)
        .options(joinedload(ProductionGroup.operator), joinedload(ProductionGroup.piecework_task), joinedload(ProductionGroup.members))
        .filter(ProductionGroup.id == group_id)
        .first()
    )
    if not gr:
        raise HTTPException(status_code=404, detail="Guruh topilmadi")
    employees = db.query(Employee).filter(Employee.is_active == True).order_by(Employee.full_name).all()
    piecework_tasks = db.query(PieceworkTask).filter(PieceworkTask.is_active == True).order_by(PieceworkTask.name).all()
    return templates.TemplateResponse("info/production_group_edit.html", {
        "request": request,
        "group": gr,
        "employees": employees,
        "piecework_tasks": piecework_tasks,
        "current_user": current_user,
        "page_title": f"Guruhni tahrirlash: {gr.name}",
    })


@app.post("/info/production-groups/edit/{group_id}", response_class=RedirectResponse)
async def info_production_groups_edit(
    group_id: int,
    request: Request,
    name: str = Form(...),
    operator_id: int = Form(...),
    piecework_task_id: Optional[int] = Form(None),
    include_qiyom: str = Form("1"),
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    """Guruhni saqlash."""
    from sqlalchemy import delete
    form = await request.form()
    member_ids_raw = form.getlist("member_ids") if hasattr(form, "getlist") else []
    member_ids = list(dict.fromkeys([int(x) for x in member_ids_raw if str(x).strip().isdigit()]))
    gr = db.query(ProductionGroup).filter(ProductionGroup.id == group_id).first()
    if not gr:
        raise HTTPException(status_code=404, detail="Guruh topilmadi")
    gr.name = name.strip()
    gr.operator_id = operator_id
    gr.piecework_task_id = int(piecework_task_id) if piecework_task_id else None
    gr.include_qiyom = (include_qiyom == "1")
    db.execute(delete(production_group_members).where(production_group_members.c.group_id == group_id))
    for eid in member_ids:
        db.execute(production_group_members.insert().values(group_id=gr.id, employee_id=eid))
    db.commit()
    return RedirectResponse(url="/info/production-groups?updated=1", status_code=303)


@app.post("/info/production-groups/delete/{group_id}", response_class=RedirectResponse)
async def info_production_groups_delete(
    group_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    """Guruhni o'chirish (is_active=False)."""
    gr = db.query(ProductionGroup).filter(ProductionGroup.id == group_id).first()
    if not gr:
        raise HTTPException(status_code=404, detail="Guruh topilmadi")
    gr.is_active = False
    db.commit()
    return RedirectResponse(url="/info/production-groups?deleted=1", status_code=303)


# --- MAHSULOT DETAIL VA BARCODE ---



@app.get("/products/barcode/{product_id}")
async def product_barcode(product_id: int, download: int = 0, db: Session = Depends(get_db)):
    product = db.query(Product).filter(Product.id == product_id).first()
    if not product or not product.barcode:
        return HTMLResponse("<h3>Shtixkod topilmadi</h3>", status_code=404)
    # Barcode rasm faylini yaratish
    barcode_path = f"app/static/images/products/barcode_{product.id}.png"
    if not os.path.exists(barcode_path):
        code128 = barcode.get('code128', product.barcode, writer=ImageWriter())
        code128.save(barcode_path[:-4])
    if download:
        return FileResponse(barcode_path, media_type="image/png", filename=f"barcode_{product.code}.png")
    return FileResponse(barcode_path, media_type="image/png")




# ==========================================
# TOVARLAR
# ==========================================
from fastapi import UploadFile, File
from fastapi.responses import StreamingResponse
import io
import openpyxl

# --- EXPORT PRODUCTS TO EXCEL ---
@app.get("/products/export")
async def export_products(db: Session = Depends(get_db), current_user: User = Depends(require_auth)):
    products = db.query(Product).all()
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Products"
    ws.append(["ID", "Kod", "Nomi", "Turi", "O'lchov", "Sotish narxi", "Olish narxi"])
    for p in products:
        ws.append([
            p.id, p.code, p.name, p.type,
            p.unit.name if p.unit else "",
            p.sale_price, p.purchase_price
        ])
    stream = io.BytesIO()
    wb.save(stream)
    stream.seek(0)
    return StreamingResponse(stream, media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", headers={"Content-Disposition": "attachment; filename=products.xlsx"})

# --- DOWNLOAD IMPORT TEMPLATE ---
@app.get("/products/template")
async def product_import_template(current_user: User = Depends(require_auth)):
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Import Template"
    
    # Headers (kategoriya yo'q)
    headers = ["ID", "Kod", "Nomi", "Turi", "O'lchov", "Sotish narxi", "Olish narxi"]
    ws.append(headers)
    # Turi ustunida quyidagilardan biri bo'lishi kerak: tayyor, yarim_tayyor, hom_ashyo
    ws.append(["", "P001", "Tayyor mahsulot", "tayyor", "dona", 15000, 10000])
    ws.append(["", "P002", "Yarim tayyor mahsulot", "yarim_tayyor", "kg", 8000, 5000])
    ws.append(["", "P003", "Xom ashyo", "hom_ashyo", "kg", 2000, 1500])
    
    # Column width adjustment
    for col in range(1, 8):
        ws.column_dimensions[openpyxl.utils.get_column_letter(col)].width = 15
    ws.column_dimensions['C'].width = 30  # Nomi ustuni
    
    stream = io.BytesIO()
    wb.save(stream)
    stream.seek(0)
    return StreamingResponse(stream, media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", headers={"Content-Disposition": "attachment; filename=tovar_andoza.xlsx"})


def _products_check_duplicates_data(db: Session):
    """Tovarlar tekshiruvi ma'lumotlari (dublikatlar)."""
    by_name: dict = {}
    for p in db.query(Product).filter(Product.is_active == True, Product.name.isnot(None)).all():
        key = (p.name or "").strip().lower()
        if key:
            by_name.setdefault(key, []).append(p)
    return [prods for prods in by_name.values() if len(prods) > 1]


@app.get("/products/check", response_class=HTMLResponse)
@app.get("/product-check", response_class=HTMLResponse)  # konfliktsiz alternativ (422 oldini olish)
async def products_check_duplicates(
    request: Request,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    """Tovarlar tekshiruvi: bir xil nomdagi tovarlar (dublikatlar) ro'yxati."""
    duplicate_groups = _products_check_duplicates_data(db)
    return templates.TemplateResponse("products/check.html", {
        "request": request,
        "current_user": current_user,
        "page_title": "Tovarlar tekshiruvi",
        "duplicate_groups": duplicate_groups,
    })


@app.get("/products/bulk-update")
async def product_bulk_update_get_redirect():
    """Brauzerda to'g'ridan-to'g'ri ochilsa tovarlar ro'yxatiga yo'naltirish (gruppavoy o'zgartirish faqat forma orqali POST)."""
    return RedirectResponse(url="/products", status_code=303)


@app.get("/products/{product_id}", response_class=HTMLResponse)
async def product_detail(request: Request, product_id: int, db: Session = Depends(get_db), current_user: User = Depends(require_auth)):
    from urllib.parse import unquote
    product = db.query(Product).filter(Product.id == product_id).first()
    if not product:
        return HTMLResponse("<h3>Mahsulot topilmadi</h3>", status_code=404)
    error_param = request.query_params.get("error")
    detail_param = unquote(request.query_params.get("detail", "") or "")
    return templates.TemplateResponse("products/detail.html", {
        "request": request, "product": product, "current_user": current_user, "page_title": product.name or "Tovar",
        "error_param": error_param, "detail_param": detail_param,
    })

# --- IMPORT PRODUCTS FROM EXCEL ---
@app.get("/products/import")
async def products_import_get(current_user: User = Depends(require_auth)):
    """Import sahifasi faqat form orqali; to'g'ridan-to'g'ri ochilsa tovarlar ro'yxatiga yo'naltirish."""
    return RedirectResponse(url="/products", status_code=303)


@app.post("/products/import")
async def import_products(
    request: Request,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth)
):
    """Excel dan tovarlarni import. Andoza: ID, Kod, Nomi, Turi, O'lchov, Sotish narxi, Olish narxi (kategoriya yo'q)."""
    from urllib.parse import quote
    from zipfile import BadZipFile
    form = await request.form()
    file = form.get("file") or form.get("excel_file")
    if not file or not getattr(file, "filename", None):
        return RedirectResponse(url="/products?error=import&detail=" + quote("Excel fayl tanlang"), status_code=303)
    try:
        contents = await file.read()
        if not contents:
            return RedirectResponse(url="/products?error=import&detail=" + quote("Fayl bo'sh"), status_code=303)
        # .xlsx fayllar ZIP formatida; boshqa format yoki buzilgan fayl xato beradi
        if contents[:2] != b"PK":
            return RedirectResponse(
                url="/products?error=import&detail=" + quote("Fayl .xlsx formati bo'lishi kerak (Excel 2007+). Eski .xls yoki boshqa format qabul qilinmaydi."),
                status_code=303,
            )
        wb = openpyxl.load_workbook(io.BytesIO(contents), read_only=False, data_only=True)
        ws = wb.active
        if ws.max_row < 2:
            return RedirectResponse(
                url="/products?import_ok=0&detail=" + quote("Excelda ma'lumot qatorlari yo'q. 1-qator sarlavha, 2-qatordan Nomi/Kod to'ldiring."),
                status_code=303,
            )
        # 2-qatordan boshlab har bir qatorni A,B,C,D,E,F,G ustunlari orqali aniq o'qish
        added = 0
        updated = 0
        skipped = 0
        for row_num in range(2, ws.max_row + 1):
            def cell(col):
                v = ws.cell(row=row_num, column=col).value
                return "" if v is None else str(v).strip()
            code = (cell(2) or cell(1)).strip()
            name = (cell(3) or cell(2)).strip()
            if not code and not name:
                skipped += 1
                continue
            if code.lower() in ("id", "kod", "nomi", "turi", "o'lchov") and (not name or name.lower() in ("id", "kod", "nomi", "turi")):
                skipped += 1
                continue
            if not code:
                code = f"P{row_num}"
            if not name:
                name = code
            raw = (cell(4) or "tayyor").replace("\xa0", " ").strip().lower()
            if raw in ("yarim tayyor", "yarim_tayyor", "yarimtayyor"):
                type_ = "yarim_tayyor"
            elif raw in ("xom ashyo", "hom_ashyo", "xom_ashyo", "xomashyo"):
                type_ = "hom_ashyo"
            else:
                type_ = "tayyor"
            unit_name = (cell(5) or "").strip() or None
            try:
                sale_price = float((cell(6) or "0").replace(" ", "").replace(",", "."))
            except (ValueError, TypeError):
                sale_price = 0
            try:
                purchase_price = float((cell(7) or "0").replace(" ", "").replace(",", "."))
            except (ValueError, TypeError):
                purchase_price = 0
            try:
                # O'lchov birligi: nomi bo'yicha registrsiz qidirish (Dona/dona bir xil)
                unit = None
                if unit_name:
                    unit = db.query(Unit).filter(func.lower(func.trim(Unit.name)) == unit_name.lower()).first()
                if not unit and unit_name:
                    base_code = (unit_name.lower().replace(" ", "_")[:10] or "u").strip("_")
                    try:
                        unit = Unit(name=unit_name, code=base_code)
                        db.add(unit)
                        db.commit()
                        db.refresh(unit)
                    except IntegrityError:
                        db.rollback()
                        unit = db.query(Unit).filter(func.lower(Unit.name) == unit_name.lower()).first()
                # Mahsulot: avval kod bo'yicha, topilmasa nom bo'yicha (bir xil nom 2 marta bo'lmasin)
                product = db.query(Product).filter(
                    Product.code.isnot(None),
                    Product.code != "",
                    func.lower(func.trim(Product.code)) == code.strip().lower()
                ).first()
                if not product:
                    existing_by_name = db.query(Product).filter(
                        Product.is_active == True,
                        Product.name.isnot(None),
                        func.lower(func.trim(Product.name)) == (name or "").strip().lower(),
                    ).first()
                    if existing_by_name:
                        product = existing_by_name
                        updated += 1
                    else:
                        product = Product(code=code.strip(), is_active=True)
                        db.add(product)
                        added += 1
                else:
                    updated += 1
                product.name = name
                product.type = type_
                product.is_active = True
                product.category_id = None
                product.unit_id = unit.id if unit else None
                product.sale_price = sale_price
                product.purchase_price = purchase_price
                db.commit()
            except Exception:
                db.rollback()
                skipped += 1
                continue
        if added == 0 and updated == 0 and skipped == 0:
            detail = "Hech qanday qator import qilinmadi. Excelda 1-qator sarlavha, 2-qatordan: A yoki B=Kod, B yoki C=Nomi to'ldiring. Andoza tugmasidan fayl yuklab tekshiring."
            return RedirectResponse(
                url="/products?import_ok=0&detail=" + quote(detail),
                status_code=303,
            )
        url = "/products?import_ok=1&added=" + str(added) + "&updated=" + str(updated)
        if skipped:
            url += "&skipped=" + str(skipped)
        return RedirectResponse(url=url, status_code=303)
    except BadZipFile:
        return RedirectResponse(
            url="/products?error=import&detail=" + quote("Fayl .xlsx formati bo'lishi kerak. Boshqa format yoki buzilgan fayl yuborilgan."),
            status_code=303,
        )
    except Exception as e:
        err_msg = str(e)[:200]
        if "zip" in err_msg.lower() or "not a zip" in err_msg.lower():
            err_msg = "Fayl .xlsx formati bo'lishi kerak. Boshqa format yoki buzilgan fayl."
        return RedirectResponse(
            url="/products?error=import&detail=" + quote(err_msg),
            status_code=303,
        )

@app.get("/products", response_class=HTMLResponse)
async def products_list(
    request: Request,
    type: str = "all",
    search: Optional[str] = None,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    """Tovarlar ro'yxati — turi va qidiruv bo'yicha filtr (tozalashguncha saqlanadi)."""
    from urllib.parse import unquote
    query = db.query(Product).filter(Product.is_active == True)
    if type == "tayyor":
        query = query.filter(Product.type == "tayyor")
    elif type == "yarim_tayyor":
        query = query.filter(Product.type == "yarim_tayyor")
    elif type == "hom_ashyo":
        query = query.filter(Product.type == "hom_ashyo")
    search_q = (search or "").strip()
    if search_q:
        like = f"%{search_q}%"
        query = query.filter(
            or_(
                Product.name.ilike(like),
                func.coalesce(Product.code, "").ilike(like),
                func.coalesce(Product.barcode, "").ilike(like),
            )
        )
    products = query.all()
    categories = db.query(Category).all()
    units = db.query(Unit).all()
    import_ok = request.query_params.get("import_ok")
    added = request.query_params.get("added")
    updated = request.query_params.get("updated")
    import_skipped = request.query_params.get("skipped")
    import_error = request.query_params.get("error") == "import"
    import_detail = unquote(request.query_params.get("detail", "") or "")
    from urllib.parse import quote as url_quote
    search_encoded = url_quote(search_q) if search_q else ""
    return templates.TemplateResponse("products/list.html", {
        "request": request,
        "products": products,
        "categories": categories,
        "units": units,
        "current_type": type,
        "search_q": search_q,
        "search_encoded": search_encoded,
        "current_user": current_user,
        "page_title": "Tovarlar",
        "import_ok": import_ok,
        "import_added": added,
        "import_updated": updated,
        "import_skipped": import_skipped,
        "import_error": import_error,
        "import_detail": import_detail,
    })



def _product_name_exists(db: Session, name: str, exclude_product_id: Optional[int] = None) -> bool:
    """Boshqa tovar shu nomda (trim, katta-kichik farqsiz) mavjudmi tekshiradi."""
    if not (name or "").strip():
        return False
    q = db.query(Product.id).filter(
        Product.is_active == True,
        func.lower(func.trim(Product.name)) == (name or "").strip().lower(),
    )
    if exclude_product_id is not None:
        q = q.filter(Product.id != exclude_product_id)
    return q.first() is not None


@app.post("/products/add")
async def product_add(
    request: Request,
    name: str = Form(...),
    type: str = Form(...),
    category_id: int = Form(None),
    unit_id: int = Form(None),
    barcode: str = Form(None),
    sale_price: float = Form(0),
    purchase_price: float = Form(0),
    image: Optional[UploadFile] = File(None),
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    """Tovar qo'shish (rasm ixtiyoriy). Nom dublikati ruxsat etilmaydi."""
    from urllib.parse import quote
    if _product_name_exists(db, name):
        return RedirectResponse(
            url="/products?error=duplicate&detail=" + quote(f"«{name.strip()}» nomli tovar allaqachon mavjud. Boshqa nom yoki kod ishlating."),
            status_code=303,
        )
    product = Product(
        name=name,
        code=None,
        type=type,
        category_id=category_id if category_id and category_id > 0 else None,
        unit_id=unit_id if unit_id and unit_id > 0 else None,
        barcode=barcode,
        sale_price=sale_price,
        purchase_price=purchase_price,
        image=None
    )
    db.add(product)
    db.commit()
    product.code = f"P{product.id:05d}"
    db.commit()

    if image and (image.filename or "").strip():
        import shutil
        ext = (image.filename or "").split(".")[-1] if "." in (image.filename or "") else "jpg"
        image_filename = f"{product.code}.{ext}"
        image_path = os.path.join("app", "static", "images", "products", image_filename)
        os.makedirs(os.path.dirname(image_path), exist_ok=True)
        with open(image_path, "wb") as buffer:
            shutil.copyfileobj(image.file, buffer)
        product.image = image_filename
        db.commit()

    return RedirectResponse(url="/products", status_code=303)


@app.post("/products/edit/{product_id}")
async def product_edit(
    product_id: int,
    name: str = Form(...),
    type: str = Form(...),
    category_id: int = Form(None),
    unit_id: int = Form(None),
    barcode: str = Form(None),
    sale_price: float = Form(0),
    purchase_price: float = Form(0),
    image: Optional[UploadFile] = File(None),
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    """Tovar tahrirlash (rasm ixtiyoriy). Nom dublikati ruxsat etilmaydi."""
    from urllib.parse import quote
    product = db.query(Product).filter(Product.id == product_id).first()
    if not product:
        raise HTTPException(status_code=404, detail="Mahsulot topilmadi")
    if _product_name_exists(db, name, exclude_product_id=product_id):
        return RedirectResponse(
            url=f"/products/{product_id}?error=duplicate&detail=" + quote(f"«{name.strip()}» nomli boshqa tovar mavjud. Nom bir xil bo'lmasin."),
            status_code=303,
        )
    product.name = name
    product.type = type
    product.category_id = category_id if category_id and category_id > 0 else None
    product.unit_id = unit_id if unit_id and unit_id > 0 else None
    product.barcode = barcode or None
    product.sale_price = sale_price
    product.purchase_price = purchase_price

    if image and (image.filename or "").strip():
        import shutil
        ext = (image.filename or "").split(".")[-1] if "." in (image.filename or "") else "jpg"
        image_filename = f"{product.code}.{ext}"
        image_path = os.path.join("app", "static", "images", "products", image_filename)
        os.makedirs(os.path.dirname(image_path), exist_ok=True)
        with open(image_path, "wb") as buffer:
            shutil.copyfileobj(image.file, buffer)
        product.image = image_filename

    db.commit()
    return RedirectResponse(url="/products", status_code=303)


@app.post("/products/delete/{product_id}")
async def product_delete(
    product_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_admin),
):
    """Faqat admin: tovarni o'chirish (soft delete: is_active=False)"""
    product = db.query(Product).filter(Product.id == product_id).first()
    if not product:
        raise HTTPException(status_code=404, detail="Mahsulot topilmadi")
    product.is_active = False
    db.commit()
    return RedirectResponse(url="/products", status_code=303)


@app.post("/products/bulk-update")
async def product_bulk_update(
    request: Request,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_admin),
):
    """Faqat admin: tanlangan tovarlarga gruppaviy o'zgartirish (turi, o'lchov birligi)."""
    from urllib.parse import quote
    form = await request.form()
    ids = form.getlist("product_ids")
    new_type = (form.get("type") or "").strip()
    unit_id_raw = form.get("unit_id")
    new_unit_id = None
    if unit_id_raw is not None and str(unit_id_raw).strip() and str(unit_id_raw).strip() != "0":
        try:
            new_unit_id = int(unit_id_raw)
        except (ValueError, TypeError):
            pass
    updated = 0
    for sid in ids:
        try:
            pid = int(sid)
            product = db.query(Product).filter(Product.id == pid, Product.is_active == True).first()
            if not product:
                continue
            changed = False
            if new_type in ("tayyor", "yarim_tayyor", "hom_ashyo"):
                product.type = new_type
                changed = True
            if new_unit_id is not None:
                product.unit_id = new_unit_id
                changed = True
            if changed:
                updated += 1
        except (ValueError, TypeError):
            pass
    if updated > 0:
        db.commit()
    msg = quote(f"Tanlangan tovarlar yangilandi ({updated} ta).") if updated else quote("O'zgartirish kiritilmadi yoki tovarlar topilmadi.")
    return RedirectResponse(url="/products?updated=" + msg, status_code=303)


@app.post("/products/delete-bulk")
async def product_delete_bulk(
    request: Request,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_admin),
):
    """Faqat admin: tanlangan tovarlarni o'chirish (is_active=False)."""
    from urllib.parse import quote
    form = await request.form()
    ids = form.getlist("product_ids")
    deleted = 0
    for sid in ids:
        try:
            pid = int(sid)
            product = db.query(Product).filter(Product.id == pid).first()
            if product:
                product.is_active = False
                deleted += 1
        except (ValueError, TypeError):
            pass
    db.commit()
    msg = quote(f"Tanlangan {deleted} ta tovar o'chirildi.") if deleted else quote("Hech narsa tanlanmadi.")
    return RedirectResponse(url="/products?deleted=" + msg, status_code=303)


@app.post("/products/{product_id}/upload-image")
async def product_upload_image(
    product_id: int,
    image: UploadFile = File(...),
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    """Mahsulot rasmini yuklash"""
    product = db.query(Product).filter(Product.id == product_id).first()
    if not product:
        raise HTTPException(status_code=404, detail="Mahsulot topilmadi")
    
    if image and image.filename:
        import shutil, os
        ext = image.filename.split('.')[-1]
        image_filename = f"{product.code}.{ext}"
        image_path = os.path.join("app/static/images/products", image_filename)
        
        # Papkani yaratish
        os.makedirs(os.path.dirname(image_path), exist_ok=True)
        
        with open(image_path, "wb") as buffer:
            shutil.copyfileobj(image.file, buffer)
        
        product.image = image_filename
        db.commit()
    
    return RedirectResponse(url="/products", status_code=303)


# ==========================================
# OMBOR
# ==========================================

def create_stock_movement(
    db: Session,
    warehouse_id: int,
    product_id: int,
    quantity_change: float,
    operation_type: str,
    document_type: str,
    document_id: int,
    document_number: str = None,
    user_id: int = None,
    note: str = None
):
    """Har bir operatsiya uchun StockMovement yozuvini yaratish.
    Chiqim (quantity_change < 0) bo'lganda avval bitta ombor+mahsulot uchun barcha Stock qatorlarini birlashtiradi."""
    rows = db.query(Stock).filter(
        Stock.warehouse_id == warehouse_id,
        Stock.product_id == product_id
    ).all()
    # Bir nechta qator bo'lsa (dublikat) — bitta qatorga yig'ib, qolganini o'chirish (ayniqsa chiqimda muhim)
    if len(rows) > 1:
        total = sum(float(r.quantity or 0) for r in rows)
        keep = rows[0]
        keep.quantity = total
        keep.updated_at = datetime.now()
        for r in rows[1:]:
            db.delete(r)
        db.flush()
        stock = keep
    elif len(rows) == 1:
        stock = rows[0]
    else:
        stock = None

    # Qoldiqni yangilash
    if stock:
        stock.quantity = (stock.quantity or 0) + quantity_change
        if stock.quantity < 0:
            stock.quantity = 0
        stock.updated_at = datetime.now()
        stock_id = stock.id
        quantity_after = stock.quantity
    else:
        # Yangi stock yaratish
        quantity_after = quantity_change if quantity_change > 0 else 0
        stock = Stock(
            warehouse_id=warehouse_id,
            product_id=product_id,
            quantity=quantity_after
        )
        db.add(stock)
        db.flush()  # ID olish uchun
        stock_id = stock.id
    
    # StockMovement yozuvini yaratish
    movement = StockMovement(
        stock_id=stock_id,
        warehouse_id=warehouse_id,
        product_id=product_id,
        operation_type=operation_type,
        document_type=document_type,
        document_id=document_id,
        document_number=document_number,
        quantity_change=quantity_change,
        quantity_after=quantity_after,
        user_id=user_id,
        note=note
    )
    db.add(movement)
    return movement


def delete_stock_movements_for_document(db: Session, document_type: str, document_id: int) -> int:
    """Hujjat tasdiqi bekor qilinganda shu hujjatga tegishli StockMovement yozuvlarini o'chiradi.
    Natijada «Mahsulot harakati tarixi» da bekor qilingan harakatlar ko'rinmaydi."""
    deleted = db.query(StockMovement).filter(
        StockMovement.document_type == document_type,
        StockMovement.document_id == document_id,
    ).delete(synchronize_session=False)
    return deleted


# ==========================================
# INVENTARIZATSIYA
# ==========================================

def _ensure_inventory_columns(db: Session) -> None:
    """SQLite: stock_adjustment_docs.warehouse_id va stock_adjustment_doc_items.previous_quantity ustunlari yo'q bo'lsa qo'shadi."""
    try:
        r = db.execute(text("PRAGMA table_info(stock_adjustment_docs)"))
        cols_doc = [row[1] for row in r.fetchall()]
        if "warehouse_id" not in cols_doc:
            db.execute(text("ALTER TABLE stock_adjustment_docs ADD COLUMN warehouse_id INTEGER REFERENCES warehouses(id)"))
            db.commit()
        r = db.execute(text("PRAGMA table_info(stock_adjustment_doc_items)"))
        cols_item = [row[1] for row in r.fetchall()]
        if "previous_quantity" not in cols_item:
            db.execute(text("ALTER TABLE stock_adjustment_doc_items ADD COLUMN previous_quantity REAL"))
            db.commit()
    except Exception:
        db.rollback()


@app.get("/inventory", response_class=HTMLResponse)
async def inventory_list_page(
    request: Request,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    """Inventarizatsiya hujjatlari ro'yxati (warehouse_id to'ldirilgan hujjatlar)."""
    if not current_user:
        return RedirectResponse(url="/login", status_code=303)
    migration_warning = None
    try:
        docs = (
            db.query(StockAdjustmentDoc)
            .filter(StockAdjustmentDoc.warehouse_id.isnot(None))
            .order_by(StockAdjustmentDoc.created_at.desc())
            .limit(200)
            .all()
        )
    except Exception:
        try:
            _ensure_inventory_columns(db)
            docs = (
                db.query(StockAdjustmentDoc)
                .filter(StockAdjustmentDoc.warehouse_id.isnot(None))
                .order_by(StockAdjustmentDoc.created_at.desc())
                .limit(200)
                .all()
            )
        except Exception:
            docs = []
            migration_warning = "Inventarizatsiya uchun bazada warehouse_id ustuni kerak. Loyiha ildizida: alembic upgrade head"
    message = request.query_params.get("message", "").strip()
    return templates.TemplateResponse("inventory/list.html", {
        "request": request,
        "docs": docs,
        "current_user": current_user,
        "page_title": "Inventarizatsiya",
        "migration_warning": migration_warning,
        "message": message,
    })


@app.get("/inventory/new", response_class=HTMLResponse)
async def inventory_new_page(
    request: Request,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    """Yangi inventarizatsiya hujjati — avval ombor tanlash, keyin bo'sh hujjat yaratib tahrirlashga yo'naltirish."""
    if not current_user:
        return RedirectResponse(url="/login", status_code=303)
    warehouses = db.query(Warehouse).filter(Warehouse.is_active == True).order_by(Warehouse.name).all()
    return templates.TemplateResponse("inventory/new.html", {
        "request": request,
        "warehouses": warehouses,
        "current_user": current_user,
        "page_title": "Inventarizatsiya — yangi hujjat",
    })


@app.post("/inventory/create-draft")
async def inventory_create_draft(
    warehouse_id: int = Form(...),
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    """Bo'sh qoralama hujjat yaratish va tahrirlash sahifasiga yo'naltirish."""
    if not current_user:
        return RedirectResponse(url="/login", status_code=303)
    wh = db.query(Warehouse).filter(Warehouse.id == warehouse_id).first()
    if not wh:
        return RedirectResponse(url="/inventory/new?message=Ombor topilmadi.", status_code=303)
    today = datetime.now()
    doc = StockAdjustmentDoc(
        number="INV-PENDING",  # Saqlanganda INV-YYYYMMDD-NNNN ga almashtiriladi
        date=today,
        warehouse_id=warehouse_id,
        user_id=current_user.id,
        status="draft",
        total_tannarx=0,
        total_sotuv=0,
    )
    db.add(doc)
    db.flush()
    doc.number = f"INV-PENDING-{doc.id}"  # Unique bo'lishi uchun
    db.commit()
    db.refresh(doc)
    return RedirectResponse(url=f"/inventory/{doc.id}/edit", status_code=303)


@app.get("/inventory/{doc_id}/edit", response_class=HTMLResponse)
async def inventory_edit_page(
    doc_id: int,
    request: Request,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    """Inventarizatsiya hujjatini tahrirlash (faqat qoralama)."""
    if not current_user:
        return RedirectResponse(url="/login", status_code=303)
    doc = db.query(StockAdjustmentDoc).filter(StockAdjustmentDoc.id == doc_id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Hujjat topilmadi")
    if doc.status != "draft":
        return RedirectResponse(url=f"/inventory/{doc_id}", status_code=303)
    if not doc.warehouse_id:
        return RedirectResponse(url="/inventory", status_code=303)
    warehouse = doc.warehouse
    wh_id = doc.warehouse_id
    # Hisobiy qoldiq = hozirgi Stock yig'indisi (har bir mahsulot uchun)
    stocks_by_product = {}
    for s in db.query(Stock).filter(Stock.warehouse_id == wh_id).all():
        pid = s.product_id
        stocks_by_product[pid] = stocks_by_product.get(pid, 0) + float(s.quantity or 0)
    by_product = {}
    for item in doc.items:
        pid = item.product_id
        if pid not in by_product:
            prod = item.product
            by_product[pid] = {
                "item_id": item.id,
                "product_id": pid,
                "product_name": (prod.name or "") if prod else "",
                "product_code": (prod.code or "") if prod else "",
                "current_quantity": stocks_by_product.get(pid, 0),
                "actual_quantity": float(item.quantity or 0),
                "cost_price": float(item.cost_price or 0),
                "sale_price": float(item.sale_price or 0),
            }
    product_ids_in_doc = set(by_product.keys())
    if product_ids_in_doc:
        products_to_add = db.query(Product).filter(Product.is_active == True).filter(~Product.id.in_(product_ids_in_doc)).order_by(Product.name).all()
    else:
        products_to_add = db.query(Product).filter(Product.is_active == True).order_by(Product.name).all()
    products_data = sorted(by_product.values(), key=lambda x: (x["product_name"].lower(), x["product_id"]))
    show_tannarx = getattr(current_user, "role", None) == "admin"
    inv_date = doc.date or datetime.now()
    doc_date_value = inv_date.strftime("%Y-%m-%dT%H:%M") if inv_date else ""
    doc_date_display = inv_date.strftime("%d.%m.%Y %H:%M") if inv_date else ""
    return templates.TemplateResponse("inventory/edit.html", {
        "request": request,
        "doc": doc,
        "warehouse": warehouse,
        "products_data": products_data,
        "products_to_add": products_to_add,
        "show_tannarx": show_tannarx,
        "current_user": current_user,
        "doc_date_value": doc_date_value,
        "doc_date_display": doc_date_display,
        "page_title": "Inventarizatsiya — tahrirlash",
    })


@app.post("/inventory/{doc_id}/load")
async def inventory_load_warehouse(
    doc_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    """Shu ombor qoldiq tovarlarini hujjatga yuklash (dublikatlarni birlashtirib)."""
    if not current_user:
        return RedirectResponse(url="/login", status_code=303)
    doc = db.query(StockAdjustmentDoc).filter(StockAdjustmentDoc.id == doc_id).first()
    if not doc or doc.status != "draft" or not doc.warehouse_id:
        return RedirectResponse(url="/inventory", status_code=303)
    wh_id = doc.warehouse_id
    stocks = (
        db.query(Stock, Product, Unit)
        .join(Product, Stock.product_id == Product.id)
        .outerjoin(Unit, Product.unit_id == Unit.id)
        .filter(Stock.warehouse_id == wh_id)
        .filter(Product.is_active == True)
        .all()
    )
    by_product = {}
    for stock, product, unit in stocks:
        pid = product.id
        qty = float(stock.quantity or 0)
        if pid not in by_product:
            by_product[pid] = qty
        else:
            by_product[pid] += qty
    existing_ids = {item.product_id for item in doc.items}
    for pid, qty in by_product.items():
        if qty is None or float(qty or 0) <= 0:
            continue
        if pid in existing_ids:
            continue
        product = db.query(Product).filter(Product.id == pid).first()
        if not product:
            continue
        cost = float(product.purchase_price or 0)
        sale = float(product.sale_price or 0)
        if (product.sale_price or 0) <= 0:
            pp = db.query(ProductPrice).filter(ProductPrice.product_id == pid).first()
            if pp:
                sale = float(pp.sale_price or 0)
        db.add(StockAdjustmentDocItem(
            doc_id=doc_id,
            product_id=pid,
            warehouse_id=wh_id,
            quantity=qty,
            cost_price=cost,
            sale_price=sale,
        ))
        existing_ids.add(pid)
    db.commit()
    return RedirectResponse(url=f"/inventory/{doc_id}/edit?message=Qoldiq tovarlar yuklandi.", status_code=303)


@app.post("/inventory/{doc_id}/add-product")
async def inventory_add_product(
    doc_id: int,
    product_id: int = Form(...),
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    """Hujjatga bitta tovar qo'shish (qoldiq 0)."""
    if not current_user:
        return RedirectResponse(url="/login", status_code=303)
    doc = db.query(StockAdjustmentDoc).filter(StockAdjustmentDoc.id == doc_id).first()
    if not doc or doc.status != "draft" or not doc.warehouse_id:
        return RedirectResponse(url="/inventory", status_code=303)
    if any(item.product_id == product_id for item in doc.items):
        return RedirectResponse(url=f"/inventory/{doc_id}/edit?message=Ushbu tovar allaqachon jadvalda.", status_code=303)
    product = db.query(Product).filter(Product.id == product_id, Product.is_active == True).first()
    if not product:
        return RedirectResponse(url=f"/inventory/{doc_id}/edit?message=Tovar topilmadi.", status_code=303)
    cost = float(product.purchase_price or 0)
    sale = float(product.sale_price or 0)
    if (product.sale_price or 0) <= 0:
        pp = db.query(ProductPrice).filter(ProductPrice.product_id == product_id).first()
        if pp:
            sale = float(pp.sale_price or 0)
    db.add(StockAdjustmentDocItem(
        doc_id=doc_id,
        product_id=product_id,
        warehouse_id=doc.warehouse_id,
        quantity=0,
        cost_price=cost,
        sale_price=sale,
    ))
    db.commit()
    return RedirectResponse(url=f"/inventory/{doc_id}/edit?message=Tovar qo'shildi.", status_code=303)


@app.post("/inventory/{doc_id}/remove-zero-balance")
async def inventory_remove_zero_balance(
    doc_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    """Hujjatdan shu omborda hozirgi qoldig'i 0 bo'lgan qatorlarni olib tashlash (o'tkazilgan tovarlar ro'yxatda qolmasin)."""
    from urllib.parse import quote
    if not current_user:
        return RedirectResponse(url="/login", status_code=303)
    doc = db.query(StockAdjustmentDoc).filter(StockAdjustmentDoc.id == doc_id).first()
    if not doc or doc.status != "draft" or not doc.warehouse_id:
        return RedirectResponse(url="/inventory", status_code=303)
    wh_id = doc.warehouse_id
    removed = 0
    for item in list(doc.items):
        total = sum(
            float(s.quantity or 0)
            for s in db.query(Stock).filter(
                Stock.warehouse_id == wh_id,
                Stock.product_id == item.product_id,
            ).all()
        )
        if total <= 0:
            db.delete(item)
            removed += 1
    db.commit()
    msg = f"0 qoldiqli {removed} ta qator olib tashlandi." if removed else "0 qoldiqli qatorlar yo'q."
    return RedirectResponse(url=f"/inventory/{doc_id}/edit?message=" + quote(msg), status_code=303)


def _parse_quantity(value) -> float:
    """Formadan miqdorni o'qiydi; vergulni nuqtaga almashtiradi."""
    if value is None or str(value).strip() == "":
        return 0.0
    try:
        return float(str(value).strip().replace(",", "."))
    except (TypeError, ValueError):
        return 0.0


def _next_inventory_number(db: Session, date_str: str) -> str:
    """Berilgan sana (YYYYMMDD) uchun keyingi bo'sh INV-YYYYMMDD-NNNN raqamini qaytaradi (max suffix + 1)."""
    prefix = f"INV-{date_str}-"
    rows = db.query(StockAdjustmentDoc.number).filter(
        StockAdjustmentDoc.number.like(f"{prefix}%")
    ).all()
    max_suffix = 0
    for (num,) in rows:
        if num and num.startswith(prefix):
            try:
                suf = int(num[len(prefix):].strip())
                if suf > max_suffix:
                    max_suffix = suf
            except (ValueError, TypeError):
                pass
    return f"{prefix}{str(max_suffix + 1).zfill(4)}"


@app.post("/inventory/{doc_id}/save")
async def inventory_save_draft(
    doc_id: int,
    request: Request,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    """Qoralama hujjatni saqlash (haqiqiy qoldiqlarni yangilash)."""
    from urllib.parse import quote
    if not current_user:
        return RedirectResponse(url="/login", status_code=303)
    doc = db.query(StockAdjustmentDoc).filter(StockAdjustmentDoc.id == doc_id).first()
    if not doc or doc.status != "draft":
        return RedirectResponse(url="/inventory", status_code=303)
    try:
        form = await request.form()
        doc_date_str = form.get("doc_date")
        if doc_date_str:
            parsed = _parse_doc_date(doc_date_str)
            if parsed:
                doc.date = parsed
        # Sana tanlanib saqlanganda haqiqiy raqam berish (INV-YYYYMMDD-NNNN)
        if doc.number and doc.number.startswith("INV-PENDING") and doc.date:
            date_str = doc.date.strftime("%Y%m%d")
            doc.number = _next_inventory_number(db, date_str)
        item_ids = form.getlist("item_id")
        quantities = form.getlist("actual_quantity")
        total_tannarx = 0.0
        total_sotuv = 0.0
        for i, iid in enumerate(item_ids):
            if not iid:
                continue
            try:
                item_id = int(iid)
            except (TypeError, ValueError):
                continue
            qty = _parse_quantity(quantities[i] if i < len(quantities) else None)
            item = db.query(StockAdjustmentDocItem).filter(
                StockAdjustmentDocItem.id == item_id,
                StockAdjustmentDocItem.doc_id == doc_id,
            ).first()
            if item:
                item.quantity = qty
                total_tannarx += qty * float(item.cost_price or 0)
                total_sotuv += qty * float(item.sale_price or 0)
        doc.total_tannarx = total_tannarx
        doc.total_sotuv = total_sotuv
        try:
            db.commit()
        except Exception as commit_err:
            db.rollback()
            # Raqam boshqa hujjatda bo'lsa — yangi raqam bilan qayta urinish
            if "UNIQUE" in str(commit_err) and "number" in str(commit_err).lower() and doc.date:
                date_str = doc.date.strftime("%Y%m%d")
                doc.number = _next_inventory_number(db, date_str)
                db.commit()
            else:
                raise commit_err
        return RedirectResponse(url=f"/inventory/{doc_id}/edit?message=Saqlandi.", status_code=303)
    except Exception as e:
        db.rollback()
        import logging
        logging.getLogger(__name__).exception("inventory save error doc_id=%s: %s", doc_id, e)
        return RedirectResponse(
            url=f"/inventory/{doc_id}/edit?error=" + quote(f"Saqlashda xatolik: {str(e)}"),
            status_code=303,
        )


@app.get("/inventory/{doc_id}", response_class=HTMLResponse)
async def inventory_view_page(
    doc_id: int,
    request: Request,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    """Inventarizatsiya hujjatini ko'rish (tasdiqlangan: faqat ko'rish + bekor qilish; qoralama: tahrirlashga yo'naltirish)."""
    if not current_user:
        return RedirectResponse(url="/login", status_code=303)
    doc = db.query(StockAdjustmentDoc).filter(StockAdjustmentDoc.id == doc_id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Hujjat topilmadi")
    if doc.status == "draft":
        return RedirectResponse(url=f"/inventory/{doc_id}/edit", status_code=303)
    if not doc.warehouse_id:
        return RedirectResponse(url="/inventory", status_code=303)
    warehouse = doc.warehouse
    rows = []
    for item in doc.items:
        prod = item.product
        rows.append({
            "product_name": (prod.name or "") if prod else "",
            "product_code": (prod.code or "") if prod else "",
            "quantity": float(item.quantity or 0),
            "cost_price": float(item.cost_price or 0),
            "sale_price": float(item.sale_price or 0),
        })
    show_tannarx = getattr(current_user, "role", None) == "admin"
    return templates.TemplateResponse("inventory/view.html", {
        "request": request,
        "doc": doc,
        "warehouse": warehouse,
        "rows": rows,
        "show_tannarx": show_tannarx,
        "current_user": current_user,
        "page_title": "Inventarizatsiya — " + (doc.number or ""),
    })


@app.post("/inventory/{doc_id}/confirm")
async def inventory_confirm(
    doc_id: int,
    request: Request,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    """Hujjatni tasdiqlash — formadagi haqiqiy qoldiqlarni saqlab, ombor qoldiqlarini yangilash."""
    if not current_user:
        return RedirectResponse(url="/login", status_code=303)
    doc = db.query(StockAdjustmentDoc).filter(StockAdjustmentDoc.id == doc_id).first()
    if not doc or doc.status != "draft" or not doc.warehouse_id:
        return RedirectResponse(url="/inventory", status_code=303)
    if not doc.items:
        return RedirectResponse(url=f"/inventory/{doc_id}/edit?message=Jadval bo'sh. Avval qoldiq tovarlarni yuklang yoki tovar qo'shing.", status_code=303)
    form = await request.form()
    doc_date_str = form.get("doc_date")
    if doc_date_str:
        parsed = _parse_doc_date(doc_date_str)
        if parsed:
            doc.date = parsed
    # Sana tanlangan holda tasdiqlashda ham raqam berish (agar hali INV-PENDING bo'lsa)
    if doc.number and doc.number.startswith("INV-PENDING") and doc.date:
        date_str = doc.date.strftime("%Y%m%d")
        doc.number = _next_inventory_number(db, date_str)
    item_ids = form.getlist("item_id")
    quantities = form.getlist("actual_quantity")
    for i, iid in enumerate(item_ids):
        if not iid:
            continue
        try:
            item_id = int(iid)
        except (TypeError, ValueError):
            continue
        qty = _parse_quantity(quantities[i] if i < len(quantities) else None)
        item = db.query(StockAdjustmentDocItem).filter(
            StockAdjustmentDocItem.id == item_id,
            StockAdjustmentDocItem.doc_id == doc_id,
        ).first()
        if item:
            item.quantity = qty
    total_tannarx = sum(float(it.quantity or 0) * float(it.cost_price or 0) for it in doc.items)
    total_sotuv = sum(float(it.quantity or 0) * float(it.sale_price or 0) for it in doc.items)
    doc.total_tannarx = total_tannarx
    doc.total_sotuv = total_sotuv
    db.commit()
    db.refresh(doc)
    pairs = set((item.warehouse_id, item.product_id) for item in doc.items)
    for wh_id, prod_id in pairs:
        rows = db.query(Stock).filter(Stock.warehouse_id == wh_id, Stock.product_id == prod_id).all()
        if len(rows) > 1:
            total = sum(float(r.quantity or 0) for r in rows)
            keep = rows[0]
            keep.quantity = total
            if hasattr(keep, "updated_at"):
                keep.updated_at = datetime.now()
            for r in rows[1:]:
                db.delete(r)
    db.commit()
    for item in doc.items:
        stocks = db.query(Stock).filter(
            Stock.warehouse_id == item.warehouse_id,
            Stock.product_id == item.product_id,
        ).all()
        old_qty = sum(float(s.quantity or 0) for s in stocks)
        new_qty = float(item.quantity or 0)
        if hasattr(item, "previous_quantity"):
            item.previous_quantity = old_qty
        quantity_change = new_qty - old_qty
        if abs(quantity_change) > 1e-9:
            create_stock_movement(
                db=db,
                warehouse_id=item.warehouse_id,
                product_id=item.product_id,
                quantity_change=quantity_change,
                operation_type="adjustment",
                document_type="StockAdjustmentDoc",
                document_id=doc.id,
                document_number=doc.number,
                user_id=current_user.id,
                note=f"Inventarizatsiya: {doc.number}",
            )
    doc.status = "confirmed"
    db.commit()
    return RedirectResponse(url=f"/inventory/{doc_id}?message=Tasdiqlandi.", status_code=303)


@app.post("/inventory/{doc_id}/revoke")
async def inventory_revoke(
    doc_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_admin),
    return_to: Optional[str] = Form(None),
):
    """Tasdiqlashni bekor qilish — qoldiqlarni oldingi holatga qaytarish (faqat admin)."""
    doc = db.query(StockAdjustmentDoc).filter(StockAdjustmentDoc.id == doc_id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Hujjat topilmadi")
    if doc.status != "confirmed":
        if return_to == "list":
            return RedirectResponse(url="/qoldiqlar/tovar/hujjat", status_code=303)
        return RedirectResponse(url=f"/inventory/{doc_id}", status_code=303)
    for item in doc.items:
        stock = db.query(Stock).filter(
            Stock.warehouse_id == item.warehouse_id,
            Stock.product_id == item.product_id,
        ).first()
        prev = getattr(item, "previous_quantity", None)
        if stock is not None and prev is not None:
            stock.quantity = prev
            if hasattr(stock, "updated_at"):
                stock.updated_at = datetime.now()
    delete_stock_movements_for_document(db, "StockAdjustmentDoc", doc_id)
    doc.status = "draft"
    db.commit()
    if return_to == "list":
        return RedirectResponse(url="/qoldiqlar/tovar/hujjat?reverted=1", status_code=303)
    return RedirectResponse(url=f"/inventory/{doc_id}/edit?message=Tasdiqlash bekor qilindi.", status_code=303)


@app.post("/inventory/{doc_id}/delete")
async def inventory_delete(
    doc_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    """Qoralama inventarizatsiya hujjatini o'chirish (faqat draft)."""
    if not current_user:
        return RedirectResponse(url="/login", status_code=303)
    doc = db.query(StockAdjustmentDoc).filter(StockAdjustmentDoc.id == doc_id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Hujjat topilmadi")
    if doc.status != "draft":
        return RedirectResponse(url="/inventory?message=Faqat qoralama hujjat o'chiriladi.", status_code=303)
    for item in list(doc.items):
        db.delete(item)
    db.delete(doc)
    db.commit()
    return RedirectResponse(url="/inventory?message=Hujjat o'chirildi.", status_code=303)


@app.get("/inventory/{doc_id}/print", response_class=HTMLResponse)
async def inventory_print_page(
    doc_id: int,
    request: Request,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    """Inventarizatsiya hujjatini chop etish uchun sahifa."""
    if not current_user:
        return RedirectResponse(url="/login", status_code=303)
    doc = db.query(StockAdjustmentDoc).filter(StockAdjustmentDoc.id == doc_id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Hujjat topilmadi")
    warehouse = doc.warehouse
    rows = []
    for item in doc.items:
        prod = item.product
        rows.append({
            "product_name": (prod.name or "") if prod else "",
            "product_code": (prod.code or "") if prod else "",
            "quantity": float(item.quantity or 0),
            "cost_price": float(item.cost_price or 0),
            "sale_price": float(item.sale_price or 0),
        })
    show_tannarx = getattr(current_user, "role", None) == "admin"
    return templates.TemplateResponse("inventory/print.html", {
        "request": request,
        "doc": doc,
        "warehouse": warehouse,
        "rows": rows,
        "show_tannarx": show_tannarx,
        "current_user": current_user,
    })


@app.get("/warehouse", response_class=HTMLResponse)
async def warehouse_list(request: Request, db: Session = Depends(get_db), current_user: User = Depends(require_auth)):
    """Qaysi omborda nima bor — ombor qoldiqlari"""
    if not current_user:
        return RedirectResponse(url="/login", status_code=303)
    warehouses = db.query(Warehouse).all()
    # Faqat qoldiq > 0 bo'lgan yozuvlarni ko'rsatamiz (0 qoldiqlar ro'yxatda chiqmasin)
    stocks = db.query(Stock).join(Product).join(Warehouse).filter(Stock.quantity > 0).all()
    
    # Har bir qoldiq uchun oxirgi hujjatni StockMovement dan olish
    stock_sources = {}
    for s in stocks:
        # Oxirgi StockMovement ni topish (har bir operatsiya uchun alohida hujjat)
        last_movement = db.query(StockMovement).filter(
            StockMovement.warehouse_id == s.warehouse_id,
            StockMovement.product_id == s.product_id
        ).order_by(StockMovement.created_at.desc()).first()
        
        if last_movement:
            # Hujjat URL ni aniqlash
            doc_url = ""
            if last_movement.document_type == "Purchase":
                doc_url = f"/purchases/edit/{last_movement.document_id}"
            elif last_movement.document_type == "Production":
                doc_url = f"/production/orders"
            elif last_movement.document_type == "WarehouseTransfer":
                doc_url = f"/warehouse/transfers/{last_movement.document_id}"
            elif last_movement.document_type == "StockAdjustmentDoc":
                doc_url = f"/qoldiqlar/tovar/hujjat/{last_movement.document_id}"
            elif last_movement.document_type == "Sale":
                doc_url = f"/sales/edit/{last_movement.document_id}"
            else:
                doc_url = "#"
            
            doc_date = last_movement.created_at.strftime("%d.%m.%Y") if last_movement.created_at else ""
            stock_sources[s.id] = [(last_movement.document_number or f"{last_movement.document_type}-{last_movement.document_id}", doc_url, doc_date)]
        else:
            # Eski tizim uchun fallback - StockMovement yo'q bo'lsa, eski usulni ishlatish
            items = []
            # 1) Tovar kirim (sotib olingan) — tasdiqlangan kirimlar
            purchases = (
                db.query(Purchase)
                .join(PurchaseItem, Purchase.id == PurchaseItem.purchase_id)
                .filter(
                    Purchase.warehouse_id == s.warehouse_id,
                    PurchaseItem.product_id == s.product_id,
                    Purchase.status == "confirmed",
                )
                .order_by(Purchase.date.desc())
                .limit(1)
                .all()
            )
            for p in purchases:
                items.append((p.number, f"/purchases/edit/{p.id}", p.date.strftime("%d.%m.%Y") if p.date else ""))
            
            # 2) Ishlab chiqarish
            out_wh_id = s.warehouse_id
            productions = (
                db.query(Production)
                .join(Recipe, Production.recipe_id == Recipe.id)
                .filter(
                    Production.status == "completed",
                    Recipe.product_id == s.product_id,
                    or_(
                        Production.output_warehouse_id == out_wh_id,
                        and_(Production.output_warehouse_id.is_(None), Production.warehouse_id == out_wh_id),
                    ),
                )
                .order_by(Production.date.desc())
                .limit(1)
                .all()
            )
            for pr in productions:
                items.append((pr.number, f"/production/orders", pr.date.strftime("%d.%m.%Y") if pr.date else ""))
            
            # 3) Qoldiq hujjati
            adj_docs = (
                db.query(StockAdjustmentDoc)
                .join(StockAdjustmentDocItem, StockAdjustmentDoc.id == StockAdjustmentDocItem.doc_id)
                .filter(
                    StockAdjustmentDoc.status == "confirmed",
                    StockAdjustmentDocItem.warehouse_id == s.warehouse_id,
                    StockAdjustmentDocItem.product_id == s.product_id,
                )
                .order_by(StockAdjustmentDoc.date.desc())
                .limit(1)
                .distinct()
                .all()
            )
            for doc in adj_docs:
                items.append((doc.number, f"/qoldiqlar/tovar/hujjat/{doc.id}", doc.date.strftime("%d.%m.%Y") if doc.date else ""))
            
            # Oxirgi hujjatni olish
            if items:
                items.sort(key=lambda x: x[2] or "", reverse=True)
                stock_sources[s.id] = items[:1]
            else:
                stock_sources[s.id] = []
    
    return templates.TemplateResponse("warehouse/list.html", {
        "request": request,
        "warehouses": warehouses,
        "stocks": stocks,
        "stock_sources": stock_sources,
        "current_user": current_user,
        "page_title": "Ombor qoldiqlari"
    })


@app.post("/warehouse/stock/{stock_id}/zero")
async def warehouse_stock_zero(
    stock_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_admin),
):
    """Qoldiqni nolga tushirish (faqat admin). Ro'yxatdan o'sha qator yo'qoladi."""
    stock = db.query(Stock).filter(Stock.id == stock_id).first()
    if not stock:
        raise HTTPException(status_code=404, detail="Qoldiq topilmadi")
    stock.quantity = 0
    db.commit()
    return RedirectResponse(url="/warehouse", status_code=303)


@app.get("/warehouse/export")
async def warehouse_export(db: Session = Depends(get_db), current_user: User = Depends(require_auth)):
    """Ombor qoldiqlarini Excelga eksport qilish."""
    stocks = db.query(Stock).join(Product).join(Warehouse).filter(Stock.quantity > 0).all()
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Qoldiqlar"
    ws.append(["Ombor nomi", "Ombor kodi", "Mahsulot kodi", "Mahsulot nomi", "Qoldiq", "Tannarx (so'm)", "Summa (so'm)"])
    for s in stocks:
        pr = s.product
        wh = s.warehouse
        tannarx = (pr.purchase_price or 0) if pr else 0
        summa = s.quantity * tannarx
        ws.append([
            wh.name if wh else "",
            wh.code if wh else "",
            pr.code if pr else "",
            pr.name if pr else "",
            s.quantity,
            tannarx,
            summa,
        ])
    stream = io.BytesIO()
    wb.save(stream)
    stream.seek(0)
    return StreamingResponse(
        stream,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": "attachment; filename=ombor_qoldiqlari.xlsx"},
    )


@app.get("/warehouse/template")
async def warehouse_template(db: Session = Depends(get_db), current_user: User = Depends(require_auth)):
    """Qoldiqlar uchun Excel andoza (Tannarx va Sotuv narxi ixtiyoriy)."""
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Andoza"
    ws.append(["Ombor nomi (yoki kodi)", "Mahsulot nomi (yoki kodi)", "Qoldiq", "Tannarx (so'm)", "Sotuv narxi (so'm)"])
    ws.append(["Xom ashyo ombori", "Yong'oq", 30, "", ""])
    ws.append(["Xom ashyo ombori", "Bodom", 100, "", ""])
    for col in range(1, 6):
        ws.column_dimensions[openpyxl.utils.get_column_letter(col)].width = 22
    stream = io.BytesIO()
    wb.save(stream)
    stream.seek(0)
    return StreamingResponse(
        stream,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": "attachment; filename=qoldiqlar_andoza.xlsx"},
    )


@app.post("/warehouse/import")
async def warehouse_import(
    request: Request,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    """Exceldan qoldiqlarni yuklash — bitta hujjat (StockAdjustmentDoc) yaratiladi, barcha mahsulotlar o'sha hujjatga yoziladi."""
    from urllib.parse import quote
    form = await request.form()
    file = form.get("file") or form.get("excel_file")
    if not file or not getattr(file, "filename", None):
        return RedirectResponse(url="/warehouse?error=import&detail=" + quote("Excel fayl tanlang"), status_code=303)
    try:
        contents = await file.read()
        if not contents:
            return RedirectResponse(url="/warehouse?error=import&detail=" + quote("Fayl bo'sh"), status_code=303)
        wb = openpyxl.load_workbook(io.BytesIO(contents), read_only=False, data_only=True)
        ws = wb.active
        rows = list(ws.iter_rows(min_row=2, values_only=True))
        items_data = []  # Barcha mahsulotlar bitta hujjatga yoziladi
        skip_no_wh = 0
        skip_no_prod = 0
        skip_empty = 0
        missing_products = []
        missing_warehouses = []
        
        for idx, row in enumerate(rows):
            if not row or (row[0] is None and row[1] is None):
                skip_empty += 1
                continue
            wh_key = str(row[0] or "").strip() if len(row) > 0 else ""
            raw_prod = row[1] if len(row) > 1 else None
            if raw_prod is not None and isinstance(raw_prod, (int, float)) and float(raw_prod) == int(float(raw_prod)):
                prod_key = str(int(float(raw_prod)))
            else:
                prod_key = str(raw_prod or "").strip()
            try:
                qty = float(row[2]) if len(row) > 2 and row[2] is not None else 0
            except (TypeError, ValueError):
                qty = 0
            tannarx = 0.0
            sotuv_narxi = 0.0
            if len(row) > 3 and row[3] is not None and row[3] != "":
                try:
                    tannarx = float(row[3])
                except (TypeError, ValueError):
                    pass
            if len(row) > 4 and row[4] is not None and row[4] != "":
                try:
                    sotuv_narxi = float(row[4])
                except (TypeError, ValueError):
                    pass
            # Bo'sh qatorlarni o'tkazib yuborish
            if not wh_key or not prod_key:
                skip_empty += 1
                continue
            # Miqdor 0 yoki manfiy bo'lsa ham, hujjatga yozish (adjustment uchun)
            # Lekin qoldiqni yangilashda 0 bo'lishi mumkin
            warehouse = db.query(Warehouse).filter(
                (func.lower(Warehouse.name) == wh_key.lower()) | (Warehouse.code == wh_key)
            ).first()
            product = db.query(Product).filter(
                or_(
                    and_(Product.code.isnot(None), Product.code != "", func.lower(func.trim(Product.code)) == prod_key.lower()),
                    and_(Product.barcode.isnot(None), Product.barcode != "", func.lower(func.trim(Product.barcode)) == prod_key.lower()),
                )
            ).first()
            if not product and prod_key:
                product = db.query(Product).filter(
                    Product.name.isnot(None),
                    func.lower(func.trim(Product.name)) == prod_key.strip().lower()
                ).first()
            if not warehouse:
                if wh_key and wh_key not in missing_warehouses:
                    missing_warehouses.append(wh_key)
                skip_no_wh += 1
                continue
            if not product:
                if prod_key and prod_key not in missing_products:
                    missing_products.append(prod_key)
                skip_no_prod += 1
                continue

            # Mahsulotni hujjatga qo'shish
            items_data.append((product.id, warehouse.id, qty, tannarx, sotuv_narxi))
            
            # Mahsulot narxlarini yangilash
            if tannarx is not None and tannarx > 0:
                product.purchase_price = tannarx
            if sotuv_narxi is not None and sotuv_narxi > 0:
                product.sale_price = sotuv_narxi
        
        if not items_data:
            detail = "Hech qanday to'g'ri qator topilmadi. Ombor va mahsulot nomi/kodi to'g'ri ekanligini tekshiring."
            if missing_products:
                detail += f" Mahsulot topilmadi: {', '.join(missing_products[:10])}"
                if len(missing_products) > 10:
                    detail += f" va yana {len(missing_products) - 10} ta"
            return RedirectResponse(url="/warehouse?error=import&detail=" + quote(detail), status_code=303)
        
        # Bitta hujjat yaratish
        today = datetime.now()
        count = db.query(StockAdjustmentDoc).filter(
            StockAdjustmentDoc.date >= today.replace(hour=0, minute=0, second=0)
        ).count()
        number = f"QLD-{today.strftime('%Y%m%d')}-{str(count + 1).zfill(4)}"
        total_tannarx = sum(qty * cp for _, _, qty, cp, _ in items_data)
        total_sotuv = sum(qty * sp for _, _, qty, _, sp in items_data)
        doc = StockAdjustmentDoc(
            number=number,
            date=today,
            user_id=current_user.id if current_user else None,
            status="draft",  # Qoralama holatida yaratiladi, keyin tasdiqlash mumkin
            total_tannarx=total_tannarx,
            total_sotuv=total_sotuv,
        )
        db.add(doc)
        db.flush()
        
        # Barcha mahsulotlarni hujjatga qo'shish
        for pid, wid, qty, cp, sp in items_data:
            db.add(StockAdjustmentDocItem(
                doc_id=doc.id,
                product_id=pid,
                warehouse_id=wid,
                quantity=qty,
                cost_price=cp,
                sale_price=sp,
            ))
        
        db.flush()  # Hujjat va qatorlarni saqlash
        
        # Hujjatni avtomatik tasdiqlash va qoldiqlarni yangilash
        for item in doc.items:
            # Eski qoldiqni olish
            stock = db.query(Stock).filter(
                Stock.warehouse_id == item.warehouse_id,
                Stock.product_id == item.product_id,
            ).first()
            old_quantity = stock.quantity if stock else 0
            
            # Yangi qoldiqni hisoblash
            new_quantity = item.quantity
            quantity_change = new_quantity - old_quantity
            
            # Qoldiqni yangilash (adjustment uchun to'g'ridan-to'g'ri belgilash)
            if stock:
                stock.quantity = new_quantity
                stock.updated_at = datetime.now()
                stock_id = stock.id
                quantity_after = new_quantity
            else:
                stock = Stock(
                    warehouse_id=item.warehouse_id,
                    product_id=item.product_id,
                    quantity=new_quantity,
                )
                db.add(stock)
                db.flush()
                stock_id = stock.id
                quantity_after = new_quantity
            
            # StockMovement yozuvini yaratish (adjustment)
            if quantity_change != 0:
                movement = StockMovement(
                    stock_id=stock_id,
                    warehouse_id=item.warehouse_id,
                    product_id=item.product_id,
                    operation_type="adjustment",
                    document_type="StockAdjustmentDoc",
                    document_id=doc.id,
                    document_number=doc.number,
                    quantity_change=quantity_change,
                    quantity_after=quantity_after,
                    user_id=current_user.id if current_user else None,
                    note=f"Exceldan yuklash: {doc.number}",
                    created_at=datetime.now()
                )
                db.add(movement)
        
        # Hujjatni tasdiqlangan holatga o'tkazish
        doc.status = "confirmed"
        
        db.commit()
        
        # Xabar tayyorlash
        detail = f"Yuklandi: {len(items_data)} ta"
        if skip_no_prod or skip_no_wh:
            detail += f", o'tkazib yuborildi: {skip_no_prod + skip_no_wh} ta"
        if missing_products:
            detail += f". Mahsulot topilmadi: {', '.join(missing_products[:5])}"
            if len(missing_products) > 5:
                detail += f" va yana {len(missing_products) - 5} ta"
        
        return RedirectResponse(
            url="/warehouse?success=import&detail=" + quote(detail) + "&doc_id=" + str(doc.id) + "&doc_number=" + quote(doc.number),
            status_code=303
        )
    except Exception as e:
        traceback.print_exc()
        return RedirectResponse(url="/warehouse?error=import&detail=" + quote(str(e)[:200]), status_code=303)


@app.get("/warehouse/transfers", response_class=HTMLResponse)
async def warehouse_transfers_list(request: Request, db: Session = Depends(get_db), current_user: User = Depends(require_auth)):
    """Ombordan omborga o'tkazish hujjatlari ro'yxati (spiska)"""
    if not current_user:
        return RedirectResponse(url="/login", status_code=303)
    from urllib.parse import unquote
    transfers = db.query(WarehouseTransfer).order_by(WarehouseTransfer.date.desc()).limit(200).all()
    error = request.query_params.get("error")
    return templates.TemplateResponse("warehouse/transfers_list.html", {
        "request": request,
        "current_user": current_user,
        "transfers": transfers,
        "page_title": "Ombordan omborga o'tkazish",
        "error_message": unquote(error) if error else None,
    })


def _parse_doc_date(s: str):
    """Sana matnini parse qiladi: YYYY-MM-DDTHH:MM, dd.mm.yyyy HH:MM, dd.mm.yyyy va boshqa formatlar."""
    if not s or not str(s).strip():
        return None
    s = str(s).strip()
    # ISO / datetime-local
    try:
        return datetime.fromisoformat(s.replace("Z", "")[:19])
    except (ValueError, TypeError):
        pass
    # dd.mm.yyyy HH:MM yoki dd.mm.yyyy H:M
    import re
    m = re.match(r"(\d{1,2})[./](\d{1,2})[./](\d{4})\s+(\d{1,2}):(\d{2})", s)
    if m:
        try:
            return datetime(int(m.group(3)), int(m.group(2)), int(m.group(1)), int(m.group(4)), int(m.group(5)))
        except (ValueError, TypeError):
            pass
    # dd.mm.yyyy
    m = re.match(r"(\d{1,2})[./](\d{1,2})[./](\d{4})", s)
    if m:
        try:
            return datetime(int(m.group(3)), int(m.group(2)), int(m.group(1)))
        except (ValueError, TypeError):
            pass
    return None


def _get_product_costs_for_transfer(db: Session):
    """Mahsulot tannarxini: Product.purchase_price, keyin so'nggi kirim narxi, keyin so'nggi qoldiq hujjati cost_price."""
    products = db.query(Product).filter(Product.is_active == True).all()
    product_cost = {str(p.id): float(p.purchase_price or 0) for p in products}
    # So'nggi kirim (PurchaseItem) narxi
    last_purchase = (
        db.query(PurchaseItem.product_id, PurchaseItem.price)
        .join(Purchase, PurchaseItem.purchase_id == Purchase.id)
        .filter((PurchaseItem.price or 0) > 0)
        .order_by(Purchase.date.desc(), Purchase.id.desc())
        .all()
    )
    seen_purchase = set()
    for row in last_purchase:
        pid = str(row.product_id)
        if (product_cost.get(pid) or 0) <= 0 and pid not in seen_purchase:
            seen_purchase.add(pid)
            product_cost[pid] = float(row.price or 0)
    # So'nggi tasdiqlangan qoldiq hujjati (StockAdjustmentDocItem) cost_price
    confirmed_ids = [r[0] for r in db.query(StockAdjustmentDoc.id).filter(StockAdjustmentDoc.status == "confirmed").all()]
    if confirmed_ids:
        qoldiq_rows = (
            db.query(StockAdjustmentDocItem.product_id, StockAdjustmentDocItem.cost_price)
            .filter(StockAdjustmentDocItem.doc_id.in_(confirmed_ids), (StockAdjustmentDocItem.cost_price or 0) > 0)
            .join(StockAdjustmentDoc, StockAdjustmentDocItem.doc_id == StockAdjustmentDoc.id)
            .order_by(StockAdjustmentDoc.date.desc())
            .all()
        )
        seen = set()
        for row in qoldiq_rows:
            pid = str(row.product_id)
            if pid not in seen and ((product_cost.get(pid) or 0) <= 0):
                seen.add(pid)
                product_cost[pid] = float(row.cost_price or 0)
    return product_cost


@app.get("/warehouse/transfers/new", response_class=HTMLResponse)
async def warehouse_transfer_new(request: Request, db: Session = Depends(get_db), current_user: User = Depends(require_auth)):
    """Yangi o'tkazish hujjati"""
    if not current_user:
        return RedirectResponse(url="/login", status_code=303)
    warehouses = db.query(Warehouse).filter(Warehouse.is_active == True).all()
    products = db.query(Product).filter(Product.is_active == True).order_by(Product.name).all()
    stocks = db.query(Stock).filter(Stock.quantity > 0).all()
    stock_by_warehouse_product = {}
    for s in stocks:
        wid, pid = str(s.warehouse_id), str(s.product_id)
        if wid not in stock_by_warehouse_product:
            stock_by_warehouse_product[wid] = {}
        stock_by_warehouse_product[wid][pid] = stock_by_warehouse_product[wid].get(pid, 0) + float(s.quantity or 0)
    product_cost = _get_product_costs_for_transfer(db)
    stock_cost_by_warehouse_product = {}
    for wid, pids in stock_by_warehouse_product.items():
        stock_cost_by_warehouse_product[wid] = {pid: product_cost.get(pid, 0) for pid in pids}
    _uc = lambda u: (u.code or "").strip().lower() if u else ""
    _is_dona = lambda u: _uc(u) in ("dona", "pc", "pcs", "ta", "шт")
    products_list = [
        {"id": p.id, "name": (p.name or ""), "code": (p.code or ""), "unit_name": (p.unit.name if p.unit else "") or "", "unit_code": _uc(p.unit), "is_dona": _is_dona(p.unit)}
        for p in products
    ]
    now = datetime.now()
    return templates.TemplateResponse("warehouse/transfer_form.html", {
        "request": request,
        "current_user": current_user,
        "transfer": None,
        "warehouses": warehouses,
        "products": products,
        "products_list": products_list,
        "stock_by_warehouse_product": stock_by_warehouse_product,
        "stock_cost_by_warehouse_product": stock_cost_by_warehouse_product,
        "now": now,
        "doc_date_value": now.strftime("%Y-%m-%dT%H:%M"),
        "doc_date_display": now.strftime("%d.%m.%Y %H:%M"),
        "page_title": "Ombordan omborga o'tkazish (yaratish)"
    })


@app.get("/warehouse/transfers/{transfer_id}", response_class=HTMLResponse)
async def warehouse_transfer_edit(request: Request, transfer_id: int, db: Session = Depends(get_db), current_user: User = Depends(require_auth)):
    """O'tkazish hujjatini tahrirlash"""
    if not current_user:
        return RedirectResponse(url="/login", status_code=303)
    transfer = db.query(WarehouseTransfer).filter(WarehouseTransfer.id == transfer_id).first()
    if not transfer:
        raise HTTPException(status_code=404, detail="Hujjat topilmadi")
    warehouses = db.query(Warehouse).filter(Warehouse.is_active == True).all()
    products = db.query(Product).filter(Product.is_active == True).order_by(Product.name).all()
    stocks = db.query(Stock).filter(Stock.quantity > 0).all()
    stock_by_warehouse_product = {}
    for s in stocks:
        wid, pid = str(s.warehouse_id), str(s.product_id)
        if wid not in stock_by_warehouse_product:
            stock_by_warehouse_product[wid] = {}
        stock_by_warehouse_product[wid][pid] = stock_by_warehouse_product[wid].get(pid, 0) + float(s.quantity or 0)
    product_cost = _get_product_costs_for_transfer(db)
    stock_cost_by_warehouse_product = {}
    for wid, pids in stock_by_warehouse_product.items():
        stock_cost_by_warehouse_product[wid] = {pid: product_cost.get(pid, 0) for pid in pids}
    from_wh_id = str(transfer.from_warehouse_id) if transfer.from_warehouse_id else None
    source_costs = {}
    if from_wh_id and from_wh_id in stock_cost_by_warehouse_product:
        for pid, cost in stock_cost_by_warehouse_product[from_wh_id].items():
            try:
                source_costs[int(pid)] = cost
            except (ValueError, TypeError):
                pass
    now_dt = transfer.date or datetime.now()
    doc_date_value = now_dt.strftime("%Y-%m-%dT%H:%M") if now_dt else datetime.now().strftime("%Y-%m-%dT%H:%M")
    doc_date_display = now_dt.strftime("%d.%m.%Y %H:%M") if now_dt else datetime.now().strftime("%d.%m.%Y %H:%M")
    _uc = lambda u: (u.code or "").strip().lower() if u else ""
    _is_dona = lambda u: _uc(u) in ("dona", "pc", "pcs", "ta", "шт")
    products_list = [
        {"id": p.id, "name": (p.name or ""), "code": (p.code or ""), "unit_name": (p.unit.name if p.unit else "") or "", "unit_code": _uc(p.unit), "is_dona": _is_dona(p.unit)}
        for p in products
    ]
    return templates.TemplateResponse("warehouse/transfer_form.html", {
        "request": request,
        "current_user": current_user,
        "transfer": transfer,
        "warehouses": warehouses,
        "products": products,
        "products_list": products_list,
        "stock_by_warehouse_product": stock_by_warehouse_product,
        "stock_cost_by_warehouse_product": stock_cost_by_warehouse_product,
        "source_costs": source_costs,
        "now": now_dt,
        "doc_date_value": doc_date_value,
        "doc_date_display": doc_date_display,
        "page_title": f"O'tkazish {transfer.number}"
    })


@app.post("/warehouse/transfers/create")
async def warehouse_transfer_create(
    request: Request,
    from_warehouse_id: int = Form(...),
    to_warehouse_id: int = Form(...),
    note: str = Form(""),
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth)
):
    """Yangi o'tkazish hujjatini saqlash (draft)"""
    from urllib.parse import quote
    if not current_user:
        return RedirectResponse(url="/login", status_code=303)
    if from_warehouse_id == to_warehouse_id:
        return RedirectResponse(url="/warehouse/transfers/new?error=" + quote("Qayerdan va qayerga bir xil bo'lmasin."), status_code=303)
    form = await request.form()
    today = datetime.now()
    doc_date = _parse_doc_date(form.get("doc_date", "")) or today
    count = db.query(WarehouseTransfer).filter(
        WarehouseTransfer.date >= doc_date.replace(hour=0, minute=0, second=0)
    ).count()
    number = f"OT-{doc_date.strftime('%Y%m%d')}-{str(count + 1).zfill(4)}"
    transfer = WarehouseTransfer(
        number=number,
        date=doc_date,
        from_warehouse_id=from_warehouse_id,
        to_warehouse_id=to_warehouse_id,
        status="draft",
        user_id=current_user.id,
        note=note or None
    )
    db.add(transfer)
    db.commit()
    db.refresh(transfer)
    for key, value in form.items():
        if key.startswith("product_id_") and value:
            try:
                pid = int(value)
                qkey = "quantity_" + key.replace("product_id_", "")
                qty = float(form.get(qkey, "0").replace(",", "."))
                if pid and qty > 0:
                    db.add(WarehouseTransferItem(transfer_id=transfer.id, product_id=pid, quantity=qty))
            except (ValueError, TypeError):
                pass
    db.commit()
    return RedirectResponse(url=f"/warehouse/transfers/{transfer.id}", status_code=303)


@app.post("/warehouse/transfers/{transfer_id}/save")
async def warehouse_transfer_save(
    request: Request,
    transfer_id: int,
    from_warehouse_id: int = Form(...),
    to_warehouse_id: int = Form(...),
    note: str = Form(""),
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth)
):
    """O'tkazish hujjatini saqlash"""
    from urllib.parse import quote
    if not current_user:
        return RedirectResponse(url="/login", status_code=303)
    transfer = db.query(WarehouseTransfer).filter(WarehouseTransfer.id == transfer_id).first()
    if not transfer:
        raise HTTPException(status_code=404, detail="Hujjat topilmadi")
    # Faqat draft va pending_approval holatida tahrirlash mumkin
    if transfer.status == "confirmed":
        return RedirectResponse(url=f"/warehouse/transfers/{transfer_id}?error=" + quote("Tasdiqlangan hujjatni tahrirlab bo'lmaydi."), status_code=303)
    if from_warehouse_id == to_warehouse_id:
        return RedirectResponse(url=f"/warehouse/transfers/{transfer_id}?error=" + quote("Qayerdan va qayerga bir xil bo'lmasin."), status_code=303)
    form = await request.form()
    parsed = _parse_doc_date(form.get("doc_date", ""))
    if parsed:
        transfer.date = parsed
    transfer.from_warehouse_id = from_warehouse_id
    transfer.to_warehouse_id = to_warehouse_id
    transfer.note = note or None
    # Faqat "Tasdiqlashga yuborish" bosilganda pending_approval; "Saqlash (qoralama)" — draft qoladi
    if form.get("submit_approval") and transfer.status == "draft":
        transfer.status = "pending_approval"
    db.query(WarehouseTransferItem).filter(WarehouseTransferItem.transfer_id == transfer_id).delete()
    for key, value in form.items():
        if key.startswith("product_id_") and value:
            try:
                pid = int(value)
                qkey = "quantity_" + key.replace("product_id_", "")
                qty = float(form.get(qkey, "0").replace(",", "."))
                if pid and qty > 0:
                    db.add(WarehouseTransferItem(transfer_id=transfer_id, product_id=pid, quantity=qty))
            except (ValueError, TypeError):
                pass
    db.commit()
    return RedirectResponse(url=f"/warehouse/transfers/{transfer_id}?saved=1", status_code=303)


@app.post("/warehouse/transfers/{transfer_id}/confirm")
async def warehouse_transfer_confirm(
    transfer_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth)
):
    """O'tkazish hujjatini tasdiqlash"""
    from urllib.parse import quote
    if not current_user:
        return RedirectResponse(url="/login", status_code=303)
    transfer = db.query(WarehouseTransfer).filter(WarehouseTransfer.id == transfer_id).first()
    if not transfer:
        raise HTTPException(status_code=404, detail="Hujjat topilmadi")
    # Faqat pending_approval holatidagi hujjatni tasdiqlash mumkin
    if transfer.status == "confirmed":
        return RedirectResponse(url=f"/warehouse/transfers/{transfer_id}?error=" + quote("Hujjat allaqachon tasdiqlangan."), status_code=303)
    if transfer.status != "pending_approval":
        return RedirectResponse(url=f"/warehouse/transfers/{transfer_id}?error=" + quote("Hujjatni avval saqlang (pending_approval holatiga o'tkazish kerak)."), status_code=303)
    items = db.query(WarehouseTransferItem).filter(WarehouseTransferItem.transfer_id == transfer_id).all()
    if not items:
        return RedirectResponse(url=f"/warehouse/transfers/{transfer_id}?error=" + quote("Kamida bitta mahsulot qo'shing."), status_code=303)
    # Qayerdan omborda bitta mahsulot uchun bir nechta Stock qatori bo'lsa — yig'indiga birlashtirib, tekshiruv va chiqim to'g'ri bo'ladi
    from_wh = transfer.from_warehouse_id
    for item in items:
        rows = db.query(Stock).filter(
            Stock.warehouse_id == from_wh,
            Stock.product_id == item.product_id
        ).all()
        if len(rows) > 1:
            total = sum(float(r.quantity or 0) for r in rows)
            keep = rows[0]
            keep.quantity = total
            keep.updated_at = datetime.now()
            for r in rows[1:]:
                db.delete(r)
            db.flush()
    # Kichik kasr farqi bo'lsa (masalan qoldiq 2806.77, so'ralgan 2807) — mavjud miqdorga tushirib tasdiqlash
    TOLERANCE_ABS = 0.5   # 0.5 birlikgacha farq bo'lsa mavjudga tushiriladi
    TOLERANCE_PCT = 0.002  # yoki so'ralganning 0.2% gacha
    for item in items:
        rows = db.query(Stock).filter(
            Stock.warehouse_id == transfer.from_warehouse_id,
            Stock.product_id == item.product_id
        ).all()
        need = float(item.quantity or 0)
        have = sum(float(r.quantity or 0) for r in rows)
        if have + 1e-6 < need:
            shortfall = need - have
            allow_cap = shortfall <= max(TOLERANCE_ABS, need * TOLERANCE_PCT)
            if allow_cap and have >= 1e-6:
                # Mavjud miqdorga tushirib tasdiqlash (hujjatda ham yangilanadi)
                item.quantity = have
                db.flush()
                continue
            prod = db.query(Product).filter(Product.id == item.product_id).first()
            name = prod.name if prod else f"#{item.product_id}"
            avail_display = "0" if abs(have) < 1e-6 else ("%.6f" % have).rstrip("0").rstrip(".")
            return RedirectResponse(
                url=f"/warehouse/transfers/{transfer_id}?error=" + quote(f"Qayerdan omborda «{name}» yetarli emas (kerak: {item.quantity}, mavjud: {avail_display})"),
                status_code=303
            )
    # Qoldiqlarni yangilash - faqat tasdiqlanganda
    for item in items:
        # Qayerdan ombordan ayirish - StockMovement yozuvini yaratish
        create_stock_movement(
            db=db,
            warehouse_id=transfer.from_warehouse_id,
            product_id=item.product_id,
            quantity_change=-item.quantity,  # Chiqim
            operation_type="transfer_out",
            document_type="WarehouseTransfer",
            document_id=transfer.id,
            document_number=transfer.number,
            user_id=current_user.id if current_user else None,
            note=f"O'tkazish (chiqim): {transfer.number}"
        )
        
        # Qayerga omborga qo'shish - StockMovement yozuvini yaratish
        create_stock_movement(
            db=db,
            warehouse_id=transfer.to_warehouse_id,
            product_id=item.product_id,
            quantity_change=item.quantity,  # Kirim
            operation_type="transfer_in",
            document_type="WarehouseTransfer",
            document_id=transfer.id,
            document_number=transfer.number,
            user_id=current_user.id if current_user else None,
            note=f"O'tkazish (kirim): {transfer.number}"
        )
    
    # Tasdiqlash ma'lumotlarini saqlash
    transfer.status = "confirmed"
    transfer.approved_by_user_id = current_user.id
    transfer.approved_at = datetime.now()
    db.commit()
    return RedirectResponse(url=f"/warehouse/transfers/{transfer_id}?confirmed=1", status_code=303)


@app.post("/warehouse/transfers/{transfer_id}/revert")
async def warehouse_transfer_revert(
    transfer_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_admin),
):
    """Tasdiqlashni bekor qilish: ombor harakatini teskari qilish, hujjat qoralamaga o'tadi (faqat admin)."""
    from urllib.parse import quote
    transfer = db.query(WarehouseTransfer).filter(WarehouseTransfer.id == transfer_id).first()
    if not transfer:
        raise HTTPException(status_code=404, detail="Hujjat topilmadi")
    if transfer.status != "confirmed":
        return RedirectResponse(url=f"/warehouse/transfers?error=" + quote("Faqat tasdiqlangan hujjatning tasdiqini bekor qilish mumkin."), status_code=303)
    items = db.query(WarehouseTransferItem).filter(WarehouseTransferItem.transfer_id == transfer_id).all()
    for item in items:
        dest = db.query(Stock).filter(
            Stock.warehouse_id == transfer.to_warehouse_id,
            Stock.product_id == item.product_id,
        ).first()
        if dest:
            dest.quantity -= item.quantity
            if dest.quantity < 0:
                dest.quantity = 0
        src = db.query(Stock).filter(
            Stock.warehouse_id == transfer.from_warehouse_id,
            Stock.product_id == item.product_id,
        ).first()
        if src:
            src.quantity += item.quantity
        else:
            db.add(Stock(warehouse_id=transfer.from_warehouse_id, product_id=item.product_id, quantity=item.quantity))
    delete_stock_movements_for_document(db, "WarehouseTransfer", transfer_id)
    transfer.status = "draft"
    db.commit()
    return RedirectResponse(url="/warehouse/transfers?reverted=1", status_code=303)


@app.post("/warehouse/transfers/{transfer_id}/delete")
async def warehouse_transfer_delete(
    transfer_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_admin),
):
    """O'tkazish hujjatini o'chirish (faqat admin). Faqat qoralama holatida o'chirish mumkin; tasdiqlangan bo'lsa avval tasdiqni bekor qilish kerak."""
    from urllib.parse import quote
    transfer = db.query(WarehouseTransfer).filter(WarehouseTransfer.id == transfer_id).first()
    if not transfer:
        raise HTTPException(status_code=404, detail="Hujjat topilmadi")
    if transfer.status == "confirmed":
        return RedirectResponse(
            url="/warehouse/transfers?error=" + quote("Tasdiqlangan hujjatni to'g'ridan-to'g'ri o'chirib bo'lmaydi. Avval tasdiqni bekor qiling."),
            status_code=303
        )
    db.delete(transfer)
    db.commit()
    return RedirectResponse(url="/warehouse/transfers?deleted=1", status_code=303)


@app.get("/warehouse/movement", response_class=HTMLResponse)
async def warehouse_movement(request: Request, current_user: User = Depends(require_auth)):
    """Spiskaga yo'naltirish"""
    return RedirectResponse(url="/warehouse/transfers", status_code=302)


@app.post("/warehouse/transfer")
async def warehouse_transfer(
    request: Request,
    from_warehouse_id: int = Form(...),
    to_warehouse_id: int = Form(...),
    product_id: int = Form(...),
    quantity: float = Form(...),
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth)
):
    """Ombordan omborga o'tkazish: bir ombordan ayirib, ikkinchisiga qo'shish"""
    from urllib.parse import quote
    if not current_user:
        return RedirectResponse(url="/login", status_code=303)
    if from_warehouse_id == to_warehouse_id:
        return RedirectResponse(url="/warehouse/movement?error=1&detail=" + quote("Qayerdan va qayerga ombor bir xil bo'lmasin."), status_code=303)
    if quantity <= 0:
        return RedirectResponse(url="/warehouse/movement?error=1&detail=" + quote("Miqdor 0 dan katta bo'lishi kerak."), status_code=303)
    source = db.query(Stock).filter(
        Stock.warehouse_id == from_warehouse_id,
        Stock.product_id == product_id
    ).first()
    need_q = float(quantity or 0)
    have_q = float(source.quantity or 0) if source else 0
    if not source or (have_q + 1e-6 < need_q):
        product = db.query(Product).filter(Product.id == product_id).first()
        name = product.name if product else f"#{product_id}"
        avail_display = "0" if abs(have_q) < 1e-6 else ("%.6f" % have_q).rstrip("0").rstrip(".")
        return RedirectResponse(url="/warehouse/movement?error=1&detail=" + quote(f"Qayerdan omborda «{name}» yetarli emas (kerak: {quantity}, mavjud: {avail_display})"), status_code=303)
    source.quantity -= quantity
    if source.quantity <= 0:
        source.quantity = 0
    dest = db.query(Stock).filter(
        Stock.warehouse_id == to_warehouse_id,
        Stock.product_id == product_id
    ).first()
    if dest:
        dest.quantity += quantity
    else:
        db.add(Stock(warehouse_id=to_warehouse_id, product_id=product_id, quantity=quantity))
    db.commit()
    return RedirectResponse(url="/warehouse/movement?success=1", status_code=303)


# ==========================================
# TOVAR KIRIMI (PURCHASE)
# ==========================================

@app.get("/purchases", response_class=HTMLResponse)
async def purchases_list(
    request: Request,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    """Tovar kirimlari ro'yxati. from=pos bo'lsa — faqat joriy foydalanuvchi kiritganlar."""
    from urllib.parse import unquote
    if not current_user:
        return RedirectResponse(url="/login", status_code=303)
    query = db.query(Purchase)
    from_pos = request.query_params.get("from") == "pos"
    if from_pos:
        query = query.filter(Purchase.user_id == current_user.id)
    purchases = query.order_by(Purchase.date.desc()).limit(100).all()
    error = request.query_params.get("error")
    error_detail = unquote(request.query_params.get("detail", "") or "")
    return templates.TemplateResponse("purchases/list.html", {
        "request": request,
        "purchases": purchases,
        "current_user": current_user,
        "page_title": "Tovar kirimlari",
        "error": error,
        "error_detail": error_detail,
        "from_pos": from_pos,
    })


@app.get("/purchases/new", response_class=HTMLResponse)
async def purchase_new(request: Request, db: Session = Depends(get_db), current_user: User = Depends(require_auth)):
    """Yangi tovar kirimi — ombor faqat foydalanuvchiga tegishli/tayinlanganlar."""
    from_pos = request.query_params.get("from") == "pos"
    products = db.query(Product).filter(Product.is_active == True).all()
    partners = db.query(Partner).filter(Partner.is_active == True).order_by(Partner.name).all()
    warehouses = get_warehouses_for_user(db, current_user)
    return templates.TemplateResponse("purchases/new.html", {
        "request": request,
        "products": products,
        "partners": partners,
        "warehouses": warehouses,
        "current_user": current_user,
        "page_title": "Yangi tovar kirimi",
        "from_pos": from_pos,
    })


@app.post("/purchases/create")
async def purchase_create(
    request: Request,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    """Tovar kirimini yaratish — tovarlar va xarajatlar bilan bir sahifada"""
    form = await request.form()
    partner_id = form.get("partner_id")
    warehouse_id = form.get("warehouse_id")
    if not partner_id or not warehouse_id:
        raise HTTPException(status_code=400, detail="Ta'minotchi va omborni tanlang")
    try:
        partner_id = int(partner_id)
        warehouse_id = int(warehouse_id)
    except (TypeError, ValueError):
        raise HTTPException(status_code=400, detail="Noto'g'ri ma'lumot")
    allowed_warehouses = get_warehouses_for_user(db, current_user)
    allowed_warehouse_ids = [w.id for w in allowed_warehouses]
    if warehouse_id not in allowed_warehouse_ids:
        raise HTTPException(status_code=403, detail="Tanlangan ombor sizga ruxsat etilmagan")

    product_ids = form.getlist("product_id")
    quantities = form.getlist("quantity")
    prices = form.getlist("price")
    expense_names = form.getlist("expense_name")
    expense_amounts = form.getlist("expense_amount")

    items_data = []
    for i, pid in enumerate(product_ids):
        if not pid or not pid.strip():
            continue
        try:
            qty = float(quantities[i]) if i < len(quantities) else 0
            pr = float(prices[i]) if i < len(prices) else 0
        except (TypeError, ValueError):
            continue
        if qty <= 0:
            continue
        try:
            items_data.append((int(pid), qty, pr))
        except ValueError:
            continue

    if not items_data:
        raise HTTPException(status_code=400, detail="Kamida bitta mahsulot qo'shing (mahsulot, miqdor va narx).")

    today = datetime.now()
    count = db.query(Purchase).filter(
        Purchase.date >= today.replace(hour=0, minute=0, second=0)
    ).count()
    number = f"P-{today.strftime('%Y%m%d')}-{str(count + 1).zfill(4)}"

    total = sum(qty * pr for _, qty, pr in items_data)
    total_expenses = 0
    for j, name in enumerate(expense_names):
        if not (name and str(name).strip()):
            continue
        try:
            amt = float(expense_amounts[j]) if j < len(expense_amounts) else 0
        except (TypeError, ValueError):
            amt = 0
        if amt > 0:
            total_expenses += amt

    purchase = Purchase(
        number=number,
        partner_id=partner_id,
        warehouse_id=warehouse_id,
        user_id=current_user.id if current_user else None,
        total=total,
        total_expenses=total_expenses,
        status="draft",
    )
    db.add(purchase)
    db.flush()

    for pid, qty, pr in items_data:
        item = PurchaseItem(
            purchase_id=purchase.id,
            product_id=pid,
            quantity=qty,
            price=pr,
            total=qty * pr,
        )
        db.add(item)

    for j, name in enumerate(expense_names):
        if not (name and str(name).strip()):
            continue
        try:
            amt = float(expense_amounts[j]) if j < len(expense_amounts) else 0
        except (TypeError, ValueError):
            amt = 0
        if amt > 0:
            db.add(PurchaseExpense(purchase_id=purchase.id, name=str(name).strip(), amount=amt))

    db.commit()
    from_pos = (form.get("from_pos") or "").strip() == "1"
    edit_url = f"/purchases/edit/{purchase.id}" + ("?from=pos" if from_pos else "")
    return RedirectResponse(url=edit_url, status_code=303)


@app.get("/purchases/edit/{purchase_id}", response_class=HTMLResponse)
async def purchase_edit(request: Request, purchase_id: int, db: Session = Depends(get_db), current_user: User = Depends(require_auth)):
    """Tovar kirimini tahrirlash (tasdiqlangan kirimni faqat admin tahrirlashi mumkin)"""
    from urllib.parse import unquote
    if not current_user:
        return RedirectResponse(url="/login", status_code=303)
    purchase = db.query(Purchase).filter(Purchase.id == purchase_id).first()
    if not purchase:
        raise HTTPException(status_code=404, detail="Tovar kirimi topilmadi")
    if purchase.status == "confirmed" and current_user.role != "admin":
        raise HTTPException(status_code=403, detail="Tasdiqlangan kirimni faqat administrator tahrirlashi mumkin")
    from_pos = request.query_params.get("from") == "pos"
    products = db.query(Product).filter(Product.is_active == True).all()
    revert_error = request.query_params.get("error") == "revert"
    revert_detail = unquote(request.query_params.get("detail", "") or "")
    return templates.TemplateResponse("purchases/edit.html", {
        "request": request,
        "purchase": purchase,
        "products": products,
        "current_user": current_user,
        "page_title": f"Tovar kirimi: {purchase.number}",
        "revert_error": revert_error,
        "revert_detail": revert_detail,
        "from_pos": from_pos,
    })


@app.post("/purchases/{purchase_id}/add-item")
async def purchase_add_item(
    purchase_id: int,
    product_id: int = Form(...),
    quantity: float = Form(...),
    price: float = Form(...),
    db: Session = Depends(get_db)
):
    """Tovar kirimiga mahsulot qo'shish"""
    purchase = db.query(Purchase).filter(Purchase.id == purchase_id).first()
    if not purchase:
        raise HTTPException(status_code=404, detail="Tovar kirimi topilmadi")
    
    total = quantity * price
    item = PurchaseItem(
        purchase_id=purchase_id,
        product_id=product_id,
        quantity=quantity,
        price=price,
        total=total
    )
    db.add(item)
    
    purchase.total = db.query(PurchaseItem).filter(
        PurchaseItem.purchase_id == purchase_id
    ).with_entities(func.sum(PurchaseItem.total)).scalar() or 0
    purchase.total += total
    
    db.commit()
    return RedirectResponse(url=f"/purchases/edit/{purchase_id}", status_code=303)


@app.post("/purchases/{purchase_id}/delete-item/{item_id}")
async def purchase_delete_item(
    purchase_id: int,
    item_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    """Tovar kirimidan mahsulot qatorini o'chirish (faqat qoralama)"""
    purchase = db.query(Purchase).filter(Purchase.id == purchase_id).first()
    if not purchase or purchase.status != "draft":
        raise HTTPException(status_code=400, detail="Faqat qoralamani tahrirlash mumkin")
    item = db.query(PurchaseItem).filter(PurchaseItem.id == item_id, PurchaseItem.purchase_id == purchase_id).first()
    if not item:
        raise HTTPException(status_code=404, detail="Qator topilmadi")
    purchase.total = (purchase.total or 0) - (item.total or 0)
    db.delete(item)
    db.commit()
    return RedirectResponse(url=f"/purchases/edit/{purchase_id}", status_code=303)


@app.post("/purchases/{purchase_id}/add-expense")
async def purchase_add_expense(
    purchase_id: int,
    request: Request,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    """Tovar kirimiga xarajat qo'shish"""
    purchase = db.query(Purchase).filter(Purchase.id == purchase_id).first()
    if not purchase or purchase.status != "draft":
        raise HTTPException(status_code=400, detail="Faqat qoralamani tahrirlash mumkin")
    form = await request.form()
    name = (form.get("name") or "").strip()
    try:
        amount = float(form.get("amount") or 0)
    except (TypeError, ValueError):
        amount = 0
    if not name or amount <= 0:
        return RedirectResponse(url=f"/purchases/edit/{purchase_id}", status_code=303)
    db.add(PurchaseExpense(purchase_id=purchase_id, name=name, amount=amount))
    purchase.total_expenses = (purchase.total_expenses or 0) + amount
    db.commit()
    return RedirectResponse(url=f"/purchases/edit/{purchase_id}", status_code=303)


@app.post("/purchases/{purchase_id}/delete-expense/{expense_id}")
async def purchase_delete_expense(
    purchase_id: int,
    expense_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    """Tovar kirimidan xarajatni o'chirish (faqat qoralama)"""
    purchase = db.query(Purchase).filter(Purchase.id == purchase_id).first()
    if not purchase or purchase.status != "draft":
        raise HTTPException(status_code=400, detail="Faqat qoralamani tahrirlash mumkin")
    expense = db.query(PurchaseExpense).filter(
        PurchaseExpense.id == expense_id,
        PurchaseExpense.purchase_id == purchase_id,
    ).first()
    if not expense:
        raise HTTPException(status_code=404, detail="Xarajat topilmadi")
    purchase.total_expenses = (purchase.total_expenses or 0) - (expense.amount or 0)
    if purchase.total_expenses < 0:
        purchase.total_expenses = 0
    db.delete(expense)
    db.commit()
    return RedirectResponse(url=f"/purchases/edit/{purchase_id}", status_code=303)


@app.post("/purchases/{purchase_id}/confirm")
async def purchase_confirm(purchase_id: int, db: Session = Depends(get_db), current_user: User = Depends(require_auth)):
    """Tovar kirimini tasdiqlash va ombor qoldiqlarini yangilash"""
    purchase = db.query(Purchase).filter(Purchase.id == purchase_id).first()
    if not purchase:
        raise HTTPException(status_code=404, detail="Tovar kirimi topilmadi")
    
    if purchase.status != "draft":
        raise HTTPException(status_code=400, detail="Faqat qoralama holatidagi kirimlarni tasdiqlash mumkin")
    
    if not purchase.items:
        raise HTTPException(status_code=400, detail="Tasdiqlash uchun kamida bitta mahsulot qo'shing. Kirimda mahsulotlar bo'lishi kerak.")
    
    total_expenses = purchase.total_expenses or 0
    items_total = purchase.total or 0
    for item in purchase.items:
        # StockMovement yozuvini yaratish - har bir operatsiya uchun alohida hujjat
        create_stock_movement(
            db=db,
            warehouse_id=purchase.warehouse_id,
            product_id=item.product_id,
            quantity_change=item.quantity,
            operation_type="purchase",
            document_type="Purchase",
            document_id=purchase.id,
            document_number=purchase.number,
            user_id=current_user.id if current_user else None,
            note=f"Tovar kirimi: {purchase.number}"
        )
        
        # Tannarx = qator narxi + xarajat ulushi (tovar kirimi summasi + xarajat = tannarx)
        product = db.query(Product).filter(Product.id == item.product_id).first()
        if product:
            cost_per_unit = item.price
            if total_expenses > 0 and items_total > 0 and item.total and item.quantity:
                expense_share = (item.total / items_total) * total_expenses
                cost_per_unit = item.price + (expense_share / item.quantity)
            product.purchase_price = cost_per_unit
    
    purchase.status = "confirmed"
    total_with_expenses = items_total + total_expenses
    if purchase.partner_id:
        partner = db.query(Partner).filter(Partner.id == purchase.partner_id).first()
        if partner:
            partner.balance -= total_with_expenses

    db.commit()
    check_low_stock_and_notify(db)
    return RedirectResponse(url=f"/purchases", status_code=303)


@app.post("/purchases/{purchase_id}/revert")
async def purchase_revert(
    purchase_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_admin)
):
    """Tasdiqni bekor qilish (faqat admin): ombor qoldig'ini qaytarish, holatni qoralamaga o'tkazish"""
    from urllib.parse import quote
    purchase = db.query(Purchase).filter(Purchase.id == purchase_id).first()
    if not purchase:
        raise HTTPException(status_code=404, detail="Tovar kirimi topilmadi")
    if purchase.status != "confirmed":
        db.rollback()
        return RedirectResponse(
            url=f"/purchases/edit/{purchase_id}?error=revert&detail=" + quote("Faqat tasdiqlangan kirimning tasdiqini bekor qilish mumkin."),
            status_code=303
        )
    for item in purchase.items:
        stock = db.query(Stock).filter(
            Stock.warehouse_id == purchase.warehouse_id,
            Stock.product_id == item.product_id
        ).first()
        if not stock:
            db.rollback()
            return RedirectResponse(
                url=f"/purchases/edit/{purchase_id}?error=revert&detail=" + quote("Ombor qoldig'i topilmadi yoki o'zgargan. Tasdiqni bekor qilish mumkin emas."),
                status_code=303
            )
        stock.quantity -= item.quantity
        if stock.quantity < 0:
            db.rollback()
            return RedirectResponse(
                url=f"/purchases/edit/{purchase_id}?error=revert&detail=" + quote("Ombor qoldig'i yetarli emas (qoldiq o'zgartirilgan). Tasdiqni bekor qilish mumkin emas."),
                status_code=303
            )
    total_with_expenses = purchase.total + (purchase.total_expenses or 0)
    if purchase.partner_id:
        partner = db.query(Partner).filter(Partner.id == purchase.partner_id).first()
        if partner:
            partner.balance += total_with_expenses
    delete_stock_movements_for_document(db, "Purchase", purchase_id)
    purchase.status = "draft"
    db.commit()
    return RedirectResponse(url=f"/purchases/edit/{purchase_id}", status_code=303)


@app.post("/purchases/{purchase_id}/delete")
async def purchase_delete(
    purchase_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_admin),
):
    """Tovar kirimini o'chirish — faqat qoralama, faqat admin"""
    purchase = db.query(Purchase).filter(Purchase.id == purchase_id).first()
    if not purchase:
        raise HTTPException(status_code=404, detail="Tovar kirimi topilmadi")
    if purchase.status != "draft":
        from urllib.parse import quote
        return RedirectResponse(
            url=f"/purchases?error=delete&detail=" + quote("Faqat qoralama holatidagi kirimni o'chirish mumkin. Avval tasdiqni bekor qiling."),
            status_code=303
        )
    db.delete(purchase)
    db.commit()
    return RedirectResponse(url="/purchases", status_code=303)


# ==========================================
# MIJOZLAR
# ==========================================

@app.get("/partners", response_class=HTMLResponse)
async def partners_list(request: Request, type: str = "all", db: Session = Depends(get_db), current_user: User = Depends(require_auth)):
    """Kontragentlar ro'yxati"""
    query = db.query(Partner)
    if type != "all":
        query = query.filter(Partner.type == type)
    
    partners = query.all()
    
    try:
        from app.config.maps_config import YANDEX_MAPS_API_KEY
        yandex_apikey = YANDEX_MAPS_API_KEY or ""
    except Exception:
        yandex_apikey = ""
    return templates.TemplateResponse("partners/list.html", {
        "request": request,
        "partners": partners,
        "current_type": type,
        "current_user": current_user,
        "page_title": "Kontragentlar",
        "yandex_maps_apikey": yandex_apikey,
    })


@app.post("/partners/add")
async def partner_add(
    request: Request,
    name: str = Form(...),
    type: str = Form(...),
    phone: str = Form(""),
    address: str = Form(""),
    credit_limit: float = Form(0),
    discount_percent: float = Form(0),
    db: Session = Depends(get_db)
):
    """Kontragent qo'shish"""
    # Dublikat tekshiruvi - nom bo'yicha
    existing_by_name = db.query(Partner).filter(Partner.name == name).first()
    if existing_by_name:
        raise HTTPException(status_code=400, detail=f"'{name}' nomli kontragent allaqachon mavjud!")
    
    # Dublikat tekshiruvi - telefon bo'yicha (agar telefon kiritilgan bo'lsa)
    if phone and phone.strip():
        existing_by_phone = db.query(Partner).filter(Partner.phone == phone).first()
        if existing_by_phone:
            raise HTTPException(status_code=400, detail=f"'{phone}' telefon raqamli kontragent allaqachon mavjud!")
    
    partner = Partner(
        name=name,
        code=None,  # Kod kerak emas
        type=type,
        phone=phone,
        address=address,
        credit_limit=credit_limit,
        discount_percent=discount_percent
    )
    db.add(partner)
    db.commit()
    return RedirectResponse(url="/partners", status_code=303)


@app.post("/partners/edit/{partner_id}")
async def partner_edit(
    partner_id: int,
    name: str = Form(...),
    type: str = Form(...),
    phone: str = Form(""),
    address: str = Form(""),
    credit_limit: float = Form(0),
    discount_percent: float = Form(0),
    db: Session = Depends(get_db)
):
    """Kontragentni tahrirlash"""
    partner = db.query(Partner).filter(Partner.id == partner_id).first()
    if not partner:
        raise HTTPException(status_code=404, detail="Kontragent topilmadi")
    
    # Dublikat tekshiruvi - nom bo'yicha (o'zidan boshqa)
    existing_by_name = db.query(Partner).filter(
        Partner.name == name,
        Partner.id != partner_id
    ).first()
    if existing_by_name:
        raise HTTPException(status_code=400, detail=f"'{name}' nomli kontragent allaqachon mavjud!")
    
    # Dublikat tekshiruvi - telefon bo'yicha (agar telefon kiritilgan bo'lsa va o'zidan boshqa)
    if phone and phone.strip():
        existing_by_phone = db.query(Partner).filter(
            Partner.phone == phone,
            Partner.id != partner_id
        ).first()
        if existing_by_phone:
            raise HTTPException(status_code=400, detail=f"'{phone}' telefon raqamli kontragent allaqachon mavjud!")
    
    partner.name = name
    # partner.code o'zgartirilmaydi - avtomatik generatsiya qilingan
    partner.type = type
    partner.phone = phone
    partner.address = address
    partner.credit_limit = credit_limit
    partner.discount_percent = discount_percent
    
    db.commit()
    return RedirectResponse(url="/partners", status_code=303)


@app.post("/partners/delete/{partner_id}")
async def partner_delete(partner_id: int, db: Session = Depends(get_db)):
    """Kontragentni o'chirish"""
    partner = db.query(Partner).filter(Partner.id == partner_id).first()
    if not partner:
        raise HTTPException(status_code=404, detail="Kontragent topilmadi")
    
    # Kontragent bilan bog'liq buyurtmalar borligini tekshirish
    has_orders = db.query(Order).filter(Order.partner_id == partner_id).first()
    has_purchases = db.query(Purchase).filter(Purchase.partner_id == partner_id).first()
    
    if has_orders or has_purchases:
        raise HTTPException(
            status_code=400, 
            detail="Bu kontragent bilan bog'liq buyurtmalar yoki kirimlar mavjud. O'chirish mumkin emas."
        )
    
    db.delete(partner)
    db.commit()
    return RedirectResponse(url="/partners", status_code=303)


# --- PARTNERS EXCEL OPERATIONS ---
@app.get("/partners/export")
async def export_partners(db: Session = Depends(get_db), current_user: User = Depends(require_auth)):
    partners = db.query(Partner).all()
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Partners"
    ws.append(["ID", "Kod", "Nomi", "Turi", "Telefon", "Manzil", "Kredit Limit", "Chegirma %"])
    for p in partners:
        ws.append([p.id, p.code, p.name, p.type, p.phone, p.address, p.credit_limit, p.discount_percent])
    stream = io.BytesIO()
    wb.save(stream)
    stream.seek(0)
    return StreamingResponse(stream, media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", headers={"Content-Disposition": "attachment; filename=kontragentlar.xlsx"})

@app.get("/partners/template")
async def template_partners():
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Template"
    ws.append(["Nomi", "Turi", "Telefon", "Manzil", "Kredit Limit", "Chegirma %"])
    ws.append(["Mijoz MCHJ", "customer", "+998901234567", "Toshkent", 1000000, 0])
    ws.append(["Yetkazib Beruvchi", "supplier", "+998909876543", "Samarqand", 0, 0])
    stream = io.BytesIO()
    wb.save(stream)
    stream.seek(0)
    return StreamingResponse(stream, media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", headers={"Content-Disposition": "attachment; filename=kontragent_andoza.xlsx"})

@app.post("/partners/import")
async def import_partners(file: UploadFile = File(...), db: Session = Depends(get_db), current_user: User = Depends(require_auth)):
    contents = await file.read()
    wb = openpyxl.load_workbook(io.BytesIO(contents))
    ws = wb.active
    rows = list(ws.iter_rows(min_row=2, values_only=True))
    for row in rows:
        if not row[0]: continue
        name, type_, phone, address, credit_limit, discount_percent = row[0:6]
        partner = db.query(Partner).filter(Partner.name == name).first()
        if not partner:
            # Kod generatsiya
            count = db.query(Partner).count()
            code = f"P{count + 1:04d}"
            partner = Partner(
                code=code,
                name=name, 
                type=type_, 
                phone=phone, 
                address=address, 
                credit_limit=credit_limit, 
                discount_percent=discount_percent
            )
            db.add(partner)
        else:
            partner.phone = phone
            partner.address = address
            partner.credit_limit = credit_limit
            partner.discount_percent = discount_percent
        db.commit()
    return RedirectResponse(url="/partners", status_code=303)


# ==========================================
# SAVDO
# ==========================================

@app.get("/sales", response_class=HTMLResponse)
async def sales_list(request: Request, db: Session = Depends(get_db), current_user: User = Depends(require_auth)):
    """Sotuvlar ro'yxati"""
    from urllib.parse import unquote
    orders = db.query(Order).filter(Order.type == "sale").order_by(Order.date.desc()).limit(100).all()
    error = request.query_params.get("error")
    error_detail = unquote(request.query_params.get("detail", "") or "")
    return templates.TemplateResponse("sales/list.html", {
        "request": request,
        "orders": orders,
        "page_title": "Sotuvlar",
        "current_user": current_user,
        "error": error,
        "error_detail": error_detail,
    })


@app.get("/sales/new", response_class=HTMLResponse)
async def sales_new(
    request: Request,
    price_type_id: Optional[int] = None,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth)
):
    """Yangi sotuv — narx turini tanlang, shu bo'yicha mahsulot narxlari ko'rsatiladi. Barcha turlar (tayyor, yarim_tayyor, xom ashyo) ombor qoldig'ida ko'rinadi."""
    products = (
        db.query(Product)
        .options(joinedload(Product.unit))
        .filter(
            Product.type.in_(["tayyor", "yarim_tayyor", "hom_ashyo", "material"]),
            Product.is_active == True,
        )
        .order_by(Product.name)
        .all()
    )
    partners = db.query(Partner).filter(Partner.is_active == True).order_by(Partner.name).all()
    warehouses = db.query(Warehouse).all()
    price_types = db.query(PriceType).filter(PriceType.is_active == True).order_by(PriceType.name).all()
    current_pt_id = price_type_id or (price_types[0].id if price_types else None)
    product_prices_by_type = {}
    if current_pt_id:
        pps = db.query(ProductPrice).filter(ProductPrice.price_type_id == current_pt_id).all()
        product_prices_by_type = {pp.product_id: pp.sale_price for pp in pps}
    # Ombor bo'yicha qoldiq: warehouse_id -> [product_id, ...] (sum(quantity) > 0) va product_id -> miqdor
    warehouse_products = {}
    warehouse_stock_quantities = {}
    for wh in warehouses:
        rows = (
            db.query(Stock.product_id)
            .filter(Stock.warehouse_id == wh.id)
            .group_by(Stock.product_id)
            .having(func.coalesce(func.sum(Stock.quantity), 0) > 0)
            .all()
        )
        warehouse_products[str(wh.id)] = [r[0] for r in rows]
        qty_rows = (
            db.query(Stock.product_id, func.coalesce(func.sum(Stock.quantity), 0).label("total"))
            .filter(Stock.warehouse_id == wh.id)
            .group_by(Stock.product_id)
            .all()
        )
        warehouse_stock_quantities[str(wh.id)] = {str(r[0]): float(r[1] or 0) for r in qty_rows}
    return templates.TemplateResponse("sales/new.html", {
        "request": request,
        "products": products,
        "partners": partners,
        "warehouses": warehouses,
        "price_types": price_types,
        "current_price_type_id": current_pt_id,
        "product_prices_by_type": product_prices_by_type,
        "warehouse_products": warehouse_products,
        "warehouse_stock_quantities": warehouse_stock_quantities,
        "current_user": current_user,
        "page_title": "Yangi sotuv"
    })


@app.post("/sales/create")
async def sales_create(
    request: Request,
    partner_id: int = Form(...),
    warehouse_id: int = Form(...),
    price_type_id: Optional[int] = Form(None),
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth)
):
    """Sotuv yaratish — narx turi saqlanadi; savatdagi mahsulotlar va narxlar (agar yuborilsa) qo'shiladi"""
    form = await request.form()
    product_ids = [int(x) for x in form.getlist("product_id") if str(x).strip().isdigit()]
    quantities_raw = form.getlist("quantity")
    prices_raw = form.getlist("price")
    quantities = []
    for q in quantities_raw:
        try:
            quantities.append(float(q))
        except (ValueError, TypeError):
            pass
    prices = []
    for p in prices_raw:
        try:
            prices.append(float(p))
        except (ValueError, TypeError):
            pass
    last_order = db.query(Order).filter(Order.type == "sale").order_by(Order.id.desc()).first()
    new_number = f"S-{datetime.now().strftime('%Y%m%d')}-{(last_order.id + 1) if last_order else 1:04d}"
    order = Order(
        number=new_number,
        type="sale",
        partner_id=partner_id,
        warehouse_id=warehouse_id,
        price_type_id=price_type_id if price_type_id else None,
        status="draft"
    )
    db.add(order)
    db.commit()
    db.refresh(order)
    for i in range(min(len(product_ids), len(quantities))):
        pid, qty = product_ids[i], float(quantities[i])
        if pid and qty > 0:
            price = prices[i] if i < len(prices) and prices[i] >= 0 else None
            if price is None or price < 0:
                pp = db.query(ProductPrice).filter(ProductPrice.product_id == pid, ProductPrice.price_type_id == order.price_type_id).first()
                if pp:
                    price = pp.sale_price or 0
                else:
                    prod = db.query(Product).filter(Product.id == pid).first()
                    price = (prod.sale_price or prod.purchase_price or 0) if prod else 0
            total_row = qty * price
            item = OrderItem(order_id=order.id, product_id=pid, quantity=qty, price=price, total=total_row)
            db.add(item)
            order.subtotal = (order.subtotal or 0) + total_row
            order.total = (order.total or 0) + total_row
    db.commit()
    return RedirectResponse(url=f"/sales/edit/{order.id}", status_code=303)


@app.get("/sales/edit/{order_id}", response_class=HTMLResponse)
async def sales_edit(
    request: Request,
    order_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth)
):
    """Sotuv tafsiloti — ko'rish va qoralama holatida tahrirlash"""
    from urllib.parse import unquote
    order = db.query(Order).filter(Order.id == order_id, Order.type == "sale").first()
    if not order:
        raise HTTPException(status_code=404, detail="Sotuv topilmadi")
    products = db.query(Product).filter(Product.type.in_(["tayyor", "yarim_tayyor"]), Product.is_active == True).order_by(Product.name).all()
    product_prices_by_type = {}
    if order.price_type_id:
        pps = db.query(ProductPrice).filter(ProductPrice.price_type_id == order.price_type_id).all()
        product_prices_by_type = {pp.product_id: pp.sale_price for pp in pps}
    error = request.query_params.get("error")
    error_detail = unquote(request.query_params.get("detail", "") or "")
    return templates.TemplateResponse("sales/edit.html", {
        "request": request,
        "order": order,
        "products": products,
        "product_prices_by_type": product_prices_by_type,
        "current_user": current_user,
        "page_title": f"Sotuv: {order.number}",
        "error": error,
        "error_detail": error_detail,
    })


@app.post("/sales/{order_id}/add-item")
async def sales_add_item(
    order_id: int,
    product_id: int = Form(...),
    quantity: float = Form(...),
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth)
):
    """Sotuvga mahsulot qo'shish"""
    order = db.query(Order).filter(Order.id == order_id, Order.type == "sale").first()
    if not order:
        raise HTTPException(status_code=404, detail="Sotuv topilmadi")
    if order.status != "draft":
        return RedirectResponse(url=f"/sales/edit/{order_id}", status_code=303)
    price = 0
    pp = db.query(ProductPrice).filter(ProductPrice.product_id == product_id, ProductPrice.price_type_id == order.price_type_id).first()
    if pp:
        price = pp.sale_price or 0
    if not price:
        prod = db.query(Product).filter(Product.id == product_id).first()
        price = (prod.sale_price or prod.purchase_price or 0) if prod else 0
    total_row = quantity * price
    item = OrderItem(order_id=order_id, product_id=product_id, quantity=quantity, price=price, total=total_row)
    db.add(item)
    order.subtotal = (order.subtotal or 0) + total_row
    order.total = (order.total or 0) + total_row
    db.commit()
    return RedirectResponse(url=f"/sales/edit/{order_id}", status_code=303)


@app.post("/sales/{order_id}/add-items")
async def sales_add_items(
    request: Request,
    order_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth)
):
    """Sotuvga savatdagi barcha mahsulotlarni bir harakatda qo'shish"""
    order = db.query(Order).filter(Order.id == order_id, Order.type == "sale").first()
    if not order:
        raise HTTPException(status_code=404, detail="Sotuv topilmadi")
    if order.status != "draft":
        return RedirectResponse(url=f"/sales/edit/{order_id}", status_code=303)
    form = await request.form()
    product_ids = [int(x) for x in form.getlist("product_id") if str(x).strip().isdigit()]
    quantities_raw = form.getlist("quantity")
    quantities = []
    for q in quantities_raw:
        try:
            quantities.append(float(q))
        except (ValueError, TypeError):
            pass
    for i in range(min(len(product_ids), len(quantities))):
        pid, qty = product_ids[i], quantities[i]
        if not pid or qty <= 0:
            continue
        price = 0
        pp = db.query(ProductPrice).filter(ProductPrice.product_id == pid, ProductPrice.price_type_id == order.price_type_id).first()
        if pp:
            price = pp.sale_price or 0
        if not price:
            prod = db.query(Product).filter(Product.id == pid).first()
            price = (prod.sale_price or prod.purchase_price or 0) if prod else 0
        total_row = qty * price
        item = OrderItem(order_id=order_id, product_id=pid, quantity=qty, price=price, total=total_row)
        db.add(item)
        order.subtotal = (order.subtotal or 0) + total_row
        order.total = (order.total or 0) + total_row
    db.commit()
    return RedirectResponse(url=f"/sales/edit/{order_id}", status_code=303)


@app.post("/sales/{order_id}/confirm")
async def sales_confirm(
    order_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth)
):
    """Sotuvni tasdiqlash — ombor qoldig'ini kamaytirish"""
    order = db.query(Order).filter(Order.id == order_id, Order.type == "sale").first()
    if not order:
        raise HTTPException(status_code=404, detail="Sotuv topilmadi")
    if order.status != "draft":
        return RedirectResponse(url=f"/sales/edit/{order_id}", status_code=303)
    for item in order.items:
        stock = db.query(Stock).filter(
            Stock.warehouse_id == order.warehouse_id,
            Stock.product_id == item.product_id
        ).first()
        if not stock or stock.quantity < item.quantity:
            from urllib.parse import quote
            name = item.product.name if item.product else f"#{item.product_id}"
            return RedirectResponse(
                url=f"/sales/edit/{order_id}?error=stock&detail=" + quote(f"Yetarli yo'q: {name}"),
                status_code=303
            )
    for item in order.items:
        stock = db.query(Stock).filter(
            Stock.warehouse_id == order.warehouse_id,
            Stock.product_id == item.product_id
        ).first()
        if stock:
            # StockMovement yozuvini yaratish (chiqim - sotuv)
            create_stock_movement(
                db=db,
                warehouse_id=order.warehouse_id,
                product_id=item.product_id,
                quantity_change=-item.quantity,  # Chiqim
                operation_type="sale",
                document_type="Sale",
                document_id=order.id,
                document_number=order.number,
                user_id=current_user.id if current_user else None,
                note=f"Sotuv: {order.number}"
            )
    order.status = "completed"
    db.commit()
    check_low_stock_and_notify(db)
    return RedirectResponse(url=f"/sales/edit/{order_id}", status_code=303)


@app.post("/sales/{order_id}/delete-item/{item_id}")
async def sales_delete_item(
    order_id: int,
    item_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth)
):
    """Sotuvdan qatorni o'chirish (faqat qoralama)"""
    order = db.query(Order).filter(Order.id == order_id, Order.type == "sale").first()
    if not order or order.status != "draft":
        return RedirectResponse(url=f"/sales/edit/{order_id}", status_code=303)
    item = db.query(OrderItem).filter(OrderItem.id == item_id, OrderItem.order_id == order_id).first()
    if item:
        order.total = (order.total or 0) - (item.total or 0)
        order.subtotal = (order.subtotal or 0) - (item.total or 0)
        db.delete(item)
        db.commit()
    return RedirectResponse(url=f"/sales/edit/{order_id}", status_code=303)


@app.post("/sales/{order_id}/revert")
async def sales_revert(
    order_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_admin),
):
    """Sotuv tasdiqini bekor qilish (faqat admin): ombor qoldig'ini qaytarish, holatni qoralamaga o'tkazish"""
    from urllib.parse import quote
    order = db.query(Order).filter(Order.id == order_id, Order.type == "sale").first()
    if not order:
        raise HTTPException(status_code=404, detail="Sotuv topilmadi")
    if order.status != "completed":
        return RedirectResponse(
            url=f"/sales/edit/{order_id}?error=revert&detail=" + quote("Faqat bajarilgan sotuvning tasdiqini bekor qilish mumkin."),
            status_code=303
        )
    for item in order.items:
        stocks = db.query(Stock).filter(
            Stock.warehouse_id == order.warehouse_id,
            Stock.product_id == item.product_id
        ).all()
        if stocks:
            # Bitta qatorga qaytarish (barcha qatorlar yig'indisiga qo'shamiz)
            stocks[0].quantity = (stocks[0].quantity or 0) + item.quantity
            stocks[0].updated_at = datetime.now()
        else:
            # Stock qatori yo'q bo'lsa (sotuvda tovar chiqarilganda) — yangi qator yaratamiz
            db.add(Stock(
                warehouse_id=order.warehouse_id,
                product_id=item.product_id,
                quantity=float(item.quantity or 0),
            ))
    delete_stock_movements_for_document(db, "Sale", order_id)
    order.status = "draft"
    db.commit()
    return RedirectResponse(url=f"/sales/edit/{order_id}", status_code=303)


@app.post("/sales/delete/{order_id}")
async def sales_delete(
    order_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_admin)
):
    """Sotuvni o'chirish (admin). Qoralama — bekor qilingan qiladi; bekor qilingan — bazadan o'chiradi."""
    from urllib.parse import quote
    order = db.query(Order).filter(Order.id == order_id, Order.type == "sale").first()
    if not order:
        raise HTTPException(status_code=404, detail="Sotuv topilmadi")
    if order.status not in ("draft", "cancelled"):
        return RedirectResponse(
            url="/sales?error=delete&detail=" + quote("Faqat qoralama yoki bekor qilingan sotuvni o'chirish mumkin. Avval tasdiqni bekor qiling."),
            status_code=303
        )
    if order.status == "draft":
        order.status = "cancelled"
        db.commit()
    else:
        db.query(OrderItem).filter(OrderItem.order_id == order_id).delete()
        db.query(Payment).filter(Payment.order_id == order_id).update({Payment.order_id: None})
        db.query(Order).filter(Order.id == order_id).delete()
        db.commit()
    return RedirectResponse(url="/sales", status_code=303)


# ---------- Sotuvchi uchun POS (Sotuv oynasi) ----------
def _get_sales_warehouse(db: Session):
    """Umumiy fallback: code/nomida 'sotuv' yoki birinchi ombor (admin tekshiruvi uchun)."""
    wh = db.query(Warehouse).filter(
        Warehouse.is_active == True,
        or_(Warehouse.code.ilike("%sotuv%"), Warehouse.name.ilike("%sotuv%"))
    ).first()
    if wh:
        return wh
    return db.query(Warehouse).filter(Warehouse.is_active == True).order_by(Warehouse.id).first()


def _get_pos_price_type(db: Session):
    """POS (chakana savdo) uchun narx turi: code='chakana' bo'lgani, yo'q bo'lsa birinchi faol narx turi."""
    pt = (
        db.query(PriceType)
        .filter(PriceType.is_active == True, PriceType.code.ilike("chakana"))
        .order_by(PriceType.id)
        .first()
    )
    if pt:
        return pt
    return db.query(PriceType).filter(PriceType.is_active == True).order_by(PriceType.id).first()


def _get_pos_warehouses_for_user(db: Session, current_user: User):
    """Foydalanuvchiga tegishli omborlar ro'yxati (POS tepada ko'rsatish va tanlash uchun)."""
    if not current_user:
        return []
    role = (current_user.role or "").strip()
    if role == "admin" or role == "manager":
        return db.query(Warehouse).filter(Warehouse.is_active == True).order_by(Warehouse.name).all()
    if role != "sotuvchi":
        return []
    from sqlalchemy.orm import joinedload
    user = db.query(User).options(
        joinedload(User.warehouses_list),
        joinedload(User.departments_list),
    ).filter(User.id == current_user.id).first()
    if not user:
        return []
    seen = set()
    result = []
    for w in (user.warehouses_list or []):
        if w and w.id not in seen and (getattr(w, "is_active", True)):
            seen.add(w.id)
            result.append(w)
    for dept in (user.departments_list or []):
        if not dept:
            continue
        for w in db.query(Warehouse).filter(
            Warehouse.department_id == dept.id,
            Warehouse.is_active == True
        ).order_by(Warehouse.name).all():
            if w.id not in seen:
                seen.add(w.id)
                result.append(w)
    if user.warehouse_id and user.warehouse_id not in seen:
        wh = db.query(Warehouse).filter(Warehouse.id == user.warehouse_id, Warehouse.is_active == True).first()
        if wh:
            seen.add(wh.id)
            result.append(wh)
    if user.department_id and user.department_id not in (getattr(d, "id", None) for d in (user.departments_list or [])):
        for w in db.query(Warehouse).filter(
            Warehouse.department_id == user.department_id,
            Warehouse.is_active == True
        ).order_by(Warehouse.name).all():
            if w.id not in seen:
                seen.add(w.id)
                result.append(w)
    return result


def _get_pos_warehouse_for_user(db: Session, current_user: User):
    """
    Sotuvchi uchun ombor: foydalanuvchining warehouses_list (yoki departments_list) bo'yicha.
    Admin/menejer: _get_sales_warehouse (umumiy).
    """
    if not current_user:
        return None
    role = (current_user.role or "").strip()
    if role == "admin" or role == "manager":
        return _get_sales_warehouse(db)
    if role != "sotuvchi":
        return None
    from sqlalchemy.orm import joinedload
    user = db.query(User).options(
        joinedload(User.warehouses_list),
        joinedload(User.departments_list),
    ).filter(User.id == current_user.id).first()
    if not user:
        return None
    if user.warehouses_list:
        return user.warehouses_list[0]
    if user.departments_list:
        dept = user.departments_list[0]
        wh = db.query(Warehouse).filter(
            Warehouse.department_id == dept.id,
            Warehouse.is_active == True
        ).order_by(Warehouse.id).first()
        if wh:
            return wh
    if user.warehouse_id:
        wh = db.query(Warehouse).filter(
            Warehouse.id == user.warehouse_id,
            Warehouse.is_active == True
        ).first()
        if wh:
            return wh
    if user.department_id:
        wh = db.query(Warehouse).filter(
            Warehouse.department_id == user.department_id,
            Warehouse.is_active == True
        ).order_by(Warehouse.id).first()
        if wh:
            return wh
    return None


def _get_pos_partner(db: Session):
    """POS uchun default mijoz (Chakana xaridor)."""
    p = db.query(Partner).filter(Partner.is_active == True, or_(Partner.code == "chakana", Partner.code == "pos")).first()
    if p:
        return p
    return db.query(Partner).filter(Partner.is_active == True).order_by(Partner.id).first()


@app.get("/sales/pos", response_class=HTMLResponse)
async def sales_pos(
    request: Request,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    """Sotuv oynasi: faqat sotuvchi (yoki admin/menejer tekshiruvi uchun). Tovarlar foydalanuvchi bo'limi/omboridan."""
    _ensure_orders_payment_due_date_column(db)
    from urllib.parse import unquote
    role = (current_user.role or "").strip()
    if role not in ("sotuvchi", "admin", "manager"):
        return RedirectResponse(url="/?error=pos_access", status_code=303)
    pos_user_warehouses = _get_pos_warehouses_for_user(db, current_user)
    sales_warehouse = _get_pos_warehouse_for_user(db, current_user)
    warehouse_id_param = request.query_params.get("warehouse_id")
    if warehouse_id_param and pos_user_warehouses:
        try:
            wid = int(warehouse_id_param)
            chosen = next((w for w in pos_user_warehouses if w.id == wid), None)
            if chosen:
                sales_warehouse = chosen
        except (TypeError, ValueError):
            pass
    from datetime import date as date_type
    today_date = date_type.today()
    pos_today_orders = db.query(Order).filter(
        Order.type == "sale",
        Order.status == "completed",
        func.date(Order.created_at) == today_date,
    ).order_by(Order.created_at.desc()).limit(100).all()
    if not sales_warehouse:
        err = "no_warehouse" if role == "sotuvchi" else "no_warehouse_admin"
        detail_msg = "Sizga ombor yoki bo'lim biriktirilmagan. Administrator bilan bog'laning." if role == "sotuvchi" else (unquote(request.query_params.get("detail", "") or "") or "Sotuv bo'limi ombori topilmadi.")
        return templates.TemplateResponse("sales/pos.html", {
            "request": request,
            "page_title": "Sotuv oynasi",
            "current_user": current_user,
            "warehouse": None,
            "pos_user_warehouses": pos_user_warehouses,
            "pos_today_orders": pos_today_orders,
            "products": [],
            "product_prices": {},
            "stock_by_product": {},
            "pos_categories": [],
            "pos_all_categories": [],
            "success": request.query_params.get("success"),
            "error": err,
            "error_detail": detail_msg,
            "number": request.query_params.get("number", ""),
        })
    product_ids_in_warehouse = [
        r[0] for r in db.query(Stock.product_id).filter(
            Stock.warehouse_id == sales_warehouse.id,
            Stock.quantity > 0
        ).distinct().all()
    ]
    if product_ids_in_warehouse:
        products = db.query(Product).filter(
            Product.id.in_(product_ids_in_warehouse),
            Product.is_active == True
        ).order_by(Product.name).all()
    else:
        products = []
    price_type = _get_pos_price_type(db)
    product_prices = {}
    if price_type:
        pps = db.query(ProductPrice).filter(ProductPrice.price_type_id == price_type.id).all()
        product_prices = {pp.product_id: (pp.sale_price or 0) for pp in pps}
    for p in products:
        if p.id not in product_prices or product_prices[p.id] == 0:
            product_prices[p.id] = float(p.sale_price or p.purchase_price or 0)
    stock_by_product = {}
    if sales_warehouse and product_ids_in_warehouse:
        for row in db.query(Stock.product_id, Stock.quantity).filter(
            Stock.warehouse_id == sales_warehouse.id,
            Stock.product_id.in_(product_ids_in_warehouse)
        ).all():
            stock_by_product[row[0]] = float(row[1] or 0)
    pos_categories = []
    if products:
        cat_ids = list({p.category_id for p in products if p.category_id})
        if cat_ids:
            for c in db.query(Category).filter(Category.id.in_(cat_ids)).order_by(Category.name).all():
                pos_categories.append({"id": c.id, "name": c.name or c.code or ""})
    pos_all_categories = [{"id": c.id, "name": c.name or c.code or ""} for c in db.query(Category).order_by(Category.name).all()]
    from urllib.parse import unquote
    success = request.query_params.get("success")
    error = request.query_params.get("error")
    error_detail = unquote(request.query_params.get("detail", "") or "")
    number = request.query_params.get("number", "")
    pos_partners = db.query(Partner).filter(Partner.is_active == True).order_by(Partner.name).all()
    default_partner = _get_pos_partner(db)
    default_partner_id = default_partner.id if default_partner else None
    return templates.TemplateResponse("sales/pos.html", {
        "request": request,
        "page_title": "Sotuv oynasi",
        "current_user": current_user,
        "warehouse": sales_warehouse,
        "pos_user_warehouses": pos_user_warehouses,
        "pos_today_orders": pos_today_orders,
        "products": products,
        "product_prices": product_prices,
        "stock_by_product": stock_by_product,
        "price_type": price_type,
        "pos_categories": pos_categories,
        "pos_all_categories": pos_all_categories,
        "pos_partners": pos_partners,
        "default_partner_id": default_partner_id,
        "success": success,
        "error": error,
        "error_detail": error_detail,
        "number": number,
    })


@app.get("/sales/pos/daily-orders")
async def sales_pos_daily_orders(
    request: Request,
    date_from: Optional[str] = None,
    date_to: Optional[str] = None,
    order_type: Optional[str] = None,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    """Kunlik / sanadan-sanagacha sotuvlar yoki qaytarishlar ro'yxati (JSON). order_type: sale (sotuvlar) yoki return_sale (qaytarishlar)."""
    from datetime import date as date_type, datetime as dt
    role = (current_user.role or "").strip()
    if role not in ("sotuvchi", "admin", "manager"):
        return []
    today = date_type.today()
    try:
        d_from = dt.strptime(date_from, "%Y-%m-%d").date() if date_from else today
    except (ValueError, TypeError):
        d_from = today
    try:
        d_to = dt.strptime(date_to, "%Y-%m-%d").date() if date_to else today
    except (ValueError, TypeError):
        d_to = today
    if d_from > d_to:
        d_from, d_to = d_to, d_from
    o_type = (order_type or "sale").strip().lower()
    if o_type != "return_sale":
        o_type = "sale"
    orders = db.query(Order).filter(
        Order.type == o_type,
        Order.status == "completed",
        func.date(Order.created_at) >= d_from,
        func.date(Order.created_at) <= d_to,
    ).order_by(Order.created_at.desc()).limit(200).all()
    out = []
    for o in orders:
        out.append({
            "id": o.id,
            "number": o.number or "",
            "type": o.type or "sale",
            "created_at": o.created_at.strftime("%H:%M") if o.created_at else "-",
            "date": o.created_at.strftime("%d.%m.%Y") if o.created_at else "-",
            "partner_name": o.partner.name if o.partner else "-",
            "warehouse_name": o.warehouse.name if o.warehouse else "-",
            "total": float(o.total or 0),
        })
    return out


@app.post("/sales/pos/draft/save")
async def sales_pos_draft_save(
    request: Request,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    """Chekni saqlash — savatdagi tovarlarni vaqtinchalik saqlab qo'yish."""
    role = (current_user.role or "").strip()
    if role not in ("sotuvchi", "admin", "manager"):
        return JSONResponse({"ok": False, "error": "Ruxsat yo'q"}, status_code=403)
    try:
        body = await request.json()
    except Exception:
        return JSONResponse({"ok": False, "error": "JSON xato"}, status_code=400)
    items = body.get("items")
    if not items or not isinstance(items, list):
        return JSONResponse({"ok": False, "error": "Savat bo'sh. Kamida bitta mahsulot qo'shing."}, status_code=400)
    warehouse = _get_pos_warehouse_for_user(db, current_user)
    name = (body.get("name") or "").strip() or None
    import json
    items_json = json.dumps(items, ensure_ascii=False)
    draft = PosDraft(
        user_id=current_user.id,
        warehouse_id=warehouse.id if warehouse else None,
        name=name,
        items_json=items_json,
    )
    db.add(draft)
    db.commit()
    db.refresh(draft)
    return JSONResponse({"ok": True, "id": draft.id, "message": "Chek saqlandi."})


@app.get("/sales/pos/drafts")
async def sales_pos_drafts_list(
    request: Request,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    """Saqlangan cheklar ro'yxati (chekni yuklash uchun)."""
    role = (current_user.role or "").strip()
    if role not in ("sotuvchi", "admin", "manager"):
        return JSONResponse([], status_code=200)
    drafts = (
        db.query(PosDraft)
        .filter(PosDraft.user_id == current_user.id)
        .order_by(PosDraft.created_at.desc())
        .limit(50)
        .all()
    )
    import json
    out = []
    for d in drafts:
        try:
            items = json.loads(d.items_json or "[]")
        except Exception:
            items = []
        total = sum((float(x.get("price") or 0) * float(x.get("quantity") or 0)) for x in items)
        out.append({
            "id": d.id,
            "name": d.name or f"Chek #{d.id}",
            "created_at": d.created_at.strftime("%d.%m.%Y %H:%M") if d.created_at else "-",
            "total": round(total, 2),
            "item_count": len(items),
        })
    return out


@app.get("/sales/pos/draft/{draft_id}")
async def sales_pos_draft_get(
    draft_id: int,
    request: Request,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    """Bitta saqlangan chekni olish (savatga yuklash uchun)."""
    role = (current_user.role or "").strip()
    if role not in ("sotuvchi", "admin", "manager"):
        return JSONResponse({"ok": False, "error": "Ruxsat yo'q"}, status_code=403)
    draft = db.query(PosDraft).filter(PosDraft.id == draft_id, PosDraft.user_id == current_user.id).first()
    if not draft:
        return JSONResponse({"ok": False, "error": "Chek topilmadi"}, status_code=404)
    import json
    try:
        items = json.loads(draft.items_json or "[]")
    except Exception:
        items = []
    return JSONResponse({"ok": True, "items": items})


def _get_pos_cash_register(db: Session, payment_type: str, department_id: Optional[int] = None):
    """POS to'lov: savdo qaysi bo'limdan bo'lsa o'sha bo'lim kassasiga. Avvalo kassaning payment_type maydoni (Kassalar sahifasida tanlangan) bo'yicha, keyin nomida so'z qidiriladi."""
    payment_type = (payment_type or "").strip().lower()
    key = payment_type if payment_type in ("naqd", "plastik", "click", "terminal") else "plastik"
    q = db.query(CashRegister).filter(CashRegister.is_active == True)
    if department_id:
        q = q.filter(CashRegister.department_id == department_id)
    active = q.order_by(CashRegister.id).all()
    if not active and department_id:
        active = db.query(CashRegister).filter(CashRegister.is_active == True).order_by(CashRegister.id).all()
    if not active:
        return None
    # 1) payment_type maydoni mos kassani tanlash (Kassalar sahifasida tanlangan)
    for c in active:
        if getattr(c, "payment_type", None) and (c.payment_type or "").strip().lower() == key:
            return c
    # 2) Eski usul: kassa nomida naqd/plastik/click/terminal bor bo'lsa
    for c in active:
        if c.name and key in (c.name or "").lower():
            return c
    if key in ("click", "terminal"):
        for c in active:
            if getattr(c, "payment_type", None) and (c.payment_type or "").strip().lower() == "plastik":
                return c
        for c in active:
            if c.name and "plastik" in (c.name or "").lower():
                return c
    return active[0]


@app.post("/sales/pos/complete")
async def sales_pos_complete(
    request: Request,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    """POS savatni sotuv qilish. Naqd mijoz → pul kassaga; boshqa kontragent → qarz, to'lov muddati."""
    _ensure_orders_payment_due_date_column(db)
    role = (current_user.role or "").strip()
    if role not in ("sotuvchi", "admin", "manager"):
        return RedirectResponse(url="/?error=pos_access", status_code=303)
    form = await request.form()
    payment_type = (form.get("payment_type") or "").strip().lower()
    if payment_type not in ("naqd", "plastik", "click", "terminal"):
        return RedirectResponse(url="/sales/pos?error=payment", status_code=303)
    warehouse = _get_pos_warehouse_for_user(db, current_user)
    wh_id_form = form.get("warehouse_id")
    if wh_id_form:
        try:
            wh_id = int(wh_id_form)
            allowed = _get_pos_warehouses_for_user(db, current_user)
            chosen = next((w for w in allowed if w.id == wh_id), None)
            if chosen:
                warehouse = chosen
        except (TypeError, ValueError):
            pass
    default_partner = _get_pos_partner(db)
    if not warehouse or not default_partner:
        return RedirectResponse(url="/sales/pos?error=config", status_code=303)
    partner_id_form = form.get("partner_id")
    partner = default_partner
    if partner_id_form and str(partner_id_form).strip().isdigit():
        try:
            pid = int(partner_id_form)
            p = db.query(Partner).filter(Partner.id == pid, Partner.is_active == True).first()
            if p:
                partner = p
        except (ValueError, TypeError):
            pass
    product_ids = [int(x) for x in form.getlist("product_id") if str(x).strip().isdigit()]
    quantities = []
    for q in form.getlist("quantity"):
        try:
            quantities.append(float(q))
        except (ValueError, TypeError):
            pass
    prices = []
    for p in form.getlist("price"):
        try:
            prices.append(float(p))
        except (ValueError, TypeError):
            pass
    if not product_ids or len(quantities) < len(product_ids):
        return RedirectResponse(url="/sales/pos?error=empty", status_code=303)
    price_type = _get_pos_price_type(db)
    last_order = db.query(Order).filter(Order.type == "sale").order_by(Order.id.desc()).first()
    new_number = f"S-{datetime.now().strftime('%Y%m%d')}-{(last_order.id + 1) if last_order else 1:04d}"
    order = Order(
        number=new_number,
        type="sale",
        partner_id=partner.id,
        warehouse_id=warehouse.id,
        price_type_id=price_type.id if price_type else None,
        user_id=current_user.id if current_user else None,
        status="draft",
        payment_type=payment_type,
    )
    db.add(order)
    db.commit()
    db.refresh(order)
    total_order = 0.0
    items_for_stock = []
    for i in range(min(len(product_ids), len(quantities))):
        pid, qty = product_ids[i], float(quantities[i])
        if not pid or qty <= 0:
            continue
        price = prices[i] if i < len(prices) and prices[i] >= 0 else None
        if price is None or price < 0:
            pp = db.query(ProductPrice).filter(ProductPrice.product_id == pid, ProductPrice.price_type_id == order.price_type_id).first()
            if pp:
                price = pp.sale_price or 0
            else:
                prod = db.query(Product).filter(Product.id == pid).first()
                price = (prod.sale_price or prod.purchase_price or 0) if prod else 0
        total_row = qty * price
        db.add(OrderItem(order_id=order.id, product_id=pid, quantity=qty, price=price, total=total_row))
        total_order += total_row
        items_for_stock.append((pid, qty))
    order.subtotal = total_order
    discount_percent = 0.0
    discount_amount = 0.0
    try:
        discount_percent = float(form.get("discount_percent") or 0)
    except (ValueError, TypeError):
        pass
    try:
        discount_amount = float(form.get("discount_amount") or 0)
    except (ValueError, TypeError):
        pass
    discount_sum = (total_order * discount_percent / 100.0) + discount_amount
    if discount_sum > total_order:
        discount_sum = total_order
    order.discount_percent = discount_percent
    order.discount_amount = discount_amount
    order.total = total_order - discount_sum
    is_cash_client = (partner.id == default_partner.id)
    if is_cash_client:
        order.paid = order.total
        order.debt = 0
    else:
        order.paid = 0
        order.debt = order.total
        due_str = (form.get("payment_due_date") or "").strip()
        if due_str:
            try:
                order.payment_due_date = datetime.strptime(due_str, "%Y-%m-%d").date()
            except (ValueError, TypeError):
                order.payment_due_date = (datetime.now() + timedelta(days=7)).date()
        else:
            order.payment_due_date = (datetime.now() + timedelta(days=7)).date()
    for pid, qty in items_for_stock:
        stock = db.query(Stock).filter(
            Stock.warehouse_id == order.warehouse_id,
            Stock.product_id == pid
        ).first()
        if not stock or (stock.quantity or 0) < qty:
            from urllib.parse import quote
            prod = db.query(Product).filter(Product.id == pid).first()
            name = prod.name if prod else f"#{pid}"
            mavjud = float(stock.quantity or 0) if stock else 0
            order.status = "cancelled"
            db.commit()
            detail = f"Yetarli yo'q: {name} (savatda: {qty}, omborda: {mavjud:.0f})"
            url = "/sales/pos?error=stock&detail=" + quote(detail)
            if warehouse and warehouse.id:
                url += "&warehouse_id=" + str(warehouse.id)
            return RedirectResponse(url=url, status_code=303)
    for pid, qty in items_for_stock:
        create_stock_movement(
            db=db,
            warehouse_id=order.warehouse_id,
            product_id=pid,
            quantity_change=-qty,
            operation_type="sale",
            document_type="Sale",
            document_id=order.id,
            document_number=order.number,
            user_id=current_user.id if current_user else None,
            note=f"Sotuv (POS {payment_type}): {order.number}"
        )
    order.status = "completed"
    db.commit()
    if is_cash_client:
        # Naqd mijoz: pul avtomatik kassaga
        department_id = getattr(warehouse, "department_id", None) if warehouse else None
        if not department_id and current_user:
            department_id = getattr(current_user, "department_id", None)
        cash_register = _get_pos_cash_register(db, payment_type, department_id)
        if cash_register and (order.total or 0) > 0:
            today_start = datetime.now().replace(hour=0, minute=0, second=0, microsecond=0)
            pay_count = db.query(Payment).filter(Payment.created_at >= today_start).count()
            pay_number = f"PAY-{datetime.now().strftime('%Y%m%d')}-{pay_count + 1:04d}"
            pay_type = "cash" if payment_type == "naqd" else ("click" if payment_type == "click" else ("terminal" if payment_type == "terminal" else "card"))
            db.add(Payment(
                number=pay_number,
                type="income",
                cash_register_id=cash_register.id,
                partner_id=order.partner_id,
                order_id=order.id,
                amount=order.total,
                payment_type=pay_type,
                category="sale",
                description=f"POS sotuv {order.number}",
                user_id=current_user.id if current_user else None,
            ))
            cash_register.balance = (cash_register.balance or 0) + (order.total or 0)
            db.commit()
    else:
        # Boshqa kontragent: qarz, hisobiga yoziladi
        partner.balance = (partner.balance or 0) + (order.total or 0)
        db.commit()
    check_low_stock_and_notify(db)
    return RedirectResponse(url="/sales/pos?success=1&number=" + order.number, status_code=303)


@app.get("/sales/pos/receipt", response_class=HTMLResponse)
async def sales_pos_receipt(
    request: Request,
    number: str = "",
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    """POS sotuv cheki — Xprinter (yoki boshqa printer) uchun chop etish sahifasi."""
    if not number or not number.strip():
        return HTMLResponse("<html><body><p>Hujjat raqami ko'rsatilmagan.</p></body></html>", status_code=400)
    order = (
        db.query(Order)
        .options(
            joinedload(Order.items).joinedload(OrderItem.product),
            joinedload(Order.partner),
            joinedload(Order.user),
        )
        .filter(Order.number == number.strip(), Order.type == "sale")
        .first()
    )
    if not order:
        return HTMLResponse("<html><body><p>Hujjat topilmadi.</p></body></html>", status_code=404)
    receipt_barcode_b64 = None
    try:
        writer = ImageWriter()
        writer.set_options({
            "module_width": 0.35,
            "module_height": 14,
            "font_size": 10,
            "dpi": 600,
        })
        buf = io.BytesIO()
        code128 = barcode.get("code128", order.number, writer=writer)
        code128.write(buf)
        buf.seek(0)
        receipt_barcode_b64 = base64.b64encode(buf.getvalue()).decode("ascii")
    except Exception:
        pass
    return templates.TemplateResponse("sales/pos_receipt.html", {
        "request": request,
        "order": order,
        "receipt_barcode_b64": receipt_barcode_b64,
    })


@app.get("/sales/returns", response_class=HTMLResponse)
async def sales_returns_list(request: Request, db: Session = Depends(get_db), current_user: User = Depends(require_auth)):
    """Savdodan qaytarish — bajarilgan sotuvlar ro'yxati (qaysi sotuvdan qaytarishni tanlash)."""
    from urllib.parse import unquote
    orders = db.query(Order).filter(
        Order.type == "sale",
        Order.status == "completed"
    ).options(
        joinedload(Order.partner),
        joinedload(Order.warehouse),
    ).order_by(Order.date.desc()).limit(200).all()
    success = request.query_params.get("success")
    number = request.query_params.get("number", "")
    warehouse_name = unquote(request.query_params.get("warehouse", "") or "")
    error = request.query_params.get("error")
    error_detail = unquote(request.query_params.get("detail", "") or "")
    return_docs = db.query(Order).filter(
        Order.type == "return_sale"
    ).options(
        joinedload(Order.partner),
        joinedload(Order.warehouse),
    ).order_by(Order.created_at.desc()).limit(100).all()
    return templates.TemplateResponse("sales/returns_list.html", {
        "request": request,
        "orders": orders,
        "return_docs": return_docs,
        "page_title": "Savdodan qaytarish",
        "current_user": current_user,
        "success": success,
        "number": number,
        "warehouse_name": warehouse_name,
        "error": error,
        "error_detail": error_detail,
    })


@app.get("/sales/return/{order_id}", response_class=HTMLResponse)
async def sales_return_form(
    request: Request,
    order_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    """Savdodan qaytarish — tanlangan sotuv bo'yicha qaytarish miqdorlarini kiritish."""
    from urllib.parse import unquote
    order = db.query(Order).filter(
        Order.id == order_id,
        Order.type == "sale",
        Order.status == "completed"
    ).first()
    if not order:
        raise HTTPException(status_code=404, detail="Sotuv topilmadi yoki bajarilmagan.")
    error = request.query_params.get("error")
    error_detail = unquote(request.query_params.get("detail", "") or "")
    return templates.TemplateResponse("sales/return_form.html", {
        "request": request,
        "order": order,
        "page_title": "Savdodan qaytarish",
        "current_user": current_user,
        "error": error,
        "error_detail": error_detail,
    })


@app.post("/sales/return/create")
async def sales_return_create(
    request: Request,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    """Savdodan qaytarishni rasmiylashtirish — qaytarish orderi va omborga qoldiq qaytarish."""
    from urllib.parse import quote
    form = await request.form()
    order_id_raw = form.get("order_id")
    if not order_id_raw or not str(order_id_raw).strip().isdigit():
        return RedirectResponse(url="/sales/returns?error=empty&detail=" + quote("Sotuv tanlanmadi."), status_code=303)
    order_id = int(order_id_raw)
    sale = db.query(Order).filter(
        Order.id == order_id,
        Order.type == "sale",
        Order.status == "completed"
    ).options(joinedload(Order.items)).first()
    if not sale:
        return RedirectResponse(url="/sales/returns?error=not_found&detail=" + quote("Sotuv topilmadi."), status_code=303)
    product_ids = [int(x) for x in form.getlist("product_id") if str(x).strip().isdigit()]
    quantities_raw = form.getlist("quantity_return")
    quantities = []
    for q in quantities_raw:
        try:
            quantities.append(float(q))
        except (ValueError, TypeError):
            quantities.append(0)
    if not product_ids or all(q <= 0 for q in quantities[:len(product_ids)]):
        return RedirectResponse(
            url="/sales/return/" + str(order_id) + "?error=empty&detail=" + quote("Kamida bitta mahsulot uchun qaytarish miqdorini kiriting."),
            status_code=303
        )
    sale_items_by_product = {item.product_id: item for item in sale.items}
    for i in range(min(len(product_ids), len(quantities))):
        pid, qty = product_ids[i], quantities[i]
        if qty <= 0:
            continue
        item = sale_items_by_product.get(pid)
        if not item:
            prod = db.query(Product).filter(Product.id == pid).first()
            name = prod.name if prod else "#" + str(pid)
            return RedirectResponse(
                url="/sales/return/" + str(order_id) + "?error=qty&detail=" + quote(f"'{name}' ushbu sotuvda yo'q."),
                status_code=303
            )
        sold_qty = item.quantity or 0
        if qty > sold_qty + 1e-6:
            name = (item.product.name if item.product else "") or ("#" + str(pid))
            return RedirectResponse(
                url="/sales/return/" + str(order_id) + "?error=qty&detail=" + quote(f"'{name}' uchun qaytarish miqdori sotilgan miqdordan oshmasin (sotilgan: {sold_qty:.3f}, kiritilgan: {qty:.3f})."),
                status_code=303
            )
    from datetime import date as date_type
    today_start = date_type.today()
    # POS sotuvda ombor bo‘lmasa, qaytarishni birinchi omborga yozamiz (qoldiq haqiqatan qaytsin)
    return_warehouse_id = sale.warehouse_id
    if not return_warehouse_id:
        default_wh = db.query(Warehouse).order_by(Warehouse.id).first()
        return_warehouse_id = default_wh.id if default_wh else None
    if not return_warehouse_id:
        return RedirectResponse(
            url="/sales/returns?error=no_warehouse&detail=" + quote("Ombor topilmadi. Avval ombor yarating."),
            status_code=303
        )
    count = db.query(Order).filter(
        Order.type == "return_sale",
        func.date(Order.created_at) == today_start
    ).count()
    new_number = f"R-{datetime.now().strftime('%Y%m%d')}-{count + 1:04d}"
    return_order = Order(
        number=new_number,
        type="return_sale",
        partner_id=sale.partner_id,
        warehouse_id=return_warehouse_id,
        price_type_id=sale.price_type_id,
        user_id=current_user.id if current_user else None,
        status="completed",
        payment_type=sale.payment_type,
        note=f"Savdodan qaytarish: {sale.number}",
    )
    db.add(return_order)
    db.commit()
    db.refresh(return_order)
    total_return = 0.0
    for i in range(min(len(product_ids), len(quantities))):
        pid, qty = product_ids[i], quantities[i]
        if not pid or qty <= 0:
            continue
        item = sale_items_by_product.get(pid)
        if not item:
            continue
        price = item.price or 0
        total_row = qty * price
        db.add(OrderItem(order_id=return_order.id, product_id=pid, quantity=qty, price=price, total=total_row))
        total_return += total_row
        create_stock_movement(
            db=db,
            warehouse_id=return_warehouse_id,
            product_id=pid,
            quantity_change=+qty,
            operation_type="return_sale",
            document_type="SaleReturn",
            document_id=return_order.id,
            document_number=return_order.number,
            user_id=current_user.id if current_user else None,
            note=f"Savdodan qaytarish: {sale.number} -> {return_order.number}",
        )
    return_order.subtotal = total_return
    return_order.total = total_return
    return_order.paid = total_return
    return_order.debt = 0
    db.commit()
    wh_name = ""
    if return_warehouse_id:
        wh = db.query(Warehouse).filter(Warehouse.id == return_warehouse_id).first()
        wh_name = (wh.name or "").strip()
    params = "success=1&number=" + quote(return_order.number)
    if wh_name:
        params += "&warehouse=" + quote(wh_name)
    return RedirectResponse(url="/sales/returns?" + params, status_code=303)


@app.get("/sales/return/document/{number}", response_class=HTMLResponse)
async def sales_return_document(
    request: Request,
    number: str,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    """Qaytarish hujjati (R-...) — ko'rish / chop etish."""
    doc = (
        db.query(Order)
        .options(
            joinedload(Order.items).joinedload(OrderItem.product),
            joinedload(Order.partner),
            joinedload(Order.warehouse),
            joinedload(Order.user),
        )
        .filter(Order.number == number.strip(), Order.type == "return_sale")
        .first()
    )
    if not doc:
        raise HTTPException(status_code=404, detail="Qaytarish hujjati topilmadi.")
    return templates.TemplateResponse("sales/return_document.html", {
        "request": request,
        "doc": doc,
        "page_title": "Qaytarish " + doc.number,
        "current_user": current_user,
    })


@app.post("/sales/return/revert/{return_order_id}")
async def sales_return_revert(
    return_order_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_admin),
):
    """Qaytarish tasdiqini bekor qilish (faqat admin): omborga qo'shilgan qoldiqni olib tashlash, holatni bekor qilingan qilish."""
    from urllib.parse import quote
    doc = (
        db.query(Order)
        .options(joinedload(Order.items))
        .filter(Order.id == return_order_id, Order.type == "return_sale")
        .first()
    )
    if not doc:
        raise HTTPException(status_code=404, detail="Qaytarish hujjati topilmadi.")
    if doc.status != "completed":
        return RedirectResponse(
            url="/sales/returns?error=revert&detail=" + quote("Faqat tasdiqlangan qaytarishning tasdiqini bekor qilish mumkin."),
            status_code=303
        )
    wh_id = doc.warehouse_id
    if not wh_id:
        return RedirectResponse(
            url="/sales/returns?error=revert&detail=" + quote("Hujjatda ombor ko'rsatilmagan."),
            status_code=303
        )
    for item in doc.items:
        create_stock_movement(
            db=db,
            warehouse_id=wh_id,
            product_id=item.product_id,
            quantity_change=-(item.quantity or 0),
            operation_type="return_sale_revert",
            document_type="SaleReturnRevert",
            document_id=doc.id,
            document_number=doc.number,
            user_id=current_user.id if current_user else None,
            note=f"Qaytarish tasdiqini bekor: {doc.number}",
        )
    doc.status = "cancelled"
    db.commit()
    return RedirectResponse(url="/sales/return/document/" + doc.number + "?reverted=1", status_code=303)


@app.post("/sales/return/delete/{return_order_id}")
async def sales_return_delete(
    return_order_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_admin),
):
    """Qaytarish hujjatini o'chirish (faqat admin). Faqat tasdiqni bekor qilingan hujjatni o'chirish mumkin."""
    from urllib.parse import quote
    doc = db.query(Order).filter(Order.id == return_order_id, Order.type == "return_sale").first()
    if not doc:
        raise HTTPException(status_code=404, detail="Qaytarish hujjati topilmadi.")
    if doc.status != "cancelled":
        return RedirectResponse(
            url="/sales/returns?error=delete&detail=" + quote("Faqat tasdiqni bekor qilgandan keyin o'chirish mumkin. Avval tasdiqni bekor qiling."),
            status_code=303
        )
    number = doc.number
    for item in list(doc.items):
        db.delete(item)
    db.delete(doc)
    db.commit()
    return RedirectResponse(url="/sales/returns?deleted=1&number=" + quote(number), status_code=303)


@app.get("/sales/return/edit/{return_order_id}", response_class=HTMLResponse)
async def sales_return_edit_form(
    request: Request,
    return_order_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    """Qaytarish hujjatini tahrirlash (faqat tasdiqni bekor qilingan hujjat)."""
    doc = (
        db.query(Order)
        .options(
            joinedload(Order.items).joinedload(OrderItem.product),
            joinedload(Order.partner),
            joinedload(Order.warehouse),
        )
        .filter(Order.id == return_order_id, Order.type == "return_sale")
        .first()
    )
    if not doc:
        raise HTTPException(status_code=404, detail="Qaytarish hujjati topilmadi.")
    if doc.status != "cancelled":
        from urllib.parse import quote
        return RedirectResponse(
            url="/sales/return/document/" + doc.number + "?error=edit&detail=" + quote("Faqat tasdiqni bekor qilingan hujjatni tahrirlash mumkin."),
            status_code=303
        )
    return templates.TemplateResponse("sales/return_edit.html", {
        "request": request,
        "doc": doc,
        "page_title": "Qaytarishni tahrirlash " + doc.number,
        "current_user": current_user,
    })


@app.post("/sales/return/update")
async def sales_return_update(
    request: Request,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    """Qaytarish hujjati qatorlarini yangilash (miqdor/narx) — faqat bekor qilingan hujjat."""
    from urllib.parse import quote
    form = await request.form()
    order_id_raw = form.get("order_id")
    if not order_id_raw or not str(order_id_raw).strip().isdigit():
        return RedirectResponse(url="/sales/returns?error=update", status_code=303)
    return_order_id = int(order_id_raw)
    doc = (
        db.query(Order)
        .options(joinedload(Order.items))
        .filter(Order.id == return_order_id, Order.type == "return_sale")
        .first()
    )
    if not doc or doc.status != "cancelled":
        return RedirectResponse(url="/sales/returns?error=update&detail=" + quote("Hujjat topilmadi yoki tahrirlash mumkin emas."), status_code=303)
    product_ids = [int(x) for x in form.getlist("product_id") if str(x).strip().isdigit()]
    quantities = []
    for q in form.getlist("quantity"):
        try:
            quantities.append(float(q))
        except (ValueError, TypeError):
            quantities.append(0)
    prices = []
    for p in form.getlist("price"):
        try:
            prices.append(float(p))
        except (ValueError, TypeError):
            prices.append(0)
    items_by_pid = {item.product_id: item for item in doc.items}
    total_return = 0.0
    for i in range(min(len(product_ids), len(quantities))):
        pid, qty = product_ids[i], quantities[i]
        if not pid or qty < 0:
            continue
        item = items_by_pid.get(pid)
        if not item:
            continue
        price = prices[i] if i < len(prices) and prices[i] >= 0 else (item.price or 0)
        item.quantity = qty
        item.price = price
        item.total = qty * price
        total_return += item.total
    doc.subtotal = total_return
    doc.total = total_return
    doc.paid = total_return
    doc.debt = 0
    db.commit()
    return RedirectResponse(url="/sales/return/document/" + doc.number + "?updated=1", status_code=303)


@app.post("/sales/return/confirm/{return_order_id}")
async def sales_return_confirm(
    return_order_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_admin),
):
    """Qaytarishni qayta tasdiqlash (faqat bekor qilingan hujjat): omborga qoldiq qo'shish."""
    from urllib.parse import quote
    doc = (
        db.query(Order)
        .options(joinedload(Order.items))
        .filter(Order.id == return_order_id, Order.type == "return_sale")
        .first()
    )
    if not doc:
        raise HTTPException(status_code=404, detail="Qaytarish hujjati topilmadi.")
    if doc.status != "cancelled":
        return RedirectResponse(
            url="/sales/returns?error=confirm&detail=" + quote("Faqat bekor qilingan hujjatni qayta tasdiqlash mumkin."),
            status_code=303
        )
    wh_id = doc.warehouse_id
    if not wh_id:
        return RedirectResponse(url="/sales/returns?error=confirm&detail=" + quote("Hujjatda ombor ko'rsatilmagan."), status_code=303)
    for item in doc.items:
        create_stock_movement(
            db=db,
            warehouse_id=wh_id,
            product_id=item.product_id,
            quantity_change=+(item.quantity or 0),
            operation_type="return_sale",
            document_type="SaleReturn",
            document_id=doc.id,
            document_number=doc.number,
            user_id=current_user.id if current_user else None,
            note=f"Qaytarish qayta tasdiqlandi: {doc.number}",
        )
    doc.status = "completed"
    db.commit()
    return RedirectResponse(url="/sales/return/document/" + doc.number + "?confirmed=1", status_code=303)


# ==========================================
# MOLIYA
# ==========================================

def _ensure_payments_status_column(db: Session) -> None:
    """Agar payments jadvalida status ustuni bo'lmasa, qo'shadi (xato bo'lmasin)."""
    try:
        db.execute(text("ALTER TABLE payments ADD COLUMN status VARCHAR(20) DEFAULT 'confirmed'"))
        db.commit()
    except OperationalError as e:
        db.rollback()
        if "duplicate column" not in str(e).lower():
            raise
    except Exception:
        db.rollback()


def _ensure_orders_payment_due_date_column(db: Session) -> None:
    """Agar orders jadvalida payment_due_date ustuni bo'lmasa, qo'shadi."""
    try:
        db.execute(text("ALTER TABLE orders ADD COLUMN payment_due_date DATE"))
        db.commit()
    except OperationalError as e:
        db.rollback()
        if "duplicate column" not in str(e).lower():
            raise
    except Exception:
        db.rollback()


@app.get("/finance", response_class=HTMLResponse)
async def finance(
    request: Request,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
    date_from: Optional[str] = None,
    date_to: Optional[str] = None,
):
    """Moliya - kassa. So'nggi to'lovlar sana bo'yicha filtrlanishi mumkin."""
    _ensure_payments_status_column(db)
    cash_registers = db.query(CashRegister).all()
    partners = db.query(Partner).filter(Partner.is_active == True).order_by(Partner.name).all()
    q = (
        db.query(Payment)
        .options(joinedload(Payment.cash_register), joinedload(Payment.partner))
        .order_by(Payment.date.desc())
    )
    if (date_from or "").strip():
        try:
            df = datetime.strptime(str(date_from).strip()[:10], "%Y-%m-%d").date()
            q = q.filter(Payment.date >= df)
        except ValueError:
            pass
    if (date_to or "").strip():
        try:
            dt = datetime.strptime(str(date_to).strip()[:10], "%Y-%m-%d").date()
            q = q.filter(Payment.date < datetime.combine(dt + timedelta(days=1), datetime.min.time()))
        except ValueError:
            pass
    payments = q.limit(200).all()
    filter_date_from = str(date_from or "").strip()[:10] if date_from else ""
    filter_date_to = str(date_to or "").strip()[:10] if date_to else ""
    # Bugungi statistika — faqat tasdiqlangan to'lovlar
    today = datetime.now().date()
    try:
        _status_ok = or_(Payment.status == "confirmed", Payment.status == None)
        today_income = db.query(Payment).filter(
            Payment.type == "income",
            Payment.date >= today,
            _status_ok
        ).all()
        today_expense = db.query(Payment).filter(
            Payment.type == "expense",
            Payment.date >= today,
            _status_ok
        ).all()
    except OperationalError:
        today_income = db.query(Payment).filter(Payment.type == "income", Payment.date >= today).all()
        today_expense = db.query(Payment).filter(Payment.type == "expense", Payment.date >= today).all()
    stats = {
        "today_income": sum(p.amount for p in today_income),
        "today_expense": sum(p.amount for p in today_expense),
    }
    return templates.TemplateResponse("finance/index.html", {
        "request": request,
        "cash_registers": cash_registers,
        "partners": partners,
        "payments": payments,
        "stats": stats,
        "filter_date_from": filter_date_from,
        "filter_date_to": filter_date_to,
        "current_user": current_user,
        "page_title": "Moliya"
    })


@app.get("/finance/kassa/{cash_register_id}", response_class=HTMLResponse)
async def finance_kassa_detail(
    cash_register_id: int,
    request: Request,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
    date_from: Optional[str] = None,
    date_to: Optional[str] = None,
):
    """Kassaning kirim/chiqimlari — qayerdan keldi, qayerga ketdi (hujjatlar ro'yxati)."""
    cash = db.query(CashRegister).filter(CashRegister.id == cash_register_id).first()
    if not cash:
        raise HTTPException(status_code=404, detail="Kassa topilmadi")
    q = (
        db.query(Payment)
        .options(joinedload(Payment.partner))
        .filter(Payment.cash_register_id == cash_register_id)
        .order_by(Payment.date.desc())
    )
    if (date_from or "").strip():
        try:
            df = datetime.strptime(str(date_from).strip()[:10], "%Y-%m-%d").date()
            q = q.filter(Payment.date >= df)
        except ValueError:
            pass
    if (date_to or "").strip():
        try:
            dt = datetime.strptime(str(date_to).strip()[:10], "%Y-%m-%d").date()
            q = q.filter(Payment.date < datetime.combine(dt + timedelta(days=1), datetime.min.time()))
        except ValueError:
            pass
    payments = q.limit(500).all()
    filter_date_from = str(date_from or "").strip()[:10] if date_from else ""
    filter_date_to = str(date_to or "").strip()[:10] if date_to else ""
    total_income = sum(p.amount or 0 for p in payments if getattr(p, "type", None) == "income")
    total_expense = sum(p.amount or 0 for p in payments if getattr(p, "type", None) == "expense")
    return templates.TemplateResponse("finance/kassa_detail.html", {
        "request": request,
        "cash": cash,
        "payments": payments,
        "filter_date_from": filter_date_from,
        "filter_date_to": filter_date_to,
        "total_income": total_income,
        "total_expense": total_expense,
        "current_user": current_user,
        "page_title": (cash.name or "Kassa") + " — kirim/chiqimlar",
    })


@app.post("/finance/payment")
async def finance_payment_post(
    request: Request,
    type: str = Form(...),
    amount: float = Form(...),
    cash_register_id: int = Form(...),
    partner_id: Optional[int] = Form(None),
    description: str = Form(""),
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    """Kassa kirim/chiqim kiritish (Moliya sahifasidagi «To'lov kiritish» formasi). Kontragent — kimga/kimdan (hisob-kitob uchun)."""
    _ensure_payments_status_column(db)
    if type not in ("income", "expense"):
        return RedirectResponse(url="/finance?error=type", status_code=303)
    cash = db.query(CashRegister).filter(CashRegister.id == cash_register_id).first()
    if not cash:
        return RedirectResponse(url="/finance?error=cash", status_code=303)
    amount = float(amount)
    if amount <= 0:
        return RedirectResponse(url="/finance?error=amount", status_code=303)
    pid = None
    if partner_id is not None and int(partner_id) > 0:
        p = db.query(Partner).filter(Partner.id == int(partner_id)).first()
        if p:
            pid = p.id
    today_start = datetime.now().replace(hour=0, minute=0, second=0, microsecond=0)
    pay_count = db.query(Payment).filter(Payment.created_at >= today_start).count()
    pay_number = f"PAY-{datetime.now().strftime('%Y%m%d')}-{pay_count + 1:04d}"
    desc = (description or "").strip() or ("Kirim" if type == "income" else "Chiqim")
    db.add(Payment(
        number=pay_number,
        type=type,
        cash_register_id=cash_register_id,
        partner_id=pid,
        order_id=None,
        amount=amount,
        payment_type="cash",
        category="other",
        description=desc,
        user_id=current_user.id if current_user else None,
        status="confirmed",
    ))
    cash.balance = (cash.balance or 0) + (amount if type == "income" else -amount)
    db.commit()
    return RedirectResponse(url="/finance?success=1", status_code=303)


def _payment_apply_balance(db: Session, payment: Payment, sign: int):
    """payment ga muvofiq kassa balansini o'zgartirish. sign=1 kirim, sign=-1 chiqim."""
    cash = db.query(CashRegister).filter(CashRegister.id == payment.cash_register_id).first()
    if cash and payment.amount:
        delta = (payment.amount or 0) * sign * (1 if payment.type == "income" else -1)
        cash.balance = (cash.balance or 0) + delta


@app.post("/finance/payment/{payment_id}/confirm")
async def finance_payment_confirm(
    payment_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_admin),
):
    """To'lovni tasdiqlash — bekor qilingan bo'lsa kassaga qayta hisoblaydi."""
    payment = db.query(Payment).filter(Payment.id == payment_id).first()
    if not payment:
        raise HTTPException(status_code=404, detail="To'lov topilmadi")
    status = getattr(payment, "status", "confirmed")
    if status == "confirmed":
        return RedirectResponse(url="/finance?msg=already_confirmed", status_code=303)
    payment.status = "confirmed"
    _payment_apply_balance(db, payment, 1)
    db.commit()
    return RedirectResponse(url="/finance?success=confirmed", status_code=303)


@app.post("/finance/payment/{payment_id}/cancel")
async def finance_payment_cancel(
    payment_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_admin),
):
    """Tasdiqni bekor qilish — kassa balansidan chiqaradi, hujjat bekor qilingan bo'ladi."""
    payment = db.query(Payment).filter(Payment.id == payment_id).first()
    if not payment:
        raise HTTPException(status_code=404, detail="To'lov topilmadi")
    status = getattr(payment, "status", "confirmed")
    if status == "cancelled":
        return RedirectResponse(url="/finance?msg=already_cancelled", status_code=303)
    payment.status = "cancelled"
    _payment_apply_balance(db, payment, -1)
    db.commit()
    return RedirectResponse(url="/finance?success=cancelled", status_code=303)


@app.get("/finance/payment/{payment_id}/edit", response_class=HTMLResponse)
async def finance_payment_edit_page(
    payment_id: int,
    request: Request,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_admin),
):
    """To'lovni tahrirlash sahifasi (admin). Faqat tasdiq bekor qilingan to'lovlar."""
    payment = db.query(Payment).filter(Payment.id == payment_id).first()
    if not payment:
        raise HTTPException(status_code=404, detail="To'lov topilmadi")
    if getattr(payment, "status", "confirmed") == "confirmed":
        return RedirectResponse(
            url="/finance?error=" + quote("Tasdiqlangan to'lovni tahrirlash mumkin emas. Avval tasdiqni bekor qiling."),
            status_code=303,
        )
    partners = db.query(Partner).filter(Partner.is_active == True).order_by(Partner.name).all()
    cash_registers = db.query(CashRegister).filter(CashRegister.is_active == True).all()
    return templates.TemplateResponse("finance/payment_edit.html", {
        "request": request,
        "payment": payment,
        "partners": partners,
        "cash_registers": cash_registers,
        "current_user": current_user,
        "page_title": "To'lovni tahrirlash",
    })


@app.post("/finance/payment/{payment_id}/edit")
async def finance_payment_edit_post(
    payment_id: int,
    type: str = Form(...),
    amount: float = Form(...),
    cash_register_id: int = Form(...),
    partner_id: Optional[int] = Form(None),
    description: str = Form(""),
    db: Session = Depends(get_db),
    current_user: User = Depends(require_admin),
):
    """To'lovni saqlash (admin). Faqat tasdiq bekor qilingan to'lovlar tahrirlanadi."""
    payment = db.query(Payment).filter(Payment.id == payment_id).first()
    if not payment:
        raise HTTPException(status_code=404, detail="To'lov topilmadi")
    if getattr(payment, "status", "confirmed") == "confirmed":
        return RedirectResponse(
            url="/finance?error=" + quote("Tasdiqlangan to'lovni tahrirlash mumkin emas. Avval tasdiqni bekor qiling."),
            status_code=303,
        )
    if type not in ("income", "expense"):
        return RedirectResponse(url=f"/finance/payment/{payment_id}/edit?error=type", status_code=303)
    amount = float(amount)
    if amount <= 0:
        return RedirectResponse(url=f"/finance/payment/{payment_id}/edit?error=amount", status_code=303)
    cash_new = db.query(CashRegister).filter(CashRegister.id == cash_register_id).first()
    if not cash_new:
        return RedirectResponse(url=f"/finance/payment/{payment_id}/edit?error=cash", status_code=303)
    pid = None
    if partner_id is not None and int(partner_id) > 0:
        p = db.query(Partner).filter(Partner.id == int(partner_id)).first()
        if p:
            pid = p.id
    # Yangi qiymatlar (faqat tasdiq bekor qilingan to'lov — balansga ta'sir qilmagan)
    payment.type = type
    payment.amount = amount
    payment.cash_register_id = cash_register_id
    payment.partner_id = pid
    payment.description = (description or "").strip() or ("Kirim" if type == "income" else "Chiqim")
    db.commit()
    return RedirectResponse(url="/finance?success=edited", status_code=303)


@app.post("/finance/payment/{payment_id}/delete")
async def finance_payment_delete(
    payment_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_admin),
):
    """To'lovni o'chirish (admin). Faqat tasdiq bekor qilingan to'lovlar o'chiriladi."""
    payment = db.query(Payment).filter(Payment.id == payment_id).first()
    if not payment:
        raise HTTPException(status_code=404, detail="To'lov topilmadi")
    if getattr(payment, "status", "confirmed") == "confirmed":
        return RedirectResponse(
            url="/finance?error=" + quote("Tasdiqlangan to'lovni o'chirish mumkin emas. Avval tasdiqni bekor qiling."),
            status_code=303,
        )
    db.delete(payment)
    db.commit()
    return RedirectResponse(url="/finance?success=deleted", status_code=303)


# ==========================================
# XODIMLAR
# ==========================================

@app.get("/employees", response_class=HTMLResponse)
async def employees_list(
    request: Request,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
    show_dismissed: bool = False,
):
    """Xodimlar ro'yxati — odatiy holda faqat faol xodimlar."""
    q = db.query(Employee).order_by(Employee.full_name)
    if not show_dismissed:
        q = q.filter(Employee.is_active == True)
    employees = q.all()
    piecework_tasks = db.query(PieceworkTask).filter(PieceworkTask.is_active == True).order_by(PieceworkTask.name).all()
    departments = db.query(Department).filter(Department.is_active == True).order_by(Department.name).all()
    positions = db.query(Position).filter(Position.is_active == True).order_by(Position.name).all()
    return templates.TemplateResponse("employees/list.html", {
        "request": request,
        "employees": employees,
        "piecework_tasks": piecework_tasks,
        "departments": departments,
        "positions": positions,
        "current_user": current_user,
        "page_title": "Xodimlar",
        "show_dismissed": show_dismissed,
    })


@app.post("/employees/add")
async def employee_add(
    request: Request,
    full_name: str = Form(...),
    code: str = Form(...),
    position: str = Form(""),
    department: str = Form(""),
    phone: str = Form(""),
    salary: float = Form(0),
    salary_type: str = Form(""),
    piecework_task_ids: List[int] = Form([]),
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    """Xodim qo'shish"""
    st = (salary_type or "").strip() or None
    if st and st not in ("oylik", "soatlik", "bo'lak", "bo'lak_oylik"):
        st = None
    task_ids = [int(x) for x in (piecework_task_ids or []) if str(x).strip().isdigit()]
    task_ids = list(dict.fromkeys(task_ids))
    employee = Employee(
        full_name=full_name,
        code=code,
        position=position,
        department=department,
        phone=phone,
        salary=salary,
        salary_type=st,
        piecework_task_id=task_ids[0] if task_ids else None,  # legacy
    )
    db.add(employee)
    db.flush()
    if st in ("bo'lak", "bo'lak_oylik") and task_ids:
        tasks = db.query(PieceworkTask).filter(PieceworkTask.id.in_(task_ids)).all()
        employee.piecework_tasks = tasks
    db.commit()
    return RedirectResponse(url="/employees", status_code=303)


@app.get("/employees/edit/{employee_id}", response_class=HTMLResponse)
async def employee_edit_page(
    request: Request,
    employee_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    """Xodim tahrirlash sahifasi"""
    emp = db.query(Employee).filter(Employee.id == employee_id).first()
    if not emp:
        return RedirectResponse(url="/employees?error=Xodim topilmadi", status_code=303)
    piecework_tasks = db.query(PieceworkTask).filter(PieceworkTask.is_active == True).order_by(PieceworkTask.name).all()
    return templates.TemplateResponse("employees/edit.html", {
        "request": request,
        "emp": emp,
        "piecework_tasks": piecework_tasks,
        "current_user": current_user,
        "page_title": "Xodimni tahrirlash"
    })


@app.post("/employees/update/{employee_id}")
async def employee_update(
    employee_id: int,
    full_name: str = Form(...),
    code: str = Form(...),
    position: str = Form(""),
    department: str = Form(""),
    phone: str = Form(""),
    salary: float = Form(0),
    salary_type: str = Form(""),
    piecework_task_ids: List[int] = Form([]),
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    """Xodim ma'lumotlarini yangilash"""
    from urllib.parse import quote
    emp = db.query(Employee).filter(Employee.id == employee_id).first()
    if not emp:
        return RedirectResponse(url="/employees?error=Xodim topilmadi", status_code=303)
    duplicate = db.query(Employee).filter(Employee.code == code, Employee.id != employee_id).first()
    if duplicate:
        return RedirectResponse(url="/employees?error=" + quote("Bunday kod boshqa xodimda mavjud: " + code), status_code=303)
    emp.full_name = full_name
    emp.code = code
    emp.position = position
    emp.department = department
    emp.phone = phone
    emp.salary = salary
    st = (salary_type or "").strip() or None
    if st and st not in ("oylik", "soatlik", "bo'lak", "bo'lak_oylik"):
        st = None
    emp.salary_type = st
    task_ids = [int(x) for x in (piecework_task_ids or []) if str(x).strip().isdigit()]
    task_ids = list(dict.fromkeys(task_ids))
    emp.piecework_task_id = task_ids[0] if task_ids else None  # legacy
    if st in ("bo'lak", "bo'lak_oylik") and task_ids:
        tasks = db.query(PieceworkTask).filter(PieceworkTask.id.in_(task_ids)).all()
        emp.piecework_tasks = tasks
    else:
        emp.piecework_tasks = []
    db.commit()
    return RedirectResponse(url="/employees?updated=1", status_code=303)


@app.post("/employees/delete/{employee_id}")
async def employee_delete(
    employee_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    """Xodimni o'chirish. Bog'liq hujjatlar bo'lsa (ishga qabul, avans, oylik, davomat va h.k.) DB xatolik beradi — foydalanuvchiga xabar."""
    from urllib.parse import quote
    emp = db.query(Employee).filter(Employee.id == employee_id).first()
    if not emp:
        return RedirectResponse(url="/employees?error=Xodim topilmadi", status_code=303)
    try:
        db.delete(emp)
        db.commit()
        return RedirectResponse(url="/employees?deleted=1", status_code=303)
    except IntegrityError:
        db.rollback()
        return RedirectResponse(
            url="/employees?error=" + quote(
                "Xodimni o'chirib bo'lmaydi: unga bog'liq yozuvlar mavjud (ishga qabul hujjati, avans, oylik, davomat va h.k.). "
                "Xodimni o'chirmasdan «Faol emas» deb belgilang yoki avval bog'liq hujjatlarni olib tashlang."
            ),
            status_code=303,
        )


# --- ISHDAN BO'SHATISH ---
DISMISSAL_REASONS = [
    ("own_will", "O'z ixtiyori bilan"),
    ("contract_end", "Shartnoma muddati tugadi"),
    ("discipline", "Mehnat intizomini buzgani"),
    ("reduction", "Loyihadan (shtatdan) qisqartirish"),
    ("agreement", "O'zaro kelishuv"),
    ("other", "Boshqa"),
]


@app.get("/employees/dismissal/create", response_class=HTMLResponse)
async def dismissal_create_page(
    request: Request,
    employee_id: int = Query(..., description="Xodim ID"),
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    """Ishdan bo'shatish hujjati yaratish — forma."""
    emp = db.query(Employee).filter(Employee.id == employee_id).first()
    if not emp:
        return RedirectResponse(url="/employees?error=Xodim topilmadi", status_code=303)
    if not emp.is_active:
        return RedirectResponse(url="/employees?error=Xodim allaqachon ishdan bo'shatilgan", status_code=303)
    default_date = datetime.now().date().strftime("%Y-%m-%d")
    return templates.TemplateResponse("employees/dismissal_form.html", {
        "request": request,
        "employee": emp,
        "reasons": DISMISSAL_REASONS,
        "default_date": default_date,
        "current_user": current_user,
        "page_title": "Ishdan bo'shatish",
    })


@app.post("/employees/dismissal/create", response_class=RedirectResponse)
async def dismissal_create_submit(
    request: Request,
    employee_id: int = Form(...),
    doc_date: str = Form(...),
    reason: str = Form(""),
    note: str = Form(""),
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    """Ishdan bo'shatish hujjatini yaratadi, xodimni faol emas qiladi."""
    from urllib.parse import quote
    emp = db.query(Employee).filter(Employee.id == employee_id).first()
    if not emp:
        return RedirectResponse(url="/employees?error=Xodim topilmadi", status_code=303)
    if not emp.is_active:
        return RedirectResponse(url="/employees?error=Xodim allaqachon ishdan bo'shatilgan", status_code=303)
    try:
        doc_d = datetime.strptime(doc_date.strip()[:10], "%Y-%m-%d").date()
    except (ValueError, TypeError):
        return RedirectResponse(url=f"/employees/dismissal/create?employee_id={employee_id}&error=Noto%27g%27ri sana", status_code=303)
    reason_label = next((r[1] for r in DISMISSAL_REASONS if r[0] == reason), reason or "—")
    count = db.query(DismissalDoc).filter(DismissalDoc.doc_date >= doc_d.replace(day=1)).count()
    number = f"IB-{doc_d.strftime('%Y%m%d')}-{count + 1:04d}"
    doc = DismissalDoc(
        number=number,
        employee_id=emp.id,
        doc_date=doc_d,
        reason=reason_label,
        note=(note or "").strip() or None,
        user_id=current_user.id if current_user else None,
    )
    db.add(doc)
    db.flush()
    emp.is_active = False
    db.commit()
    return RedirectResponse(url=f"/employees/dismissal/{doc.id}?created=1", status_code=303)


@app.get("/employees/dismissal/{doc_id}", response_class=HTMLResponse)
async def dismissal_doc_view(
    doc_id: int,
    request: Request,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    """Ishdan bo'shatish hujjati ko'rinishi."""
    doc = (
        db.query(DismissalDoc)
        .options(joinedload(DismissalDoc.employee), joinedload(DismissalDoc.user))
        .filter(DismissalDoc.id == doc_id)
        .first()
    )
    if not doc:
        raise HTTPException(status_code=404, detail="Hujjat topilmadi")
    return templates.TemplateResponse("employees/dismissal_doc.html", {
        "request": request,
        "doc": doc,
        "current_user": current_user,
        "page_title": f"Ishdan bo'shatish {doc.number}",
    })


def _build_dismissal_docx(doc, company_name: str, employer_rep_name: str):
    """Ishdan bo'shatish hujjatini Word (.docx) sifatida qaytaradi (BytesIO)."""
    d = Document()
    style = d.styles["Normal"]
    style.font.size = Pt(11)
    style.font.name = "Times New Roman"
    d.add_heading("ISHDAN BO'SHATISH HAQIDA BUYRUQ", level=0)
    h = d.paragraphs[-1]
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    d.add_paragraph()
    p = d.add_paragraph()
    p.add_run(f"№ {doc.number}").bold = True
    p.add_run(f"   Sana: {doc.doc_date.strftime('%d.%m.%Y') if doc.doc_date else '—'}")
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    d.add_paragraph()
    d.add_paragraph(f"Joy: ____________________________")
    d.add_paragraph(f"Korxona: {company_name}")
    d.add_paragraph()
    emp = doc.employee
    d.add_paragraph(
        f"1. {emp.full_name} (xodim kodi: {emp.code or '—'}), "
        f"{doc.doc_date.strftime('%d.%m.%Y')} sanadan boshlab ishdan bo'shatiladi."
    )
    d.add_paragraph(f"2. Ishdan bo'shatish sababi: {doc.reason or '—'}.")
    if doc.note:
        d.add_paragraph(f"3. Izoh: {doc.note}")
    d.add_paragraph()
    d.add_paragraph("Ish beruvchi:")
    d.add_paragraph(f"Korxona: {company_name}")
    d.add_paragraph(f"Rahbar: {employer_rep_name}")
    d.add_paragraph("Imzo: ______________________")
    d.add_paragraph()
    d.add_paragraph("Xodim bilan tanishtirildi:")
    d.add_paragraph(f"F.I.O: {emp.full_name}")
    d.add_paragraph("Imzo: ______________________")
    buf = io.BytesIO()
    d.save(buf)
    buf.seek(0)
    return buf


@app.get("/employees/dismissal/{doc_id}/export-word")
async def dismissal_doc_export_word(
    doc_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    """Ishdan bo'shatish hujjatini Word (.docx) formatida yuklab olish."""
    doc = (
        db.query(DismissalDoc)
        .options(joinedload(DismissalDoc.employee))
        .filter(DismissalDoc.id == doc_id)
        .first()
    )
    if not doc:
        raise HTTPException(status_code=404, detail="Hujjat topilmadi")
    company_name = "TOTLI HOLVA SWEETS"
    employer_rep_name = "Rahimov D.A."
    buf = _build_dismissal_docx(doc, company_name, employer_rep_name)
    safe_number = (doc.number or "ib").replace("/", "-").replace("\\", "-")
    filename = f"Ishdan_bo'shatish_{safe_number}.docx"
    return Response(
        content=buf.read(),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename*=UTF-8\'\'{quote(filename)}'},
    )


# --- ISHGA QABUL QILISH HUJJATI ---
@app.get("/employees/hiring-docs", response_class=HTMLResponse)
async def employment_docs_list(
    request: Request,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    """Ishga qabul qilish hujjatlari ro'yxati — barcha foydalanuvchilar barcha hujjatlarni ko'radi"""
    docs = (
        db.query(EmploymentDoc)
        .options(joinedload(EmploymentDoc.employee))
        .order_by(EmploymentDoc.created_at.desc())
        .all()
    )
    return templates.TemplateResponse("employees/hiring_docs_list.html", {
        "request": request,
        "docs": docs,
        "current_user": current_user,
        "page_title": "Ishga qabul qilish hujjatlari"
    })


@app.get("/employees/hiring-doc/create", response_class=HTMLResponse)
async def employment_doc_create_page(
    request: Request,
    employee_id: Optional[int] = None,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    """Ishga qabul hujjati yaratish (xodim tanlash). Har bir xodim faqat bir marta ishga qabul qilinadi."""
    employees = db.query(Employee).filter(Employee.is_active == True).order_by(Employee.full_name).all()
    emp = db.query(Employee).filter(Employee.id == employee_id).first() if employee_id else None
    # Agar tanlangan xodimda allaqachon ishga qabul hujjati bo'lsa — yangi yaratishga ruxsat yo'q
    if emp:
        existing = db.query(EmploymentDoc).filter(EmploymentDoc.employee_id == emp.id).order_by(EmploymentDoc.doc_date.desc()).first()
        if existing:
            return RedirectResponse(
                url="/employees/hiring-docs?error=" + quote(f"«{emp.full_name}» allaqachon ishga qabul qilingan. Yangi hujjat yaratib bo'lmaydi — mavjud hujjatni ko'ring yoki tahrirlang.")
                + "&existing_doc_id=" + str(existing.id),
                status_code=303,
            )
    today_str = date.today().isoformat()
    departments = db.query(Department).filter(Department.is_active == True).order_by(Department.name).all()
    positions = db.query(Position).filter(Position.is_active == True).order_by(Position.name).all()
    piecework_tasks = db.query(PieceworkTask).filter(PieceworkTask.is_active == True).order_by(PieceworkTask.name).all()
    return templates.TemplateResponse("employees/hiring_doc_form.html", {
        "request": request,
        "employees": employees,
        "emp": emp,
        "today_str": today_str,
        "departments": departments,
        "positions": positions,
        "piecework_tasks": piecework_tasks,
        "current_user": current_user,
        "page_title": "Ishga qabul hujjati yaratish"
    })


@app.post("/employees/hiring-doc/create")
async def employment_doc_create(
    employee_id: int = Form(...),
    doc_date: str = Form(...),
    hire_date: str = Form(None),
    position: str = Form(""),
    department: str = Form(""),
    salary: float = Form(0),
    salary_type: str = Form(""),
    piecework_task_ids: List[int] = Form([]),
    rest_days: List[str] = Form([]),
    probation: str = Form(""),
    contract_type: str = Form("indefinite"),
    contract_end_date: str = Form(None),
    note: str = Form(""),
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    """Ishga qabul hujjati yaratish (O'zR Mehnat kodeksi, gov.uz tamoyillari asosida). Har bir xodim faqat bir marta ishga qabul qilinadi."""
    from urllib.parse import quote
    emp = db.query(Employee).filter(Employee.id == employee_id).first()
    if not emp:
        return RedirectResponse(url="/employees/hiring-docs?error=" + quote("Xodim topilmadi"), status_code=303)
    existing = db.query(EmploymentDoc).filter(EmploymentDoc.employee_id == emp.id).first()
    if existing:
        return RedirectResponse(
            url="/employees/hiring-docs?error=" + quote(f"«{emp.full_name}» allaqachon ishga qabul qilingan. Yangi hujjat yaratib bo'lmaydi.")
            + "&existing_doc_id=" + str(existing.id),
            status_code=303,
        )
    try:
        doc_d = datetime.strptime(doc_date.strip(), "%Y-%m-%d").date()
    except (ValueError, TypeError):
        return RedirectResponse(url="/employees/hiring-doc/create?employee_id=" + str(employee_id) + "&error=" + quote("Noto'g'ri sana"), status_code=303)
    hire_d = None
    if hire_date and hire_date.strip():
        try:
            hire_d = datetime.strptime(hire_date.strip(), "%Y-%m-%d").date()
        except (ValueError, TypeError):
            pass
    end_d = None
    if contract_end_date and contract_end_date.strip() and (contract_type or "").strip() == "fixed":
        try:
            end_d = datetime.strptime(contract_end_date.strip(), "%Y-%m-%d").date()
        except (ValueError, TypeError):
            pass
    count = db.query(EmploymentDoc).filter(EmploymentDoc.doc_date >= doc_d.replace(day=1)).count()
    number = f"IQ-{doc_d.strftime('%Y%m%d')}-{count + 1:04d}"
    doc_salary = float(salary) if salary else (emp.salary or 0)
    doc_department = (department or "").strip() or (emp.department or "").strip() or None
    st = (salary_type or "").strip() or None
    if st and st not in ("oylik", "soatlik", "bo'lak", "bo'lak_oylik"):
        st = None
    task_ids = [int(x) for x in (piecework_task_ids or []) if str(x).strip().isdigit()]
    task_ids = list(dict.fromkeys(task_ids))
    rest_days_clean = [d for d in (rest_days or []) if d in ("mon","tue","wed","thu","fri","sat","sun")]
    probation_clean = (probation or "").strip() or None
    ct = (contract_type or "").strip() or "indefinite"
    if ct not in ("indefinite", "fixed", "task"):
        ct = "indefinite"
    doc = EmploymentDoc(
        number=number,
        employee_id=emp.id,
        doc_date=doc_d,
        hire_date=hire_d,
        position=(position or "").strip() or (emp.position or "").strip() or None,
        department=doc_department,
        salary=doc_salary,
        salary_type=st,
        piecework_task_ids=",".join(str(x) for x in task_ids) if (st in ("bo'lak", "bo'lak_oylik") and task_ids) else None,
        rest_days=",".join(rest_days_clean) if rest_days_clean else None,
        probation=probation_clean,
        contract_type=ct,
        contract_end_date=end_d,
        note=note or None,
        user_id=current_user.id,
        confirmed_at=datetime.now(),  # Hujjat yaratilganda avtomatik tasdiqlanadi
    )
    db.add(doc)
    db.flush()
    emp.salary = doc_salary
    if st:
        emp.salary_type = st
    if st in ("bo'lak", "bo'lak_oylik"):
        if task_ids:
            tasks = db.query(PieceworkTask).filter(PieceworkTask.id.in_(task_ids)).all()
            emp.piecework_tasks = tasks
            emp.piecework_task_id = task_ids[0]  # legacy
        else:
            emp.piecework_tasks = []
            emp.piecework_task_id = None
    if hire_d:
        emp.hire_date = hire_d
    if (position or "").strip():
        emp.position = (position or "").strip()
    if doc_department:
        emp.department = doc_department
    db.commit()
    return RedirectResponse(url=f"/employees/hiring-doc/{doc.id}?created=1", status_code=303)


@app.get("/employees/hiring-doc/{doc_id}", response_class=HTMLResponse)
async def employment_doc_view(
    request: Request,
    doc_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    """Ishga qabul hujjati ko'rish / chop etish — barcha maydonlar to'liq ko'rsatiladi."""
    doc = (
        db.query(EmploymentDoc)
        .options(joinedload(EmploymentDoc.employee))
        .filter(EmploymentDoc.id == doc_id)
        .first()
    )
    if not doc:
        return RedirectResponse(url="/employees/hiring-docs?error=Hujjat topilmadi", status_code=303)
    # Bo'lim: hujjatdagi yoki xodimdagi (matn) yoki xodimning department_id orqali
    display_department = (doc.department or "").strip() or None
    if not display_department and doc.employee:
        display_department = (doc.employee.department or "").strip() or None
        if not display_department and getattr(doc.employee, "department_id", None):
            dept = db.query(Department).filter(Department.id == doc.employee.department_id).first()
            if dept:
                display_department = dept.name
    if not display_department:
        display_department = "—"

    # Bo'lak ishlar (snapshot) — hujjatda saqlangan ro'yxat
    piecework_task_names = []
    try:
        raw = (doc.piecework_task_ids or "").strip()
        ids = [int(x) for x in raw.split(",") if x.strip().isdigit()] if raw else []
        if ids:
            tasks = db.query(PieceworkTask).filter(PieceworkTask.id.in_(ids)).order_by(PieceworkTask.name).all()
            for t in tasks:
                nm = (t.name or t.code or str(t.id))
                piecework_task_names.append(nm)
    except Exception:
        piecework_task_names = []
    return templates.TemplateResponse("employees/hiring_doc.html", {
        "request": request,
        "doc": doc,
        "display_department": display_department,
        "piecework_task_names": piecework_task_names,
        "current_user": current_user,
        "page_title": f"Ishga qabul {doc.number}"
    })


@app.get("/employees/hiring-doc/{doc_id}/contract", response_class=HTMLResponse)
async def employment_doc_contract(
    request: Request,
    doc_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    """Mehnat shartnomasi (to'liq) — namuna asosida chop etish."""
    doc = (
        db.query(EmploymentDoc)
        .options(joinedload(EmploymentDoc.employee), joinedload(EmploymentDoc.user))
        .filter(EmploymentDoc.id == doc_id)
        .first()
    )
    if not doc:
        return RedirectResponse(url="/employees/hiring-docs?error=Hujjat topilmadi", status_code=303)

    # Bo'lim ko'rsatish
    display_department = (doc.department or "").strip() or None
    if not display_department and doc.employee:
        display_department = (doc.employee.department or "").strip() or None
        if not display_department and getattr(doc.employee, "department_id", None):
            dept = db.query(Department).filter(Department.id == doc.employee.department_id).first()
            if dept:
                display_department = dept.name
    if not display_department:
        display_department = "—"

    # Tanlangan bo'lak ishlar (snapshot) — stavkalari bilan
    selected_piecework_tasks = []
    try:
        raw = (doc.piecework_task_ids or "").strip()
        ids = [int(x) for x in raw.split(",") if x.strip().isdigit()] if raw else []
        if ids:
            selected_piecework_tasks = db.query(PieceworkTask).filter(PieceworkTask.id.in_(ids)).order_by(PieceworkTask.name).all()
    except Exception:
        selected_piecework_tasks = []

    # Dam olish kunlari matni
    rest_days_display = ""
    try:
        raw_rest = (doc.rest_days or "").strip()
        codes = [x for x in raw_rest.split(",") if x]
        name_map = {
            "mon": "dushanba",
            "tue": "seshanba",
            "wed": "chorshanba",
            "thu": "payshanba",
            "fri": "juma",
            "sat": "shanba",
            "sun": "yakshanba",
        }
        names = [name_map.get(c, c) for c in codes]
        if names:
            rest_days_display = ", ".join(names)
    except Exception:
        rest_days_display = ""

    company_name = "TOTLI HOLVA SWEETS"
    employer_rep_name = "Rahimov D.A."

    return templates.TemplateResponse("employees/labor_contract.html", {
        "request": request,
        "doc": doc,
        "display_department": display_department,
        "selected_piecework_tasks": selected_piecework_tasks,
        "company_name": company_name,
        "employer_rep_name": employer_rep_name,
        "rest_days_display": rest_days_display,
        "current_user": current_user,
        "page_title": f"Mehnat shartnomasi {doc.number}",
    })


def _build_labor_contract_docx(doc, display_department, selected_piecework_tasks, rest_days_display, company_name, employer_rep_name):
    """Shartnoma matnini Word hujjati (.docx) sifatida qaytaradi (BytesIO)."""
    d = Document()
    style = d.styles["Normal"]
    style.font.size = Pt(11)
    style.font.name = "Times New Roman"

    # Sarlavha
    h = d.add_heading("MEHNAT SHARTNOMASI", level=0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    d.add_paragraph()
    p = d.add_paragraph()
    p.add_run(f"№ {doc.number}").bold = True
    p.add_run(f"   Sana: {doc.doc_date.strftime('%d.%m.%Y') if doc.doc_date else '—'}")
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    d.add_paragraph()

    d.add_paragraph("Joy: ____________________________")
    d.add_paragraph(f"Korxona: {company_name}")
    d.add_paragraph()

    d.add_paragraph(
        f"{company_name} (keyingi o'rinlarda Ish beruvchi) va fuqaro {doc.employee.full_name} "
        "(keyingi o'rinlarda Xodim), mazkur mehnat shartnomasini quyidagilar haqida tuzdilar."
    )
    d.add_paragraph()

    d.add_heading("1. UMUMIY QOIDALAR", level=1)
    hire_date_str = doc.hire_date.strftime("%d.%m.%Y") if doc.hire_date else "________________"
    pos = doc.position or (doc.employee.position if doc.employee else "") or "________________"
    d.add_paragraph(
        f"1.1. Xodim {hire_date_str} sanadan boshlab {display_department} bo'limida {pos} lavozimiga ishga qabul qilinadi."
    )
    d.add_paragraph(f"1.2. Xodimning ish joyi: {display_department}.")
    if doc.contract_type == "fixed":
        end = f" ({doc.contract_end_date.strftime('%d.%m.%Y')} gacha)" if doc.contract_end_date else ""
        d.add_paragraph(f"1.3. Mazkur shartnomaning amal qilish muddati: muayyan muddatga{end}.")
    elif doc.contract_type == "task":
        d.add_paragraph("1.3. Mazkur shartnomaning amal qilish muddati: muayyan ishni bajarish davriga.")
    else:
        d.add_paragraph("1.3. Mazkur shartnomaning amal qilish muddati: nomuayyan muddatga.")
    prob = doc.probation if doc.probation else "sinovsiz"
    d.add_paragraph(f"1.4. Sinov muddati: {prob}.")
    d.add_paragraph("1.5. Xodim lavozim yo'riqnomasi va amaldagi qonunchilikka muvofiq mehnat majburiyatlarini bajaradi.")
    d.add_paragraph()

    d.add_heading("2. TOMONLARNING HUQUQ VA MAJBURIYATLARI", level=1)
    d.add_paragraph("2.1. Ish beruvchining majburiyatlari:")
    d.add_paragraph("  • Xodimga xavfsiz va samarali mehnat qilish uchun shart-sharoitlar yaratish.")
    d.add_paragraph("  • Ichki mehnat tartibi qoidalari va lavozim yo'riqnomasi bilan tanishtirish.")
    d.add_paragraph("  • Ish haqini o'z vaqtida to'lash.")
    d.add_paragraph("2.2. Xodimning majburiyatlari:")
    d.add_paragraph("  • Mehnat intizomi va ichki tartib qoidalariga rioya qilish.")
    d.add_paragraph("  • Ish beruvchining qonuniy topshiriqlarini o'z vaqtida va aniq bajarish.")
    d.add_paragraph("  • Mehnat muhofazasi va texnika xavfsizligi talablariga rioya qilish.")
    d.add_paragraph()

    d.add_heading("3. ISH VAQTI VA DAM OLISH VAQTI", level=1)
    d.add_paragraph("3.1. Ish kuni vaqti: 09:00 dan 18:00 gacha (to'liq ish kuni asosida).")
    rest = rest_days_display if rest_days_display else "shanba va yakshanba"
    d.add_paragraph(f"3.2. Dam olish kunlari: {rest}.")
    d.add_paragraph("3.3. Qonunchilikda belgilangan tartibda dam olish/bayram kunlari ishga jalb etilishi mumkin.")
    d.add_paragraph()

    d.add_heading("4. MEHNATGA HAQ TO'LASH", level=1)
    salary_type_map = {"oylik": "Oylik", "soatlik": "Soatlik", "bo'lak": "Bo'lak", "bo'lak_oylik": "Bo'lak + oylik"}
    st = salary_type_map.get(doc.salary_type, "________________")
    d.add_paragraph(f"4.1. Ish haqi turi: {st}.")
    if doc.salary_type in ("bo'lak", "bo'lak_oylik") and selected_piecework_tasks:
        d.add_paragraph("Bo'lak ishlar va stavkalar:")
        for t in selected_piecework_tasks:
            name = t.name or t.code or str(t.id)
            price = f"{t.price_per_unit:,.0f}" if t.price_per_unit is not None else "0"
            unit = t.unit_name or "birlik"
            d.add_paragraph(f"  • {name} — {price} so'm/{unit}")
    salary_val = f"{doc.salary:,.0f}" if doc.salary else "0"
    d.add_paragraph(f"4.2. Mehnat haqi miqdori: {salary_val} so'm.")
    d.add_paragraph("4.3. Ish haqi har oyda kamida ikki marta to'lanadi.")
    d.add_paragraph()

    d.add_heading("5. XIZMAT SAFARLARI", level=1)
    d.add_paragraph("5.1. Ish zaruriyatiga ko'ra Xodim xizmat safariga yuborilishi mumkin. Xarajatlar amaldagi qonunchilikka muvofiq qoplanadi.")
    d.add_paragraph()

    d.add_heading("6. MEHNAT SHARTNOMASINI BEKOR QILISH", level=1)
    d.add_paragraph("6.1. Mehnat shartnomasi O'zbekiston Respublikasi Mehnat kodeksida belgilangan tartibda bekor qilinishi mumkin.")
    d.add_paragraph()

    d.add_heading("7. MEHNAT NIZOLARI", level=1)
    d.add_paragraph("7.1. Mehnat nizolari qonun hujjatlarida belgilangan tartibda hal qilinadi.")
    d.add_paragraph()

    d.add_heading("8. TOMONLAR REKVIZITLARI VA IMZOLARI", level=1)
    d.add_paragraph("Ish beruvchi:")
    d.add_paragraph(f"Korxona: {company_name}")
    d.add_paragraph("Manzil: O'zbekiston Respublikasi, Qo'qon shahri, Jasorat ko'chasi, 52-uy")
    d.add_paragraph("STIR: 311469106")
    d.add_paragraph("Hisob raqam: 202088409071067110001")
    d.add_paragraph('Bank: "Asaka" banki Qo\'qon filiali')
    d.add_paragraph("MFO: 00873")
    d.add_paragraph(f"Rahbar: {employer_rep_name}")
    d.add_paragraph("Imzo: ______________________")
    d.add_paragraph()
    d.add_paragraph("Xodim:")
    d.add_paragraph(f"F.I.O: {doc.employee.full_name}")
    d.add_paragraph(f"Kodi: {doc.employee.code or '—'}")
    d.add_paragraph(f"Telefon: {doc.employee.phone or '—'}")
    d.add_paragraph("Manzil: ____________________________")
    d.add_paragraph("Pasport: ____________________________")
    d.add_paragraph("Imzo: ______________________")

    buf = io.BytesIO()
    d.save(buf)
    buf.seek(0)
    return buf


@app.get("/employees/hiring-doc/{doc_id}/contract/export-word")
async def employment_doc_contract_export_word(
    doc_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    """Mehnat shartnomasini Word (.docx) formatida yuklab olish."""
    doc = (
        db.query(EmploymentDoc)
        .options(joinedload(EmploymentDoc.employee), joinedload(EmploymentDoc.user))
        .filter(EmploymentDoc.id == doc_id)
        .first()
    )
    if not doc:
        raise HTTPException(status_code=404, detail="Hujjat topilmadi")

    display_department = (doc.department or "").strip() or None
    if not display_department and doc.employee:
        display_department = (doc.employee.department or "").strip() or None
        if not display_department and getattr(doc.employee, "department_id", None):
            dept = db.query(Department).filter(Department.id == doc.employee.department_id).first()
            if dept:
                display_department = dept.name
    if not display_department:
        display_department = "—"

    selected_piecework_tasks = []
    try:
        raw = (doc.piecework_task_ids or "").strip()
        ids = [int(x) for x in raw.split(",") if x.strip().isdigit()] if raw else []
        if ids:
            selected_piecework_tasks = db.query(PieceworkTask).filter(PieceworkTask.id.in_(ids)).order_by(PieceworkTask.name).all()
    except Exception:
        selected_piecework_tasks = []

    rest_days_display = ""
    try:
        raw_rest = (doc.rest_days or "").strip()
        codes = [x for x in raw_rest.split(",") if x]
        name_map = {"mon": "dushanba", "tue": "seshanba", "wed": "chorshanba", "thu": "payshanba", "fri": "juma", "sat": "shanba", "sun": "yakshanba"}
        names = [name_map.get(c, c) for c in codes]
        if names:
            rest_days_display = ", ".join(names)
    except Exception:
        rest_days_display = ""

    company_name = "TOTLI HOLVA SWEETS"
    employer_rep_name = "Rahimov D.A."

    buf = _build_labor_contract_docx(
        doc, display_department, selected_piecework_tasks, rest_days_display, company_name, employer_rep_name
    )
    safe_number = (doc.number or "shartnoma").replace("/", "-").replace("\\", "-")
    filename = f"Mehnat_shartnomasi_{safe_number}.docx"
    return Response(
        content=buf.read(),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename*=UTF-8\'\'{quote(filename)}'},
    )


@app.post("/employees/hiring-docs/bulk-confirm")
async def employment_docs_bulk_confirm(
    request: Request,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_admin),
):
    """Tanlangan ishga qabul hujjatlarini tasdiqlash"""
    from urllib.parse import quote
    form = await request.form()
    doc_ids_raw = form.getlist("doc_ids")
    try:
        doc_ids = [int(x) for x in doc_ids_raw if str(x).strip().isdigit()]
    except (ValueError, TypeError):
        doc_ids = []
    if not doc_ids:
        return RedirectResponse(url="/employees/hiring-docs?error=" + quote("Hech qanday hujjat tanlanmagan."), status_code=303)
    confirmed = 0
    for did in doc_ids:
        doc = db.query(EmploymentDoc).filter(EmploymentDoc.id == did).first()
        if doc and not doc.confirmed_at:
            doc.confirmed_at = datetime.now()
            confirmed += 1
    db.commit()
    return RedirectResponse(url=f"/employees/hiring-docs?confirmed=1&count={confirmed}", status_code=303)


@app.post("/employees/hiring-docs/bulk-cancel-confirm")
async def employment_docs_bulk_cancel_confirm(
    request: Request,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_admin),
):
    """Tanlangan hujjatlarda tasdiqlashni bekor qilish"""
    from urllib.parse import quote
    form = await request.form()
    doc_ids_raw = form.getlist("doc_ids")
    try:
        doc_ids = [int(x) for x in doc_ids_raw if str(x).strip().isdigit()]
    except (ValueError, TypeError):
        doc_ids = []
    if not doc_ids:
        return RedirectResponse(url="/employees/hiring-docs?error=" + quote("Hech qanday hujjat tanlanmagan."), status_code=303)
    unconfirmed = 0
    for did in doc_ids:
        doc = db.query(EmploymentDoc).filter(EmploymentDoc.id == did).first()
        if doc and doc.confirmed_at:
            doc.confirmed_at = None
            unconfirmed += 1
    db.commit()
    return RedirectResponse(url=f"/employees/hiring-docs?unconfirmed=1&count={unconfirmed}", status_code=303)


@app.post("/employees/hiring-doc/{doc_id}/confirm")
async def employment_doc_confirm(
    doc_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_admin),
):
    """Ishga qabul hujjatini tasdiqlash"""
    doc = db.query(EmploymentDoc).filter(EmploymentDoc.id == doc_id).first()
    if not doc:
        return RedirectResponse(url="/employees/hiring-docs?error=Hujjat topilmadi", status_code=303)
    doc.confirmed_at = datetime.now()
    db.commit()
    return RedirectResponse(url="/employees/hiring-docs?confirmed=1", status_code=303)


@app.post("/employees/hiring-doc/{doc_id}/cancel-confirm")
async def employment_doc_cancel_confirm(
    doc_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_admin),
):
    """Ishga qabul hujjati tasdiqlashni bekor qilish"""
    doc = db.query(EmploymentDoc).filter(EmploymentDoc.id == doc_id).first()
    if not doc:
        return RedirectResponse(url="/employees/hiring-docs?error=Hujjat topilmadi", status_code=303)
    doc.confirmed_at = None
    db.commit()
    return RedirectResponse(url="/employees/hiring-docs?unconfirmed=1", status_code=303)


@app.post("/employees/hiring-doc/{doc_id}/delete")
async def employment_doc_delete(
    doc_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_admin),
):
    """Ishga qabul hujjatini o'chirish — faqat tasdiqlanmagan hujjatni o'chirish mumkin."""
    from urllib.parse import quote
    doc = db.query(EmploymentDoc).filter(EmploymentDoc.id == doc_id).first()
    if not doc:
        return RedirectResponse(url="/employees/hiring-docs?error=Hujjat topilmadi", status_code=303)
    if doc.confirmed_at:
        return RedirectResponse(
            url="/employees/hiring-docs?error=" + quote("Tasdiqlangan hujjatni o'chirish mumkin emas. Avval «Bekor qilish» orqali tasdiqlashni bekor qiling."),
            status_code=303
        )
    db.delete(doc)
    db.commit()
    return RedirectResponse(url="/employees/hiring-docs?deleted=1", status_code=303)


@app.get("/employees/hiring-doc/{doc_id}/edit", response_class=HTMLResponse)
async def employment_doc_edit_page(
    request: Request,
    doc_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_admin),
):
    """Ishga qabul hujjatini tahrirlash"""
    doc = (
        db.query(EmploymentDoc)
        .options(joinedload(EmploymentDoc.employee))
        .filter(EmploymentDoc.id == doc_id)
        .first()
    )
    if not doc:
        return RedirectResponse(url="/employees/hiring-docs?error=Hujjat topilmadi", status_code=303)
    if doc.confirmed_at:
        from urllib.parse import quote
        return RedirectResponse(
            url="/employees/hiring-docs?error=" + quote("Tasdiqlangan hujjatni tahrirlash mumkin emas. Avval «Bekor qilish» orqali tasdiqlashni bekor qiling."),
            status_code=303,
        )
    departments = db.query(Department).filter(Department.is_active == True).order_by(Department.name).all()
    positions = db.query(Position).filter(Position.is_active == True).order_by(Position.name).all()
    piecework_tasks = db.query(PieceworkTask).filter(PieceworkTask.is_active == True).order_by(PieceworkTask.name).all()
    display_department = (doc.department or "").strip() or (getattr(doc.employee, "department", "") or "").strip() or "—"
    return templates.TemplateResponse("employees/hiring_doc_edit.html", {
        "request": request,
        "doc": doc,
        "departments": departments,
        "positions": positions,
        "piecework_tasks": piecework_tasks,
        "display_department": display_department,
        "current_user": current_user,
        "page_title": f"Ishga qabul {doc.number} — tahrirlash",
    })


@app.post("/employees/hiring-doc/{doc_id}/edit")
async def employment_doc_edit_save(
    doc_id: int,
    doc_date: str = Form(...),
    hire_date: str = Form(None),
    position: str = Form(""),
    department: str = Form(""),
    salary: float = Form(0),
    salary_type: str = Form(""),
    piecework_task_ids: List[int] = Form([]),
    rest_days: List[str] = Form([]),
    probation: str = Form(""),
    contract_type: str = Form("indefinite"),
    contract_end_date: str = Form(None),
    note: str = Form(""),
    db: Session = Depends(get_db),
    current_user: User = Depends(require_admin),
):
    """Ishga qabul hujjatini saqlash (tahrirlash) — faqat tasdiqlanmagan hujjatni tahrirlash mumkin."""
    from urllib.parse import quote
    doc = db.query(EmploymentDoc).filter(EmploymentDoc.id == doc_id).first()
    if not doc:
        return RedirectResponse(url="/employees/hiring-docs?error=Hujjat topilmadi", status_code=303)
    if doc.confirmed_at:
        return RedirectResponse(
            url="/employees/hiring-docs?error=" + quote("Tasdiqlangan hujjatni tahrirlash mumkin emas. Avval «Bekor qilish» orqali tasdiqlashni bekor qiling."),
            status_code=303,
        )
    emp = db.query(Employee).filter(Employee.id == doc.employee_id).first()
    if not emp:
        return RedirectResponse(url="/employees/hiring-docs?error=" + quote("Xodim topilmadi"), status_code=303)
    try:
        doc_d = datetime.strptime(doc_date.strip(), "%Y-%m-%d").date()
    except (ValueError, TypeError):
        return RedirectResponse(url=f"/employees/hiring-doc/{doc_id}/edit?error=" + quote("Noto'g'ri sana"), status_code=303)
    hire_d = None
    if hire_date and hire_date.strip():
        try:
            hire_d = datetime.strptime(hire_date.strip(), "%Y-%m-%d").date()
        except (ValueError, TypeError):
            pass
    end_d = None
    if contract_end_date and contract_end_date.strip() and (contract_type or "").strip() == "fixed":
        try:
            end_d = datetime.strptime(contract_end_date.strip(), "%Y-%m-%d").date()
        except (ValueError, TypeError):
            pass
    st = (salary_type or "").strip() or None
    if st and st not in ("oylik", "soatlik", "bo'lak", "bo'lak_oylik"):
        st = None
    task_ids = [int(x) for x in (piecework_task_ids or []) if str(x).strip().isdigit()]
    task_ids = list(dict.fromkeys(task_ids))
    rest_days_clean = [d for d in (rest_days or []) if d in ("mon","tue","wed","thu","fri","sat","sun")]
    probation_clean = (probation or "").strip() or None
    ct = (contract_type or "").strip() or "indefinite"
    if ct not in ("indefinite", "fixed", "task"):
        ct = "indefinite"

    doc.doc_date = doc_d
    doc.hire_date = hire_d
    doc.position = (position or "").strip() or None
    doc.department = (department or "").strip() or None
    doc.salary = float(salary or 0)
    doc.salary_type = st
    doc.piecework_task_ids = ",".join(str(x) for x in task_ids) if (st in ("bo'lak", "bo'lak_oylik") and task_ids) else None
    doc.contract_type = ct
    doc.contract_end_date = end_d
    doc.note = (note or "").strip() or None
    doc.probation = probation_clean
    doc.rest_days = ",".join(rest_days_clean) if rest_days_clean else None

    # Employee snapshot yangilash (o'ylik hisoblash uchun)
    emp.salary = doc.salary
    if st:
        emp.salary_type = st
    if hire_d:
        emp.hire_date = hire_d
    if doc.position:
        emp.position = doc.position
    if doc.department:
        emp.department = doc.department
    if st in ("bo'lak", "bo'lak_oylik"):
        if task_ids:
            tasks = db.query(PieceworkTask).filter(PieceworkTask.id.in_(task_ids)).all()
            emp.piecework_tasks = tasks
            emp.piecework_task_id = task_ids[0]  # legacy
        else:
            emp.piecework_tasks = []
            emp.piecework_task_id = None

    db.commit()
    return RedirectResponse(url=f"/employees/hiring-doc/{doc.id}?edited=1", status_code=303)


# --- EMPLOYEES EXCEL OPERATIONS ---
@app.get("/employees/export")
async def export_employees(db: Session = Depends(get_db), current_user: User = Depends(require_auth)):
    employees = db.query(Employee).all()
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Employees"
    ws.append(["ID", "Kod", "F.I.SH", "Lavozim", "Bo'lim", "Telefon", "Oylik"])
    for e in employees:
        ws.append([e.id, e.code, e.full_name, e.position, e.department, e.phone, e.salary])
    stream = io.BytesIO()
    wb.save(stream)
    stream.seek(0)
    return StreamingResponse(stream, media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", headers={"Content-Disposition": "attachment; filename=xodimlar.xlsx"})

@app.get("/employees/template")
async def template_employees():
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Template"
    ws.append(["Kod", "F.I.SH", "Lavozim", "Bo'lim", "Telefon", "Oylik"])
    ws.append(["X001", "Aliyev Vali", "Ishchi", "Ishlab chiqarish", "+998901234567", 3000000])
    stream = io.BytesIO()
    wb.save(stream)
    stream.seek(0)
    return StreamingResponse(stream, media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", headers={"Content-Disposition": "attachment; filename=xodim_andoza.xlsx"})

@app.post("/employees/import")
async def import_employees(file: UploadFile = File(...), db: Session = Depends(get_db), current_user: User = Depends(require_auth)):
    contents = await file.read()
    wb = openpyxl.load_workbook(io.BytesIO(contents))
    ws = wb.active
    rows = list(ws.iter_rows(min_row=2, values_only=True))
    for row in rows:
        if not row[0]: continue
        code, full_name, position, department, phone, salary = row[0:6]
        employee = db.query(Employee).filter(Employee.code == code).first()
        if not employee:
            employee = Employee(
                code=code, 
                full_name=full_name, 
                position=position, 
                department=department, 
                phone=phone, 
                salary=salary
            )
            db.add(employee)
        else:
            employee.full_name = full_name
            employee.position = position
            employee.department = department
            employee.phone = phone
            employee.salary = salary
        db.commit()
    return RedirectResponse(url="/employees", status_code=303)


@app.post("/employees/import-from-hikvision-preview")
async def employees_import_from_hikvision_preview(
    request: Request,
    hikvision_host: str = Form(...),
    hikvision_port: str = Form("443"),
    hikvision_username: str = Form("admin"),
    hikvision_password: str = Form(""),
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    """Hikvision ulanishi va yuklanadigan shaxslar ro'yxatini ko'rsatadi; tanlanganlarni keyin yuklash mumkin."""
    from urllib.parse import quote
    try:
        port = int((hikvision_port or "").strip() or "443")
    except (ValueError, TypeError):
        port = 443
    try:
        from app.utils.hikvision import HikvisionAPI
        api = HikvisionAPI(
            host=(hikvision_host or "").strip(),
            port=port,
            username=(hikvision_username or "admin").strip(),
            password=(hikvision_password or ""),
        )
        if not api.test_connection():
            return RedirectResponse(
                url="/employees?error=" + quote(api._last_error or "Qurilma bilan bog'lanib bo'lmadi."),
                status_code=303
            )
        persons = api.get_person_list()
    except Exception as e:
        return RedirectResponse(url="/employees?error=" + quote("Hikvision: " + str(e)[:150]), status_code=303)
    return templates.TemplateResponse("employees/hikvision_import_preview.html", {
        "request": request,
        "persons": persons or [],
        "hikvision_host": (hikvision_host or "").strip(),
        "hikvision_port": str(port),
        "hikvision_username": (hikvision_username or "admin").strip(),
        "hikvision_password": hikvision_password or "",
        "current_user": current_user,
        "page_title": "Hikvision — xodimlarni tanlash"
    })


@app.get("/employees/import-from-hikvision-preview", response_class=HTMLResponse)
async def employees_import_from_hikvision_preview_get(
    request: Request,
    current_user: User = Depends(require_auth),
):
    """Preview sahifasiga to'g'ridan-to'g'ri kirilsa xodimlar ro'yxatiga yo'naltiradi."""
    return RedirectResponse(url="/employees", status_code=303)


@app.post("/employees/import-from-hikvision")
async def employees_import_from_hikvision(
    hikvision_host: str = Form(...),
    hikvision_port: str = Form("443"),
    hikvision_username: str = Form("admin"),
    hikvision_password: str = Form(""),
    employee_no: Optional[List[str]] = Form(None),
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    """Hikvision qurilmasidan tanlangan (yoki barcha) xodimlarni Employee jadvaliga qo'shadi."""
    from urllib.parse import quote
    try:
        port = int((hikvision_port or "").strip() or "443")
    except (ValueError, TypeError):
        port = 443
    employee_nos = employee_no if isinstance(employee_no, list) and employee_no else None
    try:
        from app.utils.hikvision import import_employees_from_hikvision
        result = import_employees_from_hikvision(
            (hikvision_host or "").strip(),
            port,
            (hikvision_username or "admin").strip(),
            (hikvision_password or ""),
            db,
            employee_nos=employee_nos,
        )
        err_list = result.get("errors") or []
        imported = result.get("imported", 0)
        updated = result.get("updated", 0)
        if err_list:
            msg = f"Qo'shildi: {imported}, yangilandi: {updated}. Xato: " + "; ".join(str(e) for e in err_list[:3])
            return RedirectResponse(url="/employees?warning=" + quote(msg), status_code=303)
        msg = f"Qo'shildi: {imported}, yangilandi: {updated}."
        return RedirectResponse(url="/employees?imported=1&msg=" + quote(msg), status_code=303)
    except Exception as e:
        return RedirectResponse(url="/employees?error=" + quote("Hikvision: " + str(e)[:150]), status_code=303)


# ==========================================
# DAVOMAT (KUNLIK TABELLAR)
# ==========================================

@app.get("/employees/attendance", response_class=HTMLResponse)
async def attendance_docs_list(
    request: Request,
    start_date: Optional[str] = None,
    end_date: Optional[str] = None,
    sort: Optional[str] = None,
    order: Optional[str] = None,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    """Kunlik tabel hujjatlari ro'yxati — saralash: number, date, count, confirmed_at"""
    today = date.today()
    start_date = start_date or (today - timedelta(days=30)).strftime("%Y-%m-%d")
    end_date = end_date or today.strftime("%Y-%m-%d")
    sort = (sort or "date").strip().lower()
    order = (order or "asc").strip().lower()
    if order not in ("asc", "desc"):
        order = "asc"
    query = (
        db.query(AttendanceDoc)
        .filter(AttendanceDoc.date >= start_date, AttendanceDoc.date <= end_date)
    )
    if sort == "number":
        query = query.order_by(AttendanceDoc.number.desc() if order == "desc" else AttendanceDoc.number.asc())
    elif sort == "date":
        query = query.order_by(AttendanceDoc.date.desc() if order == "desc" else AttendanceDoc.date.asc())
    elif sort == "confirmed_at":
        query = query.order_by(
            AttendanceDoc.confirmed_at.desc() if order == "desc" else AttendanceDoc.confirmed_at.asc()
        )
    else:
        query = query.order_by(AttendanceDoc.date.desc())
    docs = query.all()
    count_by_doc = {}
    for doc in docs:
        count_by_doc[doc.id] = db.query(Attendance).filter(Attendance.date == doc.date).count()
    if sort == "count":
        reverse = order == "desc"
        docs = sorted(docs, key=lambda d: count_by_doc.get(d.id, 0), reverse=reverse)
    return templates.TemplateResponse("employees/attendance_docs_list.html", {
        "request": request,
        "docs": docs,
        "count_by_doc": count_by_doc,
        "start_date": start_date,
        "end_date": end_date,
        "sort": sort,
        "order": order,
        "current_user": current_user,
        "page_title": "Kunlik tabellar",
    })


@app.get("/employees/attendance/form", response_class=HTMLResponse)
async def attendance_form(
    request: Request,
    date_param: Optional[str] = Query(None, alias="date"),
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    """Tabel formasi — sana tanlash, shu kundagi yozuvlar, Hikvision yuklash. Sana query ?date=YYYY-MM-DD orqali olinadi."""
    today = date.today()
    form_date_str = (date_param or "").strip() or today.strftime("%Y-%m-%d")
    try:
        form_date = datetime.strptime(form_date_str, "%Y-%m-%d").date()
    except ValueError:
        form_date = today
        form_date_str = form_date.strftime("%Y-%m-%d")
    attendances = (
        db.query(Attendance)
        .filter(Attendance.date == form_date)
        .order_by(Attendance.employee_id)
        .all()
    )
    attendance_by_employee = {a.employee_id: a for a in attendances}
    # Barcha faol xodimlar + shu kunda davomat bo‘lgan (lekin ro‘yxatda bo‘lmagan) xodimlar
    employees_active = (
        db.query(Employee)
        .filter(Employee.is_active == True)
        .order_by(Employee.full_name)
        .all()
    )
    employee_ids_in_rows = {e.id for e in employees_active}
    attendance_rows = [{"employee": e, "attendance": attendance_by_employee.get(e.id)} for e in employees_active]
    for att in attendances:
        if att.employee_id not in employee_ids_in_rows:
            emp = db.query(Employee).filter(Employee.id == att.employee_id).first()
            if emp:
                attendance_rows.append({"employee": emp, "attendance": att})
                employee_ids_in_rows.add(emp.id)
    doc = db.query(AttendanceDoc).filter(AttendanceDoc.date == form_date).first()
    return templates.TemplateResponse("employees/attendance_form.html", {
        "request": request,
        "form_date": form_date,
        "form_date_str": form_date_str,
        "attendances": attendances,
        "attendance_rows": attendance_rows,
        "doc": doc,
        "current_user": current_user,
        "page_title": "Tabel formasi",
    })


@app.post("/employees/attendance/sync-hikvision")
async def attendance_sync_hikvision(
    request: Request,
    start_date: str = Form(...),
    end_date: str = Form(...),
    hikvision_host: str = Form(...),
    hikvision_port: str = Form("443"),
    hikvision_username: str = Form("admin"),
    hikvision_password: str = Form(""),
    redirect_url: str = Form("/employees/attendance/form"),
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    """Hikvision'dan davomat yuklash"""
    from urllib.parse import quote
    sep = "&" if "?" in (redirect_url or "") else "?"
    base_redirect = (redirect_url or "/employees/attendance/form").strip()
    try:
        start_d = datetime.strptime((start_date or "").strip(), "%Y-%m-%d").date()
        end_d = datetime.strptime((end_date or "").strip(), "%Y-%m-%d").date()
    except (ValueError, TypeError):
        return RedirectResponse(url=base_redirect + sep + "error=" + quote("Noto'g'ri sana"), status_code=303)
    try:
        port = int(hikvision_port.strip() or "443")
    except (ValueError, TypeError):
        port = 443
    try:
        from app.utils.hikvision import sync_hikvision_attendance
        result = sync_hikvision_attendance(
            (hikvision_host or "").strip(),
            port,
            (hikvision_username or "admin").strip(),
            (hikvision_password or ""),
            start_d,
            end_d,
            db,
        )
        err_list = result.get("errors") or []
        events_count = result.get("events_count", 0)
        imported = result.get("imported", 0)
        msg = f"Hodisa: {events_count} ta, yuklangan: {imported} ta. Xato: {len(err_list)} ta."
        if err_list:
            msg += " " + "; ".join(str(e) for e in err_list[:3])
        return RedirectResponse(url=base_redirect + sep + "synced=1&msg=" + quote(msg), status_code=303)
    except Exception as e:
        err_msg = str(e)[:200] if e else "Noma'lum xato"
        traceback.print_exc()
        return RedirectResponse(url=base_redirect + sep + "error=" + quote("Hikvision yuklash: " + err_msg), status_code=303)


def _parse_time(s: str):
    """'09:00' yoki '09:00:00' dan time object qaytaradi, bo'sh bo'lsa None."""
    if not s or not str(s).strip():
        return None
    s = str(s).strip()
    for fmt in ("%H:%M", "%H:%M:%S"):
        try:
            from datetime import time as dt_time
            t = datetime.strptime(s, fmt).time()
            return t
        except ValueError:
            continue
    return None


@app.post("/employees/attendance/form/bulk-time")
async def attendance_form_bulk_time(
    request: Request,
    date_param: str = Form(..., alias="date"),
    check_in_time: str = Form("09:00"),
    check_out_time: str = Form("18:00"),
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    """Barcha faol xodimlarga tanlangan kun uchun Keldi/Ketdi/Soat (9:00–18:00) yuklash va saqlash."""
    from urllib.parse import quote
    try:
        doc_date = datetime.strptime(date_param.strip(), "%Y-%m-%d").date()
    except (ValueError, AttributeError):
        return RedirectResponse(url="/employees/attendance/form?error=" + quote("Noto'g'ri sana"), status_code=303)
    t_in = _parse_time((check_in_time or "09:00").strip())
    t_out = _parse_time((check_out_time or "18:00").strip())
    if not t_in:
        t_in = datetime.strptime("09:00", "%H:%M").time()
    if not t_out:
        t_out = datetime.strptime("18:00", "%H:%M").time()
    check_in_dt = datetime.combine(doc_date, t_in)
    check_out_dt = datetime.combine(doc_date, t_out)
    delta = check_out_dt - check_in_dt
    if delta.total_seconds() < 0:
        delta += timedelta(days=1)
    hours_worked = round(delta.total_seconds() / 3600 * 2) / 2
    form = await request.form()
    employee_ids_param = form.getlist("employee_ids")
    if employee_ids_param:
        try:
            emp_ids = [int(x) for x in employee_ids_param if str(x).strip().isdigit()]
        except (ValueError, TypeError):
            emp_ids = []
        employees = db.query(Employee).filter(Employee.id.in_(emp_ids), Employee.is_active == True).all() if emp_ids else []
    else:
        employees = db.query(Employee).filter(Employee.is_active == True).all()
    saved = 0
    for emp in employees:
        att = db.query(Attendance).filter(Attendance.employee_id == emp.id, Attendance.date == doc_date).first()
        if not att:
            att = Attendance(employee_id=emp.id, date=doc_date)
            db.add(att)
        att.check_in = check_in_dt
        att.check_out = check_out_dt
        att.hours_worked = hours_worked
        att.status = "present"
        saved += 1
    db.commit()
    msg = f"{saved} ta xodimga vaqt yuklandi (Keldi {check_in_time or '09:00'}, Ketdi {check_out_time or '18:00'})."
    return RedirectResponse(
        url=f"/employees/attendance/form?date={doc_date.strftime('%Y-%m-%d')}&saved={saved}&msg=" + quote(msg),
        status_code=303,
    )


@app.post("/employees/attendance/form/save")
async def attendance_form_save(
    request: Request,
    date_param: str = Form(..., alias="date"),
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    """Tabelni qo'lda to'ldirish — har bir xodim uchun Keldi/Ketdi/Soat/Holat/Izoh saqlanadi."""
    from urllib.parse import quote
    try:
        doc_date = datetime.strptime(date_param.strip(), "%Y-%m-%d").date()
    except (ValueError, AttributeError):
        return RedirectResponse(url="/employees/attendance/form?error=" + quote("Noto'g'ri sana"), status_code=303)
    form = await request.form()
    # employee_ids = form.getlist("employee_id"), check_in_1=..., check_out_1=..., hours_1=..., status_1=..., note_1=...
    employee_ids = form.getlist("employee_id")
    saved = 0
    for i, emp_id_str in enumerate(employee_ids):
        try:
            emp_id = int(emp_id_str)
        except (ValueError, TypeError):
            continue
        emp = db.query(Employee).filter(Employee.id == emp_id).first()
        if not emp:
            continue
        check_in_str = (form.get(f"check_in_{emp_id}") or "").strip()
        check_out_str = (form.get(f"check_out_{emp_id}") or "").strip()
        hours_str = (form.get(f"hours_{emp_id}") or "").strip().replace(",", ".")
        status_val = (form.get(f"status_{emp_id}") or "").strip() or "present"
        note_val = (form.get(f"note_{emp_id}") or "").strip() or None
        if status_val not in ("present", "absent", "leave"):
            status_val = "present"
        if not check_in_str and not check_out_str:
            status_val = "absent"
        try:
            hours_worked = float(hours_str) if hours_str else None
        except ValueError:
            hours_worked = None
        check_in_time = _parse_time(check_in_str)
        check_out_time = _parse_time(check_out_str)
        check_in_dt = datetime.combine(doc_date, check_in_time) if check_in_time else None
        check_out_dt = datetime.combine(doc_date, check_out_time) if check_out_time else None
        if hours_worked is None and check_in_dt and check_out_dt:
            delta = check_out_dt - check_in_dt
            if delta.total_seconds() < 0:
                delta += timedelta(days=1)
            hours_worked = round(delta.total_seconds() / 3600 * 2) / 2
        att = db.query(Attendance).filter(Attendance.employee_id == emp_id, Attendance.date == doc_date).first()
        if not att:
            att = Attendance(employee_id=emp_id, date=doc_date)
            db.add(att)
        att.check_in = check_in_dt
        att.check_out = check_out_dt
        att.hours_worked = hours_worked if hours_worked is not None else (att.hours_worked if att.hours_worked is not None else 0)
        att.status = status_val
        att.note = note_val
        saved += 1
    db.commit()
    return RedirectResponse(
        url=f"/employees/attendance/form?date={doc_date.strftime('%Y-%m-%d')}&saved={saved}&msg=" + quote("Tabel qo'lda saqlandi."),
        status_code=303,
    )


@app.post("/employees/attendance/form/confirm")
async def attendance_form_confirm(
    request: Request,
    date_param: str = Form(..., alias="date"),
    db: Session = Depends(get_db),
    current_user: User = Depends(require_admin),
):
    """Kunni tasdiqlash — AttendanceDoc yaratiladi"""
    try:
        doc_date = datetime.strptime(date_param, "%Y-%m-%d").date()
    except ValueError:
        return RedirectResponse(url="/employees/attendance/form?error=Noto'g'ri sana", status_code=303)
    existing = db.query(AttendanceDoc).filter(AttendanceDoc.date == doc_date).first()
    if existing:
        if existing.confirmed_at:
            return RedirectResponse(url="/employees/attendance?already=1", status_code=303)
        existing.confirmed_at = datetime.now()
        existing.user_id = current_user.id
        db.commit()
        return RedirectResponse(url="/employees/attendance?confirmed=1", status_code=303)
    count = db.query(AttendanceDoc).filter(AttendanceDoc.date >= doc_date.replace(day=1)).count()
    number = f"TBL-{doc_date.strftime('%Y%m%d')}-{count + 1:04d}"
    doc = AttendanceDoc(number=number, date=doc_date, user_id=current_user.id, confirmed_at=datetime.now())
    db.add(doc)
    db.commit()
    return RedirectResponse(url="/employees/attendance?confirmed=1", status_code=303)


@app.get("/employees/attendance/doc/{doc_id}", response_class=HTMLResponse)
async def attendance_doc_view(
    request: Request,
    doc_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    """Kunlik tabel hujjati ko'rinishi"""
    doc = db.query(AttendanceDoc).filter(AttendanceDoc.id == doc_id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Hujjat topilmadi")
    rows = db.query(Attendance).filter(Attendance.date == doc.date).order_by(Attendance.employee_id).all()
    return templates.TemplateResponse("employees/attendance_doc.html", {
        "request": request,
        "doc": doc,
        "rows": rows,
        "current_user": current_user,
        "page_title": f"Tabel {doc.number}",
    })


@app.get("/employees/attendance/records", response_class=HTMLResponse)
async def attendance_records(
    request: Request,
    start_date: Optional[str] = None,
    end_date: Optional[str] = None,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    """Barcha davomat yozuvlari (sana oralig'i) — qo'lda qo'shish/tahrirlash"""
    today = date.today()
    start_date = start_date or (today - timedelta(days=7)).strftime("%Y-%m-%d")
    end_date = end_date or today.strftime("%Y-%m-%d")
    records = (
        db.query(Attendance)
        .filter(Attendance.date >= start_date, Attendance.date <= end_date)
        .order_by(Attendance.date.desc(), Attendance.employee_id)
        .all()
    )
    employees = db.query(Employee).filter(Employee.is_active == True).order_by(Employee.full_name).all()
    return templates.TemplateResponse("employees/attendance_records.html", {
        "request": request,
        "records": records,
        "employees": employees,
        "start_date": start_date,
        "end_date": end_date,
        "current_user": current_user,
        "page_title": "Davomat yozuvlari",
    })


@app.post("/employees/attendance/doc/{doc_id}/delete")
async def attendance_doc_delete(
    doc_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_admin),
):
    """Tabel hujjatini ro'yxatdan butunlay o'chirish (AttendanceDoc jadvaldan o'chiriladi; davomat yozuvlari saqlanadi)."""
    doc = db.query(AttendanceDoc).filter(AttendanceDoc.id == doc_id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Hujjat topilmadi")
    db.delete(doc)
    db.commit()
    return RedirectResponse(url="/employees/attendance?deleted=1", status_code=303)


@app.post("/employees/attendance/doc/{doc_id}/cancel-confirm")
async def attendance_doc_cancel_confirm(
    doc_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_admin),
):
    """Tasdiqlashni bekor qilish"""
    doc = db.query(AttendanceDoc).filter(AttendanceDoc.id == doc_id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Hujjat topilmadi")
    doc.confirmed_at = None
    db.commit()
    return RedirectResponse(url="/employees/attendance?unconfirmed=1", status_code=303)


@app.post("/employees/attendance/records/add")
async def attendance_record_add(
    request: Request,
    employee_id: int = Form(...),
    att_date: str = Form(...),
    check_in: Optional[str] = Form(None),
    check_out: Optional[str] = Form(None),
    hours_worked: float = Form(0),
    note: str = Form(""),
    start_date: str = Form(""),
    end_date: str = Form(""),
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    """Davomat yozuvi qo'shish (qo'lda)"""
    try:
        att_d = datetime.strptime(att_date, "%Y-%m-%d").date()
    except ValueError:
        return RedirectResponse(url=f"/employees/attendance/records?start_date={start_date}&end_date={end_date}&error=Noto'g'ri sana", status_code=303)
    check_in_dt = None
    if check_in:
        try:
            check_in_dt = datetime.strptime(f"{att_date} {check_in}", "%Y-%m-%d %H:%M")
        except ValueError:
            pass
    check_out_dt = None
    if check_out:
        try:
            check_out_dt = datetime.strptime(f"{att_date} {check_out}", "%Y-%m-%d %H:%M")
        except ValueError:
            pass
    att = Attendance(
        employee_id=employee_id,
        date=att_d,
        check_in=check_in_dt,
        check_out=check_out_dt,
        hours_worked=hours_worked or 0,
        status="present",
        note=note or None,
    )
    db.add(att)
    db.commit()
    return RedirectResponse(url=f"/employees/attendance/records?start_date={start_date}&end_date={end_date}&added=1", status_code=303)


@app.get("/employees/attendance/records/edit/{record_id}", response_class=HTMLResponse)
async def attendance_record_edit_page(
    request: Request,
    record_id: int,
    start_date: str = Query(""),
    end_date: str = Query(""),
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    """Davomat yozuvini tahrirlash sahifasi"""
    att = db.query(Attendance).filter(Attendance.id == record_id).first()
    if not att:
        raise HTTPException(status_code=404, detail="Yozuv topilmadi")
    return templates.TemplateResponse("employees/attendance_record_edit.html", {
        "request": request,
        "record": att,
        "start_date": start_date or att.date.strftime("%Y-%m-%d"),
        "end_date": end_date or att.date.strftime("%Y-%m-%d"),
        "current_user": current_user,
        "page_title": "Davomat yozuvini tahrirlash",
    })


@app.post("/employees/attendance/records/edit/{record_id}")
async def attendance_record_edit_save(
    record_id: int,
    check_in: Optional[str] = Form(None),
    check_out: Optional[str] = Form(None),
    hours_worked: float = Form(None),
    status: str = Form("present"),
    note: str = Form(""),
    start_date: str = Form(""),
    end_date: str = Form(""),
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    """Davomat yozuvini saqlash"""
    att = db.query(Attendance).filter(Attendance.id == record_id).first()
    if not att:
        raise HTTPException(status_code=404, detail="Yozuv topilmadi")
    att_date_str = att.date.strftime("%Y-%m-%d")
    check_in_dt = None
    if check_in and str(check_in).strip():
        try:
            check_in_dt = datetime.strptime(f"{att_date_str} {check_in.strip()}", "%Y-%m-%d %H:%M")
        except ValueError:
            pass
    check_out_dt = None
    if check_out and str(check_out).strip():
        try:
            check_out_dt = datetime.strptime(f"{att_date_str} {check_out.strip()}", "%Y-%m-%d %H:%M")
        except ValueError:
            pass
    att.check_in = check_in_dt
    att.check_out = check_out_dt
    if hours_worked is not None:
        att.hours_worked = float(hours_worked)
    elif check_in_dt and check_out_dt:
        delta = check_out_dt - check_in_dt
        if delta.total_seconds() < 0:
            delta += timedelta(days=1)
        att.hours_worked = round(delta.total_seconds() / 3600 * 2) / 2
    if status in ("present", "absent", "leave"):
        att.status = status
    att.note = (note or "").strip() or None
    db.commit()
    return RedirectResponse(url=f"/employees/attendance/records?start_date={start_date}&end_date={end_date}&updated=1", status_code=303)


@app.post("/employees/attendance/records/bulk-time")
async def attendance_records_bulk_time(
    request: Request,
    start_date: str = Form(""),
    end_date: str = Form(""),
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    """Tanlangan yozuvlarga Keldi 9:00, Ketdi 18:00, Soat 9 qo'llash"""
    from urllib.parse import quote
    form = await request.form()
    record_ids_raw = form.getlist("record_ids")
    try:
        record_ids = [int(x) for x in record_ids_raw if str(x).strip().isdigit()]
    except (ValueError, TypeError):
        record_ids = []
    if not record_ids:
        return RedirectResponse(url=f"/employees/attendance/records?start_date={start_date}&end_date={end_date}&error=" + quote("Hech qanday yozuv tanlanmagan"), status_code=303)
    check_in_dt = datetime.strptime("09:00", "%H:%M").time()
    check_out_dt = datetime.strptime("18:00", "%H:%M").time()
    hours_worked = 9.0
    updated = 0
    for rid in record_ids:
        att = db.query(Attendance).filter(Attendance.id == rid).first()
        if not att:
            continue
        att.check_in = datetime.combine(att.date, check_in_dt)
        att.check_out = datetime.combine(att.date, check_out_dt)
        att.hours_worked = hours_worked
        att.status = "present"
        updated += 1
    db.commit()
    return RedirectResponse(url=f"/employees/attendance/records?start_date={start_date}&end_date={end_date}&updated={updated}&msg=" + quote("Vaqt yuklandi (9:00–18:00)."), status_code=303)


@app.post("/employees/attendance/records/bulk-time-all")
async def attendance_records_bulk_time_all(
    start_date: str = Form(...),
    end_date: str = Form(...),
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    """Tanlangan sana oralig'idagi har bir kun uchun barcha faol xodimlarga Keldi 9:00, Ketdi 18:00, Soat 9 yuklash."""
    from urllib.parse import quote
    try:
        d_start = datetime.strptime(start_date.strip()[:10], "%Y-%m-%d").date()
        d_end = datetime.strptime(end_date.strip()[:10], "%Y-%m-%d").date()
    except (ValueError, TypeError):
        return RedirectResponse(url=f"/employees/attendance/records?error=" + quote("Noto'g'ri sana"), status_code=303)
    if d_end < d_start:
        d_end = d_start
    check_in_t = datetime.strptime("09:00", "%H:%M").time()
    check_out_t = datetime.strptime("18:00", "%H:%M").time()
    hours_worked = 9.0
    employees = db.query(Employee).filter(Employee.is_active == True).all()
    saved = 0
    d = d_start
    while d <= d_end:
        check_in_dt = datetime.combine(d, check_in_t)
        check_out_dt = datetime.combine(d, check_out_t)
        for emp in employees:
            att = db.query(Attendance).filter(Attendance.employee_id == emp.id, Attendance.date == d).first()
            if not att:
                att = Attendance(employee_id=emp.id, date=d)
                db.add(att)
            att.check_in = check_in_dt
            att.check_out = check_out_dt
            att.hours_worked = hours_worked
            att.status = "present"
            saved += 1
        d += timedelta(days=1)
    db.commit()
    msg = quote(f"Barcha xodimlar uchun vaqt yuklandi: {saved} ta yozuv (9:00–18:00).")
    return RedirectResponse(url=f"/employees/attendance/records?start_date={start_date}&end_date={end_date}&updated={saved}&msg={msg}", status_code=303)


@app.post("/employees/attendance/records/delete/{record_id}")
async def attendance_record_delete(
    record_id: int,
    start_date: str = Form(""),
    end_date: str = Form(""),
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    """Davomat yozuvini o'chirish"""
    att = db.query(Attendance).filter(Attendance.id == record_id).first()
    if att:
        db.delete(att)
        db.commit()
    return RedirectResponse(url=f"/employees/attendance/records?start_date={start_date}&end_date={end_date}", status_code=303)


@app.post("/employees/attendance/records/bulk-delete")
async def attendance_records_bulk_delete(
    request: Request,
    start_date: str = Form(""),
    end_date: str = Form(""),
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    """Tanlangan davomat yozuvlarini o'chirish."""
    from urllib.parse import quote
    form = await request.form()
    record_ids_raw = form.getlist("record_ids")
    try:
        record_ids = [int(x) for x in record_ids_raw if str(x).strip().isdigit()]
    except (ValueError, TypeError):
        record_ids = []
    if not record_ids:
        return RedirectResponse(url=f"/employees/attendance/records?start_date={start_date}&end_date={end_date}&error=" + quote("Hech qanday yozuv tanlanmagan."), status_code=303)
    deleted = 0
    for rid in record_ids:
        att = db.query(Attendance).filter(Attendance.id == rid).first()
        if att:
            db.delete(att)
            deleted += 1
    db.commit()
    msg = quote(f"Tanlangan {deleted} ta yozuv o'chirildi.")
    return RedirectResponse(url=f"/employees/attendance/records?start_date={start_date}&end_date={end_date}&deleted={deleted}&msg={msg}", status_code=303)


# ==========================================
# AVANS BERISH
# ==========================================

def _advances_list_redirect_params(form_or_params, key_from="date_from", key_to="date_to"):
    """Filtr parametrlarini redirect URL ga qo'shish."""
    parts = []
    if hasattr(form_or_params, "get"):
        df, dt = form_or_params.get(key_from) or "", form_or_params.get(key_to) or ""
    else:
        df = form_or_params.get(key_from, "") or ""
        dt = form_or_params.get(key_to, "") or ""
    if (df or "").strip():
        parts.append("date_from=" + quote(str(df).strip()[:10]))
    if (dt or "").strip():
        parts.append("date_to=" + quote(str(dt).strip()[:10]))
    return "&".join(parts)


@app.get("/employees/advances", response_class=HTMLResponse)
async def employee_advances_list(
    request: Request,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
    date_from: Optional[str] = None,
    date_to: Optional[str] = None,
):
    """Xodim avanslari ro'yxati — sana bo'yicha filtrlash."""
    q = db.query(EmployeeAdvance).options(joinedload(EmployeeAdvance.cash_register)).order_by(EmployeeAdvance.advance_date.desc())
    if (date_from or "").strip():
        try:
            df = datetime.strptime(str(date_from).strip()[:10], "%Y-%m-%d").date()
            q = q.filter(EmployeeAdvance.advance_date >= df)
        except ValueError:
            pass
    if (date_to or "").strip():
        try:
            dt = datetime.strptime(str(date_to).strip()[:10], "%Y-%m-%d").date()
            q = q.filter(EmployeeAdvance.advance_date <= dt)
        except ValueError:
            pass
    advances = q.all()
    employees = db.query(Employee).filter(Employee.is_active == True).order_by(Employee.full_name).all()
    cash_registers = db.query(CashRegister).filter(CashRegister.is_active == True).order_by(CashRegister.name).all()
    default_date = date.today().strftime("%Y-%m-%d")
    filter_date_from = str(date_from or "").strip()[:10] if date_from else ""
    filter_date_to = str(date_to or "").strip()[:10] if date_to else ""
    return templates.TemplateResponse("employees/advances_list.html", {
        "request": request,
        "advances": advances,
        "employees": employees,
        "cash_registers": cash_registers,
        "default_date": default_date,
        "filter_date_from": filter_date_from,
        "filter_date_to": filter_date_to,
        "current_user": current_user,
        "page_title": "Avans berish",
    })


@app.post("/employees/advances/add")
async def employee_advance_add(
    request: Request,
    employee_id: int = Form(...),
    amount: float = Form(...),
    advance_date: str = Form(...),
    cash_register_id: Optional[int] = Form(None),
    note: str = Form(""),
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    """Avans qo'shish; tanlangan kassadan chiqim yoziladi."""
    try:
        adv_date = datetime.strptime(advance_date, "%Y-%m-%d").date()
    except ValueError:
        return RedirectResponse(url="/employees/advances?error=Noto'g'ri sana", status_code=303)
    if amount <= 0:
        return RedirectResponse(url="/employees/advances?error=Summa 0 dan katta bo'lishi kerak", status_code=303)
    emp = db.query(Employee).filter(Employee.id == employee_id).first()
    if not emp:
        return RedirectResponse(url="/employees/advances?error=Xodim topilmadi", status_code=303)
    cash = None
    if cash_register_id:
        cash = db.query(CashRegister).filter(CashRegister.id == cash_register_id, CashRegister.is_active == True).first()
    if not cash:
        return RedirectResponse(url="/employees/advances?error=Kassani tanlang", status_code=303)
    adv = EmployeeAdvance(
        employee_id=employee_id,
        amount=amount,
        advance_date=adv_date,
        cash_register_id=cash.id,
        note=note or None,
    )
    db.add(adv)
    db.flush()
    # Kassadan chiqim — kassa balansini kamaytirish
    current_balance = float(cash.balance or 0)
    new_balance = current_balance - amount
    today = datetime.now()
    count = db.query(CashBalanceDoc).filter(
        CashBalanceDoc.date >= today.replace(hour=0, minute=0, second=0)
    ).count()
    doc_number = f"KLD-{today.strftime('%Y%m%d')}-{str(count + 1).zfill(4)}"
    doc = CashBalanceDoc(
        number=doc_number,
        date=today,
        user_id=current_user.id if current_user else None,
        status="confirmed",
    )
    db.add(doc)
    db.flush()
    db.add(CashBalanceDocItem(doc_id=doc.id, cash_register_id=cash.id, balance=new_balance, previous_balance=current_balance))
    cash.balance = new_balance
    adv.confirmed_at = today
    db.commit()
    return RedirectResponse(url="/employees/advances?added=1", status_code=303)


@app.get("/employees/advances/edit/{advance_id}", response_class=HTMLResponse)
async def employee_advance_edit_page(
    advance_id: int,
    request: Request,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    """Avansni tahrirlash sahifasi — faqat tasdiqlanmagan avanslar."""
    adv = db.query(EmployeeAdvance).options(
        joinedload(EmployeeAdvance.employee),
        joinedload(EmployeeAdvance.cash_register),
    ).filter(EmployeeAdvance.id == advance_id).first()
    if not adv:
        return RedirectResponse(url="/employees/advances?error=Avans topilmadi", status_code=303)
    if adv.confirmed_at:
        return RedirectResponse(
            url="/employees/advances?error=" + quote("Tasdiqlangan avansni tahrirlash mumkin emas. Avval tasdiqni bekor qiling."),
            status_code=303,
        )
    employees = db.query(Employee).filter(Employee.is_active == True).order_by(Employee.full_name).all()
    if adv.employee and not any(e.id == adv.employee_id for e in employees):
        employees = [adv.employee] + list(employees)
    cash_registers = db.query(CashRegister).filter(CashRegister.is_active == True).order_by(CashRegister.name).all()
    next_ids = (request.query_params.get("next_ids") or "").strip()
    next_count = len([x for x in next_ids.split(",") if x.strip()]) if next_ids else 0
    return templates.TemplateResponse("employees/advance_edit.html", {
        "request": request,
        "advance": adv,
        "employees": employees,
        "cash_registers": cash_registers,
        "current_user": current_user,
        "page_title": "Avansni tahrirlash",
        "next_ids": next_ids,
        "next_count": next_count,
    })


@app.post("/employees/advances/edit/{advance_id}")
async def employee_advance_edit_save(
    advance_id: int,
    request: Request,
    employee_id: int = Form(...),
    amount: float = Form(...),
    advance_date: str = Form(...),
    cash_register_id: Optional[int] = Form(None),
    note: str = Form(""),
    next_ids: str = Form(""),
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    """Avansni saqlash (tahrirlash) — faqat tasdiqlanmagan avanslar."""
    adv = db.query(EmployeeAdvance).filter(EmployeeAdvance.id == advance_id).first()
    if not adv:
        return RedirectResponse(url="/employees/advances?error=Avans topilmadi", status_code=303)
    if adv.confirmed_at:
        return RedirectResponse(
            url="/employees/advances?error=" + quote("Tasdiqlangan avansni tahrirlash mumkin emas."),
            status_code=303,
        )
    try:
        adv_date = datetime.strptime(advance_date, "%Y-%m-%d").date()
    except ValueError:
        return RedirectResponse(url=f"/employees/advances/edit/{advance_id}?error=Noto'g'ri sana", status_code=303)
    if amount <= 0:
        return RedirectResponse(url=f"/employees/advances/edit/{advance_id}?error=Summa 0 dan katta bo'lishi kerak", status_code=303)
    emp = db.query(Employee).filter(Employee.id == employee_id).first()
    if not emp:
        return RedirectResponse(url=f"/employees/advances/edit/{advance_id}?error=Xodim topilmadi", status_code=303)
    if cash_register_id:
        cash = db.query(CashRegister).filter(CashRegister.id == cash_register_id, CashRegister.is_active == True).first()
        adv.cash_register_id = cash.id if cash else adv.cash_register_id
    else:
        adv.cash_register_id = None
    adv.employee_id = employee_id
    adv.amount = amount
    adv.advance_date = adv_date
    adv.note = note or None
    adv.confirmed_at = datetime.now()  # Tahrirlashda saqlash = tasdiqlash
    db.commit()
    # Ketma-ket tahrirlash: keyingi avansga yo'naltirish
    next_param = (next_ids or "").strip()
    if next_param:
        rest = [x.strip() for x in next_param.split(",") if x.strip()]
        if rest:
            try:
                next_id = int(rest[0])
                remaining = ",".join(rest[1:])
                url = f"/employees/advances/edit/{next_id}"
                if remaining:
                    url += "?next_ids=" + remaining
                return RedirectResponse(url=url, status_code=303)
            except (ValueError, TypeError):
                pass
    return RedirectResponse(url="/employees/advances?edited=1", status_code=303)


@app.post("/employees/advances/confirm/{advance_id}")
async def employee_advance_confirm(
    advance_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    """Avansni tasdiqlash"""
    adv = db.query(EmployeeAdvance).filter(EmployeeAdvance.id == advance_id).first()
    if not adv:
        return RedirectResponse(url="/employees/advances?error=Avans topilmadi", status_code=303)
    adv.confirmed_at = datetime.now()
    db.commit()
    return RedirectResponse(url="/employees/advances?confirmed=1", status_code=303)


@app.post("/employees/advances/unconfirm/{advance_id}")
async def employee_advance_unconfirm(
    advance_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    """Avans tasdiqini bekor qilish"""
    adv = db.query(EmployeeAdvance).filter(EmployeeAdvance.id == advance_id).first()
    if not adv:
        return RedirectResponse(url="/employees/advances?error=Avans topilmadi", status_code=303)
    adv.confirmed_at = None
    db.commit()
    return RedirectResponse(url="/employees/advances?unconfirmed=1", status_code=303)


@app.post("/employees/advances/delete/{advance_id}")
async def employee_advance_delete(
    advance_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    """Avansni ro'yxatdan o'chirish — faqat tasdiqlanmagan (tasdiq bekor qilingan) avanslar."""
    adv = db.query(EmployeeAdvance).filter(EmployeeAdvance.id == advance_id).first()
    if not adv:
        return RedirectResponse(url="/employees/advances?error=Avans topilmadi", status_code=303)
    if adv.confirmed_at:
        return RedirectResponse(
            url="/employees/advances?error=" + quote("Tasdiqlangan avansni o'chirish mumkin emas. Avval tasdiqni bekor qiling."),
            status_code=303,
        )
    db.delete(adv)
    db.commit()
    return RedirectResponse(url="/employees/advances?deleted=1", status_code=303)


@app.post("/employees/advances/bulk-edit", response_class=RedirectResponse)
async def employee_advances_bulk_edit(
    request: Request,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    """Tanlangan tasdiqlanmagan avanslarni ketma-ket tahrirlash — birinchisiga yo'naltiradi."""
    form = await request.form()
    raw = form.getlist("advance_ids")
    ids = []
    for x in raw:
        try:
            ids.append(int(x))
        except (TypeError, ValueError):
            pass
    if not ids:
        return RedirectResponse(url="/employees/advances?error=" + quote("Hech qaysi avans tanlanmagan."), status_code=303)
    unconfirmed = (
        db.query(EmployeeAdvance.id)
        .filter(EmployeeAdvance.id.in_(ids), EmployeeAdvance.confirmed_at.is_(None))
        .order_by(EmployeeAdvance.id)
        .all()
    )
    unconfirmed_ids = [r[0] for r in unconfirmed]
    if not unconfirmed_ids:
        return RedirectResponse(url="/employees/advances?error=" + quote("Tanlangan avanslar tasdiqlangan. Faqat tasdiqlanmagan avanslarni tahrirlash mumkin."), status_code=303)
    first_id = unconfirmed_ids[0]
    next_ids = unconfirmed_ids[1:]
    next_param = ",".join(str(i) for i in next_ids) if next_ids else ""
    url = f"/employees/advances/edit/{first_id}"
    if next_param:
        url += "?next_ids=" + next_param
    return RedirectResponse(url=url, status_code=303)


@app.post("/employees/advances/bulk-unconfirm", response_class=RedirectResponse)
async def employee_advances_bulk_unconfirm(
    request: Request,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    """Tanlangan tasdiqlangan avanslarning tasdiqini bekor qilish"""
    form = await request.form()
    raw = form.getlist("advance_ids")
    ids = []
    for x in raw:
        try:
            ids.append(int(x))
        except (TypeError, ValueError):
            pass
    if not ids:
        return RedirectResponse(url="/employees/advances?error=" + quote("Hech qaysi avans tanlanmagan."), status_code=303)
    updated = db.query(EmployeeAdvance).filter(EmployeeAdvance.id.in_(ids), EmployeeAdvance.confirmed_at.isnot(None)).update({EmployeeAdvance.confirmed_at: None}, synchronize_session=False)
    db.commit()
    base = "/employees/advances?bulk_unconfirmed=" + str(updated)
    extra = _advances_list_redirect_params(form)
    return RedirectResponse(url=base + ("&" + extra if extra else ""), status_code=303)


@app.post("/employees/advances/bulk-confirm", response_class=RedirectResponse)
async def employee_advances_bulk_confirm(
    request: Request,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    """Tanlangan tasdiqlanmagan avanslarni tasdiqlash"""
    form = await request.form()
    raw = form.getlist("advance_ids")
    ids = []
    for x in raw:
        try:
            ids.append(int(x))
        except (TypeError, ValueError):
            pass
    if not ids:
        return RedirectResponse(url="/employees/advances?error=" + quote("Hech qaysi avans tanlanmagan."), status_code=303)
    now = datetime.now()
    updated = db.query(EmployeeAdvance).filter(EmployeeAdvance.id.in_(ids), EmployeeAdvance.confirmed_at.is_(None)).update({EmployeeAdvance.confirmed_at: now}, synchronize_session=False)
    db.commit()
    base = "/employees/advances?bulk_confirmed=" + str(updated)
    extra = _advances_list_redirect_params(form)
    return RedirectResponse(url=base + ("&" + extra if extra else ""), status_code=303)


@app.post("/employees/advances/bulk-delete", response_class=RedirectResponse)
async def employee_advances_bulk_delete(
    request: Request,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    """Tanlangan tasdiqlanmagan avanslarni o'chirish (tasdiqlanganlarni o'chirish mumkin emas)."""
    form = await request.form()
    raw = form.getlist("advance_ids")
    ids = []
    for x in raw:
        try:
            ids.append(int(x))
        except (TypeError, ValueError):
            pass
    if not ids:
        base = "/employees/advances?error=" + quote("Hech qaysi avans tanlanmagan.")
        extra = _advances_list_redirect_params(form)
        return RedirectResponse(url=base + ("&" + extra if extra else ""), status_code=303)
    deleted = db.query(EmployeeAdvance).filter(EmployeeAdvance.id.in_(ids), EmployeeAdvance.confirmed_at.is_(None)).delete(synchronize_session=False)
    db.commit()
    if ids and deleted == 0:
        base = "/employees/advances?error=" + quote("Tanlangan avanslar tasdiqlangan. O'chirish uchun avval tasdiqni bekor qiling.")
    else:
        base = "/employees/advances?bulk_deleted=" + str(deleted)
    extra = _advances_list_redirect_params(form)
    return RedirectResponse(url=base + ("&" + extra if extra else ""), status_code=303)


# ==========================================
# OYLIK HISOBLASH
# ==========================================

@app.get("/employees/salary", response_class=HTMLResponse)
async def employee_salary_page(
    request: Request,
    year: Optional[int] = None,
    month: Optional[int] = None,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    """Oylik hisoblash — oy tanlash, xodimlar ro'yxati (base, bonus, deduction, avans, total)"""
    today = date.today()
    year = year or today.year
    month = month or today.month
    if not (1 <= month <= 12):
        month = today.month
    if year < 2020 or year > 2030:
        year = today.year
    # Faqat ishga qabul hujjati bor xodimlar ro'yxatda ko'rinadi
    hired_employee_ids = db.query(EmploymentDoc.employee_id).distinct().all()
    hired_ids = [r[0] for r in hired_employee_ids if r[0]]
    if not hired_ids:
        employees = []
    else:
        employees = (
            db.query(Employee)
            .filter(Employee.is_active == True, Employee.id.in_(hired_ids))
            .order_by(Employee.full_name)
            .all()
        )
    salaries = {s.employee_id: s for s in db.query(Salary).filter(Salary.year == year, Salary.month == month).all()}
    # Ishga qabul hujjatidagi oylik — avval tasdiqlangan, keyin har qanday oxirgi hujjat (qadoqlovchilar va b. uchun)
    emp_ids = [e.id for e in employees]
    latest_doc_salary = {}
    if emp_ids:
        # Har bir xodim uchun tasdiqlangan hujjatlar ichidan eng oxirgi (max doc_date) dagi oylik
        subq_conf = (
            db.query(EmploymentDoc.employee_id, func.max(EmploymentDoc.doc_date).label("max_date"))
            .filter(EmploymentDoc.employee_id.in_(emp_ids), EmploymentDoc.confirmed_at.isnot(None))
            .group_by(EmploymentDoc.employee_id)
        ).subquery()
        docs_confirmed = (
            db.query(EmploymentDoc.employee_id, EmploymentDoc.salary)
            .join(subq_conf, (EmploymentDoc.employee_id == subq_conf.c.employee_id) & (EmploymentDoc.doc_date == subq_conf.c.max_date))
            .filter(EmploymentDoc.employee_id.in_(emp_ids), EmploymentDoc.confirmed_at.isnot(None))
            .all()
        )
        for row in docs_confirmed:
            if (row.salary or 0) > 0:
                latest_doc_salary[row.employee_id] = float(row.salary)
        # Asos hali 0 bo'lgan xodimlar uchun: har bir xodimning eng oxirgi hujjatidagi oylik (max doc_date)
        missing = [eid for eid in emp_ids if eid not in latest_doc_salary]
        if missing:
            subq = (
                db.query(EmploymentDoc.employee_id, func.max(EmploymentDoc.doc_date).label("max_date"))
                .filter(EmploymentDoc.employee_id.in_(missing))
                .group_by(EmploymentDoc.employee_id)
            ).subquery()
            docs_latest = (
                db.query(EmploymentDoc.employee_id, EmploymentDoc.salary)
                .join(subq, (EmploymentDoc.employee_id == subq.c.employee_id) & (EmploymentDoc.doc_date == subq.c.max_date))
                .filter(EmploymentDoc.employee_id.in_(missing))
                .all()
            )
            for row in docs_latest:
                if (row.salary or 0) > 0:
                    latest_doc_salary[row.employee_id] = float(row.salary)
    # Avanslar (shu oy berilgan) — hisoblash uchun
    advance_sums = {}
    from calendar import monthrange
    _, last_day = monthrange(year, month)
    start_d = date(year, month, 1)
    end_d = date(year, month, last_day)
    for a in db.query(EmployeeAdvance).filter(
        EmployeeAdvance.advance_date >= start_d,
        EmployeeAdvance.advance_date <= end_d,
    ).all():
        advance_sums[a.employee_id] = advance_sums.get(a.employee_id, 0) + a.amount
    # Oylik turi xodimlar uchun: tabel bo'yicha ishlagan kunlar — kunlik = oylik/oydagi kunlar, asos = kunlik * ishlagan kun
    worked_days_by_emp = {}
    if emp_ids:
        try:
            worked_rows = (
                db.query(Attendance.employee_id, func.count(func.distinct(Attendance.date)).label("days"))
                .filter(
                    Attendance.employee_id.in_(emp_ids),
                    Attendance.date >= start_d,
                    Attendance.date <= end_d,
                    or_(
                        Attendance.status == "present",
                        Attendance.check_in.isnot(None),
                    ),
                )
                .group_by(Attendance.employee_id)
                .all()
            )
            for r in worked_rows:
                worked_days_by_emp[r.employee_id] = int(r.days or 0)
        except Exception:
            pass
    days_in_month = last_day
    # Bo'lak ish haqi: ishlab chiqarilgan miqdor * (bitta bo'lak narxi). Bitta stavka ishlatiladi (min), yig'indi emas.
    piecework_calculated = {}
    emp_by_id = {e.id: e for e in employees}
    group_member_ids = set()
    # Ishlab chiqarish guruhlari (qiyomchilar): operator ishi kunlik tabel bo'yicha kелgan a'zolar orasida teng bo'linadi
    production_groups = (
        db.query(ProductionGroup)
        .options(joinedload(ProductionGroup.members), joinedload(ProductionGroup.piecework_task))
        .filter(ProductionGroup.is_active == True, ProductionGroup.operator_id.in_(emp_ids) if emp_ids else False)
        .all()
    )
    for gr in production_groups:
        member_ids = [m.id for m in gr.members if m.id in emp_ids] if hasattr(gr, "members") and gr.members else []
        if not member_ids or gr.operator_id not in emp_ids:
            continue
        group_member_ids.update(member_ids)
        rate = float(gr.piecework_task.price_per_unit or 0) if gr.piecework_task else 0
        if rate <= 0:
            continue
        prod_list = (
            db.query(Production)
            .options(joinedload(Production.recipe))
            .filter(
                Production.operator_id == gr.operator_id,
                Production.status == "completed",
                func.date(Production.date) >= start_d,
                func.date(Production.date) <= end_d,
            )
            .all()
        )
        day_kg = {}
        for p in prod_list:
            if not getattr(gr, "include_qiyom", True) and _is_qiyom_recipe(p.recipe):
                continue
            kg = (float(p.quantity or 0) * _kg_per_unit_from_recipe(p.recipe)) if p.recipe else 0
            if kg <= 0:
                continue
            d = p.date.date() if hasattr(p.date, "date") else p.date
            day_kg[d] = day_kg.get(d, 0) + kg
        attendances = (
            db.query(Attendance.employee_id, Attendance.date, Attendance.status, Attendance.check_in)
            .filter(Attendance.employee_id.in_(member_ids), Attendance.date >= start_d, Attendance.date <= end_d)
            .all()
        )
        present_by_date = {}
        for row in attendances:
            d = row.date
            if d not in present_by_date:
                present_by_date[d] = set()
            if (row.status or "").strip() == "present" or (getattr(row, "check_in", None) is not None):
                present_by_date[d].add(row.employee_id)
        member_kg = {mid: 0.0 for mid in member_ids}
        for d, kg in day_kg.items():
            present = present_by_date.get(d, set()) & set(member_ids)
            cnt = len(present)
            if cnt <= 0:
                continue
            per_person = kg / cnt
            for mid in present:
                member_kg[mid] = member_kg.get(mid, 0) + per_person
        for mid in member_ids:
            piecework_calculated[mid] = member_kg.get(mid, 0) * rate
    # Har bir xodim uchun bitta bo'lak narxi (min stavka — bir xil ish uchun bitta narx)
    piece_rate_sum = {}
    if emp_ids:
        rows_rates = (
            db.query(employee_piecework_tasks.c.employee_id, func.min(PieceworkTask.price_per_unit).label("rate"))
            .join(PieceworkTask, PieceworkTask.id == employee_piecework_tasks.c.task_id)
            .filter(employee_piecework_tasks.c.employee_id.in_(emp_ids))
            .filter(PieceworkTask.price_per_unit > 0)
            .group_by(employee_piecework_tasks.c.employee_id)
            .all()
        )
        for eid, rate in rows_rates:
            piece_rate_sum[int(eid)] = float(rate or 0)
    # Bo'lak va Bo'lak+oylik xodimlar, lekin employee_piecework_tasks da yo'q: stavkani ishga qabul hujjatidagi piecework_task_ids dan olaymiz
    boalak_emp_ids = [e.id for e in employees if getattr(e, "salary_type", None) in ("bo'lak", "bo'lak_oylik")]
    if boalak_emp_ids:
        docs_with_tasks = (
            db.query(EmploymentDoc.employee_id, EmploymentDoc.piecework_task_ids)
            .filter(EmploymentDoc.employee_id.in_(boalak_emp_ids), EmploymentDoc.confirmed_at.isnot(None))
            .order_by(EmploymentDoc.doc_date.desc())
            .all()
        )
        for row in docs_with_tasks:
            eid = row.employee_id
            if piece_rate_sum.get(eid, 0) > 0:
                continue
            raw = (row.piecework_task_ids or "").strip()
            ids = [int(x) for x in raw.split(",") if x.strip().isdigit()] if raw else []
            if not ids:
                continue
            first_task = db.query(PieceworkTask).filter(PieceworkTask.id == ids[0], PieceworkTask.price_per_unit > 0).first()
            if first_task:
                piece_rate_sum[eid] = float(first_task.price_per_unit)
    boalak_employees = [e for e in employees if getattr(e, "salary_type", None) in ("bo'lak", "bo'lak_oylik") and piece_rate_sum.get(e.id, 0) > 0]
    emp_by_id = {e.id: e for e in employees}
    user_to_employee_id = {}
    for e in employees:
        if e.user_id:
            user_to_employee_id[e.user_id] = e.id
    # Bo'lak ish haqi: Operator bo'yicha ishlab chiqarishdagi kabi KG hisoblanadi (qiyom hisobga olinmaydi), keyin kg * bo'lak narxi
    if boalak_employees:
        productions_for_salary = (
            db.query(Production)
            .options(joinedload(Production.recipe))
            .filter(
                Production.status == "completed",
                func.date(Production.date) >= start_d,
                func.date(Production.date) <= end_d,
            )
            .all()
        )
        total_kg_by_emp_id = {}
        for p in productions_for_salary:
            if _is_qiyom_recipe(p.recipe):
                continue
            kg = (float(p.quantity or 0) * _kg_per_unit_from_recipe(p.recipe)) if p.recipe else 0
            if kg <= 0:
                continue
            emp_id = None
            if p.operator_id and p.operator_id in emp_by_id:
                emp_id = p.operator_id
            elif p.user_id and p.user_id in user_to_employee_id:
                emp_id = user_to_employee_id[p.user_id]
            if emp_id and emp_id not in group_member_ids:
                total_kg_by_emp_id[emp_id] = total_kg_by_emp_id.get(emp_id, 0) + kg
        for emp in boalak_employees:
            if emp.id in group_member_ids:
                continue
            total_kg = total_kg_by_emp_id.get(emp.id, 0)
            rate = piece_rate_sum.get(emp.id, 0)
            if total_kg > 0 and rate > 0:
                piecework_calculated[emp.id] = total_kg * rate
    rows = []
    for emp in employees:
        s = salaries.get(emp.id)
        piecework_amount = float(piecework_calculated.get(emp.id, 0) or 0)
        base_source = ""  # "oylik" | "bo'lak" — asos qaysi manbadan olingan
        # Guruh a'zosi (qiyomchilar): guruh bo'lak ulushi bo'yicha, asos = max(mehnat haqi, bo'lak)
        if emp.id in group_member_ids and emp.id in piecework_calculated:
            mehnat_haqi = float(latest_doc_salary.get(emp.id, 0) or 0) or float(emp.salary or 0)
            piece_total = piecework_amount
            base = max(mehnat_haqi, piece_total)
            base_source = "bo'lak" if piece_total >= mehnat_haqi and piece_total > 0 else "oylik"
        elif getattr(emp, "salary_type", None) == "bo'lak":
            base = piecework_amount
            base_source = "bo'lak" if piecework_amount > 0 else ""
        elif getattr(emp, "salary_type", None) == "bo'lak_oylik":
            mehnat_haqi = float(latest_doc_salary.get(emp.id, 0) or 0) or float(emp.salary or 0)
            piece_total = piecework_amount
            base = max(mehnat_haqi, piece_total)
            base_source = "bo'lak" if piece_total >= mehnat_haqi and piece_total > 0 else "oylik"
        else:
            base = (s.base_salary if s else 0) or (emp.salary or 0) or latest_doc_salary.get(emp.id, 0)
            if not base and emp.id in piecework_calculated:
                base = piecework_calculated[emp.id]
            base = float(base or 0)
            if getattr(emp, "salary_type", None) in ("oylik", "soatlik") or not getattr(emp, "salary_type", None):
                base_source = "oylik" if base > 0 else ""
        base = float(base or 0)
        # Hisoblangan oylik (tabel bo'yicha): Oylik turi — doim; Bo'lak+oylik / guruh a'zosi — faqat asos "oylikdan" bo'lsa
        calculated_base = None
        if days_in_month and days_in_month > 0:
            contract_monthly = float(latest_doc_salary.get(emp.id, 0) or 0) or float(emp.salary or 0)
            worked_days = worked_days_by_emp.get(emp.id, 0) or 0
            if getattr(emp, "salary_type", None) == "oylik":
                calculated_base = round((contract_monthly / days_in_month) * worked_days, 2)
            elif base_source == "oylik" and contract_monthly > 0:
                # Bo'lak+oylik yoki guruh a'zosi, asos oylikdan — tabel bo'yicha hisoblangan oylik
                calculated_base = round((contract_monthly / days_in_month) * worked_days, 2)
        amount_for_total = calculated_base if calculated_base is not None else base
        bonus = float(s.bonus if s and s.bonus is not None else 0) or 0
        deduction = float(s.deduction if s and s.deduction is not None else 0) or 0
        adv_ded = None
        if s and getattr(s, "advance_deduction", None) is not None:
            adv_ded = float(s.advance_deduction)
        if adv_ded is None:
            adv_ded = float(advance_sums.get(emp.id, 0) or 0)
        total = amount_for_total + bonus - deduction - adv_ded
        total = round(total, 2)
        paid = float(s.paid if s and s.paid is not None else 0) or 0
        status = (s.status if s else "pending") or "pending"
        if total == 0 and paid == 0:
            status = "pending"
        elif (total or 0) > 0 and paid >= total:
            status = "paid"
        elif (total or 0) > 0:
            status = "pending"
        rows.append({
            "employee": emp,
            "salary_row": s,
            "base_salary": base,
            "calculated_base": calculated_base,
            "piecework_amount": piecework_amount,
            "base_source": base_source,
            "bonus": bonus,
            "deduction": deduction,
            "advance_deduction": adv_ded,
            "total": total,
            "paid": paid,
            "status": status,
        })
    cash_doc_id = request.query_params.get("cash_doc")
    try:
        cash_doc_id = int(cash_doc_id) if cash_doc_id else None
    except (TypeError, ValueError):
        cash_doc_id = None
    no_cash_warn = request.query_params.get("no_cash") == "1"
    return templates.TemplateResponse("employees/salary_list.html", {
        "request": request,
        "year": year,
        "month": month,
        "rows": rows,
        "current_user": current_user,
        "page_title": "Oylik hisoblash",
        "cash_doc_id": cash_doc_id,
        "no_cash_warn": no_cash_warn,
    })


@app.post("/employees/salary/save")
async def employee_salary_save(
    request: Request,
    year: int = Form(...),
    month: int = Form(...),
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    """Oylik yozuvlarini saqlash; jami to'lov summasiga kassadan chiqim hujjati yaratiladi va tasdiqlanadi. Faqat ishga qabul qilingan xodimlar."""
    if not (1 <= month <= 12) or year < 2020 or year > 2030:
        return RedirectResponse(url="/employees/salary?error=Noto'g'ri oy yoki yil", status_code=303)
    form = await request.form()
    hired_ids = [r[0] for r in db.query(EmploymentDoc.employee_id).distinct().all() if r[0]]
    if not hired_ids:
        employees = []
    else:
        employees = db.query(Employee).filter(Employee.is_active == True, Employee.id.in_(hired_ids)).all()
    total_payroll = 0.0
    for emp in employees:
        base = float(form.get(f"base_{emp.id}", 0) or 0)
        bonus = float(form.get(f"bonus_{emp.id}", 0) or 0)
        deduction = float(form.get(f"deduction_{emp.id}", 0) or 0)
        advance_deduction = float(form.get(f"advance_{emp.id}", 0) or 0)
        total = base + bonus - deduction - advance_deduction
        total_payroll += max(0, float(total))  # Kassadan faqat musbat to'lovlar chiqadi
        s = db.query(Salary).filter(Salary.employee_id == emp.id, Salary.year == year, Salary.month == month).first()
        if not s:
            s = Salary(employee_id=emp.id, year=year, month=month)
            db.add(s)
        s.base_salary = base
        s.bonus = bonus
        s.deduction = deduction
        s.advance_deduction = advance_deduction
        s.total = total
        if s.paid is None:
            s.paid = 0
        s.status = "paid" if (s.paid or 0) >= total else "pending"
    db.commit()
    # Oylik vedomost bo'yicha kassaga yo'naltirish — kassadan chiqim hujjati yaratish va tasdiqlash
    cash_doc_id = None
    no_cash_warn = False
    if total_payroll > 0:
        cash = db.query(CashRegister).filter(CashRegister.is_active == True).order_by(CashRegister.id).first()
        if cash:
            current_balance = float(cash.balance or 0)
            new_balance = current_balance - total_payroll
            today = datetime.now()
            count = db.query(CashBalanceDoc).filter(
                CashBalanceDoc.date >= today.replace(hour=0, minute=0, second=0)
            ).count()
            number = f"KLD-{today.strftime('%Y%m%d')}-{str(count + 1).zfill(4)}"
            doc = CashBalanceDoc(
                number=number,
                date=today,
                user_id=current_user.id if current_user else None,
                status="draft",
            )
            db.add(doc)
            db.flush()
            item = CashBalanceDocItem(doc_id=doc.id, cash_register_id=cash.id, balance=new_balance, previous_balance=current_balance)
            db.add(item)
            db.flush()
            cash.balance = new_balance
            doc.status = "confirmed"
            db.commit()
            cash_doc_id = doc.id
        else:
            no_cash_warn = True
    params = f"year={year}&month={month}&saved=1"
    if cash_doc_id:
        params += f"&cash_doc={cash_doc_id}"
    if no_cash_warn:
        params += "&no_cash=1"
    return RedirectResponse(url=f"/employees/salary?{params}", status_code=303)


@app.post("/employees/salary/mark-paid/{employee_id}")
async def employee_salary_mark_paid(
    request: Request,
    employee_id: int,
    year: int = Form(...),
    month: int = Form(...),
    paid_amount: float = Form(...),
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    """Oylik to'langanligini belgilash"""
    s = db.query(Salary).filter(
        Salary.employee_id == employee_id,
        Salary.year == year,
        Salary.month == month,
    ).first()
    if not s:
        s = Salary(employee_id=employee_id, year=year, month=month, base_salary=0, total=0, paid=0)
        db.add(s)
    s.paid = paid_amount
    s.status = "paid" if paid_amount >= (s.total or 0) else "pending"
    db.commit()
    return RedirectResponse(url=f"/employees/salary?year={year}&month={month}", status_code=303)


# ==========================================
# ISHLAB CHIQARISH
# ==========================================

@app.get("/production", response_class=HTMLResponse)
async def production_index_page(request: Request, db: Session = Depends(get_db), current_user: User = Depends(require_auth)):
    """Ishlab chiqarish bosh sahifasi — retseptni yozib qidirish va tanlash uchun to'liq ro'yxat."""
    warehouses = db.query(Warehouse).all()
    recipes = (
        db.query(Recipe)
        .options(
            joinedload(Recipe.product).joinedload(Product.unit),
        )
        .filter(Recipe.is_active == True)
        .all()
    )

    today = datetime.now().date()
    total_recipes = db.query(Recipe).filter(Recipe.is_active == True).count()
    # Operator/ishlab chiqarish rollari uchun faqat o'zining; admin/manager uchun barcha buyurtmalar
    _operator_roles = ("production", "qadoqlash", "rahbar", "raxbar", "operator")
    role = (getattr(current_user, "role", None) or "").strip().lower() if current_user else ""
    filter_by_user = role in _operator_roles
    current_user_employee = None
    if current_user and getattr(current_user, "id", None):
        current_user_employee = db.query(Employee).filter(Employee.user_id == current_user.id).first()
    user_or_operator_filter = None
    if filter_by_user and current_user and getattr(current_user, "id", None):
        if current_user_employee:
            user_or_operator_filter = or_(
                Production.user_id == current_user.id,
                Production.operator_id == current_user_employee.id,
            )
        else:
            user_or_operator_filter = Production.user_id == current_user.id

    today_start = datetime.combine(today, datetime.min.time())
    today_productions = (
        db.query(Production)
        .options(joinedload(Production.recipe))
        .filter(
            Production.date >= today_start,
            Production.status == "completed"
        )
    )
    if user_or_operator_filter is not None:
        today_productions = today_productions.filter(user_or_operator_filter)
    today_productions = today_productions.all()
    # Qiyom (oralama mahsulot) hisobga olinmaydi — faqat yarim tayyor / tayyor / dona; kg ga o'tkazamiz
    today_quantity = 0.0
    for p in today_productions:
        if _is_qiyom_recipe(p.recipe):
            continue
        kg_per = _kg_per_unit_from_recipe(p.recipe) if p.recipe else 1.0
        today_quantity += (p.quantity or 0) * (kg_per if kg_per and kg_per > 0 else 1.0)
    pending_productions = db.query(Production).filter(Production.status == "draft")
    if user_or_operator_filter is not None:
        pending_productions = pending_productions.filter(user_or_operator_filter)
    pending_productions = pending_productions.count()
    recent_qry = (
        db.query(Production)
        .options(
            joinedload(Production.recipe).joinedload(Recipe.product).joinedload(Product.unit),
        )
    )
    if user_or_operator_filter is not None:
        recent_qry = recent_qry.filter(user_or_operator_filter)
    recent_productions = recent_qry.order_by(Production.date.desc()).limit(10).all()

    # Operator / ishlab chiqarish / qadoqlash uchun "Retseptlar" ro'yxati blokini yashirish
    _operator_type_roles = ("operator", "production", "qadoqlash")
    show_recipes_section = not (
        current_user
        and getattr(current_user, "role", None)
        and (current_user.role or "").strip().lower() in _operator_type_roles
    )

    now = datetime.now()
    resp = templates.TemplateResponse("production/index.html", {
        "request": request,
        "current_user": current_user,
        "total_recipes": total_recipes,
        "today_quantity": today_quantity,
        "pending_productions": pending_productions,
        "recent_productions": recent_productions,
        "recipes": recipes,
        "warehouses": warehouses,
        "show_recipes_section": show_recipes_section,
        "page_title": "Ishlab chiqarish",
        "now": now,
    })
    resp.headers["Cache-Control"] = "no-store, no-cache, must-revalidate"
    return resp


@app.get("/production/api/quick-recipes")
async def production_api_quick_recipes(
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    """Tezkor ishlab chiqarish uchun retseptlar ro'yxati (JSON). Sahifada quickRecipes bo'sh bo'lsa brauzer shu API orqali yuklaydi."""
    recipes = (
        db.query(Recipe)
        .options(
            joinedload(Recipe.product).joinedload(Product.unit),
        )
        .filter(Recipe.is_active == True)
        .all()
    )
    out = []
    for r in recipes:
        unit = "kg"
        if r.product and getattr(r.product, "unit", None):
            u = r.product.unit
            unit = (getattr(u, "name", None) or getattr(u, "code", None) or "kg") or "kg"
        out.append({
            "id": r.id,
            "name": r.name or "",
            "unit": unit,
            "wh": str(r.default_warehouse_id) if r.default_warehouse_id else "",
            "whOut": str(r.default_output_warehouse_id) if r.default_output_warehouse_id else "",
        })
    return out


@app.get("/production/recipes", response_class=HTMLResponse)
async def production_recipes(request: Request, db: Session = Depends(get_db), current_user: User = Depends(require_auth)):
    """Retseptlar ro'yxati — retsept nomi bo'yicha qidiruv (q)."""
    import json
    q = (request.query_params.get("q") or "").strip()
    warehouses = db.query(Warehouse).all()
    query = db.query(Recipe).options(
        joinedload(Recipe.product).joinedload(Product.unit),
        joinedload(Recipe.items).joinedload(RecipeItem.product).joinedload(Product.unit),
    )
    if q:
        search = f"%{q}%"
        query = query.filter(Recipe.name.ilike(search))
    recipes = query.order_by(Recipe.name).all()
    products = db.query(Product).filter(Product.type.in_(["tayyor", "yarim_tayyor"])).all()
    materials = db.query(Product).filter(Product.type == "hom_ashyo").all()
    recipe_products_json = json.dumps([
        {"id": p.id, "name": (p.name or ""), "unit": (p.unit.name if p.unit else "kg")}
        for p in products
    ]).replace("<", "\\u003c")

    return templates.TemplateResponse("production/recipes.html", {
        "request": request,
        "current_user": current_user,
        "recipes": recipes,
        "products": products,
        "recipe_products_json": recipe_products_json,
        "materials": materials,
        "warehouses": warehouses,
        "page_title": "Retseptlar",
        "search_q": q,
    })


def _recipe_detail_redirect_url(recipe_id: int, return_q: Optional[str] = None) -> str:
    """Retsept tafsilotiga redirect URL — qidiruv parametri q bo'lsa saqlanadi."""
    url = f"/production/recipes/{recipe_id}"
    if return_q and str(return_q).strip():
        from urllib.parse import quote
        url += "?q=" + quote(str(return_q).strip())
    return url


@app.get("/production/recipes/{recipe_id}", response_class=HTMLResponse)
async def production_recipe_detail(request: Request, recipe_id: int, db: Session = Depends(get_db), current_user: User = Depends(get_current_user)):
    """Retsept tafsilotlari — omborlar ro'yxati «Ishlab chiqarishda ishlatiladigan omborlar» uchun."""
    recipe = db.query(Recipe).filter(Recipe.id == recipe_id).first()
    if not recipe:
        raise HTTPException(status_code=404, detail="Retsept topilmadi")
    search_q = (request.query_params.get("q") or "").strip()
    materials = db.query(Product).filter(Product.type.in_(["hom_ashyo", "yarim_tayyor", "tayyor"])).all()
    try:
        recipe_stages = sorted(recipe.stages, key=lambda s: s.stage_number) if recipe.stages else []
    except Exception:
        recipe_stages = []
    warehouses = db.query(Warehouse).filter(Warehouse.is_active == True).all()
    return templates.TemplateResponse("production/recipe_detail.html", {
        "request": request,
        "current_user": current_user,
        "recipe": recipe,
        "materials": materials,
        "recipe_stages": recipe_stages,
        "warehouses": warehouses,
        "page_title": f"Retsept: {recipe.name}",
        "search_q": search_q,
    })


@app.post("/production/recipes/add")
async def add_recipe(
    request: Request,
    name: str = Form(...),
    product_id: int = Form(...),
    output_quantity: float = Form(1),
    description: str = Form(""),
    return_q: str = Form(""),
    db: Session = Depends(get_db)
):
    """Yangi retsept qo'shish"""
    return_q = (return_q or "").strip()
    recipe = Recipe(
        name=name,
        product_id=product_id,
        output_quantity=output_quantity,
        description=description,
        is_active=True
    )
    db.add(recipe)
    db.commit()
    return RedirectResponse(url=_recipe_detail_redirect_url(recipe.id, return_q), status_code=303)


@app.post("/production/recipes/{recipe_id}/add-item")
async def add_recipe_item(
    recipe_id: int,
    product_id: int = Form(...),
    quantity: float = Form(...),
    return_q: str = Form(""),
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth)
):
    """Retseptga xom ashyo qo'shish"""
    return_q = (return_q or "").strip()
    recipe = db.query(Recipe).filter(Recipe.id == recipe_id).first()
    if not recipe:
        raise HTTPException(status_code=404, detail="Retsept topilmadi")
    item = RecipeItem(
        recipe_id=recipe_id,
        product_id=product_id,
        quantity=quantity
    )
    db.add(item)
    db.commit()
    return RedirectResponse(url=_recipe_detail_redirect_url(recipe_id, return_q), status_code=303)


@app.post("/production/recipes/{recipe_id}/set-name")
async def set_recipe_name(
    recipe_id: int,
    request: Request,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    """Retsept nomini o'zgartirish."""
    form = await request.form()
    name = (form.get("name") or "").strip()
    return_q = (form.get("return_q") or "").strip()
    recipe = db.query(Recipe).filter(Recipe.id == recipe_id).first()
    if not recipe:
        raise HTTPException(status_code=404, detail="Retsept topilmadi")
    if name:
        recipe.name = name
    db.commit()
    return RedirectResponse(url=_recipe_detail_redirect_url(recipe_id, return_q), status_code=303)


@app.post("/production/recipes/{recipe_id}/set-warehouses")
async def set_recipe_warehouses(
    recipe_id: int,
    request: Request,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    """Retsept uchun omborlarni saqlash."""
    form = await request.form()
    return_q = (form.get("return_q") or "").strip()
    default_warehouse_id = form.get("default_warehouse_id")
    default_output_warehouse_id = form.get("default_output_warehouse_id")
    if default_warehouse_id is not None and str(default_warehouse_id).strip() == "":
        default_warehouse_id = None
    if default_output_warehouse_id is not None and str(default_output_warehouse_id).strip() == "":
        default_output_warehouse_id = None
    if default_warehouse_id is not None:
        try:
            default_warehouse_id = int(default_warehouse_id)
        except (TypeError, ValueError):
            default_warehouse_id = None
    if default_output_warehouse_id is not None:
        try:
            default_output_warehouse_id = int(default_output_warehouse_id)
        except (TypeError, ValueError):
            default_output_warehouse_id = None
    recipe = db.query(Recipe).filter(Recipe.id == recipe_id).first()
    if not recipe:
        raise HTTPException(status_code=404, detail="Retsept topilmadi")
    recipe.default_warehouse_id = default_warehouse_id
    recipe.default_output_warehouse_id = default_output_warehouse_id
    db.commit()
    return RedirectResponse(url=_recipe_detail_redirect_url(recipe_id, return_q), status_code=303)


@app.post("/production/recipes/{recipe_id}/edit-item/{item_id}")
async def edit_recipe_item(
    recipe_id: int,
    item_id: int,
    product_id: int = Form(...),
    quantity: float = Form(...),
    return_q: str = Form(""),
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth)
):
    """Retsept tarkibidagi qatorni tahrirlash"""
    return_q = (return_q or "").strip()
    item = db.query(RecipeItem).filter(
        RecipeItem.id == item_id,
        RecipeItem.recipe_id == recipe_id
    ).first()
    if not item:
        raise HTTPException(status_code=404, detail="Tarkib qatori topilmadi")
    item.product_id = product_id
    item.quantity = quantity
    db.commit()
    return RedirectResponse(url=_recipe_detail_redirect_url(recipe_id, return_q), status_code=303)


@app.post("/production/recipes/{recipe_id}/add-stage")
async def add_recipe_stage(
    recipe_id: int,
    stage_number: int = Form(...),
    name: str = Form(...),
    return_q: str = Form(""),
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    """Retseptga bosqich qo'shish"""
    return_q = (return_q or "").strip()
    recipe = db.query(Recipe).filter(Recipe.id == recipe_id).first()
    if not recipe:
        raise HTTPException(status_code=404, detail="Retsept topilmadi")
    stage = RecipeStage(recipe_id=recipe_id, stage_number=stage_number, name=(name or "").strip())
    db.add(stage)
    db.commit()
    return RedirectResponse(url=_recipe_detail_redirect_url(recipe_id, return_q), status_code=303)


@app.post("/production/recipes/{recipe_id}/delete-stage/{stage_id}")
async def delete_recipe_stage(
    recipe_id: int,
    stage_id: int,
    request: Request,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    """Retsept bosqichini o'chirish"""
    form = await request.form()
    return_q = (form.get("return_q") or "").strip()
    stage = db.query(RecipeStage).filter(
        RecipeStage.id == stage_id,
        RecipeStage.recipe_id == recipe_id,
    ).first()
    if not stage:
        raise HTTPException(status_code=404, detail="Bosqich topilmadi")
    db.delete(stage)
    db.commit()
    return RedirectResponse(url=_recipe_detail_redirect_url(recipe_id, return_q), status_code=303)


@app.post("/production/recipes/{recipe_id}/delete-item/{item_id}")
async def delete_recipe_item(
    recipe_id: int,
    item_id: int,
    request: Request,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth)
):
    """Retsept tarkibidagi qatorni o'chirish"""
    form = await request.form()
    return_q = (form.get("return_q") or "").strip()
    item = db.query(RecipeItem).filter(
        RecipeItem.id == item_id,
        RecipeItem.recipe_id == recipe_id
    ).first()
    if not item:
        raise HTTPException(status_code=404, detail="Tarkib qatori topilmadi")
    db.delete(item)
    db.commit()
    return RedirectResponse(url=_recipe_detail_redirect_url(recipe_id, return_q), status_code=303)


@app.get("/production/{prod_id}/materials", response_class=HTMLResponse)
async def production_edit_materials(
    prod_id: int,
    request: Request,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth)
):
    """Kutilmoqdagi buyurtma uchun xom ashyo miqdorlarini tahrirlash. Yakunlangan buyurtmani faqat admin ko'ra oladi (faqat ko'rish)."""
    if not current_user:
        return RedirectResponse(url="/login", status_code=303)
    production = db.query(Production).filter(Production.id == prod_id).first()
    if not production:
        return RedirectResponse(
            url="/production/orders?error=not_found&detail=" + quote("Buyurtma topilmadi yoki o'chirilgan."),
            status_code=303,
        )
    if production.status == "completed" and current_user.role != "admin":
        return RedirectResponse(url="/production/orders?error=forbidden&detail=" + quote("Yakunlangan buyurtmani faqat administrator ko'ra oladi."), status_code=303)
    if production.status not in ("draft", "completed"):
        return RedirectResponse(
            url="/production/orders?error=invalid&detail=" + quote("Faqat kutilmoqdagi yoki yakunlangan buyurtmani ko'rish mumkin."),
            status_code=303,
        )
    recipe = db.query(Recipe).filter(Recipe.id == production.recipe_id).first()
    if not recipe:
        return RedirectResponse(
            url="/production/orders?error=not_found&detail=" + quote("Retsept topilmadi."),
            status_code=303,
        )
    # Agar production_items bo'sh bo'lsa, retseptdan yaratib olaylik
    if not production.production_items:
        for item in recipe.items:
            pi = ProductionItem(
                production_id=production.id,
                product_id=item.product_id,
                quantity=item.quantity * production.quantity
            )
            db.add(pi)
        db.commit()
        db.refresh(production)
    read_only = production.status == "completed"
    return templates.TemplateResponse("production/edit_materials.html", {
        "request": request,
        "current_user": current_user,
        "production": production,
        "recipe": recipe,
        "read_only": read_only,
        "page_title": f"Xom ashyo: {production.number}",
    })


@app.post("/production/{prod_id}/materials")
async def production_save_materials(
    prod_id: int,
    request: Request,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth)
):
    """Xom ashyo miqdorlarini saqlash."""
    production = db.query(Production).filter(Production.id == prod_id).first()
    if not production or production.status != "draft":
        raise HTTPException(status_code=404, detail="Buyurtma topilmadi yoki tahrirlab bo'lmaydi")
    form = await request.form()
    for key, value in form.items():
        if key.startswith("qty_"):
            try:
                item_id = int(key.replace("qty_", ""))
                qty = float(value.replace(",", "."))
            except (ValueError, TypeError):
                continue
            pi = db.query(ProductionItem).filter(
                ProductionItem.id == item_id,
                ProductionItem.production_id == prod_id
            ).first()
            if pi and qty >= 0:
                pi.quantity = qty
    db.commit()
    return RedirectResponse(url="/production/orders", status_code=303)


@app.get("/production/{prod_id}/movements", response_class=HTMLResponse)
async def production_movements_page(
    prod_id: int,
    request: Request,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    """Shu ishlab chiqarish buyurtmasi uchun ombor harakati tarixi (xom ashyo chiqimi, tayyor mahsulot kirimi)."""
    if not current_user:
        return RedirectResponse(url="/login", status_code=303)
    production = db.query(Production).filter(Production.id == prod_id).first()
    if not production:
        return RedirectResponse(
            url="/production/orders?error=not_found&detail=" + quote("Buyurtma topilmadi."),
            status_code=303,
        )
    movements = (
        db.query(StockMovement)
        .filter(
            StockMovement.document_type == "Production",
            StockMovement.document_id == prod_id,
        )
        .order_by(StockMovement.created_at.asc())
        .all()
    )
    rows = []
    for m in movements:
        wh_name = (m.warehouse.name if m.warehouse else "") or "—"
        prod_name = (m.product.name if m.product else "") or "—"
        code = (m.product.code if m.product else "") or ""
        qty = float(m.quantity_change or 0)
        rows.append({
            "warehouse_name": wh_name,
            "product_name": prod_name,
            "product_code": code,
            "quantity_change": qty,
            "quantity_after": float(m.quantity_after or 0),
            "created_at": m.created_at.strftime("%d.%m.%Y %H:%M") if m.created_at else "—",
        })
    return templates.TemplateResponse("production/movements.html", {
        "request": request,
        "current_user": current_user,
        "production": production,
        "rows": rows,
        "page_title": f"Harakat tarixi — {production.number}",
    })


@app.get("/production/orders", response_class=HTMLResponse)
async def production_orders(
    request: Request,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
    date_from: Optional[str] = None,
    date_to: Optional[str] = None,
    q: Optional[str] = None,
):
    """Ishlab chiqarish buyurtmalari. Sana sukutda bugun; faqat joriy foydalanuvchi o'zi ishlab chiqarganlari ko'rinadi."""
    from sqlalchemy.orm import joinedload
    from sqlalchemy import func
    from urllib.parse import unquote, quote

    if not current_user:
        return RedirectResponse(url="/login?next=/production/orders", status_code=303)
    try:
        today = date.today()
        role = (getattr(current_user, "role", None) or "").strip().lower()
        # Admin/rahbar/raxbar uchun sukutda oy boshi–bugun (ko'proq buyurtma ko'rinsin); boshqalar uchun bugun
        if not (date_from or "").strip():
            date_from = (today.replace(day=1).strftime("%Y-%m-%d") if role in ("admin", "rahbar", "raxbar") else today.strftime("%Y-%m-%d"))
        if not (date_to or "").strip():
            date_to = today.strftime("%Y-%m-%d")
        operator_id_raw = request.query_params.get("operator_id")
        operator_id = None
        if operator_id_raw is not None and str(operator_id_raw).strip():
            try:
                operator_id = int(operator_id_raw)
            except (ValueError, TypeError):
                operator_id = None

        qry = (
            db.query(Production)
            .options(
                joinedload(Production.recipe).joinedload(Recipe.stages),
                joinedload(Production.recipe).joinedload(Recipe.product).joinedload(Product.unit),
                joinedload(Production.production_items),
                joinedload(Production.operator),
                joinedload(Production.user),
                joinedload(Production.warehouse),
                joinedload(Production.output_warehouse),
                joinedload(Production.machine),
            )
            .order_by(Production.date.desc())
        )
        if role in ("admin", "rahbar", "raxbar"):
            # Admin/rahbar/raxbar: operator tanlangan bo'lsa — operator_id yoki shu xodimga bog'langan user orqali kiritilgan buyurtmalar
            if operator_id and operator_id > 0:
                emp = db.query(Employee).filter(Employee.id == operator_id).first()
                if emp and getattr(emp, "user_id", None):
                    qry = qry.filter(
                        or_(
                            Production.operator_id == operator_id,
                            Production.user_id == emp.user_id,
                        )
                    )
                else:
                    qry = qry.filter(Production.operator_id == operator_id)
            # operator tanlanmagan: barcha buyurtmalar ko'rinadi
        else:
            qry = qry.filter(Production.user_id == current_user.id)

        if (date_from or "").strip():
            try:
                d_from = datetime.strptime(str(date_from).strip()[:10], "%Y-%m-%d").date()
                qry = qry.filter(func.date(Production.date) >= d_from)
            except (ValueError, TypeError):
                pass
        if (date_to or "").strip():
            try:
                d_to = datetime.strptime(str(date_to).strip()[:10], "%Y-%m-%d").date()
                qry = qry.filter(func.date(Production.date) <= d_to)
            except (ValueError, TypeError):
                pass
        if (q or "").strip():
            search = "%" + str(q).strip().lower() + "%"
            try:
                ids_num = [r[0] for r in db.query(Production.id).filter(func.lower(Production.number).like(search)).all()]
                ids_recipe = [r[0] for r in db.query(Production.id).join(Recipe, Production.recipe_id == Recipe.id).filter(func.lower(Recipe.name).like(search)).all()]
                match_ids = list(set(ids_num) | set(ids_recipe))
                if match_ids:
                    qry = qry.filter(Production.id.in_(match_ids))
                else:
                    qry = qry.filter(Production.id == -1)
            except Exception:
                qry = qry.filter(func.lower(Production.number).like(search))
        productions = qry.all()
        # T.kg — faqat tayyor mahsulot chiqadigan qatorlarda; Y.t.kg — faqat yarim tayyor chiqadigan qatorlarda
        total_output_kg = 0.0
        total_yarim_tayyor_kg = 0.0
        for p in productions:
            out_wh_name = (getattr(p.output_warehouse, "name", None) or "").lower()
            is_yarim_tayyor_output = "yarim" in out_wh_name  # 2-ombor yarim tayyor ombori bo'lsa
            out_kg = _kg_per_unit_from_recipe(p.recipe) * (float(p.quantity or 0))
            inp_kg = sum(float(pi.quantity or 0) for pi in (p.production_items or []))
            completed_only = getattr(p, "status", None) == "completed"  # faqat yakunlanganlar jamiga
            if is_yarim_tayyor_output:
                p.output_kg = 0.0
                if _is_qiyom_recipe(p.recipe):
                    p.yarim_tayyor_kg = 0.0  # qiyom retseptlari Y.t.kg da ko'rsatilmaydi va jamiga kiritilmaydi
                else:
                    p.yarim_tayyor_kg = inp_kg
                    if completed_only:
                        total_yarim_tayyor_kg += inp_kg
            else:
                p.output_kg = out_kg
                p.yarim_tayyor_kg = 0.0
                if completed_only:
                    total_output_kg += out_kg
        machines = db.query(Machine).filter(Machine.is_active == True).all()
        employees = db.query(Employee).filter(Employee.is_active == True).all()
        current_user_employee = db.query(Employee).filter(Employee.user_id == current_user.id).first() if current_user else None
        error = request.query_params.get("error")
        detail = unquote(request.query_params.get("detail", "") or "")
        return templates.TemplateResponse("production/orders.html", {
            "request": request,
            "current_user": current_user,
            "productions": productions,
            "total_output_kg": total_output_kg,
            "total_yarim_tayyor_kg": total_yarim_tayyor_kg,
            "machines": machines,
            "employees": employees,
            "current_user_employee_id": current_user_employee.id if current_user_employee else None,
            "page_title": "Ishlab chiqarish buyurtmalari",
            "error": error,
            "error_detail": detail,
            "stage_names": PRODUCTION_STAGE_NAMES,
            "filter_date_from": (date_from or "").strip()[:10] if date_from else today.strftime("%Y-%m-%d"),
            "filter_date_to": (date_to or "").strip()[:10] if date_to else today.strftime("%Y-%m-%d"),
            "filter_operator_id": operator_id if (operator_id and operator_id > 0) else None,
            "filter_q": (q or "").strip(),
        })
    except Exception as e:
        import traceback
        traceback.print_exc()
        return RedirectResponse(url=f"/production?error=orders&detail={quote(str(e)[:80])}", status_code=303)


@app.post("/production/orders/bulk-revert")
async def production_orders_bulk_revert(
    request: Request,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_admin),
):
    """Tanlangan yakunlangan buyurtmalar uchun tasdiqni bekor qilish."""
    from app.routes.production import _production_revert_one
    form = await request.form()
    raw_ids = form.getlist("prod_ids")
    prod_ids = [int(x) for x in raw_ids if str(x).strip().isdigit()]
    if not prod_ids:
        return RedirectResponse(
            url="/production/orders?error=revert&detail=" + quote("Hech qaysi buyurtma tanlanmagan."),
            status_code=303,
        )
    reverted = 0
    for pid in prod_ids:
        production = db.query(Production).filter(Production.id == pid).first()
        if not production:
            continue
        err = _production_revert_one(db, production)
        if err:
            db.rollback()
            return RedirectResponse(
                url="/production/orders?error=revert&detail=" + quote(f"{production.number}: {err}"),
                status_code=303,
            )
        reverted += 1
    db.commit()
    return RedirectResponse(url="/production/orders?bulk_reverted=" + str(reverted), status_code=303)


@app.post("/production/orders/bulk-complete")
async def production_orders_bulk_complete(
    request: Request,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    """Tanlangan buyurtmalarni yakunlash."""
    from app.routes.production import _do_complete_production_stock, _recipe_max_stage
    from app.utils.notifications import check_low_stock_and_notify
    from app.utils.production_order import notify_managers_production_ready
    form = await request.form()
    raw_ids = form.getlist("prod_ids")
    prod_ids = [int(x) for x in raw_ids if str(x).strip().isdigit()]
    if not prod_ids:
        return RedirectResponse(
            url="/production/orders?error=complete&detail=" + quote("Hech qaysi buyurtma tanlanmagan."),
            status_code=303,
        )
    completed = 0
    for pid in prod_ids:
        production = db.query(Production).filter(Production.id == pid).first()
        if not production or production.status not in ("draft", "in_progress"):
            continue
        recipe = db.query(Recipe).filter(Recipe.id == production.recipe_id).first()
        if not recipe:
            return RedirectResponse(
                url="/production/orders?error=complete&detail=" + quote(f"{production.number}: Retsept topilmadi."),
                status_code=303,
            )
        err = _do_complete_production_stock(db, production, recipe)
        if err:
            db.rollback()
            return err
        production.status = "completed"
        production.current_stage = _recipe_max_stage(recipe)
        completed += 1
    db.commit()
    check_low_stock_and_notify(db)
    for pid in prod_ids:
        production = db.query(Production).filter(Production.id == pid).first()
        if production and production.status == "completed":
            notify_managers_production_ready(db, production)
    return RedirectResponse(url="/production/orders?bulk_completed=" + str(completed), status_code=303)


def _is_qiyom_recipe(recipe) -> bool:
    """Retsept 'qiyom' (oralama mahsulot) bo'lsa — jamida hisoblanmasin (shablon: qiyom hisobga olinmaydi)."""
    if not recipe or not getattr(recipe, "name", None):
        return False
    return "qiyom" in (recipe.name or "").lower()


def _kg_per_unit_from_recipe(recipe) -> float:
    """Recipe nomi yoki output_quantity dan kg per birlik (dona) ni hisoblaydi."""
    if not recipe:
        return 1.0
    name = (recipe.name or "").lower()
    if "250gr" in name or "250 gr" in name:
        return 0.25
    if "400gr" in name or "400 gr" in name:
        return 0.4
    if "5kg" in name or "5 kg" in name:
        return 5.0
    if "4kg" in name or "4 kg" in name:
        return 4.0
    if "3kg" in name or "3 kg" in name:
        return 3.0
    if "2kg" in name or "2 kg" in name:
        return 2.0
    if "1kg" in name or "1 kg" in name:
        return 1.0
    if getattr(recipe, "output_quantity", None) and recipe.output_quantity:
        return float(recipe.output_quantity)
    return 1.0


@app.get("/production/by-operator", response_class=HTMLResponse)
async def production_by_operator(
    request: Request,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
    date_from: Optional[str] = None,
    date_to: Optional[str] = None,
):
    """Operator bo'yicha ishlab chiqarish (yakunlanganlar, sana bo'yicha)."""
    from sqlalchemy.orm import joinedload
    from sqlalchemy import func
    from collections import defaultdict

    if not current_user:
        return RedirectResponse(url="/login?next=/production/by-operator", status_code=303)
    today = date.today()
    if not (date_from or "").strip():
        date_from = today.replace(day=1).strftime("%Y-%m-%d")  # oy boshi
    if not (date_to or "").strip():
        date_to = today.strftime("%Y-%m-%d")
    d_from = d_to = None
    try:
        d_from = datetime.strptime(str(date_from).strip()[:10], "%Y-%m-%d").date()
    except (ValueError, TypeError):
        d_from = today.replace(day=1)
    try:
        d_to = datetime.strptime(str(date_to).strip()[:10], "%Y-%m-%d").date()
    except (ValueError, TypeError):
        d_to = today

    qry = (
        db.query(Production)
        .options(
            joinedload(Production.recipe).joinedload(Recipe.product).joinedload(Product.unit),
            joinedload(Production.operator),
            joinedload(Production.user),
        )
        .filter(Production.status == "completed")
        .filter(func.date(Production.date) >= d_from)
        .filter(func.date(Production.date) <= d_to)
        .order_by(Production.date.desc())
    )
    # joinedload bilan ba'zi driverlarda bir xil qator takrorlanishi mumkin — id bo'yicha bitta qoldiramiz
    seen = {}
    for p in qry.all():
        if p.id not in seen:
            seen[p.id] = p
    productions = list(seen.values())
    productions.sort(key=lambda x: (x.date or datetime.min), reverse=True)

    # User_id -> employee_id: user orqali ism ko'rinadigan operatorlarni ham link qilish uchun
    user_ids_from_productions = {p.user_id for p in productions if p.user_id and not p.operator_id}
    user_to_employee = {}
    if user_ids_from_productions:
        for emp in db.query(Employee).filter(Employee.user_id.in_(user_ids_from_productions)).all():
            if emp.user_id:
                user_to_employee[emp.user_id] = emp.id

    # Operator bo'yicha jami (kg) — har bir buyurtma faqat bir marta, qiyom hisobga olinmaydi
    totals_by_name = defaultdict(float)
    name_to_employee_id = {}
    for p in productions:
        # Qiyom (oralama mahsulot) — jamida ko'rsatilmasin (yarim tayyor / tayyor / dona hisoblanadi)
        if _is_qiyom_recipe(p.recipe):
            continue
        op_name = "—"
        if p.operator:
            op_name = getattr(p.operator, "full_name", None) or str(p.operator_id)
            if p.operator_id and op_name != "—":
                name_to_employee_id[op_name] = p.operator_id
        elif p.user:
            op_name = getattr(p.user, "full_name", None) or str(p.user_id)
            # User orqali ko'rinadigan ism uchun ham employee_id topilsa link qilamiz
            if p.user_id and op_name != "—":
                emp_id = user_to_employee.get(p.user_id)
                if emp_id:
                    name_to_employee_id[op_name] = emp_id
        kg = (p.quantity or 0) * _kg_per_unit_from_recipe(p.recipe)
        totals_by_name[op_name] += kg
    operator_totals = sorted(totals_by_name.items(), key=lambda x: -x[1])

    return templates.TemplateResponse("production/by_operator.html", {
        "request": request,
        "current_user": current_user,
        "productions": productions,
        "operator_totals": operator_totals,
        "name_to_employee_id": name_to_employee_id,
        "user_to_employee": user_to_employee,
        "filter_date_from": (date_from or "").strip()[:10],
        "filter_date_to": (date_to or "").strip()[:10],
        "page_title": "Operator bo'yicha ishlab chiqarish",
    })


@app.get("/production/new", response_class=HTMLResponse)
async def production_new(request: Request, db: Session = Depends(get_db), current_user: User = Depends(require_auth)):
    """Yangi ishlab chiqarish"""
    warehouses = db.query(Warehouse).all()
    recipes = db.query(Recipe).filter(Recipe.is_active == True).all()
    machines = db.query(Machine).filter(Machine.is_active == True).all()
    employees = db.query(Employee).filter(Employee.is_active == True).all()

    return templates.TemplateResponse("production/new_order.html", {
        "request": request,
        "current_user": current_user,
        "recipes": recipes,
        "warehouses": warehouses,
        "machines": machines,
        "employees": employees,
        "page_title": "Yangi ishlab chiqarish"
    })



from typing import List

@app.get("/production/create", response_class=HTMLResponse)
async def production_create_get(
    request: Request,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    """GET /production/create — yangi ishlab chiqarish formasi (production/new bilan bir xil)."""
    if not current_user:
        return RedirectResponse(url="/login?next=/production/create", status_code=303)
    try:
        warehouses = db.query(Warehouse).all()
        recipes = (
            db.query(Recipe)
            .options(
                joinedload(Recipe.product).joinedload(Product.unit),
                joinedload(Recipe.items),
            )
            .filter(Recipe.is_active == True)
            .all()
        )
        machines = db.query(Machine).filter(Machine.is_active == True).all()
        employees = db.query(Employee).filter(Employee.is_active == True).all()
        current_user_employee = db.query(Employee).filter(Employee.user_id == current_user.id).first() if current_user else None
        return templates.TemplateResponse("production/new_order.html", {
            "request": request,
            "current_user": current_user,
            "recipes": recipes,
            "warehouses": warehouses,
            "machines": machines,
            "employees": employees,
            "current_user_employee_id": current_user_employee.id if current_user_employee else None,
            "page_title": "Yangi ishlab chiqarish",
        })
    except Exception as e:
        import traceback
        traceback.print_exc()
        return RedirectResponse(url="/production/new", status_code=303)


def _safe_int(val, default=None):
    if val is None or (isinstance(val, str) and not val.strip()):
        return default
    try:
        return int(float(val)) if isinstance(val, str) else int(val)
    except (ValueError, TypeError):
        return default


@app.post("/production/create")
async def create_production(
    request: Request,
    recipe_id: int = Form(...),
    warehouse_id: int = Form(...),
    output_warehouse_id: Optional[int] = Form(None),
    quantity: float = Form(...),
    note: str = Form(""),
    machine_id: Optional[str] = Form(None),
    operator_id: Optional[str] = Form(None),
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """Ishlab chiqarish yaratish: 1-ombor (xom ashyo) dan oladi, 2-ombor (yarim tayyor) ga yozadi."""
    from urllib.parse import quote
    try:
        if output_warehouse_id is None:
            output_warehouse_id = warehouse_id
        from sqlalchemy.orm import joinedload
        recipe = db.query(Recipe).options(joinedload(Recipe.stages), joinedload(Recipe.items)).filter(Recipe.id == recipe_id).first()
        if not recipe:
            raise HTTPException(status_code=404, detail="Retsept topilmadi")
        max_stage = _recipe_max_stage(recipe)
        today = datetime.now()
        prefix = f"PR-{today.strftime('%Y%m%d')}-"
        last = db.query(Production.number).filter(Production.number.like(prefix + "%")).all()
        next_n = 1
        for (num,) in last:
            if num and num.startswith(prefix):
                try:
                    n = int(num[len(prefix):].strip())
                    if n >= next_n:
                        next_n = n + 1
                except ValueError:
                    pass
        number = f"{prefix}{str(next_n).zfill(3)}"
        mid = _safe_int(machine_id)
        oid = _safe_int(operator_id)
        production = Production(
            number=number,
            recipe_id=recipe_id,
            warehouse_id=warehouse_id,
            output_warehouse_id=output_warehouse_id,
            quantity=quantity,
            note=note or "",
            status="draft",
            current_stage=1,
            max_stage=max_stage,
            user_id=current_user.id if current_user else None,
            machine_id=mid,
            operator_id=oid,
        )
        db.add(production)
        db.commit()
        db.refresh(production)
        for stage_num in range(1, max_stage + 1):
            stage = ProductionStage(production_id=production.id, stage_number=stage_num)
            db.add(stage)
        db.commit()
        if recipe.items:
            for item in recipe.items:
                pi = ProductionItem(
                    production_id=production.id,
                    product_id=item.product_id,
                    quantity=item.quantity * quantity
                )
                db.add(pi)
            db.commit()
        return RedirectResponse(url="/production/orders", status_code=303)
    except HTTPException:
        raise
    except IntegrityError as e:
        if "productions.number" in str(e.orig) if getattr(e, "orig", None) else "number" in str(e):
            try:
                prefix = f"PR-{datetime.now().strftime('%Y%m%d')}-"
                last = db.query(Production.number).filter(Production.number.like(prefix + "%")).all()
                next_n = 1
                for (num,) in last:
                    if num and num.startswith(prefix):
                        try:
                            n = int(num[len(prefix):].strip())
                            if n >= next_n:
                                next_n = n + 1
                        except ValueError:
                            pass
                number = f"{prefix}{str(next_n).zfill(3)}"
                production = Production(
                    number=number,
                    recipe_id=recipe_id,
                    warehouse_id=warehouse_id,
                    output_warehouse_id=output_warehouse_id,
                    quantity=quantity,
                    note=note or "",
                    status="draft",
                    current_stage=1,
                    max_stage=max_stage,
                    user_id=current_user.id if current_user else None,
                    machine_id=mid,
                    operator_id=oid,
                )
                db.add(production)
                db.commit()
                db.refresh(production)
                for stage_num in range(1, max_stage + 1):
                    stage = ProductionStage(production_id=production.id, stage_number=stage_num)
                    db.add(stage)
                db.commit()
                if recipe.items:
                    for item in recipe.items:
                        pi = ProductionItem(production_id=production.id, product_id=item.product_id, quantity=item.quantity * quantity)
                        db.add(pi)
                    db.commit()
                return RedirectResponse(url="/production/orders", status_code=303)
            except Exception:
                pass
        import traceback
        traceback.print_exc()
        msg = quote(str(e)[:120], safe="")
        return RedirectResponse(url=f"/production/new?error=create&detail={msg}", status_code=303)
    except Exception as e:
        import traceback
        traceback.print_exc()
        msg = quote(str(e)[:120], safe="")
        return RedirectResponse(url=f"/production/new?error=create&detail={msg}", status_code=303)


def _do_complete_production_stock(db, production, recipe):
    """Xom ashyo ayirish, tayyor mahsulot qo'shish. RedirectResponse qaytaradi xato bo'lsa."""
    from urllib.parse import quote
    wh_id = production.warehouse_id
    if not wh_id:
        return RedirectResponse(
            url="/production/orders?error=insufficient_stock&detail=" + quote("Buyurtmada 1-ombor (xom ashyo) tanlanmagan."),
            status_code=303
        )
    wh = db.query(Warehouse).filter(Warehouse.id == wh_id).first()
    wh_name = wh.name if wh else f"#{wh_id}"
    if production.production_items:
        items_to_use = [(pi.product_id, pi.quantity) for pi in production.production_items]
    else:
        items_to_use = [(item.product_id, item.quantity * production.quantity) for item in recipe.items]
    for product_id, required in items_to_use:
        stock = db.query(Stock).filter(
            Stock.warehouse_id == wh_id,
            Stock.product_id == product_id
        ).first()
        available = float(stock.quantity or 0) if stock else 0
        if available < required:
            product_name = db.query(Product).filter(Product.id == product_id).first()
            name = product_name.name if product_name else f"#{product_id}"
            msg = quote(
                f"«{wh_name}» da {name} yetarli emas: kerak {required}, mavjud {available}. Kerakli mahsulotni shu omborga kiriting.",
                safe=""
            )
            return RedirectResponse(url=f"/production/orders?error=insufficient_stock&detail={msg}", status_code=303)
    # Xom ashyolarni ayirish va StockMovement yozuvlarini yaratish
    for product_id, required in items_to_use:
        stock = db.query(Stock).filter(
            Stock.warehouse_id == production.warehouse_id,
            Stock.product_id == product_id
        ).first()
        if stock:
            # StockMovement yozuvini yaratish (chiqim)
            create_stock_movement(
                db=db,
                warehouse_id=production.warehouse_id,
                product_id=product_id,
                quantity_change=-required,  # Chiqim
                operation_type="production_consumption",
                document_type="Production",
                document_id=production.id,
                document_number=production.number,
                user_id=production.user_id,
                note=f"Ishlab chiqarish (xom ashyo): {production.number}"
            )
    
    total_material_cost = 0.0
    for product_id, required in items_to_use:
        product = db.query(Product).filter(Product.id == product_id).first()
        if product and getattr(product, "purchase_price", None) is not None:
            total_material_cost += required * (product.purchase_price or 0)
    output_units = production.quantity * (recipe.output_quantity or 1)
    cost_per_unit = (total_material_cost / output_units) if output_units > 0 else 0
    out_wh_id = production.output_warehouse_id if production.output_warehouse_id else production.warehouse_id
    
    # Tayyor mahsulotni qo'shish va StockMovement yozuvini yaratish
    product_stock = db.query(Stock).filter(
        Stock.warehouse_id == out_wh_id,
        Stock.product_id == recipe.product_id
    ).first()
    if product_stock:
        product_stock.quantity += output_units
    else:
        db.add(Stock(warehouse_id=out_wh_id, product_id=recipe.product_id, quantity=output_units))
    
    # StockMovement yozuvini yaratish (kirim - tayyor mahsulot)
    create_stock_movement(
        db=db,
        warehouse_id=out_wh_id,
        product_id=recipe.product_id,
        quantity_change=output_units,  # Kirim
        operation_type="production_output",
        document_type="Production",
        document_id=production.id,
        document_number=production.number,
        user_id=production.user_id,
        note=f"Ishlab chiqarish (tayyor mahsulot): {production.number}"
    )
    output_product = db.query(Product).filter(Product.id == recipe.product_id).first()
    if output_product:
        product_stock = db.query(Stock).filter(Stock.warehouse_id == out_wh_id, Stock.product_id == recipe.product_id).first()
        old_price = output_product.purchase_price or 0
        old_qty = (product_stock.quantity - output_units) if product_stock else 0
        if old_qty > 0 and old_price > 0 and output_units > 0:
            output_product.purchase_price = (old_qty * old_price + output_units * cost_per_unit) / (old_qty + output_units)
        elif cost_per_unit > 0:
            output_product.purchase_price = cost_per_unit
    return None


def _recipe_max_stage(recipe) -> int:
    """Retseptdagi oxirgi bosqich raqami; bo'lmasa 2 (tezkor ishlab chiqarish uchun)."""
    if not recipe or not recipe.stages:
        return 2
    return max(s.stage_number for s in recipe.stages)


def _production_orders_redirect_url(form, base="/production/orders", extra_params=None):
    """Buyurtmalar sahifasiga qaytishda sana/operator/q filtrlarni saqlash."""
    from urllib.parse import urlencode
    params = {}
    for key in ("redirect_date_from", "redirect_date_to", "redirect_operator_id", "redirect_q"):
        val = form.get(key) if hasattr(form, "get") else (form or {}).get(key)
        if val is not None and str(val).strip():
            qkey = key.replace("redirect_", "")
            params[qkey] = str(val).strip()
    if extra_params:
        params.update(extra_params)
    return base + ("?" + urlencode(params) if params else "")


@app.post("/production/{prod_id}/complete-stage")
async def complete_production_stage(
    request: Request,
    prod_id: int,
    stage_number: int = Form(...),
    machine_id: Optional[int] = Form(None),
    operator_id: Optional[int] = Form(None),
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    """Bosqichni yakunlash. Oxirgi bosqichda ombor harakati qiladi va buyurtma yakunlanadi (retseptdagi bosqichlar soniga qarab)."""
    form = await request.form()
    def _redirect(extra=None):
        return RedirectResponse(url=_production_orders_redirect_url(form, extra_params=extra), status_code=303)
    production = db.query(Production).filter(Production.id == prod_id).first()
    if not production:
        raise HTTPException(status_code=404, detail="Topilmadi")
    from sqlalchemy.orm import joinedload
    recipe = (
        db.query(Recipe)
        .options(joinedload(Recipe.stages))
        .filter(Recipe.id == production.recipe_id)
        .first()
    )
    if not recipe:
        raise HTTPException(status_code=404, detail="Retsept topilmadi")
    # Har doim retseptdagi bosqichlar soniga qarab (eski buyurtmalarda max_stage=4 qolgan bo'lsa ham)
    max_stage = _recipe_max_stage(recipe)
    if stage_number < 1 or stage_number > max_stage:
        raise HTTPException(status_code=400, detail=f"Bosqich 1–{max_stage} oralig'ida bo'lishi kerak")
    if production.status == "completed":
        return _redirect()
    current = getattr(production, "current_stage", None) or 1
    # Eski buyurtma 4 bosqichda qolgan, retsept endi 2 bosqich — bosqichni bosganda darhol yakunlash
    if current > max_stage:
        err = _do_complete_production_stock(db, production, recipe)
        if err:
            return err
        production.status = "completed"
        production.current_stage = max_stage
        db.commit()
        check_low_stock_and_notify(db)
        return _redirect()
    if stage_number != current:
        return _redirect(extra_params={"error": "stage", "detail": f"Keyingi bosqich {current}"})
    stage_row = db.query(ProductionStage).filter(
        ProductionStage.production_id == prod_id,
        ProductionStage.stage_number == stage_number,
    ).first()
    now = datetime.now()
    if stage_row:
        if not stage_row.started_at:
            stage_row.started_at = now
        stage_row.completed_at = now
        stage_row.machine_id = int(machine_id) if machine_id else None
        stage_row.operator_id = int(operator_id) if operator_id else None
    if stage_number < max_stage:
        production.current_stage = stage_number + 1
        production.status = "in_progress"
        db.commit()
        return _redirect()
    err = _do_complete_production_stock(db, production, recipe)
    if err:
        return err
    production.status = "completed"
    production.current_stage = max_stage
    db.commit()
    check_low_stock_and_notify(db)
    return _redirect()


@app.post("/production/{prod_id}/complete")
async def complete_production(
    request: Request,
    prod_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    """Ishlab chiqarishni bir martada yakunlash (barcha bosqichlarsiz) — tezkor ishlab chiqarishga qaytadi."""
    form = await request.form()
    redirect_url = _production_orders_redirect_url(form)
    production = db.query(Production).filter(Production.id == prod_id).first()
    if not production:
        raise HTTPException(status_code=404, detail="Topilmadi")
    recipe = db.query(Recipe).filter(Recipe.id == production.recipe_id).first()
    if not recipe:
        raise HTTPException(status_code=404, detail="Retsept topilmadi")
    err = _do_complete_production_stock(db, production, recipe)
    if err:
        return err
    production.status = "completed"
    production.current_stage = _recipe_max_stage(recipe)
    db.commit()
    check_low_stock_and_notify(db)
    return RedirectResponse(url=redirect_url, status_code=303)


@app.post("/production/{prod_id}/revert")
async def production_revert(
    request: Request,
    prod_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_admin)
):
    """Tasdiqni bekor qilish (faqat admin): xom ashyoni qaytarish, tayyor mahsulotni olib tashlash, holatni qoralamaga o'tkazish"""
    from urllib.parse import quote
    form = await request.form()

    def _redirect(extra=None):
        return RedirectResponse(
            url=_production_orders_redirect_url(form, extra_params=extra),
            status_code=303
        )

    production = db.query(Production).filter(Production.id == prod_id).first()
    if not production:
        raise HTTPException(status_code=404, detail="Topilmadi")
    if production.status != "completed":
        return _redirect(extra_params={"error": "revert", "detail": "Faqat yakunlangan buyurtmaning tasdiqini bekor qilish mumkin."})
    recipe = db.query(Recipe).filter(Recipe.id == production.recipe_id).first()
    if not recipe:
        return _redirect(extra_params={"error": "revert", "detail": "Retsept topilmadi."})
    items_to_use = [(pi.product_id, pi.quantity) for pi in production.production_items] if production.production_items else [(item.product_id, item.quantity * production.quantity) for item in recipe.items]
    output_units = production.quantity * (recipe.output_quantity or 1)
    out_wh_id = production.output_warehouse_id if production.output_warehouse_id else production.warehouse_id
    # Tayyor mahsulotni 2-ombordan ayirish
    product_stock = db.query(Stock).filter(
        Stock.warehouse_id == out_wh_id,
        Stock.product_id == recipe.product_id
    ).first()
    current_qty = float(product_stock.quantity or 0) if product_stock else 0
    if not product_stock or current_qty < output_units:
        out_wh = db.query(Warehouse).filter(Warehouse.id == out_wh_id).first()
        out_product = db.query(Product).filter(Product.id == recipe.product_id).first()
        wh_name = (out_wh.name if out_wh else "2-ombor") or "2-ombor"
        prod_name = (out_product.name if out_product else "tayyor mahsulot") or "tayyor mahsulot"
        detail = f"«{wh_name}» da «{prod_name}» dan kerak: {output_units:,.1f}, mavjud: {current_qty:,.1f}. Mahsulot sotilgan yoki ko'chirilgan bo'lishi mumkin — tasdiqni bekor qilish uchun 2-omborda shu miqdorda qoldiq bo'lishi kerak."
        return _redirect(extra_params={"error": "revert", "detail": detail})
    product_stock.quantity -= output_units
    # Xom ashyolarni 1-omborga qaytarish
    for product_id, required in items_to_use:
        stock = db.query(Stock).filter(
            Stock.warehouse_id == production.warehouse_id,
            Stock.product_id == product_id
        ).first()
        if stock:
            stock.quantity += required
        else:
            db.add(Stock(warehouse_id=production.warehouse_id, product_id=product_id, quantity=required))
    delete_stock_movements_for_document(db, "Production", prod_id)
    production.status = "draft"
    db.commit()
    return _redirect()


@app.post("/production/{prod_id}/cancel")
async def cancel_production(
    request: Request,
    prod_id: int,
    db: Session = Depends(get_db),
):
    """Ishlab chiqarishni bekor qilish"""
    form = await request.form()
    production = db.query(Production).filter(Production.id == prod_id).first()
    if not production:
        raise HTTPException(status_code=404, detail="Topilmadi")
    production.status = "cancelled"
    db.commit()
    return RedirectResponse(url=_production_orders_redirect_url(form), status_code=303)


@app.post("/production/{prod_id}/delete")
async def delete_production(
    request: Request,
    prod_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_admin),
):
    """Ishlab chiqarish buyurtmasini o'chirish — faqat admin."""
    form = await request.form()
    production = db.query(Production).filter(Production.id == prod_id).first()
    if not production:
        raise HTTPException(status_code=404, detail="Buyurtma topilmadi")
    db.delete(production)
    db.commit()
    return RedirectResponse(url=_production_orders_redirect_url(form), status_code=303)


# ==========================================
# API (Telegram bot uchun)
# ==========================================

@app.get("/api/stats")
async def api_stats(db: Session = Depends(get_db)):
    """Statistika API"""
    today = datetime.now().date()
    
    today_sales = db.query(Order).filter(
        Order.type == "sale",
        Order.date >= today
    ).all()
    
    cash = db.query(CashRegister).first()
    
    return {
        "today_sales": sum(o.total for o in today_sales),
        "today_orders": len(today_sales),
        "cash_balance": cash.balance if cash else 0,
        "products_count": db.query(Product).count(),
        "partners_count": db.query(Partner).count(),
    }


@app.get("/api/products")
async def api_products(db: Session = Depends(get_db)):
    """Tovarlar API"""
    products = db.query(Product).filter(Product.is_active == True).all()
    return [{"id": p.id, "name": p.name, "code": p.code, "price": p.sale_price} for p in products]


@app.get("/api/partners")
async def api_partners(db: Session = Depends(get_db)):
    """Kontragentlar API"""
    partners = db.query(Partner).filter(Partner.is_active == True).all()
    return [{"id": p.id, "name": p.name, "balance": p.balance} for p in partners]


@app.get("/api/notifications/unread")
async def api_notifications_unread(
    db: Session = Depends(get_db),
    current_user: Optional[User] = Depends(require_auth),
):
    """O'qilmagan bildirishnomalar soni va oxirgisi (ovozli push / UI uchun)."""
    if not current_user:
        return {"unread_count": 0, "last": None}
    count = get_unread_count(db, current_user.id)
    last_list = get_user_notifications(db, current_user.id, unread_only=True, limit=1)
    last = None
    if last_list:
        n = last_list[0]
        last = {
            "id": n.id,
            "title": n.title or "",
            "message": n.message or "",
            "priority": n.priority or "normal",
        }
    return {"unread_count": count, "last": last}


# ==========================================
# AGENTLAR
# ==========================================

@app.get("/agents", response_class=HTMLResponse)
async def agents_list(request: Request, db: Session = Depends(get_db)):
    """Agentlar ro'yxati"""
    agents = db.query(Agent).all()
    
    # Har bir agent uchun oxirgi lokatsiya
    for agent in agents:
        last_loc = db.query(AgentLocation).filter(
            AgentLocation.agent_id == agent.id
        ).order_by(AgentLocation.recorded_at.desc()).first()
        agent.last_location = last_loc
        
        # Bugungi tashriflar
        today = datetime.now().date()
        agent.today_visits = db.query(Visit).filter(
            Visit.agent_id == agent.id,
            Visit.visit_date >= today
        ).count()
    
    return templates.TemplateResponse("agents/list.html", {
        "request": request,
        "agents": agents,
        "page_title": "Agentlar"
    })


@app.post("/agents/add")
async def agent_add(
    request: Request,
    full_name: str = Form(...),
    phone: str = Form(""),
    region: str = Form(""),
    telegram_id: str = Form(""),
    db: Session = Depends(get_db)
):
    """Agent qo'shish"""
    agent = Agent(
        full_name=full_name,
        phone=phone,
        region=region,
        telegram_id=telegram_id
    )
    db.add(agent)
    db.commit()
    return RedirectResponse(url="/agents", status_code=303)


@app.get("/agents/{agent_id}", response_class=HTMLResponse)
async def agent_detail(request: Request, agent_id: int, db: Session = Depends(get_db)):
    """Agent tafsilotlari"""
    agent = db.query(Agent).filter(Agent.id == agent_id).first()
    if not agent:
        raise HTTPException(status_code=404, detail="Agent topilmadi")
    
    # Oxirgi lokatsiyalar
    locations = db.query(AgentLocation).filter(
        AgentLocation.agent_id == agent_id
    ).order_by(AgentLocation.recorded_at.desc()).limit(50).all()
    
    # Tashriflar
    visits = db.query(Visit).filter(
        Visit.agent_id == agent_id
    ).order_by(Visit.visit_date.desc()).limit(30).all()
    
    return templates.TemplateResponse("agents/detail.html", {
        "request": request,
        "agent": agent,
        "locations": locations,
        "visits": visits,
        "page_title": f"Agent: {agent.full_name}"
    })


# ==========================================
# YETKAZIB BERISH
# ==========================================

@app.get("/delivery", response_class=HTMLResponse)
async def delivery_list(request: Request, db: Session = Depends(get_db)):
    """Yetkazib berish ro'yxati"""
    drivers = db.query(Driver).all()
    
    # Har bir haydovchi uchun statistika
    for driver in drivers:
        last_loc = db.query(DriverLocation).filter(
            DriverLocation.driver_id == driver.id
        ).order_by(DriverLocation.recorded_at.desc()).first()
        driver.last_location = last_loc
        
        # Bugungi yetkazilganlar
        today = datetime.now().date()
        driver.today_deliveries = db.query(Delivery).filter(
            Delivery.driver_id == driver.id,
            Delivery.created_at >= today
        ).count()
        
        driver.pending_deliveries = db.query(Delivery).filter(
            Delivery.driver_id == driver.id,
            Delivery.status == "pending"
        ).count()
    
    # Barcha yetkazishlar
    deliveries = db.query(Delivery).order_by(Delivery.created_at.desc()).limit(50).all()
    
    return templates.TemplateResponse("delivery/list.html", {
        "request": request,
        "drivers": drivers,
        "deliveries": deliveries,
        "page_title": "Yetkazib berish"
    })


@app.post("/drivers/add")
async def driver_add(
    request: Request,
    full_name: str = Form(...),
    phone: str = Form(""),
    vehicle_number: str = Form(""),
    vehicle_type: str = Form(""),
    telegram_id: str = Form(""),
    db: Session = Depends(get_db)
):
    """Haydovchi qo'shish"""
    driver = Driver(
        full_name=full_name,
        phone=phone,
        vehicle_number=vehicle_number,
        vehicle_type=vehicle_type,
        telegram_id=telegram_id
    )
    db.add(driver)
    db.commit()
    return RedirectResponse(url="/delivery", status_code=303)


@app.get("/delivery/{driver_id}", response_class=HTMLResponse)
async def driver_detail(request: Request, driver_id: int, db: Session = Depends(get_db)):
    """Haydovchi tafsilotlari"""
    driver = db.query(Driver).filter(Driver.id == driver_id).first()
    if not driver:
        raise HTTPException(status_code=404, detail="Haydovchi topilmadi")
    
    # Lokatsiya tarixi
    locations = db.query(DriverLocation).filter(
        DriverLocation.driver_id == driver_id
    ).order_by(DriverLocation.recorded_at.desc()).limit(100).all()
    
    # Yetkazishlar
    deliveries = db.query(Delivery).filter(
        Delivery.driver_id == driver_id
    ).order_by(Delivery.created_at.desc()).limit(30).all()
    
    return templates.TemplateResponse("delivery/detail.html", {
        "request": request,
        "driver": driver,
        "locations": locations,
        "deliveries": deliveries,
        "page_title": f"Haydovchi: {driver.full_name}"
    })


# ==========================================
# XARITA
# ==========================================

@app.get("/map", response_class=HTMLResponse)
async def map_view(request: Request, db: Session = Depends(get_db)):
    """Xarita - agentlar, haydovchilar, mijozlar joylashuvi"""
    
    # Agentlar
    agents = db.query(Agent).filter(Agent.is_active == True).all()
    agent_markers = []
    for agent in agents:
        last_loc = db.query(AgentLocation).filter(
            AgentLocation.agent_id == agent.id
        ).order_by(AgentLocation.recorded_at.desc()).first()
        if last_loc:
            agent_markers.append({
                "id": agent.id,
                "name": agent.full_name,
                "type": "agent",
                "lat": last_loc.latitude,
                "lng": last_loc.longitude,
                "time": last_loc.recorded_at.strftime("%H:%M")
            })
    
    # Haydovchilar
    drivers = db.query(Driver).filter(Driver.is_active == True).all()
    driver_markers = []
    for driver in drivers:
        last_loc = db.query(DriverLocation).filter(
            DriverLocation.driver_id == driver.id
        ).order_by(DriverLocation.recorded_at.desc()).first()
        if last_loc:
            driver_markers.append({
                "id": driver.id,
                "name": driver.full_name,
                "type": "driver",
                "lat": last_loc.latitude,
                "lng": last_loc.longitude,
                "time": last_loc.recorded_at.strftime("%H:%M"),
                "vehicle": driver.vehicle_number
            })
    
    # Mijozlar
    partners = db.query(Partner).filter(Partner.is_active == True).all()
    partner_locations = db.query(PartnerLocation).all()
    partner_markers = []
    for loc in partner_locations:
        partner = db.query(Partner).filter(Partner.id == loc.partner_id).first()
        if partner and loc.latitude and loc.longitude:
            partner_markers.append({
                "id": loc.partner_id,
                "name": partner.name,
                "type": "partner",
                "lat": loc.latitude,
                "lng": loc.longitude,
                "address": loc.address
            })
    
    try:
        from app.config.maps_config import MAP_PROVIDER
        map_provider = MAP_PROVIDER
    except Exception:
        map_provider = "yandex"
    try:
        from app.config.maps_config import YANDEX_MAPS_API_KEY
        yandex_apikey = YANDEX_MAPS_API_KEY or ""
    except Exception:
        yandex_apikey = ""
    return templates.TemplateResponse("map/index.html", {
        "request": request,
        "agents": agents,
        "drivers": drivers,
        "partner_locations": partner_locations,
        "agent_markers": agent_markers,
        "driver_markers": driver_markers,
        "partner_markers": partner_markers,
        "region_markers": [],
        "map_provider": map_provider,
        "yandex_maps_apikey": yandex_apikey,
        "page_title": "Xarita",
    })


# ==========================================
# SUPERVAYZER DASHBOARD
# ==========================================

@app.get("/supervisor", response_class=HTMLResponse)
async def supervisor_dashboard(request: Request, db: Session = Depends(get_db)):
    """Supervayzer dashboard"""
    today = datetime.now().date()
    
    # Agentlar statistikasi
    total_agents = db.query(Agent).filter(Agent.is_active == True).count()
    active_agents = 0
    for agent in db.query(Agent).filter(Agent.is_active == True).all():
        last_loc = db.query(AgentLocation).filter(
            AgentLocation.agent_id == agent.id,
            AgentLocation.recorded_at >= today
        ).first()
        if last_loc:
            active_agents += 1
    
    # Bugungi tashriflar
    today_visits = db.query(Visit).filter(Visit.visit_date >= today).count()
    
    # Bugungi buyurtmalar
    today_orders = db.query(Order).filter(
        Order.type == "sale",
        Order.date >= today
    ).all()
    today_sales_sum = sum(o.total for o in today_orders)
    
    # Yetkazib berish statistikasi
    total_drivers = db.query(Driver).filter(Driver.is_active == True).count()
    pending_deliveries = db.query(Delivery).filter(Delivery.status == "pending").count()
    today_delivered = db.query(Delivery).filter(
        Delivery.status == "delivered",
        Delivery.delivered_at >= today
    ).count()
    
    # Agent reytingi (bugungi savdo bo'yicha)
    agent_stats = []
    for agent in db.query(Agent).filter(Agent.is_active == True).all():
        visits = db.query(Visit).filter(
            Visit.agent_id == agent.id,
            Visit.visit_date >= today
        ).count()
        
        # So'nggi lokatsiya
        last_loc = db.query(AgentLocation).filter(
            AgentLocation.agent_id == agent.id
        ).order_by(AgentLocation.recorded_at.desc()).first()
        
        agent_stats.append({
            "agent": agent,
            "visits": visits,
            "last_seen": last_loc.recorded_at if last_loc else None,
            "is_online": last_loc and (datetime.now() - last_loc.recorded_at).seconds < 600 if last_loc else False
        })
    
    # Eng faol agentlar
    agent_stats.sort(key=lambda x: x["visits"], reverse=True)
    
    stats = {
        "total_agents": total_agents,
        "active_agents": active_agents,
        "today_visits": today_visits,
        "today_orders": len(today_orders),
        "today_sales": today_sales_sum,
        "total_drivers": total_drivers,
        "pending_deliveries": pending_deliveries,
        "today_delivered": today_delivered,
    }
    
    # Barcha agentlar va haydovchilar
    agents = db.query(Agent).filter(Agent.is_active == True).all()
    drivers = db.query(Driver).filter(Driver.is_active == True).all()
    
    # Oxirgi tashriflar va yetkazishlar
    recent_visits = db.query(Visit).order_by(Visit.visit_date.desc()).limit(10).all()
    recent_deliveries = db.query(Delivery).order_by(Delivery.created_at.desc()).limit(10).all()
    
    return templates.TemplateResponse("supervisor/dashboard.html", {
        "request": request,
        "stats": stats,
        "agents": agents,
        "drivers": drivers,
        "agent_stats": agent_stats[:10],
        "recent_visits": recent_visits,
        "recent_deliveries": recent_deliveries,
        "page_title": "Supervayzer",
        "now": datetime.now()
    })


# ==========================================
# POST Routes - Agent va Driver qo'shish
# ==========================================

@app.post("/agents/add")
async def add_agent(
    request: Request,
    full_name: str = Form(...),
    phone: str = Form(None),
    region: str = Form(None),
    telegram_id: str = Form(None),
    db: Session = Depends(get_db)
):
    """Yangi agent qo'shish"""
    # Kod generatsiya
    last_agent = db.query(Agent).order_by(Agent.id.desc()).first()
    code = f"AG{str((last_agent.id if last_agent else 0) + 1).zfill(3)}"
    
    agent = Agent(
        code=code,
        full_name=full_name,
        phone=phone,
        region=region,
        telegram_id=telegram_id,
        is_active=True
    )
    db.add(agent)
    db.commit()
    return RedirectResponse(url="/agents", status_code=303)


@app.post("/delivery/add-driver")
async def add_driver(
    request: Request,
    full_name: str = Form(...),
    phone: str = Form(None),
    vehicle_type: str = Form(None),
    vehicle_number: str = Form(None),
    telegram_id: str = Form(None),
    db: Session = Depends(get_db)
):
    """Yangi haydovchi qo'shish"""
    # Kod generatsiya
    last_driver = db.query(Driver).order_by(Driver.id.desc()).first()
    code = f"DR{str((last_driver.id if last_driver else 0) + 1).zfill(3)}"
    
    driver = Driver(
        code=code,
        full_name=full_name,
        phone=phone,
        vehicle_type=vehicle_type,
        vehicle_number=vehicle_number,
        telegram_id=telegram_id,
        is_active=True
    )
    db.add(driver)
    db.commit()
    return RedirectResponse(url="/delivery", status_code=303)


@app.post("/delivery/add-order")
async def add_delivery_order(
    request: Request,
    driver_id: int = Form(...),
    order_number: str = Form(...),
    delivery_address: str = Form(...),
    notes: str = Form(None),
    db: Session = Depends(get_db)
):
    """Yangi yetkazish buyurtmasi qo'shish"""
    delivery = Delivery(
        driver_id=driver_id,
        order_number=order_number,
        delivery_address=delivery_address,
        notes=notes,
        status="pending"
    )
    db.add(delivery)
    db.commit()
    return RedirectResponse(url=f"/delivery/{driver_id}", status_code=303)


# ==========================================
# GPS API (Mobil ilova uchun) - MOVED TO PWA API SECTION
# ==========================================

# OLD API REMOVED - See PWA API section below for new implementation




@app.post("/api/driver/location")
async def update_driver_location(
    driver_code: str = Form(...),
    latitude: float = Form(...),
    longitude: float = Form(...),
    speed: float = Form(0),
    db: Session = Depends(get_db)
):
    """Haydovchi lokatsiyasini yangilash"""
    driver = db.query(Driver).filter(Driver.code == driver_code).first()
    if not driver:
        raise HTTPException(status_code=404, detail="Haydovchi topilmadi")
    
    location = DriverLocation(
        driver_id=driver.id,
        latitude=latitude,
        longitude=longitude,
        speed=speed
    )
    db.add(location)
    db.commit()
    return {"status": "ok", "message": "Lokatsiya saqlandi"}


@app.get("/api/agents/locations")
async def get_agents_locations(db: Session = Depends(get_db)):
    """Barcha agentlarning oxirgi joylashuvi"""
    agents = db.query(Agent).filter(Agent.is_active == True).all()
    result = []
    for agent in agents:
        last_loc = db.query(AgentLocation).filter(
            AgentLocation.agent_id == agent.id
        ).order_by(AgentLocation.recorded_at.desc()).first()
        if last_loc:
            result.append({
                "id": agent.id,
                "name": agent.full_name,
                "code": agent.code,
                "lat": last_loc.latitude,
                "lng": last_loc.longitude,
                "time": last_loc.recorded_at.isoformat(),
                "battery": last_loc.battery
            })
    return result


@app.get("/api/drivers/locations")
async def get_drivers_locations(db: Session = Depends(get_db)):
    """Barcha haydovchilarning oxirgi joylashuvi"""
    drivers = db.query(Driver).filter(Driver.is_active == True).all()
    result = []
    for driver in drivers:
        last_loc = db.query(DriverLocation).filter(
            DriverLocation.driver_id == driver.id
        ).order_by(DriverLocation.recorded_at.desc()).first()
        if last_loc:
            result.append({
                "id": driver.id,
                "name": driver.full_name,
                "code": driver.code,
                "vehicle": driver.vehicle_number,
                "lat": last_loc.latitude,
                "lng": last_loc.longitude,
                "time": last_loc.recorded_at.isoformat(),
                "speed": last_loc.speed
            })
    return result



# ==========================================
# PWA API ENDPOINTS
# ==========================================

@app.post("/api/agent/login")
async def agent_login(
    username: str = Form(...),
    password: str = Form(...),
    db: Session = Depends(get_db)
):
    """Agent login API"""
    try:
        agent = db.query(Agent).filter(Agent.phone == username).first()
        
        if not agent or not agent.is_active:
            return {"success": False, "error": "Agent topilmadi yoki faol emas"}
        
        # Oddiy parol tekshiruvi (hozircha telefon = parol)
        if password != agent.phone:
            return {"success": False, "error": "Parol noto'g'ri"}
        
        # Session token yaratish
        token = create_session_token(agent.id, "agent")
        return {
            "success": True,
            "agent": {
                "id": agent.id,
                "code": agent.code,
                "full_name": agent.full_name,
                "phone": agent.phone,
            },
            "token": token
        }
    except Exception as e:
        return {"success": False, "error": str(e)}


@app.post("/api/driver/login")
async def driver_login(
    username: str = Form(...),
    password: str = Form(...),
    db: Session = Depends(get_db)
):
    """Driver login API"""
    try:
        driver = db.query(Driver).filter(Driver.phone == username).first()
        
        if not driver or not driver.is_active:
            return {"success": False, "error": "Haydovchi topilmadi yoki faol emas"}
        
        # Oddiy parol tekshiruvi
        if password != driver.phone:
            return {"success": False, "error": "Parol noto'g'ri"}
        
        token = create_session_token(driver.id, "driver")
        return {
            "success": True,
            "driver": {
                "id": driver.id,
                "code": driver.code,
                "full_name": driver.full_name,
                "phone": driver.phone,
                "vehicle_number": driver.vehicle_number,
            },
            "token": token
        }
    except Exception as e:
        return {"success": False, "error": str(e)}


@app.post("/api/agent/location_OLD_DISABLED")
async def agent_location_update_OLD(
    latitude: float = Form(...),
    longitude: float = Form(...),
    accuracy: float = Form(None),
    battery: int = Form(None),
    token: str = Form(...),
    db: Session = Depends(get_db)
):
    """Agent location update"""
    try:
        user_data = get_user_from_token(token)
        if not user_data or user_data.get("role") != "agent":
            return {"success": False, "error": "Invalid token"}
        
        agent_id = user_data["user_id"]
        
        location = AgentLocation(
            agent_id=agent_id,
            latitude=latitude,
            longitude=longitude,
            accuracy=accuracy,
            battery=battery,
        )
        db.add(location)
        db.commit()
        
        return {"success": True, "location_id": location.id}
    except Exception as e:
        db.rollback()
        return {"success": False, "error": str(e)}


@app.post("/api/driver/location")
async def driver_location_update(
    latitude: float = Form(...),
    longitude: float = Form(...),
    accuracy: float = Form(None),
    battery: int = Form(None),
    token: str = Form(...),
    db: Session = Depends(get_db)
):
    """Driver location update"""
    try:
        user_data = get_user_from_token(token)
        if not user_data or user_data.get("role") != "driver":
            return {"success": False, "error": "Invalid token"}
        
        driver_id = user_data["user_id"]
        
        location = DriverLocation(
            driver_id=driver_id,
            latitude=latitude,
            longitude=longitude,
            accuracy=accuracy,
            battery=battery,
        )
        db.add(location)
        db.commit()
        
        return {"success": True, "location_id": location.id}
    except Exception as e:
        db.rollback()
        return {"success": False, "error": str(e)}


@app.get("/api/agent/orders")
async def agent_orders(token: str, db: Session = Depends(get_db)):
    """Agent orders list"""
    try:
        user_data = get_user_from_token(token)
        if not user_data:
            return {"success": False, "error": "Invalid token"}
        
        # Hozircha bo'sh ro'yxat qaytaramiz
        return {"success": True, "orders": []}
    except Exception as e:
        return {"success": False, "error": str(e)}


@app.get("/api/agent/partners")
async def agent_partners(token: str, db: Session = Depends(get_db)):
    """Agent partners list"""
    try:
        user_data = get_user_from_token(token)
        if not user_data:
            return {"success": False, "error": "Invalid token"}
        
        partners = db.query(Partner).filter(Partner.is_active == True).all()
        return {
            "success": True,
            "partners": [
                {
                    "id": p.id,
                    "name": p.name,
                    "phone": p.phone,
                    "address": p.address,
                }
                for p in partners
            ]
        }
    except Exception as e:
        return {"success": False, "error": str(e)}


# ==========================================
# STARTUP
# ==========================================

# ==========================================
# PWA API ENDPOINTS
# ==========================================

@app.post("/api/agent/location")
async def agent_location_update(
    latitude: float = Form(...),
    longitude: float = Form(...),
    accuracy: float = Form(None),
    battery: int = Form(None),
    token: str = Form(...),
    db: Session = Depends(get_db)
):
    """Agent location update"""
    try:
        # Test mode - agent_id = 1
        agent_id = 1
        
        location = AgentLocation(
            agent_id=agent_id,
            latitude=latitude,
            longitude=longitude,
            accuracy=accuracy,
            battery=battery,
        )
        db.add(location)
        db.commit()
        
        return {"success": True, "location_id": location.id}
    except Exception as e:
        db.rollback()
        return {"success": False, "error": str(e)}



# ==========================================
# HUDUDLAR (REGIONS)
# ==========================================

@app.get("/test/regions", response_class=HTMLResponse)
async def regions_test_page(request: Request, db: Session = Depends(get_db), current_user: User = Depends(require_admin)):
    """Hududlar test sahifasi, faqat admin"""
    regions = db.query(Region).all()
    # Fake user for testing
    fake_user = {"username": "test", "role": "admin"}
    return templates.TemplateResponse("info/regions.html", {
        "request": request,
        "page_title": "Hududlar",
        "user": fake_user,
        "regions": regions
    })


@app.get("/info/regions", response_class=HTMLResponse)
async def info_regions(request: Request, db: Session = Depends(get_db), current_user: User = Depends(require_auth)):
    """Hududlar sahifasi"""
    if not current_user:
        return RedirectResponse(url="/login", status_code=303)
    from urllib.parse import unquote
    regions = db.query(Region).filter(Region.is_active == True).all()
    import_ok = request.query_params.get("import_ok")
    import_added = request.query_params.get("added")
    import_updated = request.query_params.get("updated")
    import_error = request.query_params.get("error") == "import"
    import_detail = unquote(request.query_params.get("detail", "") or "")
    return templates.TemplateResponse("info/regions.html", {
        "request": request,
        "current_user": current_user,
        "page_title": "Hududlar",
        "regions": regions,
        "import_ok": import_ok,
        "import_added": import_added,
        "import_updated": import_updated,
        "import_error": import_error,
        "import_detail": import_detail,
    })


@app.post("/info/regions/add")
async def region_add(
    code: str = Form(...),
    name: str = Form(...),
    description: str = Form(""),
    db: Session = Depends(get_db)
):
    """Hudud qo'shish"""
    existing = db.query(Region).filter(Region.code == code).first()
    if existing:
        raise HTTPException(status_code=400, detail=f"'{code}' kodli hudud allaqachon mavjud!")
    
    region = Region(code=code, name=name, description=description)
    db.add(region)
    db.commit()
    return RedirectResponse(url="/info/regions", status_code=303)


@app.post("/info/regions/edit/{region_id}")
async def region_edit(
    region_id: int,
    code: str = Form(...),
    name: str = Form(...),
    description: str = Form(""),
    db: Session = Depends(get_db)
):
    """Hududni tahrirlash"""
    region = db.query(Region).filter(Region.id == region_id).first()
    if not region:
        raise HTTPException(status_code=404, detail="Hudud topilmadi")
    
    existing = db.query(Region).filter(
        Region.code == code,
        Region.id != region_id
    ).first()
    if existing:
        raise HTTPException(status_code=400, detail=f"'{code}' kodli hudud allaqachon mavjud!")
    
    region.code = code
    region.name = name
    region.description = description
    db.commit()
    return RedirectResponse(url="/info/regions", status_code=303)


@app.post("/info/regions/delete/{region_id}")
async def region_delete(region_id: int, db: Session = Depends(get_db)):
    """Hududni o'chirish"""
    region = db.query(Region).filter(Region.id == region_id).first()
    if not region:
        raise HTTPException(status_code=404, detail="Hudud topilmadi")
    
    db.delete(region)
    db.commit()
    return RedirectResponse(url="/info/regions", status_code=303)


@app.get("/info/regions/export")
async def export_regions(db: Session = Depends(get_db), current_user: User = Depends(require_auth)):
    """Hududlarni Excel ga eksport."""
    regions = db.query(Region).filter(Region.is_active == True).all()
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Hududlar"
    ws.append(["ID", "Kod", "Nomi", "Tavsif"])
    for r in regions:
        ws.append([r.id, r.code, r.name, r.description or ""])
    stream = io.BytesIO()
    wb.save(stream)
    stream.seek(0)
    return StreamingResponse(
        stream,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": "attachment; filename=hududlar.xlsx"},
    )


@app.get("/info/regions/template")
async def template_regions(current_user: User = Depends(require_auth)):
    """Hududlar import andozasi."""
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Hududlar"
    ws.append(["Kod", "Nomi", "Tavsif"])
    ws.append(["QOQON", "Qo'qon", "Markaziy hudud"])
    ws.append(["RISHTON", "Rishton", "Janubiy hudud"])
    stream = io.BytesIO()
    wb.save(stream)
    stream.seek(0)
    return StreamingResponse(
        stream,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": "attachment; filename=hudud_andoza.xlsx"},
    )


@app.get("/info/regions/import")
async def regions_import_get(current_user: User = Depends(require_auth)):
    """Import sahifasi to'g'ridan-to'g'ri ochilsa hududlar ro'yxatiga yo'naltirish."""
    return RedirectResponse(url="/info/regions", status_code=303)


@app.post("/info/regions/import")
async def import_regions(
    request: Request,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    """Excel dan hududlarni import. Ustunlar: Kod, Nomi, Tavsif."""
    from urllib.parse import quote
    from zipfile import BadZipFile
    form = await request.form()
    file = form.get("file") or form.get("excel_file")
    if not file or not getattr(file, "filename", None):
        return RedirectResponse(url="/info/regions?error=import&detail=" + quote("Excel fayl tanlang"), status_code=303)
    try:
        contents = await file.read()
        if not contents:
            return RedirectResponse(url="/info/regions?error=import&detail=" + quote("Fayl bo'sh"), status_code=303)
        if contents[:2] != b"PK":
            return RedirectResponse(
                url="/info/regions?error=import&detail=" + quote("Fayl .xlsx formati bo'lishi kerak."),
                status_code=303,
            )
        wb = openpyxl.load_workbook(io.BytesIO(contents), read_only=False, data_only=True)
        ws = wb.active
        if ws.max_row < 2:
            return RedirectResponse(
                url="/info/regions?error=import&detail=" + quote("Excelda ma'lumot qatorlari yo'q."),
                status_code=303,
            )
        added = 0
        updated = 0
        for row_num in range(2, ws.max_row + 1):
            def cell(col):
                v = ws.cell(row=row_num, column=col).value
                return "" if v is None else str(v).strip()
            code = cell(1) or cell(2)
            name = cell(2) or cell(1)
            desc = cell(3) or ""
            if not code and not name:
                continue
            if (code or "").lower() in ("id", "kod", "nomi", "tavsif") and (not name or (name or "").lower() in ("id", "kod", "nomi")):
                continue
            if not code:
                code = f"R{row_num}"
            if not name:
                name = code
            region = db.query(Region).filter(Region.code == code).first()
            if not region:
                region = Region(code=code, name=name, description=desc or None)
                db.add(region)
                added += 1
            else:
                region.name = name
                region.description = desc or None
                updated += 1
            db.commit()
        return RedirectResponse(
            url="/info/regions?import_ok=1&added=" + str(added) + "&updated=" + str(updated),
            status_code=303,
        )
    except BadZipFile:
        return RedirectResponse(
            url="/info/regions?error=import&detail=" + quote("Fayl .xlsx formati bo'lishi kerak."),
            status_code=303,
        )
    except Exception as e:
        return RedirectResponse(
            url="/info/regions?error=import&detail=" + quote(str(e)[:150]),
            status_code=303,
        )


# ==========================================
# USKUNALAR (MACHINES)
# ==========================================

@app.get("/info/machines", response_class=HTMLResponse)
async def info_machines(request: Request, db: Session = Depends(get_db), current_user: User = Depends(require_auth)):
    """Uskunalar ro'yxati"""
    if not current_user:
        return RedirectResponse(url="/login", status_code=303)
    machines = db.query(Machine).filter(Machine.is_active == True).order_by(Machine.created_at.desc()).all()
    employees = db.query(Employee).filter(Employee.is_active == True).all()
    warehouses = db.query(Warehouse).all()
    return templates.TemplateResponse("info/machines.html", {
        "request": request,
        "current_user": current_user,
        "page_title": "Uskunalar",
        "machines": machines,
        "employees": employees,
        "warehouses": warehouses,
    })


@app.post("/info/machines/add")
async def machine_add(
    code: str = Form(...),
    name: str = Form(...),
    machine_type: str = Form(""),
    capacity: float = Form(0),
    efficiency: float = Form(100.0),
    status: str = Form("idle"),
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    """Uskuna qo'shish"""
    if not current_user:
        return RedirectResponse(url="/login", status_code=303)
    existing = db.query(Machine).filter(Machine.code == code).first()
    if existing:
        raise HTTPException(status_code=400, detail=f"'{code}' kodli uskuna allaqachon mavjud!")
    machine = Machine(
        code=code.strip(),
        name=name.strip(),
        machine_type=machine_type.strip() or "boshqa",
        capacity=float(capacity),
        efficiency=float(efficiency),
        status=status,
    )
    db.add(machine)
    db.commit()
    return RedirectResponse(url="/info/machines", status_code=303)


@app.post("/info/machines/edit/{machine_id}")
async def machine_edit(
    machine_id: int,
    code: str = Form(...),
    name: str = Form(...),
    machine_type: str = Form(""),
    capacity: float = Form(0),
    efficiency: float = Form(100.0),
    status: str = Form("idle"),
    operator_id: Optional[int] = Form(None),
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    """Uskunani tahrirlash"""
    if not current_user:
        return RedirectResponse(url="/login", status_code=303)
    machine = db.query(Machine).filter(Machine.id == machine_id).first()
    if not machine:
        raise HTTPException(status_code=404, detail="Uskuna topilmadi")
    existing = db.query(Machine).filter(Machine.code == code, Machine.id != machine_id).first()
    if existing:
        raise HTTPException(status_code=400, detail=f"'{code}' kodli uskuna allaqachon mavjud!")
    machine.code = code.strip()
    machine.name = name.strip()
    machine.machine_type = machine_type.strip() or "boshqa"
    machine.capacity = float(capacity)
    machine.efficiency = float(efficiency)
    machine.status = status
    machine.operator_id = int(operator_id) if operator_id else None
    db.commit()
    return RedirectResponse(url="/info/machines", status_code=303)


@app.post("/info/machines/delete/{machine_id}")
async def machine_delete(machine_id: int, db: Session = Depends(get_db), current_user: User = Depends(require_auth)):
    """Uskunani o'chirish (soft: is_active=False)"""
    if not current_user:
        return RedirectResponse(url="/login", status_code=303)
    machine = db.query(Machine).filter(Machine.id == machine_id).first()
    if not machine:
        raise HTTPException(status_code=404, detail="Uskuna topilmadi")
    machine.is_active = False
    db.commit()
    return RedirectResponse(url="/info/machines", status_code=303)


@app.on_event("startup")
async def startup():
    """Dastur ishga tushganda"""
    init_db()
    try:
        from app.models.database import ensure_attendance_advance_tables
        ensure_attendance_advance_tables()
    except Exception as e:
        print("[Startup] ensure_attendance_advance_tables:", e)
    try:
        from app.utils.scheduler import start_scheduler
        start_scheduler()
    except Exception as e:
        print("[Startup] Scheduler ishga tushmadi:", e)
    print("TOTLI HOLVA Business System ishga tushdi!")
    _mp = os.path.abspath(__file__)
    print("  main.py:", _mp)
    try:
        for _dir in [os.path.dirname(_mp), os.getcwd(), r"C:\Users\ELYOR\.cursor\worktrees\business_system\pwp"]:
            if _dir:
                with open(os.path.join(_dir, "server_started.txt"), "w", encoding="utf-8") as f:
                    f.write("main.py: %s\ncwd: %s\n" % (_mp, os.getcwd()))
                break
    except Exception:
        pass


if __name__ == "__main__":
    uvicorn.run("main:app", host="0.0.0.0", port=8080, reload=False)

